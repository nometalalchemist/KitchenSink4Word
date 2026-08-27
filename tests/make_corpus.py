"""Generate a synthetic test corpus, structurally equivalent to the private
real-world corpus the suite was developed against.

Any file already present in tests/corpus/ is left alone, so a local real
corpus always wins; CI and contributors get generated stand-ins with the same
structural properties (footnote counts, revision counts, threading, TOC with
cached entries, merged tables, images) but zero real content.

Run directly:  python -X utf8 tests/make_corpus.py
Or automatically via conftest when corpus files are missing.
"""

from __future__ import annotations

import random
import struct
import sys
import zlib
from pathlib import Path

CORPUS = Path(__file__).parent / "corpus"

_WORDS = (
    "the framework extends across cases while compliance rises and the "
    "institution adapts its mandate toward renewed cooperation under "
    "conditions of uncertainty where actors weigh obligations against "
    "expected consequences and prior commitments shape present choices"
).split()


def _sentence(rng: random.Random, n: int = 14) -> str:
    ws = [rng.choice(_WORDS) for _ in range(n)]
    return (" ".join(ws)).capitalize() + "."


def _paragraph(rng: random.Random, sentences: int = 4) -> str:
    return " ".join(_sentence(rng, rng.randint(9, 18)) for _ in range(sentences))


def _png(path: Path, w: int = 60, h: int = 30, rgb=(120, 60, 40)) -> Path:
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(
            ">I", zlib.crc32(c) & 0xFFFFFFFF
        )

    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    return path


def _base_doc(path: Path, rng: random.Random, paragraphs: int, headings: int = 0):
    from docx import Document

    d = Document()
    made = 0
    h = 0
    while made < paragraphs:
        if headings and h < headings and made % max(paragraphs // headings, 1) == 0:
            level = 1 if h % 3 == 0 else 2
            d.add_heading(f"Section {h + 1}: {_sentence(rng, 4)[:-1]}", level)
            h += 1
        d.add_paragraph(_paragraph(rng, rng.randint(2, 5)))
        made += 1
    d.save(str(path))
    return path


def _pkg(path):
    sys.path.insert(0, str(Path(__file__).parents[1] / "src"))
    from word_mcp.core.package import DocxPackage

    return DocxPackage(path)


def gen_ch4(path: Path):
    """~173 paragraphs, 2 tables, APA-style citations + reference list
    (>=40 parseable refs, >=40 in-text citations), 'alliance'/'Delta Model'
    vocabulary, lastRenderedPageBreak artifacts."""
    rng = random.Random(41)
    from docx import Document

    d = Document()
    surnames = [
        "Hurd", "Kelman", "Bordin", "Walt", "Snyder", "Jervis", "Lake",
        "Wendt", "Keohane", "Nye", "Mearsheimer", "Waltz", "Axelrod",
        "Fearon", "Schelling", "Putnam", "Krasner", "Ruggie", "Katzen",
        "Finnemore", "Sikkink", "Checkel", "Adler", "Barnett", "Hopf",
        "Tannenwald", "Price", "Risse", "Moravcsik", "Milner", "Gourevitch",
        "Gilpin", "Cox", "Ashley", "Campbell", "Weldes", "Doty", "Hansen",
        "Buzan", "Waever", "Booth", "Wyn", "Herz", "Deutsch", "Haas",
        "Mitrany", "Rosenau", "Singer", "Small", "Organski",
    ]
    years = [str(1959 + i) for i in range(len(surnames))]
    d.add_paragraph(
        "The Delta Model frames the alliance as a mediated structure, and "
        "the Mutual Defense Treaty anchors the argument throughout."
    )
    for i in range(150):
        text = _paragraph(rng, 3)
        if i % 3 == 0:
            k = rng.randrange(len(surnames))
            text += f" This point follows prior work ({surnames[k]}, {years[k]})."
        if i % 2 == 0:
            text += " The alliance holds where the alliance is mediated."
        d.add_paragraph(text)
    d.add_heading("References", 1)
    for k in range(len(surnames)):
        d.add_paragraph(
            f"{surnames[k]}, A. ({years[k]}). {_sentence(rng, 6)[:-1]}. "
            f"Journal of Synthetic Studies, {k + 1}(2), {k + 10}-{k + 30}."
        )
    d.save(str(path))

    pkg = _pkg(path)
    from word_mcp.core.package import qn
    from word_mcp.ops import tables as tb

    tb.create_table(pkg, [["Element", "Count"], ["Records", "210"]], at_end=True)
    tb.create_table(
        pkg, [["Case", "Code", "Note"], ["One", "G", "x"], ["Two", "T", "y"]],
        at_end=True,
    )
    # lastRenderedPageBreak artifacts inside text runs (the v1.2 HIGH trigger).
    from lxml import etree

    count = 0
    for p in pkg.root().iter(qn("w:p")):
        r = p.find(qn("w:r"))
        if r is not None and count < 48 and count % 1 == 0:
            t = r.find(qn("w:t"))
            if t is not None:
                lrpb = etree.Element(qn("w:lastRenderedPageBreak"))
                t.addprevious(lrpb)
                count += 1
    pkg.mark_dirty()
    pkg.save(do_backup=False)


def gen_ch5(path: Path):
    rng = random.Random(51)
    _base_doc(path, rng, paragraphs=91)


def gen_ch13(path: Path):
    """677+ paragraphs, 47+ headings, >5000 words, multiple sections, a real
    Word-style TOC (sdt, updated with cached TOC1/TOC2 entries), page breaks."""
    rng = random.Random(13)
    from docx import Document
    from docx.enum.text import WD_BREAK

    d = Document()
    d.add_paragraph("Front matter title page.")
    p = d.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.add_run("Dedication")
    for h in range(48):
        d.add_heading(f"Heading {h + 1}: {_sentence(rng, 3)[:-1]}", 1 if h % 4 == 0 else 2)
        for _ in range(12):
            d.add_paragraph(_paragraph(rng, rng.randint(2, 4)))
    d.save(str(path))

    pkg = _pkg(path)
    from lxml import etree

    from word_mcp.core.package import qn
    from word_mcp.ops import furniture as fu, toc

    fu.add_section_break(pkg, after_index=30, break_type="nextPage")
    toc.insert_toc(pkg, at_start=True, update_on_open=False)
    # Fabricate cached entries inside the TOC field (as if Word updated it).
    field_p = toc._find_toc_fields(pkg)[0]
    sep_run = None
    for r in field_p.iter(qn("w:r")):
        fc = r.find(qn("w:fldChar"))
        if fc is not None and fc.get(qn("w:fldCharType")) == "separate":
            sep_run = r
    anchor = sep_run
    for i in range(3):
        entry = etree.Element(qn("w:p"))
        ppr = etree.SubElement(entry, qn("w:pPr"))
        etree.SubElement(ppr, qn("w:pStyle")).set(qn("w:val"), f"TOC{1 + i % 2}")
        r = etree.SubElement(entry, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = f"Heading {i + 1}\t{3 + i}"
        field_p.addnext(entry) if i == 0 else anchor.addnext(entry)
        anchor = entry
    pkg.mark_dirty()
    pkg.save(do_backup=False)


def gen_codebook(path: Path):
    rng = random.Random(77)
    _base_doc(path, rng, paragraphs=170)
    pkg = _pkg(path)
    from word_mcp.ops import tables as tb

    for t in range(9):
        rows = [[f"H{c}" for c in range(3 + t % 3)]]
        for r in range(2 + (15 if t == 4 else t % 5)):
            rows.append([_sentence(rng, 3)[:-1] for _ in rows[0]])
        tb.create_table(pkg, rows, at_end=True)
    pkg.save(do_backup=False)


def gen_unitar(path: Path, tmp: Path):
    rng = random.Random(88)
    _base_doc(path, rng, paragraphs=100)
    pkg = _pkg(path)
    from word_mcp.ops import media

    png1 = _png(tmp / "syn1.png", rgb=(60, 90, 140))
    png2 = _png(tmp / "syn2.png", rgb=(140, 90, 60))
    media.add_image(pkg, str(png1), at_end=True, width_pt=120)
    media.add_image(pkg, str(png2), at_end=True, width_pt=120)
    pkg.save(do_backup=False)


def gen_niu(path: Path):
    """171 real footnotes referenced in body order."""
    rng = random.Random(171)
    _base_doc(path, rng, paragraphs=350)
    pkg = _pkg(path)
    from word_mcp.core.package import qn
    from word_mcp.ops import notes

    # Unique anchors: stamp a marker word into 171 paragraphs, then note them.
    from lxml import etree

    paras = [p for p in pkg.body().findall(qn("w:p"))]
    step = max(len(paras) // 171, 1)
    added = 0
    for i in range(0, len(paras), step):
        if added >= 171:
            break
        marker = f"anchorpoint{added:03d}"
        r = etree.SubElement(paras[i], qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = f" {marker}"
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        added += 1
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    pkg = _pkg(path)
    for k in range(171):
        notes.add_note(
            pkg,
            "footnote",
            anchor_text=f"anchorpoint{k:03d}",
            note_text=f"Synthetic footnote {k + 1}: {_sentence(rng, 8)}",
        )
    pkg.save(do_backup=False)

    # A few reviewer comments like the original.
    pkg = _pkg(path)
    from word_mcp.ops import comments as cm

    for k in range(3):
        cm.add_comment(
            pkg,
            anchor_text=f"anchorpoint{k * 40:03d}",
            text=f"Reviewer note {k + 1}.",
            author="Synthetic Reviewer",
        )
    pkg.save(do_backup=False)


def gen_ejir(path: Path):
    """~126 tracked revisions (insertions AND long deletions) by one author,
    plus 12 comments."""
    rng = random.Random(126)
    _base_doc(path, rng, paragraphs=140)
    pkg = _pkg(path)
    from word_mcp.ops import comments as cm, read, text as tx

    # 63 tracked replacements = 126 revision elements (del+ins pairs).
    paras = [p for p in read.get_paragraphs(pkg) if len(p["text"]) > 80]
    done = 0
    for p in paras:
        if done >= 63:
            break
        words = p["text"].split()
        find = " ".join(words[2:5])
        try:
            r = tx.search_and_replace(
                pkg,
                [{"find": find, "replace": _sentence(rng, 4)[:-1].lower()}],
                track=True,
                author="Synthetic Reviewer One",
            )
            done += r["total"]
        except Exception:
            continue
    pkg.save(do_backup=False)

    pkg = _pkg(path)
    paras = [p for p in read.get_paragraphs(pkg) if len(p["text"]) > 60]
    for k in range(12):
        words = paras[k * 3 % len(paras)]["text"].split()
        try:
            cm.add_comment(
                pkg,
                anchor_text=" ".join(words[1:4]),
                text=f"Comment {k + 1}: reconsider.",
                author="Synthetic Reviewer One",
            )
        except Exception:
            continue
    pkg.save(do_backup=False)


def gen_outline(path: Path):
    """33 comments with 7+ threaded replies."""
    rng = random.Random(33)
    _base_doc(path, rng, paragraphs=200, headings=10)
    pkg = _pkg(path)
    from word_mcp.ops import comments as cm, read

    paras = [p for p in read.get_paragraphs(pkg) if len(p["text"]) > 60]
    roots = []
    made = 0
    for k in range(26):
        words = paras[(k * 5) % len(paras)]["text"].split()
        try:
            r = cm.add_comment(
                pkg,
                anchor_text=" ".join(words[2:5]),
                text=f"Outline note {k + 1}.",
                author="Synthetic Author",
            )
            roots.append(r["comment_id"])
            made += 1
        except Exception:
            continue
    for k in range(7):
        cm.reply_to_comment(
            pkg,
            comment_id=roots[k],
            text=f"Reply {k + 1}.",
            author="Synthetic Author",
        )
    pkg.save(do_backup=False)


def gen_ch4_chair(path: Path):
    rng = random.Random(499)
    _base_doc(path, rng, paragraphs=200, headings=6)
    pkg = _pkg(path)
    from word_mcp.ops import tables as tb

    for _ in range(3):
        tb.create_table(pkg, [["a", "b"], ["c", "d"]], at_end=True)
    pkg.save(do_backup=False)


GENERATORS = {
    "ch4.docx": lambda p, tmp: gen_ch4(p),
    "ch5.docx": lambda p, tmp: gen_ch5(p),
    "ch1-3.docx": lambda p, tmp: gen_ch13(p),
    "codebook.docx": lambda p, tmp: gen_codebook(p),
    "unitar.docx": lambda p, tmp: gen_unitar(p, tmp),
    "ch4_chair.docx": lambda p, tmp: gen_ch4_chair(p),
    "niu.docx": lambda p, tmp: gen_niu(p),
    "ejir_rw.docx": lambda p, tmp: gen_ejir(p),
    "outline.docx": lambda p, tmp: gen_outline(p),
}


def generate_missing(verbose: bool = True) -> list[str]:
    import tempfile

    CORPUS.mkdir(exist_ok=True)
    tmp = Path(tempfile.mkdtemp())
    made = []
    for name, gen in GENERATORS.items():
        target = CORPUS / name
        if target.exists():
            continue
        if verbose:
            print(f"generating synthetic {name} ...")
        gen(target, tmp)
        made.append(name)
    return made


if __name__ == "__main__":
    made = generate_missing()
    print(f"generated {len(made)} synthetic corpus file(s): {made or 'none needed'}")
