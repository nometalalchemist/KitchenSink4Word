"""Review-cycle bundle: reviewer matrix (single + multi-file), revision
analytics, and the structured no-Word diff. Documents are built from scratch
with python-docx and then reviewed/revised through the server's own ops
(add_comment / reply_to_comment / resolve_comment, tracked search_and_replace,
insert_columns) so the analytics read exactly what the tools write."""

import shutil
from pathlib import Path

from docx import Document

from word_mcp.core.package import DocxPackage
from word_mcp.ops import comments as cm, reviewcycle as rc, tables as tb, text as tx


def _build_base(path: Path) -> None:
    doc = Document()
    doc.add_heading("Introduction", level=1)                              # 0
    doc.add_paragraph(
        "The alliance framework rests on legitimacy and authority "
        "working together."
    )                                                                     # 1
    doc.add_paragraph(
        "Carter announced the withdrawal plan in nineteen seventy seven."
    )                                                                     # 2
    doc.add_heading("Methods", level=1)                                   # 3
    doc.add_paragraph(
        "The coding protocol uses three categories drawn from Kelman."
    )                                                                     # 4
    doc.add_paragraph(
        "Each document was coded twice for reliability purposes overall."
    )                                                                     # 5
    doc.save(str(path))


# ------------------------------------------------------------ comment_report


def _commented_doc(tmp_path: Path) -> Path:
    p = tmp_path / "commented.docx"
    _build_base(p)
    pkg = DocxPackage(p)
    r1 = cm.add_comment(
        pkg,
        anchor_text="legitimacy and authority",
        text="Define these terms first.",
        author="Reviewer One",
        initials="R1",
    )
    cm.reply_to_comment(
        pkg,
        comment_id=r1["comment_id"],
        text="Definitions added in section 2.",
        author="Author",
    )
    r2 = cm.add_comment(
        pkg,
        anchor_text="coding protocol",
        text="Cite the codebook here.",
        author="Reviewer Two",
    )
    cm.resolve_comment(pkg, comment_id=r2["comment_id"])
    pkg.save(do_backup=False)
    return p


def test_comment_report_matrix(tmp_path):
    p = _commented_doc(tmp_path)
    rep = rc.comment_report(DocxPackage(p))

    assert rep["summary"]["threads"] == 2
    assert rep["summary"]["total_comments"] == 3
    assert rep["summary"]["open_threads"] == 1
    assert rep["summary"]["resolved_threads"] == 1
    assert rep["summary"]["by_author"] == {
        "Reviewer One": 1,
        "Author": 1,
        "Reviewer Two": 1,
    }
    assert rep["summary"]["by_section"] == {"Introduction": 1, "Methods": 1}

    t1 = next(c for c in rep["comments"] if c["author"] == "Reviewer One")
    assert t1["initials"] == "R1"
    assert t1["anchored_text"] == "legitimacy and authority"
    assert t1["text"] == "Define these terms first."
    assert t1["resolved"] is False
    assert t1["paragraph_index"] == 1
    assert t1["heading_path"] == ["Introduction"]
    assert t1["section"] == "Introduction"
    assert len(t1["replies"]) == 1
    reply = t1["replies"][0]
    assert reply["author"] == "Author"
    assert reply["text"] == "Definitions added in section 2."
    assert reply["replies"] == []

    t2 = next(c for c in rep["comments"] if c["author"] == "Reviewer Two")
    assert t2["resolved"] is True
    assert t2["paragraph_index"] == 4
    assert t2["heading_path"] == ["Methods"]
    assert t2["date"]  # w:date is written by add_comment


def test_comment_report_excludes_resolved(tmp_path):
    p = _commented_doc(tmp_path)
    rep = rc.comment_report(DocxPackage(p), include_resolved=False)
    assert rep["summary"]["threads"] == 1
    assert rep["resolved_threads_excluded"] == 1
    assert all(not c["resolved"] for c in rep["comments"])
    # Replies still ride along with their surviving thread.
    assert rep["summary"]["total_comments"] == 2


# ------------------------------------------------------ comment_report_multi


def test_comment_report_multi_collision(tmp_path):
    paths = []
    for name in ("copy_a.docx", "copy_b.docx", "copy_c.docx"):
        p = tmp_path / name
        _build_base(p)
        paths.append(p)

    def comment(path, anchor, text, author):
        pkg = DocxPackage(path)
        cm.add_comment(pkg, anchor_text=anchor, text=text, author=author)
        pkg.save(do_backup=False)

    comment(paths[0], "coding protocol", "Needs a citation.", "Reviewer One")
    comment(paths[1], "coding protocol", "Too vague as written.", "Reviewer Two")
    comment(paths[1], "withdrawal plan", "Give the exact date.", "Reviewer Two")
    comment(paths[2], "reliability purposes", "Report the kappa.", "Reviewer Three")

    rep = rc.comment_report_multi([str(p) for p in paths])

    assert rep["summary"]["files"] == 3
    assert rep["summary"]["total_threads"] == 4
    assert rep["summary"]["merged_entries"] == 4
    assert rep["summary"]["by_author"] == {
        "Reviewer One": 1,
        "Reviewer Two": 2,
        "Reviewer Three": 1,
    }
    assert [f["path"] for f in rep["files"]] == [str(p) for p in paths]
    assert rep["files"][1]["threads"] == 2

    # Every occurrence carries its file provenance.
    for entry in rep["merged"]:
        for occ in entry["occurrences"]:
            assert occ["file"] in {str(p) for p in paths}
            assert occ["comment_id"]

    # The deliberate same-span collision: two reviewers on "coding protocol".
    assert rep["summary"]["collision_count"] == 1
    col = rep["collisions"][0]
    assert col["anchored_text"] == "coding protocol"
    assert col["authors"] == ["Reviewer One", "Reviewer Two"]
    files_in_collision = {e["file"] for e in col["entries"]}
    assert files_in_collision == {str(paths[0]), str(paths[1])}


# --------------------------------------------------------- revision_analytics


def test_revision_analytics_two_authors(tmp_path):
    p = tmp_path / "revised.docx"
    _build_base(p)
    pkg = DocxPackage(p)
    tx.search_and_replace(
        pkg,
        [{"find": "withdrawal plan", "replace": "phased withdrawal schedule"}],
        track=True,
        author="Alice",
    )
    tx.search_and_replace(
        pkg,
        [{"find": "three categories", "replace": "four categories"}],
        track=True,
        author="Bob",
    )
    pkg.save(do_backup=False)

    ana = rc.revision_analytics(DocxPackage(p))
    assert ana["total_revisions"] == 4  # one ins + one del per author

    alice = ana["by_author"]["Alice"]
    assert alice["insertions"] == 1
    assert alice["deletions"] == 1
    assert alice["words_added"] == 3   # "phased withdrawal schedule"
    assert alice["words_removed"] == 2  # "withdrawal plan"
    assert alice["by_section"] == {"Introduction": 2}
    assert alice["date_range"] is not None
    assert alice["date_range"]["first"] <= alice["date_range"]["last"]

    bob = ana["by_author"]["Bob"]
    assert bob["words_added"] == 2
    assert bob["words_removed"] == 2
    assert bob["by_section"] == {"Methods": 2}

    heavy = ana["heaviest_paragraphs"]
    assert {h["paragraph_index"] for h in heavy} == {2, 4}
    # Alice's paragraph churns more characters than Bob's.
    assert heavy[0]["paragraph_index"] == 2
    assert heavy[0]["authors"] == ["Alice"]
    assert heavy[0]["section"] == "Introduction"
    assert heavy[0]["chars_changed"] > 0
    assert "Carter announced" in heavy[0]["text_preview"]


def test_revision_analytics_clean_document(tmp_path):
    p = tmp_path / "clean.docx"
    _build_base(p)
    ana = rc.revision_analytics(DocxPackage(p))
    assert ana["total_revisions"] == 0
    assert ana["by_author"] == {}
    assert ana["heaviest_paragraphs"] == []


# ------------------------------------------------------------ structured_diff

_OLD_PARAS = [
    "The first finding concerns alliance durability under external pressure.",
    "The second finding concerns domestic audiences and their expectations.",
    "The third finding concerns institutional memory across administrations.",
    "This paragraph will be deleted in the next draft entirely.",
    "This paragraph moves to a new location in the revised draft.",
]


def _table_2x2(doc):
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Case"
    table.cell(0, 1).text = "Outcome"
    table.cell(1, 0).text = "Korea"
    table.cell(1, 1).text = "Renewal"


def _build_old(path: Path) -> None:
    doc = Document()
    doc.add_heading("Results", level=1)      # 0
    for text in _OLD_PARAS:                  # 1..5
        doc.add_paragraph(text)
    _table_2x2(doc)
    doc.save(str(path))


def _build_new(path: Path) -> None:
    doc = Document()
    doc.add_heading("Results", level=1)      # 0
    doc.add_paragraph(
        "The first finding concerns alliance durability under sustained "
        "external pressure."
    )                                        # 1: modified
    doc.add_paragraph(_OLD_PARAS[4])         # 2: moved up from old index 5
    doc.add_paragraph(_OLD_PARAS[1])         # 3
    doc.add_paragraph(_OLD_PARAS[2])         # 4
    # old index 4 deliberately deleted
    _table_2x2(doc)
    doc.save(str(path))
    pkg = DocxPackage(path)
    tb.insert_columns(pkg, 0, at=2)          # 2x2 -> 2x3
    pkg.save(do_backup=False)


def test_structured_diff_full_scenario(tmp_path):
    old, new = tmp_path / "old.docx", tmp_path / "new.docx"
    _build_old(old)
    _build_new(new)

    diff = rc.structured_diff(str(old), str(new))

    assert diff["counts"]["unchanged"] == 3  # heading + two "finding" paras
    assert diff["counts"]["modified"] == 1
    assert diff["counts"]["moved"] == 1
    assert diff["counts"]["deleted"] == 1
    assert diff["counts"]["inserted"] == 0
    assert diff["counts"]["old_paragraphs"] == 6
    assert diff["counts"]["new_paragraphs"] == 5
    assert diff["identical"] is False
    assert diff["detail_capped"] is False

    mod = diff["modified"][0]
    assert (mod["old_index"], mod["new_index"]) == (1, 1)
    assert mod["section"] == "Results"
    assert any("sustained" in c["new"] for c in mod["changes"])

    mv = diff["moved"][0]
    assert (mv["old_index"], mv["new_index"]) == (5, 2)
    assert "moves to a new location" in mv["text"]

    dele = diff["deleted"][0]
    assert dele["old_index"] == 4
    assert "will be deleted" in dele["text"]
    assert dele["section"] == "Results"

    assert diff["by_section"]["Results"] == {
        "modified": 1,
        "inserted": 0,
        "deleted": 1,
        "moved": 1,
    }

    # Table: same table, one grid column inserted on the right edge.
    assert diff["tables"]["table_count"] == {"old": 1, "new": 1, "delta": 0}
    tchange = diff["tables"]["changed_tables"][0]
    assert tchange["dimensions_changed"] is True
    assert tchange["old_dims"]["columns"] == 2
    assert tchange["new_dims"]["columns"] == 3
    assert tchange["changed_cells"] == []  # surviving cells untouched

    assert diff["footnotes"]["delta"] == 0
    assert diff["headings"] == {"added": [], "removed": []}


def test_structured_diff_identical(tmp_path):
    old, copy = tmp_path / "old.docx", tmp_path / "copy.docx"
    _build_old(old)
    shutil.copy(old, copy)

    diff = rc.structured_diff(str(old), str(copy))
    assert diff["identical"] is True
    assert diff["counts"]["modified"] == 0
    assert diff["counts"]["inserted"] == 0
    assert diff["counts"]["deleted"] == 0
    assert diff["counts"]["moved"] == 0
    assert diff["counts"]["unchanged"] == 6
    assert diff["modified"] == diff["inserted"] == diff["deleted"] == []
    assert diff["tables"]["changed_tables"] == []
    assert diff["detail_capped"] is False


def test_structured_diff_detail_cap(tmp_path):
    """Cap honesty: counts stay complete while detail lists shrink."""
    old, new = tmp_path / "cap_old.docx", tmp_path / "cap_new.docx"
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"Original sentence number {i} in the early draft.")
    doc.save(str(old))
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f"Rewritten sentence number {i} in the later draft.")
    doc.save(str(new))

    diff = rc.structured_diff(str(old), str(new), detail_cap=5)
    assert diff["counts"]["modified"] == 12
    assert diff["detail_capped"] is True
    assert diff["detail_cap"] == 5
    shown = (
        len(diff["modified"])
        + len(diff["moved"])
        + len(diff["inserted"])
        + len(diff["deleted"])
    )
    assert shown == 5


# ---------------------------------------------------------------- read-only


def test_report_tools_never_mutate(tmp_path):
    """Every analytics call leaves the file byte-for-byte untouched."""
    p = _commented_doc(tmp_path)
    pkg = DocxPackage(p)
    tx.search_and_replace(
        pkg,
        [{"find": "seventy seven", "replace": "seventy eight"}],
        track=True,
        author="Alice",
    )
    pkg.save(do_backup=False)
    before = p.read_bytes()

    rc.comment_report(DocxPackage(p))
    rc.comment_report(DocxPackage(p), include_resolved=False)
    rc.revision_analytics(DocxPackage(p))
    rc.comment_report_multi([str(p)])
    other = tmp_path / "other.docx"
    shutil.copy(p, other)
    rc.structured_diff(str(p), str(other))

    assert p.read_bytes() == before
    assert not list(tmp_path.glob("*.bak-*")), "no backups: nothing was saved"
    from word_mcp.core import safesave

    d = safesave.slot_dir(p)
    assert not (d.exists() and any(d.glob("*.docx"))), (
        "no slot rotation: nothing was saved"
    )
