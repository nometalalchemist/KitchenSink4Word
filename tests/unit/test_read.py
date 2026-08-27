"""Phase 1 gate: read layer verified against python-docx as independent oracle."""

from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import read

CORPUS = Path(__file__).resolve().parents[1] / "corpus"
CORPUS_DOCS = sorted(CORPUS.glob("*.docx"))
IDS = [p.name for p in CORPUS_DOCS]


@pytest.fixture(params=CORPUS_DOCS, ids=IDS)
def both(request):
    return DocxPackage(request.param), Document(str(request.param))


def _norm(text: str) -> str:
    # python-docx does not render footnote/endnote reference markers; strip ours.
    import re

    return re.sub(r"\[(fn|en):[^\]]*\]", "", text)


def test_paragraph_texts_match_python_docx(both):
    """Oracle comparison, valid only where python-docx's semantics are correct:
    python-docx omits runs inside w:ins (tracked insertions), which ARE visible
    text, so revision-bearing paragraphs are checked separately below."""
    pkg, oracle = both
    ours = [_norm(e["text"]) for e in read.get_paragraphs(pkg)]
    theirs = [p.text for p in oracle.paragraphs]
    assert len(ours) == len(theirs)
    rev_tags = {qn("w:ins"), qn("w:del"), qn("w:moveFrom"), qn("w:moveTo")}
    mismatches = []
    for i, (a, b, o_para) in enumerate(zip(ours, theirs, oracle.paragraphs)):
        if any(o_para._p.iter(t) for t in rev_tags) and any(
            list(o_para._p.iter(t)) for t in rev_tags
        ):
            continue
        if a != b:
            mismatches.append((i, a, b))
    assert not mismatches, f"first mismatch: {mismatches[0]}"


def test_revision_paragraph_semantics(both):
    """Where a paragraph has tracked changes: inserted text is included in our
    extraction, deleted text is excluded (matching what Word displays)."""
    pkg, oracle = both
    from word_mcp.ops.read import body_items, paragraph_text, run_text

    checked = 0
    for kind, idx, el in body_items(pkg):
        if kind != "paragraph":
            continue
        ins_els = list(el.iter(qn("w:ins")))
        del_els = list(el.iter(qn("w:del")))
        if not ins_els and not del_els:
            continue
        visible = paragraph_text(el)
        for ins in ins_els:
            ins_text = "".join(run_text(r) for r in ins.iter(qn("w:r")))
            if ins_text:
                assert ins_text in visible, (
                    f"para {idx}: inserted text missing: {ins_text!r}"
                )
        deleted_only = "".join(
            run_text(r, include_deleted=True)
            for d in del_els
            for r in d.iter(qn("w:r"))
        )
        without_deleted = paragraph_text(el, include_deleted=False)
        with_deleted_len = len(paragraph_text(el, include_deleted=True))
        assert len(without_deleted) + len(deleted_only) == with_deleted_len
        checked += 1
    if checked:
        print(f"revision-bearing paragraphs checked: {checked}")


def test_table_contents_match_python_docx(both):
    pkg, oracle = both
    tables = read.list_tables(pkg)
    assert len(tables) == len(oracle.tables)
    for t_meta, o_tbl in zip(tables, oracle.tables):
        ours = read.get_table(pkg, t_meta["index"])
        assert ours["rows"] == len(o_tbl.rows)
        for r_i, o_row in enumerate(o_tbl.rows):
            our_row = ours["cells"][r_i]
            # python-docx repeats merged cells; our reader lists real w:tc only,
            # so compare against the deduplicated oracle row.
            o_cells = []
            seen = set()
            for c in o_row.cells:
                if id(c._tc) not in seen:
                    seen.add(id(c._tc))
                    o_cells.append(c)
            assert len(our_row) == len(o_cells), (
                f"table {t_meta['index']} row {r_i}: {len(our_row)} vs {len(o_cells)}"
            )
            for c_i, o_cell in enumerate(o_cells):
                assert _norm(our_row[c_i]["text"]) == o_cell.text


def test_outline_levels_are_sane(both):
    pkg, _ = both
    outline = read.get_outline(pkg)
    for h in outline:
        assert 1 <= h["level"] <= 9
        assert h["text"].strip()


def test_document_info_consistency(both):
    pkg, oracle = both
    info = read.get_document_info(pkg)
    assert info["paragraphs"] == len(oracle.paragraphs)
    assert info["tables"] == len(oracle.tables)
    assert info["footnotes"] == len(read.list_footnotes(pkg))
    assert info["comments"] == len(read.get_comments(pkg))


def test_footnote_positions_monotonic(both):
    pkg, _ = both
    notes = read.list_footnotes(pkg)
    positions = [n["position"] for n in notes if n["position"] is not None]
    assert positions == sorted(positions)


def test_comment_threading_consistency(both):
    pkg, _ = both
    comments = read.get_comments(pkg)
    ids = {c["id"] for c in comments}
    for c in comments:
        if c["reply_to"] is not None:
            assert c["reply_to"] in ids


def test_find_text_locates_known_paragraph(both):
    pkg, oracle = both
    # Pick the longest paragraph and search for a distinctive slice of it.
    texts = [(i, p.text) for i, p in enumerate(oracle.paragraphs) if len(p.text) > 80]
    if not texts:
        pytest.skip("no long paragraphs")
    idx, text = max(texts, key=lambda t: len(t[1]))
    needle = text[20:60]
    hits = read.find_text(pkg, needle)
    assert any(h.get("paragraph_index") == idx for h in hits)


def test_revision_summary_counts(both):
    pkg, _ = both
    summary = read.revision_summary(pkg)
    assert summary["total"] == len(read.get_tracked_changes(pkg))
    assert sum(summary["by_author"].values()) == summary["total"]


def test_styles_readable(both):
    pkg, _ = both
    styles = read.list_styles(pkg)
    assert styles, "styles.xml should list styles"
    assert all(s["id"] for s in styles)
