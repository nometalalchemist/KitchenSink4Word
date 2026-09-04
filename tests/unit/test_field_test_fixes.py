"""Regression tests for the 2026-09-03 field-test fix batch.

Findings credit: the author's two live Opus field-test sessions (the
3-agent stress test and the real dissertation merge; reports under
Developer Feedback/General Feedback/Beta Test Report - 2026-09-03*).
Each test reproduces the reported defect first, then asserts the fix.
"""

from __future__ import annotations

from lxml import etree

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import assembly as asm
from word_mcp.ops import batch as bt
from word_mcp.ops import diagnostics as dg
from word_mcp.ops import read as rd
from word_mcp.ops import tables as tb
from word_mcp.ops import text as tx


# --------------------------------------------------------------- helpers


def _build(path, texts=("Alpha", "Bravo", "Charlie")):
    d = Document()
    for t in texts:
        d.add_paragraph(t)
    d.save(str(path))
    return path


def _texts(path):
    pkg = DocxPackage(path)
    return [
        rd.paragraph_text(el)
        for k, _i, el in rd.body_items(pkg)
        if k == "paragraph"
    ]


def _set_docdefaults_spacing(path, *, line=None, after=None):
    """Stamp docDefaults/pPrDefault/pPr/spacing on a saved .docx."""
    pkg = DocxPackage(path)
    root = pkg.root("word/styles.xml")
    dd = root.find(qn("w:docDefaults"))
    if dd is None:
        dd = etree.Element(qn("w:docDefaults"))
        root.insert(0, dd)
    ppd = dd.find(qn("w:pPrDefault"))
    if ppd is None:
        ppd = etree.SubElement(dd, qn("w:pPrDefault"))
    ppr = ppd.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.SubElement(ppd, qn("w:pPr"))
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(ppr, qn("w:spacing"))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    if after is not None:
        sp.set(qn("w:after"), str(after))
    pkg.mark_dirty("word/styles.xml")
    pkg.save(do_backup=False)


def _body_paras(path):
    pkg = DocxPackage(path)
    return [
        el for k, _i, el in rd.body_items(pkg) if k == "paragraph"
    ], pkg


def _spacing_of(p):
    sp = p.find(f"{qn('w:pPr')}/{qn('w:spacing')}")
    if sp is None:
        return None
    return {
        "line": sp.get(qn("w:line")),
        "after": sp.get(qn("w:after")),
        "lineRule": sp.get(qn("w:lineRule")),
    }


# ==================================================================
# CRITICAL: insert_document formatting='source' loses source
# document-defaults (the half-single-spaced dissertation bug)
# ==================================================================


def test_insert_document_bakes_source_docdefaults(tmp_path):
    """Double-spaced-source (via docDefaults) into a target with different
    defaults: inserted paragraphs must carry the source spacing explicitly
    and the result must warn that the defaults differ."""
    source = _build(tmp_path / "src.docx", ("Ch5 body one.", "Ch5 body two."))
    _set_docdefaults_spacing(source, line=480, after=0)
    target = _build(tmp_path / "tgt.docx", ("T1", "T2"))
    _set_docdefaults_spacing(target, line=240, after=240)

    out = srv.insert_document(str(target), str(source))
    dd = out["document_defaults"]
    assert dd["differ"] is True
    assert "pPr.spacing.line" in dd["differing_properties"]
    assert "pPr.spacing.after" in dd["differing_properties"]
    assert dd["paragraphs_baked"] >= 2

    paras, _pkg = _body_paras(target)
    # Target's own paragraphs untouched.
    assert _spacing_of(paras[0]) is None
    assert _spacing_of(paras[1]) is None
    # Inserted paragraphs resolved to the SOURCE defaults, explicitly.
    for p in paras[2:]:
        sp = _spacing_of(p)
        assert sp is not None, "inserted paragraph lost source defaults"
        assert sp["line"] == "480"
        assert sp["after"] == "0"


def test_insert_document_same_defaults_bakes_nothing(tmp_path):
    """Identical docDefaults: no baking, no warning, clean paragraphs."""
    source = _build(tmp_path / "src.docx", ("Body.",))
    target = _build(tmp_path / "tgt.docx", ("T1",))
    out = srv.insert_document(str(target), str(source))
    assert "document_defaults" not in out
    paras, _pkg = _body_paras(target)
    assert _spacing_of(paras[-1]) is None


def test_insert_document_explicit_and_style_values_not_overwritten(tmp_path):
    """Direct spacing wins; a style-chain-defined attribute is left to the
    documented by-name style reconciliation, not baked."""
    source = _build(tmp_path / "src.docx", ("Direct.", "Styled."))
    _set_docdefaults_spacing(source, line=480, after=0)
    # Paragraph 0: explicit spacing line=360. Paragraph 1: a style that
    # defines line spacing, referenced via pStyle.
    pkg = DocxPackage(source)
    paras = [
        el for k, _i, el in rd.body_items(pkg) if k == "paragraph"
    ]
    tx.set_paragraph_format(pkg, [0], {"line_spacing": 1.5})
    sroot = pkg.root("word/styles.xml")
    style = etree.SubElement(sroot, qn("w:style"))
    style.set(qn("w:type"), "paragraph")
    style.set(qn("w:styleId"), "SpacedStyle")
    etree.SubElement(style, qn("w:name")).set(qn("w:val"), "Spaced Style")
    sppr = etree.SubElement(style, qn("w:pPr"))
    ssp = etree.SubElement(sppr, qn("w:spacing"))
    ssp.set(qn("w:line"), "300")
    ppr = etree.SubElement(paras[1], qn("w:pPr"))
    paras[1].remove(ppr)
    paras[1].insert(0, ppr)
    etree.SubElement(ppr, qn("w:pStyle")).set(qn("w:val"), "SpacedStyle")
    pkg.mark_dirty()
    pkg.mark_dirty("word/styles.xml")
    pkg.save(do_backup=False)

    target = _build(tmp_path / "tgt.docx", ("T1",))
    srv.insert_document(str(target), str(source))
    tparas, _pkg = _body_paras(target)
    direct, styled = tparas[1], tparas[2]
    assert _spacing_of(direct)["line"] == "360"  # 1.5 * 240, untouched
    # Style-defined line spacing not baked over; after (chain-undefined)
    # still baked from the source docDefaults.
    st = _spacing_of(styled)
    assert st["line"] is None or st["line"] != "480"
    assert st["after"] == "0"


# ==================================================================
# HIGH: insert_document honors position 'before'/'after'
# ==================================================================


def test_insert_document_position_before(tmp_path):
    source = _build(tmp_path / "src.docx", ("Xray",))
    target = _build(tmp_path / "tgt.docx", ("Alpha", "Bravo", "Charlie"))
    srv.insert_document(
        str(target), str(source),
        location={"paragraph": 1, "position": "before"},
    )
    assert _texts(target) == ["Alpha", "Xray", "Bravo", "Charlie"]


def test_insert_document_position_before_first(tmp_path):
    """The report had no way to insert ahead of the first body item."""
    source = _build(tmp_path / "src.docx", ("Xray",))
    target = _build(tmp_path / "tgt.docx", ("Alpha", "Bravo"))
    out = srv.insert_document(
        str(target), str(source),
        location={"paragraph": 0, "position": "before"},
    )
    assert _texts(target) == ["Xray", "Alpha", "Bravo"]
    assert out["position"]["starts_at_body_item"] == 0


def test_insert_document_position_after_default(tmp_path):
    source = _build(tmp_path / "src.docx", ("Xray",))
    target = _build(tmp_path / "tgt.docx", ("Alpha", "Bravo"))
    srv.insert_document(str(target), str(source), location={"paragraph": 0})
    assert _texts(target) == ["Alpha", "Xray", "Bravo"]


# ==================================================================
# MEDIUM: row insert inside a vertical merge + diagnose detection
# ==================================================================


def _table_doc(tmp_path, rows=4, cols=2):
    path = tmp_path / "table.docx"
    d = Document()
    d.add_table(rows=rows, cols=cols)
    d.save(str(path))
    pkg = DocxPackage(path)
    for r in range(rows):
        tb.set_cells(
            pkg, 0,
            [{"row": r, "cell": c, "text": f"r{r}c{c}"} for c in range(cols)],
        )
    pkg.save(do_backup=False)
    return path


def test_insert_row_inside_vertical_merge_continues_chain(tmp_path):
    path = _table_doc(tmp_path)
    pkg = DocxPackage(path)
    tb.merge_cells(pkg, 0, start_row=1, end_row=3, start_col=0, end_col=0)
    res = tb.insert_rows(pkg, 0, at=2)
    assert res.get("cells_continued_vertical_merges") == 1
    pkg.save(do_backup=False)
    # The inserted row's first cell continues the merge.
    pkg2 = DocxPackage(path)
    tbl = pkg2.root().find(f".//{qn('w:tbl')}")
    row2 = tbl.findall(qn("w:tr"))[2]
    vm = row2.findall(qn("w:tc"))[0].find(
        f"{qn('w:tcPr')}/{qn('w:vMerge')}"
    )
    assert vm is not None
    assert vm.get(qn("w:val"), "continue") == "continue"
    # And the document diagnoses clean (no orphaned continuation).
    report = dg.diagnose_document(pkg2)
    assert not [
        p for p in report["problems"] if p["category"] == "tables"
    ]


def test_insert_row_outside_merge_stays_standalone(tmp_path):
    path = _table_doc(tmp_path)
    pkg = DocxPackage(path)
    tb.merge_cells(pkg, 0, start_row=1, end_row=3, start_col=0, end_col=0)
    res = tb.insert_rows(pkg, 0, at=1)  # boundary: above the restart
    assert "cells_continued_vertical_merges" not in res
    tbl = pkg.root().find(f".//{qn('w:tbl')}")
    row1 = tbl.findall(qn("w:tr"))[1]
    assert row1.findall(qn("w:tc"))[0].find(
        f"{qn('w:tcPr')}/{qn('w:vMerge')}"
    ) is None


def test_diagnose_detects_orphaned_vmerge_continue(tmp_path):
    path = _table_doc(tmp_path)
    pkg = DocxPackage(path)
    tb.merge_cells(pkg, 0, start_row=1, end_row=3, start_col=0, end_col=0)
    # Break the chain by hand: remove the restart marker.
    tbl = pkg.root().find(f".//{qn('w:tbl')}")
    restart_tcpr = tbl.findall(qn("w:tr"))[1].findall(qn("w:tc"))[0].find(
        qn("w:tcPr")
    )
    restart_tcpr.remove(restart_tcpr.find(qn("w:vMerge")))
    problems = [
        p
        for p in dg.diagnose_document(pkg)["problems"]
        if p["category"] == "tables"
    ]
    assert problems and "orphaned" in problems[0]["detail"]


# ==================================================================
# MEDIUM: merge reports the ACTUAL absorbed range
# ==================================================================


def test_merge_reports_actual_range_when_absorbing(tmp_path):
    path = _table_doc(tmp_path, rows=5)
    pkg = DocxPackage(path)
    tb.merge_cells(pkg, 0, start_row=3, end_row=4, start_col=0, end_col=0)
    res = tb.merge_cells(
        pkg, 0, start_row=1, end_row=3, start_col=0, end_col=0
    )
    assert res["merged"]["rows"] == [1, 4]
    assert res["requested"]["rows"] == [1, 3]
    assert "absorbed" in res["note"]


def test_merge_plain_range_shape_unchanged(tmp_path):
    path = _table_doc(tmp_path)
    pkg = DocxPackage(path)
    res = tb.merge_cells(
        pkg, 0, start_row=1, end_row=2, start_col=0, end_col=0
    )
    assert res["merged"]["rows"] == [1, 2]
    assert "requested" not in res


# ==================================================================
# MEDIUM: set_paragraph_format numeric range validation
# ==================================================================


@pytest.mark.parametrize(
    "formatting",
    [
        {"indent_left_pt": -72},
        {"indent_right_pt": -1},
        {"space_after_pt": 99999},
        {"space_before_pt": -5},
        {"line_spacing": 0},
        {"line_spacing": 5000},
        {"first_line_indent_pt": -99999},
    ],
)
def test_paragraph_format_out_of_range_refused(tmp_path, formatting):
    path = _build(tmp_path / "fmt.docx")
    with pytest.raises(WordMcpError):
        srv.set_paragraph_format(str(path), [0], formatting)


def test_paragraph_format_valid_values_pass(tmp_path):
    path = _build(tmp_path / "fmt.docx")
    out = srv.set_paragraph_format(
        str(path), [0],
        {"line_spacing": 2, "space_after_pt": 0,
         "first_line_indent_pt": -18},
    )
    assert out["formatted_paragraphs"] == [0]


# ==================================================================
# MEDIUM: orphaned comment anchor after tracked delete
# ==================================================================


def test_comment_anchor_deleted_reported(tmp_path):
    path = _build(tmp_path / "cmt.docx", ("Keep me.", "Comment target here."))
    srv.manage_comment(
        str(path), "add", text="note",
        location={"search": {"text": "target"}},
    )
    srv.delete_paragraphs(str(path), start=1, track=True)
    comments = srv.get_comments(str(path))
    assert comments[0]["anchor_deleted"] is True
    assert "target" in comments[0]["anchored_text"]


# ==================================================================
# MEDIUM: outline + structure view on direct-formatted documents
# ==================================================================


def _direct_formatted_doc(tmp_path):
    path = tmp_path / "direct.docx"
    d = Document()
    h = d.add_paragraph()
    run = h.add_run("Chapter Four")
    run.bold = True
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(3):
        d.add_paragraph(
            f"Body paragraph {i} long enough to never look like a heading "
            "to any heuristic, with a period at the end."
        )
    d.save(str(path))
    return path


def test_get_outline_fallback_note_and_counts(tmp_path):
    path = _direct_formatted_doc(tmp_path)
    out = srv.get_outline(str(path))
    assert out["headings"] == []
    assert "direct formatting" in out["note"]
    assert out["structure"]["paragraphs"] == 4
    assert out["structure"]["approx_words"] > 0


def test_get_outline_detect_formatted_heuristic(tmp_path):
    path = _direct_formatted_doc(tmp_path)
    out = srv.get_outline(str(path), detect_formatted=True)
    assert isinstance(out, list) and len(out) == 1
    assert out[0]["text"] == "Chapter Four"
    assert out[0]["level"] == 1
    assert out[0]["detected_via"] == "formatting_heuristic"


def test_structure_view_flat_fallback(tmp_path):
    path = _direct_formatted_doc(tmp_path)
    out = srv.get_document_view(str(path), detail="structure")
    assert out["blocks"] == 0
    assert out["note"]
    assert "4 paragraphs" in out["view"]


# ==================================================================
# MEDIUM: copy_format_from excludes outline_level
# ==================================================================


def test_copy_format_from_excludes_outline_level(tmp_path):
    path = _build(tmp_path / "clone.docx", ("Heading-ish", "Body"))
    srv.set_paragraph_format(
        str(path), [0], {"outline_level": 2, "line_spacing": 2}
    )
    srv.insert_paragraphs(
        str(path), [{"text": "Research question text."}],
        location={"paragraph": 1}, copy_format_from=0,
    )
    pkg = DocxPackage(path)
    paras = [
        el for k, _i, el in rd.body_items(pkg) if k == "paragraph"
    ]
    new_p = paras[2]
    assert new_p.find(f"{qn('w:pPr')}/{qn('w:outlineLvl')}") is None
    assert _spacing_of(new_p)["line"] == "480"  # visual clone still works
    assert srv.get_outline(str(path))[0]["paragraph_index"] == 0


# ==================================================================
# LOW: set_paragraph_format start/end range syntax
# ==================================================================


def test_set_paragraph_format_range_syntax(tmp_path):
    path = _build(tmp_path / "range.docx", ("A", "B", "C", "D"))
    out = srv.set_paragraph_format(
        str(path), start=1, end=3, formatting={"line_spacing": 2}
    )
    assert out["formatted_paragraphs"] == [1, 2, 3]
    with pytest.raises(WordMcpError):
        srv.set_paragraph_format(
            str(path), [0], start=1, formatting={"line_spacing": 2}
        )
    with pytest.raises(WordMcpError):
        srv.set_paragraph_format(str(path), formatting={"line_spacing": 2})
    with pytest.raises(WordMcpError):
        srv.set_paragraph_format(
            str(path), start=3, end=1, formatting={"line_spacing": 2}
        )


# ==================================================================
# LOW: DEL 0x7F joins control-character validation
# ==================================================================


def test_del_char_refused_in_paragraph_and_cell(tmp_path):
    path = _build(tmp_path / "del.docx")
    with pytest.raises(WordMcpError, match="0x7F"):
        srv.insert_paragraphs(str(path), [{"text": "before\x7fafter"}])
    d = Document()
    d.add_table(rows=1, cols=1)
    tpath = tmp_path / "delt.docx"
    d.save(str(tpath))
    with pytest.raises(WordMcpError, match="0x7F"):
        srv.set_cells(
            str(tpath), 0, [{"row": 0, "cell": 0, "text": "x\x7fy"}]
        )


# ==================================================================
# LOW: copy_document refuses source == dest
# ==================================================================


def test_copy_document_self_copy_refused(tmp_path):
    path = _build(tmp_path / "self.docx")
    with pytest.raises(WordMcpError, match="same file"):
        srv.copy_document(str(path), str(path), overwrite=True)


# ==================================================================
# LOW: markdown list nesting cap matches insert_list (0-8)
# ==================================================================


def test_markdown_nesting_over_limit_refused():
    md = "- top\n" + " " * 20 + "- ten deep"
    with pytest.raises(WordMcpError, match="level"):
        bt.parse_markdown(md)


def test_markdown_nesting_level_8_allowed():
    md = "- top\n" + " " * 16 + "- eight deep"
    segs = bt.parse_markdown(md)
    assert segs[0]["items"][1][1] == 8


# ==================================================================
# LOW: manage_comment defines the CommentReference style
# ==================================================================


def test_comment_reference_style_created(tmp_path):
    path = _build(tmp_path / "cstyle.docx", ("Anchor text here.",))
    srv.manage_comment(
        str(path), "add", text="note",
        location={"search": {"text": "Anchor"}},
    )
    pkg = DocxPackage(path)
    ids = {
        s.get(qn("w:styleId"))
        for s in pkg.root("word/styles.xml").findall(qn("w:style"))
    }
    assert "CommentReference" in ids
    # And diagnose no longer flags the dangling style reference.
    assert not [
        p
        for p in dg.diagnose_document(pkg)["problems"]
        if p["category"] == "styles" and "CommentReference" in p["detail"]
    ]


# ==================================================================
# LOW: large row insert warning
# ==================================================================


def test_large_row_insert_warns(tmp_path):
    path = _table_doc(tmp_path, rows=2, cols=1)
    pkg = DocxPackage(path)
    res = tb.insert_rows(pkg, 0, at=2, count=1001)
    assert "warning" in res
    res_small = tb.insert_rows(pkg, 0, at=0, count=2)
    assert "warning" not in res_small


# ==================================================================
# POLISH: zero-width-only regex refuses instead of silent 0
# ==================================================================


def test_zero_width_only_regex_refused(tmp_path):
    path = _build(tmp_path / "zw.docx", ("foo bar baz",))
    pkg = DocxPackage(path)
    with pytest.raises(WordMcpError, match="zero-width"):
        tx.search_and_replace(
            pkg, [{"find": r"(?<=foo)", "replace": "X", "regex": True}]
        )
    # A lookbehind that consumes text still works.
    out = tx.search_and_replace(
        pkg, [{"find": r"(?<=foo )bar", "replace": "BAR", "regex": True}]
    )
    assert out["total"] == 1


# ==================================================================
# Docstring wiring sanity for the batch
# ==================================================================


def test_task_map_in_enable_tools_description():
    assert "Task map:" in srv.enable_tools.__doc__
    assert "media-forms" in srv.enable_tools.__doc__


def test_insert_document_op_before_first_direct(tmp_path):
    source = _build(tmp_path / "s.docx", ("X",))
    target = _build(tmp_path / "t.docx", ("A",))
    pkg = DocxPackage(target)
    asm.insert_document(pkg, str(source), before_first=True)
    pkg.save(do_backup=False)
    assert _texts(target) == ["X", "A"]
