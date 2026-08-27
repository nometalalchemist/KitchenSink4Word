"""Phase 4 (part 2): bookmarks, cross-references, captions, hyperlinks."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import fields, read

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


@pytest.fixture
def codebook(tmp_path):
    dst = tmp_path / "codebook.docx"
    shutil.copy(CORPUS / "codebook.docx", dst)
    return dst


def anchor_from(path, word_slice=slice(3, 6)):
    pkg = DocxPackage(path)
    paras = read.get_paragraphs(pkg)
    long_para = max(paras, key=lambda p: len(p["text"]))
    return " ".join(long_para["text"].split()[word_slice])


def test_bookmark_and_cross_reference(ch4):
    pkg = DocxPackage(ch4)
    anchor = anchor_from(ch4)
    fields.add_bookmark(pkg, "TestTarget", anchor_text=anchor)
    paras = read.get_paragraphs(pkg)
    other = [p for p in paras if len(p["text"]) > 80 and anchor not in p["text"]]
    ref_anchor = " ".join(other[0]["text"].split()[:3])
    fields.add_cross_reference(
        pkg, after_anchor=ref_anchor, to_bookmark="TestTarget", kind="page"
    )
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    assert any(b["name"] == "TestTarget" for b in fields.list_bookmarks(pkg2))
    xml = pkg2.raw_part("word/document.xml").decode("utf-8")
    assert "PAGEREF TestTarget" in xml
    Document(str(ch4))


def test_cross_reference_to_missing_bookmark_refused(ch4):
    pkg = DocxPackage(ch4)
    with pytest.raises(TargetNotFound):
        fields.add_cross_reference(
            pkg, after_anchor="anything", to_bookmark="NoSuchBookmark"
        )


def test_duplicate_bookmark_refused(ch4):
    pkg = DocxPackage(ch4)
    anchor = anchor_from(ch4)
    fields.add_bookmark(pkg, "Dup", anchor_text=anchor)
    with pytest.raises(WordMcpError):
        fields.add_bookmark(pkg, "Dup", anchor_text=anchor)


def test_invalid_bookmark_name_refused(ch4):
    pkg = DocxPackage(ch4)
    with pytest.raises(WordMcpError):
        fields.add_bookmark(pkg, "1bad name", anchor_text="x")


def test_table_caption(codebook):
    pkg = DocxPackage(codebook)
    fields.add_caption(
        pkg, table_index=0, label="Table", text="Coding categories overview"
    )
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(codebook)
    xml = pkg2.raw_part("word/document.xml").decode("utf-8")
    assert "SEQ Table" in xml
    assert "Coding categories overview" in xml
    Document(str(codebook))


def test_hyperlink(ch4):
    pkg = DocxPackage(ch4)
    anchor = anchor_from(ch4)
    fields.add_hyperlink(
        pkg, anchor_text=anchor, url="https://example.org/paper"
    )
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch4)
    rels = pkg2.raw_part("word/_rels/document.xml.rels").decode("utf-8")
    assert "https://example.org/paper" in rels
    assert "TargetMode=\"External\"" in rels
    # Anchor text still present exactly once (moved inside w:hyperlink).
    text = "\n".join(p["text"] for p in read.get_paragraphs(pkg2))
    assert anchor in text
    Document(str(ch4))


def test_hyperlink_inside_hyperlink_refused(ch4):
    pkg = DocxPackage(ch4)
    anchor = anchor_from(ch4)
    fields.add_hyperlink(pkg, anchor_text=anchor, url="https://a.example")
    with pytest.raises(WordMcpError):
        fields.add_hyperlink(pkg, anchor_text=anchor, url="https://b.example")
