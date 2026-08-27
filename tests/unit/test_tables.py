"""Phase 3 gate: table operations — columns through merges, bulk cells, real
codebook tables."""

import shutil
import time
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import read, tables as tb, text as t

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def codebook(tmp_path):
    dst = tmp_path / "codebook.docx"
    shutil.copy(CORPUS / "codebook.docx", dst)
    return dst


@pytest.fixture
def fresh(tmp_path):
    """A doc with a known 4x4 table built by our own create_table."""
    dst = tmp_path / "fresh.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    doc.save(str(dst))
    pkg = DocxPackage(dst)
    data = [
        ["H1", "H2", "H3", "H4"],
        ["a1", "b1", "c1", "d1"],
        ["a2", "b2", "c2", "d2"],
        ["a3", "b3", "c3", "d3"],
    ]
    tb.create_table(pkg, data, at_end=True)
    pkg.save(do_backup=False)
    return dst


def grid(path, index=0):
    return read.get_table(DocxPackage(path), index)


# ------------------------------------------------------------------ column ops


def test_delete_middle_column(fresh):
    pkg = DocxPackage(fresh)
    tb.delete_columns(pkg, 0, [1])
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["grid_columns"] == 3
    assert [c["text"] for c in g["cells"][0]] == ["H1", "H3", "H4"]
    assert [c["text"] for c in g["cells"][2]] == ["a2", "c2", "d2"]


def test_delete_multiple_columns(fresh):
    pkg = DocxPackage(fresh)
    tb.delete_columns(pkg, 0, [0, 2])
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["grid_columns"] == 2
    assert [c["text"] for c in g["cells"][0]] == ["H2", "H4"]


def test_delete_all_columns_refused(fresh):
    pkg = DocxPackage(fresh)
    with pytest.raises(WordMcpError):
        tb.delete_columns(pkg, 0, [0, 1, 2, 3])


def test_delete_column_through_horizontal_merge(fresh):
    """Merge H2:H3 horizontally in row 0, then delete grid column 2: the
    merged cell must shrink, other rows lose their cell."""
    pkg = DocxPackage(fresh)
    tb.merge_cells(pkg, 0, start_row=0, end_row=0, start_col=1, end_col=2)
    pkg.save(do_backup=False)

    pkg = DocxPackage(fresh)
    tb.delete_columns(pkg, 0, [2])
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["grid_columns"] == 3
    # Row 0: merged cell shrank from span 2 to span 1.
    assert [c["grid_span"] for c in g["cells"][0]] == [1, 1, 1]
    assert "H2" in g["cells"][0][1]["text"]
    # Other rows lost their column-2 cell.
    assert [c["text"] for c in g["cells"][1]] == ["a1", "b1", "d1"]


def test_delete_column_through_vertical_merge(fresh):
    """Vertically merge rows 1-3 of grid column 0, then delete column 0:
    the whole chain must vanish cleanly in every row."""
    pkg = DocxPackage(fresh)
    tb.merge_cells(pkg, 0, start_row=1, end_row=3, start_col=0, end_col=0)
    pkg.save(do_backup=False)

    pkg = DocxPackage(fresh)
    tb.delete_columns(pkg, 0, [0])
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["grid_columns"] == 3
    assert not g["has_merges"]
    assert [c["text"] for c in g["cells"][0]] == ["H2", "H3", "H4"]
    assert [c["text"] for c in g["cells"][1]] == ["b1", "c1", "d1"]


def test_insert_column_middle_and_edge(fresh):
    pkg = DocxPackage(fresh)
    tb.insert_columns(pkg, 0, at=1)
    tb.insert_columns(pkg, 0, at=5)  # append at right edge
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["grid_columns"] == 6
    row0 = [c["text"] for c in g["cells"][0]]
    assert row0 == ["H1", "", "H2", "H3", "H4", ""]


def test_insert_column_inside_merge_widens_it(fresh):
    pkg = DocxPackage(fresh)
    tb.merge_cells(pkg, 0, start_row=0, end_row=0, start_col=1, end_col=2)
    pkg.save(do_backup=False)
    pkg = DocxPackage(fresh)
    tb.insert_columns(pkg, 0, at=2)  # inside the merged span 1-2
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["grid_columns"] == 5
    # Row 0 merged cell widened to span 3; other rows gained a plain cell.
    assert g["cells"][0][1]["grid_span"] == 3
    assert [c["text"] for c in g["cells"][1]] == ["a1", "b1", "", "c1", "d1"]


# --------------------------------------------------------------------- row ops


def test_insert_and_delete_rows(fresh):
    pkg = DocxPackage(fresh)
    tb.insert_rows(pkg, 0, at=2, count=2)
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["rows"] == 6
    assert all(c["text"] == "" for c in g["cells"][2])
    assert all(c["text"] == "" for c in g["cells"][3])

    pkg = DocxPackage(fresh)
    tb.delete_rows(pkg, 0, 2, 3)
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["rows"] == 4
    assert [c["text"] for c in g["cells"][2]] == ["a2", "b2", "c2", "d2"]


def test_delete_row_rerooots_vertical_merge(fresh):
    pkg = DocxPackage(fresh)
    tb.merge_cells(pkg, 0, start_row=1, end_row=3, start_col=0, end_col=0)
    pkg.save(do_backup=False)
    pkg = DocxPackage(fresh)
    tb.delete_rows(pkg, 0, 1)  # the restart row
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["rows"] == 3
    # New restart in what is now row 1, content moved down.
    assert g["cells"][1][0]["vmerge"] == "restart"
    assert "a1" in g["cells"][1][0]["text"]
    assert g["cells"][2][0]["vmerge"] == "continue"


# ---------------------------------------------------------------- bulk cells


def test_bulk_cell_edit_20_cells_under_5s(codebook):
    pkg = DocxPackage(codebook)
    tables = read.list_tables(pkg)
    target = max(tables, key=lambda x: x["rows"])
    g = read.get_table(pkg, target["index"])
    edits = []
    count = 0
    for r in range(g["rows"]):
        for c in range(len(g["cells"][r])):
            if count >= 20:
                break
            edits.append({"row": r, "cell": c, "text": f"EDIT-{count}"})
            count += 1
    start = time.perf_counter()
    result = tb.set_cells(pkg, target["index"], edits)
    pkg.save(do_backup=False)
    elapsed = time.perf_counter() - start
    assert result["cells_written"] == 20
    assert elapsed < 5.0, f"20-cell edit took {elapsed:.2f}s"
    g2 = grid(codebook, target["index"])
    flat = [c["text"] for row in g2["cells"] for c in row]
    for i in range(20):
        assert f"EDIT-{i}" in flat


def test_set_cells_grid_block(fresh):
    pkg = DocxPackage(fresh)
    tb.set_cells_grid(
        pkg, 0, origin_row=1, origin_cell=1,
        data=[["X1", "Y1"], ["X2", "Y2"]],
    )
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["cells"][1][1]["text"] == "X1"
    assert g["cells"][2][2]["text"] == "Y2"


def test_set_cells_out_of_range(fresh):
    pkg = DocxPackage(fresh)
    with pytest.raises(TargetNotFound):
        tb.set_cells(pkg, 0, [{"row": 99, "cell": 0, "text": "x"}])


def test_format_cells_shading_and_bold(fresh):
    pkg = DocxPackage(fresh)
    tb.format_cells(
        pkg, 0, [{"row": 0}], {"shading": "D9E2F3", "bold": True, "valign": "center"}
    )
    pkg.save(do_backup=False)
    doc = Document(str(fresh))
    tbl = doc.tables[0]
    for cell in tbl.rows[0].cells:
        tcpr = cell._tc.tcPr
        shd = tcpr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
        )
        assert shd is not None
        assert shd.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
        ) == "D9E2F3"


# ----------------------------------------------------------- merge / unmerge


def test_merge_and_unmerge_roundtrip(fresh):
    pkg = DocxPackage(fresh)
    tb.merge_cells(pkg, 0, start_row=1, end_row=2, start_col=1, end_col=2)
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert g["has_merges"]
    merged = g["cells"][1][1]
    assert merged["grid_span"] == 2
    assert merged["vmerge"] == "restart"
    assert "b1" in merged["text"] and "c1" in merged["text"]

    pkg = DocxPackage(fresh)
    tb.unmerge_cells(pkg, 0, row=1, cell=1)
    pkg.save(do_backup=False)
    g = grid(fresh)
    assert not g["has_merges"]
    assert len(g["cells"][1]) == 4


def test_misaligned_merge_refused(fresh):
    from word_mcp.core.errors import UnsupportedStructure

    pkg = DocxPackage(fresh)
    tb.merge_cells(pkg, 0, start_row=0, end_row=0, start_col=0, end_col=1)
    pkg.save(do_backup=False)
    pkg = DocxPackage(fresh)
    # Rect col 1..1 slices through row 0's merged 0..1 cell.
    with pytest.raises(UnsupportedStructure):
        tb.merge_cells(pkg, 0, start_row=0, end_row=1, start_col=1, end_col=1)


# ------------------------------------------------------- real codebook tables


def test_codebook_all_tables_survive_column_ops(codebook):
    """Column round-trip on every codebook table: insert one, delete it,
    content identical afterward."""
    pkg = DocxPackage(codebook)
    n_tables = len(read.list_tables(pkg))
    before = [read.get_table(pkg, i) for i in range(n_tables)]

    for i in range(n_tables):
        g = before[i]
        tb.insert_columns(pkg, i, at=g["grid_columns"])
        tb.delete_columns(pkg, i, [g["grid_columns"]])
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(codebook)
    for i in range(n_tables):
        after = read.get_table(pkg2, i)
        assert after["grid_columns"] == before[i]["grid_columns"]
        assert after["rows"] == before[i]["rows"]
        assert [
            [c["text"] for c in row] for row in after["cells"]
        ] == [[c["text"] for c in row] for row in before[i]["cells"]]


def test_codebook_create_and_delete_table(codebook):
    pkg = DocxPackage(codebook)
    n = len(read.list_tables(pkg))
    tb.create_table(
        pkg, [["Col A", "Col B"], ["1", "2"]], at_end=True
    )
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(codebook)
    assert len(read.list_tables(pkg2)) == n + 1
    tb.delete_table(pkg2, n)
    pkg2.save(do_backup=False)
    assert len(read.list_tables(DocxPackage(codebook))) == n
