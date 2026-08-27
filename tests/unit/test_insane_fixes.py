"""Regression tests for the insane-mode stress findings (v1.0.3)."""

import time

import pytest
from docx import Document

from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import read, tables as tb, text as tx


@pytest.fixture
def doc(tmp_path):
    dst = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("a" * 38 + "b")  # ReDoS bait for (a+)+c
    d.add_paragraph("Normal second paragraph for other tests.")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tb.create_table(pkg, [["x", "y"], ["1", "2"]], at_end=True)
    pkg.save(do_backup=False)
    return dst


# HIGH: ReDoS guard
#
# Two layers: the `regex` engine internally defuses many pathological patterns
# (including (a+)+c, the one that hung stdlib re in the stress test), and a
# hard timeout catches the rest ((a|aa)+c still backtracks exponentially).

REDOS_PATTERN = "(a|aa)+c"  # not optimizable; guaranteed to hit the timeout


def test_original_attack_pattern_now_harmless(doc):
    """The exact pattern from the stress report completes instantly."""
    pkg = DocxPackage(doc)
    start = time.perf_counter()
    result = tx.search_and_replace(
        pkg, [{"find": "(a+)+c", "replace": "X", "regex": True}]
    )
    assert time.perf_counter() - start < 5
    assert result["total"] == 0


def test_redos_pattern_times_out_cleanly(doc):
    pkg = DocxPackage(doc)
    start = time.perf_counter()
    with pytest.raises(WordMcpError, match="backtracking"):
        tx.search_and_replace(
            pkg, [{"find": REDOS_PATTERN, "replace": "X", "regex": True}]
        )
    elapsed = time.perf_counter() - start
    assert elapsed < 15, f"guard took {elapsed:.1f}s — timeout not effective"


def test_redos_guard_in_find_text(doc):
    pkg = DocxPackage(doc)
    start = time.perf_counter()
    with pytest.raises(WordMcpError, match="backtracking"):
        read.find_text(pkg, REDOS_PATTERN, regex=True)
    assert time.perf_counter() - start < 15


def test_legitimate_regex_still_works(doc):
    pkg = DocxPackage(doc)
    result = tx.search_and_replace(
        pkg, [{"find": r"[Nn]ormal", "replace": "NORMAL", "regex": True}]
    )
    assert result["total"] == 1
    assert read.find_text(pkg, r"second\s+paragraph", regex=True)


def test_invalid_regex_clean_error(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="invalid regex"):
        tx.search_and_replace(
            pkg, [{"find": "(unclosed", "replace": "X", "regex": True}]
        )


# MEDIUM: unknown formatting keys rejected


def test_unknown_char_format_key_rejected(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="unknown character-formatting"):
        tx.format_text(pkg, find="Normal", formatting={"explode": True})


def test_typo_format_key_rejected(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="bolded"):
        tx.format_text(pkg, find="Normal", formatting={"bolded": True})


def test_unknown_paragraph_format_key_rejected(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="unknown paragraph-formatting"):
        tx.set_paragraph_format(pkg, [0], {"allignment": "center"})


def test_unknown_cell_format_key_rejected(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="unknown cell-formatting"):
        tb.format_cells(pkg, 0, [{"row": 0}], {"shade": "FF0000"})


# MEDIUM: negative widths rejected


def test_negative_column_widths_rejected(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="positive"):
        tb.set_column_widths(pkg, 0, [-50, 100])
    with pytest.raises(WordMcpError, match="positive"):
        tb.insert_columns(pkg, 0, at=0, width_pt=-10)
    with pytest.raises(WordMcpError, match="positive"):
        tb.create_table(pkg, [["a"]], at_end=True, width_pt=0)


# MEDIUM: backup collisions within one second


def test_rapid_edits_produce_distinct_backups(doc):
    for i in range(5):
        pkg = DocxPackage(doc)
        tx.insert_paragraphs(pkg, [{"text": f"edit {i}"}], at_end=True)
        pkg.save()  # backup on
    baks = sorted(doc.parent.glob("t.bak-*.docx"))
    assert len(baks) == 5, f"expected 5 distinct backups, got {len(baks)}"


# LOW: negative start rejected consistently


def test_negative_start_rejected(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match=">= 0"):
        read.get_paragraphs(pkg, start=-1)
