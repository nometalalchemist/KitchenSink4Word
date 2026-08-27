"""Phase 8 gauntlet: adversarial inputs, foreign-origin files, fuzz sequences.

Every failure mode must be a CLEAN typed error — never a corrupted output file,
never a stack trace masquerading as a result.
"""

import io
import random
import shutil
import zipfile
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import (
    DocumentCorrupt,
    DocumentNotFound,
    DocumentProtected,
    WordMcpError,
)
from word_mcp.core.package import DocxPackage
from word_mcp.ops import notes, read, tables as tb, text as tx

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


# -------------------------------------------------------------- adversarial


def test_missing_file():
    with pytest.raises(DocumentNotFound):
        DocxPackage(r"C:\nowhere\nothing.docx")


def test_zero_byte_file(tmp_path):
    f = tmp_path / "empty.docx"
    f.write_bytes(b"")
    with pytest.raises(DocumentCorrupt):
        DocxPackage(f)


def test_text_file_pretending_to_be_docx(tmp_path):
    f = tmp_path / "fake.docx"
    f.write_text("This is not a docx at all.")
    with pytest.raises(DocumentCorrupt):
        DocxPackage(f)


def test_truncated_zip(tmp_path):
    src = CORPUS / "ch4.docx"
    data = src.read_bytes()
    f = tmp_path / "truncated.docx"
    f.write_bytes(data[: len(data) // 2])
    with pytest.raises(DocumentCorrupt):
        DocxPackage(f)


def test_ole_encrypted_signature(tmp_path):
    f = tmp_path / "protected.docx"
    f.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)
    with pytest.raises(DocumentProtected):
        DocxPackage(f)


def test_zip_without_document_xml(tmp_path):
    f = tmp_path / "hollow.docx"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("hello.txt", "nothing here")
    f.write_bytes(buf.getvalue())
    with pytest.raises(DocumentCorrupt):
        DocxPackage(f)


def test_failed_validation_leaves_original_untouched(tmp_path):
    """Force a validation failure mid-save; the target file must be intact."""
    src = tmp_path / "victim.docx"
    shutil.copy(CORPUS / "ch4.docx", src)
    original = src.read_bytes()
    pkg = DocxPackage(src)
    # Sabotage: make a parsed tree unserializable-to-valid by injecting a
    # bogus raw part, then attempt save.
    pkg.set_raw_part("word/bogus.xml", b"<unclosed")
    with pytest.raises(Exception):
        pkg.save(do_backup=False)
    assert src.read_bytes() == original


# ------------------------------------------------------------- foreign files


def test_python_docx_minimal_file(tmp_path):
    """python-docx output lacks rsids and much of Word's noise."""
    f = tmp_path / "pdx.docx"
    doc = Document()
    doc.add_heading("Generated Heading", 1)
    doc.add_paragraph("Body text with enough words to anchor onto here.")
    doc.add_table(rows=2, cols=2)
    doc.save(str(f))

    pkg = DocxPackage(f)
    assert read.get_document_info(pkg)["tables"] == 1
    tx.search_and_replace(pkg, [{"find": "Body", "replace": "BODY"}])
    tb.set_cells(pkg, 0, [{"row": 0, "cell": 0, "text": "filled"}])
    notes.add_note(
        pkg, "footnote", anchor_text="anchor onto", note_text="works here too"
    )
    pkg.save(do_backup=False)
    assert "BODY" in Document(str(f)).paragraphs[1].text


def test_handbuilt_minimal_ooxml(tmp_path):
    """A bare-minimum hand-built package (Google-Docs-style: no styles.xml
    extras, no settings, no themes) must be readable, and mutating tools must
    either work or fail cleanly — never corrupt."""
    f = tmp_path / "minimal.docx"
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    document = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:document xmlns:w="{w}"><w:body>'
        f"<w:p><w:r><w:t>Minimal paragraph one.</w:t></w:r></w:p>"
        f"<w:p><w:r><w:t>Minimal paragraph two.</w:t></w:r></w:p>"
        f"</w:body></w:document>"
    )
    ct = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )
    with zipfile.ZipFile(f, "w") as zf:
        zf.writestr("[Content_Types].xml", ct)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document)

    pkg = DocxPackage(f)
    paras = read.get_paragraphs(pkg)
    assert len(paras) == 2
    tx.search_and_replace(pkg, [{"find": "one", "replace": "ONE"}])
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(f)
    assert "ONE" in read.get_paragraphs(pkg2)[0]["text"]


# --------------------------------------------------------------------- fuzzer


OPS = [
    lambda pkg, rng: tx.search_and_replace(
        pkg, [{"find": rng.choice(["the", "and", "of"]), "replace": "X-Y"}]
    ),
    lambda pkg, rng: tx.insert_paragraphs(
        pkg, [{"text": f"Fuzz {rng.randint(0, 999)}"}], at_end=True
    ),
    lambda pkg, rng: tx.set_paragraph_format(
        pkg, [rng.randint(0, 5)], {"alignment": rng.choice(["left", "center"])}
    ),
    lambda pkg, rng: tx.add_heading(
        pkg, f"Fuzz H{rng.randint(0, 99)}", rng.randint(1, 3), at_end=True
    ),
    lambda pkg, rng: _fuzz_table_op(pkg, rng),
    lambda pkg, rng: _fuzz_note_op(pkg, rng),
]


def _fuzz_table_op(pkg, rng):
    tabs = read.list_tables(pkg)
    if not tabs:
        return tb.create_table(pkg, [["a", "b"], ["c", "d"]], at_end=True)
    t = rng.choice(tabs)
    g = read.get_table(pkg, t["index"])
    choice = rng.randint(0, 3)
    if choice == 0 and g["grid_columns"] > 2:
        return tb.delete_columns(pkg, t["index"], [g["grid_columns"] - 1])
    if choice == 1:
        return tb.insert_columns(pkg, t["index"], at=g["grid_columns"])
    if choice == 2:
        return tb.insert_rows(pkg, t["index"], at=g["rows"])
    return tb.set_cells(
        pkg,
        t["index"],
        [{"row": 0, "cell": 0, "text": f"fz{rng.randint(0, 99)}"}],
    )


def _fuzz_note_op(pkg, rng):
    paras = [p for p in read.get_paragraphs(pkg) if len(p["text"]) > 60]
    if not paras:
        return {}
    p = rng.choice(paras)
    words = p["text"].split()
    i = rng.randint(0, max(0, len(words) - 4))
    anchor = " ".join(words[i : i + 3])
    try:
        return notes.add_note(
            pkg, "footnote", anchor_text=anchor, note_text=f"fuzz {i}"
        )
    except WordMcpError:
        return {}  # ambiguous/multiple anchors are legitimate refusals


@pytest.mark.parametrize("doc_name", ["ch4.docx", "codebook.docx", "outline.docx"])
def test_fuzz_sequences(doc_name, tmp_path):
    """30 random ops per doc; after each save the package must round-trip and
    note integrity must hold."""
    rng = random.Random(20260827)
    src = tmp_path / doc_name
    shutil.copy(CORPUS / doc_name, src)
    for step in range(30):
        pkg = DocxPackage(src)
        op = rng.choice(OPS)
        try:
            op(pkg, rng)
        except WordMcpError:
            continue  # clean refusals are fine
        pkg.save(do_backup=False)
        # Integrity after every mutation.
        check = DocxPackage(src)
        v = notes.validate_notes(check)
        assert v["footnotes"]["ok"], f"step {step}: {v}"
        Document(str(src))
