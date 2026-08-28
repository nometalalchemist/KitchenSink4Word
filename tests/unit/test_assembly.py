"""DOCUMENT ASSEMBLY: insert_document (positional whole-document insertion).

Self-contained: every document is built on the fly with python-docx plus the
package's own ops. A rich source document exercises every resource class
(styles incl. name-collision, numbered lists, images, footnotes+endnotes,
merged-cell tables, hyperlinks, bookmarks, a generated chart, a comment);
insertion results are validated structurally (no dangling styleId / numId /
noteId / rId, docPr unique, python-docx round-trip). Word-render validation
runs in the deferred live round.
"""

import shutil
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from word_mcp.core.errors import (
    AmbiguousTarget,
    TargetNotFound,
    UnsupportedStructure,
    WordMcpError,
)
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import assembly as asm
from word_mcp.ops import charts as ch
from word_mcp.ops import comments as cm
from word_mcp.ops import dataio as dio
from word_mcp.ops import fields as fl
from word_mcp.ops import lists as ls
from word_mcp.ops import media
from word_mcp.ops import notes
from word_mcp.ops import tables as tb
from word_mcp.ops.read import body_items, paragraph_text

R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

CORPUS = Path(__file__).parent.parent / "corpus"

CAT_DATA = {
    "categories": ["Alpha", "Beta"],
    "series": [{"name": "S1", "values": [4.3, 2.5]}],
}


# ------------------------------------------------------------------- builders


def _png(width: int = 4, height: int = 4) -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x10\x20\x30" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _add_style(pkg: DocxPackage, style_id: str, name: str, color: str) -> None:
    root = pkg.root("word/styles.xml")
    s = etree.SubElement(root, qn("w:style"))
    s.set(qn("w:type"), "paragraph")
    s.set(qn("w:styleId"), style_id)
    etree.SubElement(s, qn("w:name")).set(qn("w:val"), name)
    etree.SubElement(s, qn("w:basedOn")).set(qn("w:val"), "Normal")
    rpr = etree.SubElement(s, qn("w:rPr"))
    etree.SubElement(rpr, qn("w:color")).set(qn("w:val"), color)
    pkg.mark_dirty("word/styles.xml")


def _add_paragraph(pkg: DocxPackage, text: str, style: str | None = None):
    p = etree.Element(qn("w:p"))
    if style:
        ppr = etree.SubElement(p, qn("w:pPr"))
        etree.SubElement(ppr, qn("w:pStyle")).set(qn("w:val"), style)
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    body = pkg.body()
    sectpr = body.find(qn("w:sectPr"))
    if sectpr is not None:
        sectpr.addprevious(p)
    else:
        body.append(p)
    pkg.mark_dirty()
    return p


def _fresh(tmp_dir: Path, name: str, texts: list[str]) -> Path:
    f = tmp_dir / name
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    doc.save(str(f))
    return f


def _build_rich_source(tmp_dir: Path) -> Path:
    """Source exercising every carried resource class."""
    f = _fresh(tmp_dir, "source.docx", ["S0 opening paragraph"])
    pkg = DocxPackage(f)
    # Styles: name-collision-different-formatting + source-only style.
    _add_style(pkg, "SrcSpecial", "Special Style", "0000FF")
    _add_style(pkg, "SrcOnly", "Source Only Style", "00FF00")
    _add_paragraph(pkg, "Styled special paragraph", style="SrcSpecial")
    _add_paragraph(pkg, "Cloned style paragraph", style="SrcOnly")
    # Two independent numbered lists (restart semantics per instance).
    ls.add_list(pkg, ["one", "two"], kind="number", at_end=True)
    ls.add_list(pkg, ["uno", "dos"], kind="number", at_end=True)
    # Image.
    img = tmp_dir / "pic.png"
    img.write_bytes(_png())
    media.add_image(pkg, str(img), at_end=True)
    # Notes.
    _add_paragraph(pkg, "Sentence with FNANCHOR inside.")
    notes.add_note(
        pkg, "footnote", anchor_text="FNANCHOR", note_text="A carried footnote."
    )
    _add_paragraph(pkg, "Sentence with ENANCHOR inside.")
    notes.add_note(
        pkg, "endnote", anchor_text="ENANCHOR", note_text="A carried endnote."
    )
    # Table with merged cells.
    tb.create_table(
        pkg,
        [["A", "B", "C"], ["D", "E", "F"], ["G", "H", "I"]],
        at_end=True,
        header_row=False,
    )
    tb.merge_cells(pkg, 0, start_row=0, end_row=1, start_col=0, end_col=0)
    # Hyperlink.
    _add_paragraph(pkg, "Visit LINKANCHOR for details.")
    fl.add_hyperlink(
        pkg, anchor_text="LINKANCHOR", url="https://example.com/x"
    )
    # Bookmarks: one name that collides with the target, one that does not.
    _add_paragraph(pkg, "Bookmark hosts: SHAREDTEXT and SRCONLYTEXT.")
    fl.add_bookmark(pkg, "SharedBM", anchor_text="SHAREDTEXT")
    fl.add_bookmark(pkg, "SrcOnlyBM", anchor_text="SRCONLYTEXT")
    # Chart.
    ch.add_chart(pkg, "column", CAT_DATA, at_end=True)
    # Comment (must be stripped on insertion).
    _add_paragraph(pkg, "COMMENTANCHOR gets a comment.")
    cm.add_comment(
        pkg, anchor_text="COMMENTANCHOR", text="reviewer note", author="T"
    )
    _add_paragraph(pkg, "S last paragraph")
    pkg.save(do_backup=False)
    return f


def _build_target(tmp_dir: Path, name: str = "target.docx") -> Path:
    f = _fresh(tmp_dir, name, ["T0 intro", "Chapter 4", "T2 closing"])
    pkg = DocxPackage(f)
    _add_style(pkg, "TgtSpecial", "Special Style", "FF0000")
    fl.add_bookmark(pkg, "SharedBM", anchor_text="T0")
    ls.add_list(pkg, ["target item"], kind="number", at_end=True)
    pkg.save(do_backup=False)
    return f


# ----------------------------------------------------------- shared insertion


@pytest.fixture(scope="module")
def rich(tmp_path_factory):
    """(source_path, result_dict, merged_target_path) for the at_end insert
    of the full-featured source — shared by the resource-class assertions."""
    tmp_dir = tmp_path_factory.mktemp("assembly")
    source = _build_rich_source(tmp_dir)
    target = _build_target(tmp_dir)
    pkg = DocxPackage(target)
    nums_before = {
        n.get(qn("w:numId"))
        for n in pkg.root("word/numbering.xml").findall(qn("w:num"))
    }
    result = asm.insert_document(pkg, str(source), at_end=True)
    pkg.save(do_backup=False)
    return source, result, target, nums_before


def _story_parts(pkg):
    for part in (
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
    ):
        if pkg.has_part(part):
            yield part


def _assert_integrity(pkg: DocxPackage) -> None:
    """No dangling styleId / numId / noteId / rId; docPr ids unique;
    every internal rel target part exists."""
    style_ids = {
        s.get(qn("w:styleId"))
        for s in pkg.root("word/styles.xml").findall(qn("w:style"))
    }
    num_ids = {"0"}
    if pkg.has_part("word/numbering.xml"):
        num_ids |= {
            n.get(qn("w:numId"))
            for n in pkg.root("word/numbering.xml").findall(qn("w:num"))
        }
    docpr_ids = []
    for part in _story_parts(pkg):
        root = pkg.root(part)
        for tag in ("w:pStyle", "w:rStyle", "w:tblStyle"):
            for el in root.iter(qn(tag)):
                assert el.get(qn("w:val")) in style_ids, (
                    f"dangling style ref {el.get(qn('w:val'))} in {part}"
                )
        for el in root.iter(qn("w:numId")):
            assert el.get(qn("w:val")) in num_ids, (
                f"dangling numId {el.get(qn('w:val'))} in {part}"
            )
        rels_name = part.replace("word/", "word/_rels/") + ".rels"
        rels = {}
        if pkg.has_part(rels_name):
            rels = {r.get("Id"): r for r in pkg.root(rels_name)}
        for node in root.iter():
            for key, val in node.attrib.items():
                if key.startswith("{" + R_NS + "}"):
                    assert val in rels, f"dangling rId {val} in {part}"
        import posixpath

        for rid, rel in rels.items():
            if rel.get("TargetMode") == "External":
                continue
            resolved = posixpath.normpath(
                posixpath.join("word", rel.get("Target", ""))
            )
            assert pkg.has_part(resolved), (
                f"{rels_name} {rid} targets missing part {resolved}"
            )
        for d in root.iter(f"{{{WP}}}docPr"):
            docpr_ids.append(d.get("id"))
    assert len(docpr_ids) == len(set(docpr_ids)), "docPr ids not unique"
    vn = notes.validate_notes(pkg)
    for kind, report in vn.items():
        assert report["ok"], f"{kind} integrity broken: {report}"


# ------------------------------------------------------------------ the tests


def test_exactly_one_positioner_required(tmp_path):
    source = _fresh(tmp_path, "s.docx", ["content"])
    target = _build_target(tmp_path)
    pkg = DocxPackage(target)
    with pytest.raises(WordMcpError, match="exactly one positioner"):
        asm.insert_document(pkg, str(source))
    with pytest.raises(WordMcpError, match="exactly one positioner"):
        asm.insert_document(pkg, str(source), after_index=0, at_end=True)


def test_insert_into_self_refused(tmp_path):
    target = _build_target(tmp_path)
    pkg = DocxPackage(target)
    with pytest.raises(WordMcpError, match="into itself"):
        asm.insert_document(pkg, str(target), at_end=True)


def test_empty_source_refused(tmp_path):
    f = _fresh(tmp_path, "empty.docx", [])
    spkg = DocxPackage(f)
    body = spkg.body()
    for child in list(body):
        if etree.QName(child).localname != "sectPr":
            body.remove(child)
    spkg.mark_dirty()
    spkg.save(do_backup=False)
    target = _build_target(tmp_path)
    pkg = DocxPackage(target)
    with pytest.raises(TargetNotFound, match="no body content"):
        asm.insert_document(pkg, str(f), at_end=True)


def test_counts_reported(rich):
    _source, result, _target, _nums = rich
    assert result["paragraphs"] >= 10
    assert result["tables"] == 1
    assert result["images_carried"] == 1
    assert result["charts_carried"] == 1
    assert result["hyperlinks_carried"] == 1
    assert result["footnotes_carried"] == 1
    assert result["endnotes_carried"] == 1
    assert result["lists_carried"] == 2
    assert result["comments_stripped"] == 1
    assert result["bookmarks_carried"] >= 2
    n_items = result["paragraphs"] + result["tables"]
    start, end = result["body_item_range"]
    assert end - start + 1 == n_items
    assert result["position"]["mode"] == "at_end"
    # target had 3 paragraphs + 1 target list item = 4 body items before
    assert start == 4


def test_merged_result_parses_with_full_integrity(rich):
    _source, _result, target, _nums = rich
    _assert_integrity(DocxPackage(target))


def test_python_docx_roundtrip_and_order(rich):
    _source, _result, target, _nums = rich
    doc = Document(str(target))
    texts = [p.text for p in doc.paragraphs]
    assert "T0 intro" in texts and "S0 opening paragraph" in texts
    assert texts.index("T2 closing") < texts.index("S0 opening paragraph")
    assert texts.index("S0 opening paragraph") < texts.index(
        "S last paragraph"
    )
    assert len(doc.tables) == 1  # the carried merged-cell table


def test_style_name_collision_remaps_to_target(rich):
    _source, result, target, _nums = rich
    pkg = DocxPackage(target)
    styled = next(
        el
        for kind, _i, el in body_items(pkg)
        if kind == "paragraph"
        and paragraph_text(el) == "Styled special paragraph"
    )
    pstyle = styled.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
    assert pstyle.get(qn("w:val")) == "TgtSpecial"  # target formatting wins
    assert result["styles"]["remapped_ids"].get("SrcSpecial") == "TgtSpecial"
    named = [
        s
        for s in pkg.root("word/styles.xml").findall(qn("w:style"))
        if (s.find(qn("w:name")) is not None)
        and s.find(qn("w:name")).get(qn("w:val")) == "Special Style"
    ]
    assert len(named) == 1  # never cloned a duplicate of a name match


def test_unmatched_style_cloned(rich):
    _source, result, target, _nums = rich
    pkg = DocxPackage(target)
    cloned_names = {c["name"] for c in result["styles"]["cloned"]}
    assert "Source Only Style" in cloned_names
    ids = {
        s.get(qn("w:styleId"))
        for s in pkg.root("word/styles.xml").findall(qn("w:style"))
    }
    assert "SrcOnly" in ids
    styled = next(
        el
        for kind, _i, el in body_items(pkg)
        if kind == "paragraph"
        and paragraph_text(el) == "Cloned style paragraph"
    )
    assert (
        styled.find(f"{qn('w:pPr')}/{qn('w:pStyle')}").get(qn("w:val"))
        == "SrcOnly"
    )


def test_numbering_fresh_instances(rich):
    _source, _result, target, nums_before = rich
    pkg = DocxPackage(target)
    root = pkg.root("word/numbering.xml")
    nums = root.findall(qn("w:num"))
    # exactly the two carried instances were added, both with FRESH numIds
    # (the python-docx template ships built-in nums, so assert the delta)
    new_nids = {n.get(qn("w:numId")) for n in nums} - nums_before
    assert len(new_nids) == 2
    # The two inserted lists reference two DISTINCT numIds, both new.
    inserted_nids = set()
    for kind, _i, el in body_items(pkg):
        if kind != "paragraph":
            continue
        if paragraph_text(el) in ("one", "two", "uno", "dos"):
            nid = el.find(
                f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}"
            ).get(qn("w:val"))
            inserted_nids.add(nid)
    target_nid = next(
        el.find(f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}").get(
            qn("w:val")
        )
        for kind, _i, el in body_items(pkg)
        if kind == "paragraph" and paragraph_text(el) == "target item"
    )
    assert inserted_nids == new_nids
    assert target_nid not in inserted_nids


def test_notes_point_at_transplanted_definitions(rich):
    _source, _result, target, _nums = rich
    pkg = DocxPackage(target)
    for kind, part, ref_tag, note_tag, text in (
        (
            "footnote",
            "word/footnotes.xml",
            "w:footnoteReference",
            "w:footnote",
            "A carried footnote.",
        ),
        (
            "endnote",
            "word/endnotes.xml",
            "w:endnoteReference",
            "w:endnote",
            "A carried endnote.",
        ),
    ):
        refs = [
            r.get(qn("w:id")) for r in pkg.root().iter(qn(ref_tag))
        ]
        assert len(refs) == 1, f"expected one {kind} reference"
        defs = {
            n.get(qn("w:id")): n
            for n in pkg.root(part).findall(qn(note_tag))
        }
        assert refs[0] in defs
        carried = defs[refs[0]]
        note_text = "".join(
            t.text or "" for t in carried.iter(qn("w:t"))
        )
        assert text in note_text


def test_image_carried_and_rel_resolves(rich):
    _source, _result, target, _nums = rich
    pkg = DocxPackage(target)
    blips = list(
        pkg.root().iter(
            "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        )
    )
    assert len(blips) == 1
    rid = blips[0].get(f"{{{R_NS}}}embed")
    rels = {
        r.get("Id"): r for r in pkg.root("word/_rels/document.xml.rels")
    }
    assert rid in rels
    assert pkg.has_part("word/" + rels[rid].get("Target"))


def test_chart_carried_with_subtree(rich):
    _source, _result, target, _nums = rich
    pkg = DocxPackage(target)
    chart_parts = [
        p for p in pkg.part_names() if p.startswith("word/charts/chart")
        and p.endswith(".xml") and "_rels" not in p
    ]
    assert len(chart_parts) == 1
    chart = chart_parts[0]
    # its rels resolve (embedded workbook copied)
    rels_name = chart.rsplit("/", 1)[0] + "/_rels/" + chart.rsplit("/", 1)[1] + ".rels"
    assert pkg.has_part(rels_name)
    import posixpath

    for rel in pkg.root(rels_name):
        if rel.get("TargetMode") == "External":
            continue
        resolved = posixpath.normpath(
            posixpath.join(chart.rsplit("/", 1)[0], rel.get("Target"))
        )
        assert pkg.has_part(resolved)
    # content-type override present for the copied chart part
    overrides = {
        o.get("PartName")
        for o in pkg.root("[Content_Types].xml").findall(
            f"{{{CT_NS}}}Override"
        )
    }
    assert "/" + chart in overrides


def test_hyperlink_rel_external(rich):
    _source, _result, target, _nums = rich
    pkg = DocxPackage(target)
    link = next(pkg.root().iter(qn("w:hyperlink")))
    rid = link.get(qn("r:id"))
    rel = next(
        r
        for r in pkg.root("word/_rels/document.xml.rels")
        if r.get("Id") == rid
    )
    assert rel.get("TargetMode") == "External"
    assert rel.get("Target") == "https://example.com/x"


def test_bookmark_collision_renamed_and_ids_unique(rich):
    _source, result, target, _nums = rich
    pkg = DocxPackage(target)
    renames = result["bookmarks_renamed"]
    assert any(r["from"] == "SharedBM" for r in renames)
    starts = list(pkg.root().iter(qn("w:bookmarkStart")))
    names = [b.get(qn("w:name")) for b in starts]
    ids = [b.get(qn("w:id")) for b in starts]
    assert names.count("SharedBM") == 1  # the target's original only
    assert "SrcOnlyBM" in names  # non-colliding name kept
    assert len(ids) == len(set(ids)), "bookmark ids must be unique"
    # every start has a matching end
    end_ids = {b.get(qn("w:id")) for b in pkg.root().iter(qn("w:bookmarkEnd"))}
    assert set(ids) <= end_ids


def test_comments_stripped_cleanly(rich):
    _source, _result, target, _nums = rich
    pkg = DocxPackage(target)
    assert not list(pkg.root().iter(qn("w:commentReference")))
    assert not list(pkg.root().iter(qn("w:commentRangeStart")))
    # the commented text itself survives
    assert any(
        "COMMENTANCHOR" in paragraph_text(el)
        for kind, _i, el in body_items(pkg)
        if kind == "paragraph"
    )


def test_table_merges_carried(rich):
    _source, _result, target, _nums = rich
    pkg = DocxPackage(target)
    tbl = next(
        el for kind, _i, el in body_items(pkg) if kind == "table"
    )
    _grid, merges, _nested = dio._table_grid_data(tbl)
    assert {"row": 0, "col": 0, "rowspan": 2, "colspan": 1} in merges


def test_sectpr_never_carried(tmp_path):
    source = _fresh(tmp_path, "sect_src.docx", ["before break", "after break"])
    spkg = DocxPackage(source)
    # distinctive trailing sectPr + a mid-content section break
    trailing = spkg.body().find(qn("w:sectPr"))
    pgsz = trailing.find(qn("w:pgSz"))
    if pgsz is None:
        pgsz = etree.SubElement(trailing, qn("w:pgSz"))
    pgsz.set(qn("w:w"), "99999")
    first_p = next(iter(spkg.body().iter(qn("w:p"))))
    ppr = etree.SubElement(first_p, qn("w:pPr"))
    first_p.insert(0, ppr)
    etree.SubElement(ppr, qn("w:sectPr"))
    spkg.mark_dirty()
    spkg.save(do_backup=False)

    target = _build_target(tmp_path)
    before = DocxPackage(target)
    tgt_sect = before.body().find(qn("w:sectPr"))
    tgt_w = (
        tgt_sect.find(qn("w:pgSz")).get(qn("w:w"))
        if tgt_sect.find(qn("w:pgSz")) is not None
        else None
    )
    pkg = DocxPackage(target)
    result = asm.insert_document(pkg, str(source), at_end=True)
    pkg.save(do_backup=False)
    assert result["section_breaks_stripped"] == 1

    merged = DocxPackage(target)
    sects = merged.root().findall(f".//{qn('w:sectPr')}")
    assert len(sects) == 1, "only the target's own sectPr may remain"
    got = sects[0].find(qn("w:pgSz"))
    assert (got.get(qn("w:w")) if got is not None else None) == tgt_w
    assert not any(
        s.get(qn("w:w")) == "99999"
        for s in merged.root().iter(qn("w:pgSz"))
    )


def test_after_index_positions_by_body_item(tmp_path):
    source = _fresh(tmp_path, "pos_src.docx", ["INSERTED"])
    target = _build_target(tmp_path)
    pkg = DocxPackage(target)
    result = asm.insert_document(pkg, str(source), after_index=0)
    pkg.save(do_backup=False)
    doc = Document(str(target))
    texts = [p.text for p in doc.paragraphs]
    assert texts.index("INSERTED") == texts.index("T0 intro") + 1
    assert result["body_item_range"] == [1, 1]


def test_after_index_out_of_range(tmp_path):
    source = _fresh(tmp_path, "oor_src.docx", ["x"])
    pkg = DocxPackage(_build_target(tmp_path))
    with pytest.raises(TargetNotFound, match="body items"):
        asm.insert_document(pkg, str(source), after_index=99)


def test_after_anchor_unique_match(tmp_path):
    source = _fresh(tmp_path, "anchor_src.docx", ["CHAPTER BODY"])
    target = _build_target(tmp_path)
    pkg = DocxPackage(target)
    asm.insert_document(pkg, str(source), after_anchor="Chapter 4")
    pkg.save(do_backup=False)
    texts = [p.text for p in Document(str(target)).paragraphs]
    assert texts.index("CHAPTER BODY") == texts.index("Chapter 4") + 1


def test_after_anchor_ambiguity_refused_with_locations(tmp_path):
    source = _fresh(tmp_path, "amb_src.docx", ["x"])
    target = _fresh(
        tmp_path,
        "amb_tgt.docx",
        ["Chapter 4", "filler", "Chapter 4"],
    )
    pkg = DocxPackage(target)
    n_before = len(list(pkg.body()))
    with pytest.raises(AmbiguousTarget) as exc:
        asm.insert_document(pkg, str(source), after_anchor="Chapter 4")
    msg = str(exc.value)
    assert "2 paragraphs" in msg
    assert "'body_item_index': 0" in msg and "'body_item_index': 2" in msg
    assert len(list(pkg.body())) == n_before  # nothing half-applied


def test_after_anchor_requires_exact_paragraph_text(tmp_path):
    # substring matches must NOT count ("Chapter 4" inside body prose)
    source = _fresh(tmp_path, "exact_src.docx", ["x"])
    target = _fresh(
        tmp_path,
        "exact_tgt.docx",
        ["Chapter 4", "Chapter 4 applies the model to the case."],
    )
    pkg = DocxPackage(target)
    asm.insert_document(pkg, str(source), after_anchor="Chapter 4")
    pkg.save(do_backup=False)
    texts = [p.text for p in Document(str(target)).paragraphs]
    assert texts.index("x") == 1  # right after the exact-match heading


def test_ole_content_refused_before_any_mutation(tmp_path):
    source = _fresh(tmp_path, "ole_src.docx", ["host paragraph"])
    spkg = DocxPackage(source)
    p = next(iter(spkg.body().iter(qn("w:p"))))
    r = etree.SubElement(p, qn("w:r"))
    etree.SubElement(r, qn("w:object"))
    spkg.mark_dirty()
    spkg.save(do_backup=False)

    target = _build_target(tmp_path)
    pkg = DocxPackage(target)
    n_before = len(list(pkg.body()))
    styles_before = len(pkg.root("word/styles.xml").findall(qn("w:style")))
    with pytest.raises(UnsupportedStructure) as exc:
        asm.insert_document(pkg, str(source), at_end=True)
    msg = str(exc.value)
    assert "OLE" in msg and "Strip these from the source" in msg
    assert len(list(pkg.body())) == n_before
    assert (
        len(pkg.root("word/styles.xml").findall(qn("w:style")))
        == styles_before
    ), "refusal must precede all target mutation"


def test_unknown_embedded_part_refused(tmp_path):
    """A body element referencing a relationship of an unknown internal type
    (synthetic stand-in for SmartArt/OLE-package parts) refuses whole."""
    source = _fresh(tmp_path, "unk_src.docx", ["host paragraph"])
    spkg = DocxPackage(source)
    spkg.set_raw_part("word/weird/blob.bin", b"\x00\x01")
    rels_root = spkg.root("word/_rels/document.xml.rels")
    rel = etree.SubElement(rels_root, f"{{{REL_NS}}}Relationship")
    rel.set("Id", "rId999")
    rel.set("Type", "http://example.com/relationships/mystery")
    rel.set("Target", "weird/blob.bin")
    spkg.mark_dirty("word/_rels/document.xml.rels")
    p = next(iter(spkg.body().iter(qn("w:p"))))
    r = etree.SubElement(p, qn("w:r"))
    weird = etree.SubElement(r, qn("w:t"))
    weird.text = "x"
    p.set(f"{{{R_NS}}}dummy", "rId999")  # r:-namespaced pointer at the blob
    spkg.mark_dirty()
    spkg.save(do_backup=False)

    pkg = DocxPackage(_build_target(tmp_path))
    with pytest.raises(UnsupportedStructure, match="mystery"):
        asm.insert_document(pkg, str(source), at_end=True)


def test_docpr_ids_unique_after_double_insert(tmp_path):
    """Inserting a source with drawings into a target that already has
    drawings (a prior insert of the same source) must renumber docPr ids."""
    img = tmp_path / "p.png"
    img.write_bytes(_png())
    source = _fresh(tmp_path, "img_src.docx", ["pic host"])
    spkg = DocxPackage(source)
    media.add_image(spkg, str(img), at_end=True)
    spkg.save(do_backup=False)

    target = _build_target(tmp_path)
    for _ in range(2):
        pkg = DocxPackage(target)
        asm.insert_document(pkg, str(source), at_end=True)
        pkg.save(do_backup=False)
    merged = DocxPackage(target)
    ids = [d.get("id") for d in merged.root().iter(f"{{{WP}}}docPr")]
    assert len(ids) == 2
    assert len(ids) == len(set(ids))
    _assert_integrity(merged)


def test_full_integrity_after_positioned_insert(rich, tmp_path):
    """Insert the rich source MID-document (not at_end) and re-verify."""
    source, _result, _target, _nums = rich
    target = _build_target(tmp_path, "mid_target.docx")
    pkg = DocxPackage(target)
    result = asm.insert_document(pkg, str(source), after_index=1)
    pkg.save(do_backup=False)
    merged = DocxPackage(target)
    _assert_integrity(merged)
    texts = [p.text for p in Document(str(target)).paragraphs]
    assert texts.index("S0 opening paragraph") < texts.index("T2 closing")
    assert result["position"]["starts_at_body_item"] == 2


@pytest.mark.skipif(
    not (CORPUS / "ch1-3.docx").exists(),
    reason="tests/corpus/ch1-3.docx not present",
)
def test_corpus_insert_into_real_document(tmp_path):
    """Insert a small synthetic doc into a copy of a real corpus document
    (corpus files are never modified) and assert the validate_document
    criteria: package opens, notes consistent, fields balanced."""
    target = tmp_path / "corpus_copy.docx"
    shutil.copyfile(CORPUS / "ch1-3.docx", target)
    source = _fresh(
        tmp_path, "small_src.docx", ["Inserted synthetic paragraph."]
    )
    spkg = DocxPackage(source)
    notes.add_note(
        pkg=spkg,
        kind="footnote",
        anchor_text="synthetic",
        note_text="corpus-test footnote",
    )
    spkg.save(do_backup=False)

    pkg = DocxPackage(target)
    result = asm.insert_document(pkg, str(source), at_end=True)
    pkg.save(do_backup=False)
    assert result["paragraphs"] >= 1

    merged = DocxPackage(target)  # package_ok
    vn = notes.validate_notes(merged)
    for kind, report in vn.items():
        assert report["ok"], f"{kind} broken after corpus insert: {report}"
    xml = merged.raw_part("word/document.xml").decode(
        "utf-8", errors="replace"
    )
    assert xml.count('w:fldCharType="begin"') == xml.count(
        'w:fldCharType="end"'
    )
    Document(str(target))  # python-docx round-trip
