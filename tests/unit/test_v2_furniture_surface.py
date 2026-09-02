"""STAGED v2-surface rewrite (Wave D): structure/breaks/sections (2.5),
headers/footers/page furniture (2.6), images (2.7 slice), watermark and
protection (2.14).

Exercises the RENAMED/MERGED server surface per
integration/v2_briefs/wave_D.md: insert_break, set_section_properties
absorption, insert_list, set_header_footer, set_page_numbers, insert_image,
set_image, export_images, delete_element(type:"image"), set_watermark,
set_document_protection(protection:"none").

Skips wholesale until the Phase 2 integrator registers the v2 tools; the
ops layer beneath is unchanged and stays covered by tests/unit. Assertions
target reparsed document state, not envelope internals, so envelope
finalization does not churn this file. Original v1 test file untouched at
tests/unit/test_furniture_toc_media.py.
"""

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage, qn

pytestmark = pytest.mark.skipif(
    not hasattr(srv, "insert_break"),
    reason="v2 surface not yet registered (staged for the Phase 2 "
    "integrator)",
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_png(path: Path, w=8, h=4, rgb=(20, 90, 200)):
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(
            ">I", zlib.crc32(c) & 0xFFFFFFFF
        )

    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    return path


def fresh(tmp_path: Path, name="doc.docx") -> str:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Alpha paragraph.")
    doc.add_paragraph("Beta paragraph.")
    doc.save(str(f))
    return str(f)


def ok(result) -> bool:
    return not (isinstance(result, dict) and result.get("ok") is False)


# ----------------------------------------------------------- 2.5 insert_break


def test_insert_break_page(tmp_path):
    f = fresh(tmp_path)
    r = srv.insert_break(f, type="page", location={"paragraph": 0})
    assert ok(r)
    pkg = DocxPackage(f)
    brs = [
        b
        for b in pkg.root().iter(qn("w:br"))
        if b.get(qn("w:type")) == "page"
    ]
    assert len(brs) == 1


def test_insert_break_section_continuous(tmp_path):
    f = fresh(tmp_path)
    r = srv.insert_break(
        f, type="section_continuous", location={"paragraph": 0}
    )
    assert ok(r)
    pkg = DocxPackage(f)
    # A new sectPr appears inside a paragraph's pPr.
    in_ppr = [
        s
        for s in pkg.root().iter(qn("w:sectPr"))
        if s.getparent().tag == qn("w:pPr")
    ]
    assert len(in_ppr) == 1


def test_insert_break_bad_type_refuses(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(WordMcpError):
        srv.insert_break(f, type="column")


# --------------------------------------------- 2.5 set_section_properties


def test_section_properties_geometry_and_full_state(tmp_path):
    f = fresh(tmp_path)
    r = srv.set_section_properties(f, section=0, orientation="landscape")
    assert ok(r)
    r2 = srv.set_section_properties(f, section=0)  # read-back call
    assert ok(r2)


def test_section_properties_absorbs_columns(tmp_path):
    f = fresh(tmp_path)
    r = srv.set_section_properties(
        f, section=0, columns={"count": 2, "space_pt": 24}
    )
    assert ok(r)
    pkg = DocxPackage(f)
    cols = pkg.body().find(qn("w:sectPr") + "/" + qn("w:cols"))
    assert cols is not None and cols.get(qn("w:num")) == "2"


def test_section_properties_absorbs_line_numbering(tmp_path):
    f = fresh(tmp_path)
    r = srv.set_section_properties(
        f, section=0, line_numbering={"count_by": 5}
    )
    assert ok(r)
    pkg = DocxPackage(f)
    ln = pkg.body().find(qn("w:sectPr") + "/" + qn("w:lnNumType"))
    assert ln is not None and ln.get(qn("w:countBy")) == "5"

    r2 = srv.set_section_properties(f, section=0, line_numbering="none")
    assert ok(r2)
    pkg2 = DocxPackage(f)
    assert pkg2.body().find(qn("w:sectPr") + "/" + qn("w:lnNumType")) is None


# --------------------------------------------------------- 2.5 insert_list


def test_insert_list_at_location(tmp_path):
    f = fresh(tmp_path)
    r = srv.insert_list(
        f, ["one", "two"], kind="number", location={"paragraph": 0}
    )
    assert ok(r)
    pkg = DocxPackage(f)
    assert len(list(pkg.root().iter(qn("w:numPr")))) == 2


# ---------------------------------------------------- 2.6 set_header_footer


def test_set_header_footer_both_parts(tmp_path):
    f = fresh(tmp_path)
    assert ok(srv.set_header_footer(f, part="header", text="My Header"))
    assert ok(
        srv.set_header_footer(
            f, part="footer", text="My Footer", include_page_number=True
        )
    )
    pkg = DocxPackage(f)
    header_parts = [n for n in pkg.part_names() if "header" in n and n.endswith(".xml")]
    footer_parts = [n for n in pkg.part_names() if "footer" in n and n.endswith(".xml")]
    assert header_parts and footer_parts
    header_text = "".join(
        t.text or "" for t in pkg.root(header_parts[0]).iter(qn("w:t"))
    )
    assert "My Header" in header_text


def test_set_header_footer_bad_part_refuses(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(WordMcpError):
        srv.set_header_footer(f, part="margin", text="x")


# ----------------------------------------------------- 2.6 set_page_numbers


def test_set_page_numbers_insert_and_format(tmp_path):
    f = fresh(tmp_path)
    r = srv.set_page_numbers(
        f,
        position="footer",
        format={"number_format": "lowerRoman", "start_at": 1},
    )
    assert ok(r)
    pkg = DocxPackage(f)
    footer_parts = [n for n in pkg.part_names() if "footer" in n and n.endswith(".xml")]
    assert footer_parts
    pg = pkg.body().find(qn("w:sectPr") + "/" + qn("w:pgNumType"))
    assert pg is not None and pg.get(qn("w:fmt")) == "lowerRoman"


def test_set_page_numbers_needs_position_or_format(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(WordMcpError):
        srv.set_page_numbers(f)


# ------------------------------------------------ 2.7 images (Wave D slice)


def test_image_lifecycle_insert_set_export_delete(tmp_path):
    f = fresh(tmp_path)
    png = make_png(tmp_path / "pic.png")
    assert ok(srv.insert_image(f, str(png), location={"paragraph": 1}))

    r = srv.set_image(f, image_id=0, width_pt=120, alt_text="a blue box")
    assert ok(r)
    pkg = DocxPackage(f)
    extents = list(
        pkg.root().iter(
            "{http://schemas.openxmlformats.org/drawingml/2006/"
            "wordprocessingDrawing}extent"
        )
    )
    assert extents and int(extents[0].get("cx")) == 120 * 12700

    out = tmp_path / "exported"
    out.mkdir()
    r = srv.export_images(f, str(out))
    assert ok(r)
    assert list(out.glob("image0*"))

    r = srv.delete_element(f, type="image", id=0)
    assert ok(r)
    pkg2 = DocxPackage(f)
    assert not any(
        n.startswith("word/media/") for n in pkg2.part_names()
    )


def test_set_image_nothing_to_set_refuses(tmp_path):
    f = fresh(tmp_path)
    png = make_png(tmp_path / "pic.png")
    srv.insert_image(f, str(png))
    with pytest.raises(WordMcpError):
        srv.set_image(f, image_id=0)


# --------------------------------------------------------- 2.14 watermark


def test_set_watermark_and_none(tmp_path):
    f = fresh(tmp_path)
    srv.set_header_footer(f, part="header", text="keep me")
    assert ok(srv.set_watermark(f, watermark={"text": "DRAFT"}))
    pkg = DocxPackage(f)
    header_parts = [n for n in pkg.part_names() if "header" in n and n.endswith(".xml")]
    assert any(
        "DRAFT" in (pkg.raw_part(n).decode("utf-8", "ignore"))
        for n in header_parts
    )
    assert ok(srv.set_watermark(f, watermark="none"))
    pkg2 = DocxPackage(f)
    assert not any(
        "PowerPlusWaterMarkObject" in pkg2.raw_part(n).decode("utf-8", "ignore")
        or "DRAFT" in pkg2.raw_part(n).decode("utf-8", "ignore")
        for n in [m for m in pkg2.part_names() if "header" in m and m.endswith(".xml")]
    )


# -------------------------------------------------------- 2.14 protection


def test_protection_set_and_none(tmp_path):
    f = fresh(tmp_path)
    assert ok(srv.set_document_protection(f, protection="readOnly"))
    state = srv.get_protection(f)
    assert state["protected"] is True and state["edit"] == "readOnly"
    assert ok(srv.set_document_protection(f, protection="none"))
    state2 = srv.get_protection(f)
    assert state2["protected"] is False


def test_protection_none_with_password_refuses(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(WordMcpError):
        srv.set_document_protection(f, protection="none", password="x")
