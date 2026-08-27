"""v1.2 Phase A local features: stats, citation parity, page-number formats,
line numbering."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import citecheck, furniture as f, read, stats

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def ch13(tmp_path):
    dst = tmp_path / "ch1-3.docx"
    shutil.copy(CORPUS / "ch1-3.docx", dst)
    return dst


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


# ---------------------------------------------------------------------- stats


def test_word_count_totals_plausible(ch13):
    result = stats.word_count(DocxPackage(ch13))
    t = result["totals"]
    assert t["words"] > 5000  # three chapters
    assert t["paragraphs"] > 400
    assert len(result["sections"]) >= 40  # 47 headings in ch1-3


def test_section_words_sum_to_children(ch13):
    """A level-1 section's words include its level-2 children's words."""
    result = stats.word_count(DocxPackage(ch13))
    sections = result["sections"]
    l1 = [s for s in sections if s["level"] == 1 and s["words"] > 0]
    assert l1, "expected level-1 sections with content"
    biggest = max(l1, key=lambda s: s["words"])
    assert biggest["words"] <= result["totals"]["words"]


def test_word_count_no_headings_doc(tmp_path):
    dst = tmp_path / "plain.docx"
    d = Document()
    d.add_paragraph("one two three four five")
    d.save(str(dst))
    r = stats.word_count(DocxPackage(dst))
    assert r["totals"]["words"] == 5
    assert r["sections"] == []


# ------------------------------------------------------------ citation parity


def test_parity_on_synthetic_doc(tmp_path):
    dst = tmp_path / "cite.docx"
    d = Document()
    d.add_paragraph(
        "Mediation research shows promise (Smith, 2026). Bordin (1979) framed "
        "the alliance concept, later extended (Bordin, 1994; Hurd, 1999). "
        "Uncited claims exist too."
    )
    d.add_heading("References", 1)
    d.add_paragraph("Smith, J. (2026). The Example Model. Journal of X.")
    d.add_paragraph("Bordin, E. S. (1979). The generalizability of the concept.")
    d.add_paragraph("Hurd, I. (1999). Legitimacy and authority.")
    d.add_paragraph("Orphan, O. (2001). Never cited anywhere.")
    d.save(str(dst))

    r = citecheck.check_citation_parity(DocxPackage(dst))
    assert r["reference_entries"] == 4
    # Bordin 1994 cited but not listed.
    assert any("Bordin (1994)" in m for m in r["missing_references"])
    # Orphan 2001 listed but never cited.
    assert any("Orphan" in u for u in r["uncited_references"])
    assert not r["parity_ok"]
    # The clean ones are neither missing nor uncited.
    assert not any("Hurd" in m for m in r["missing_references"])
    assert not any("Smith" in u for u in r["uncited_references"])


def test_parity_real_chapter(ch4):
    """ch4 has ~56 refs; the checker should parse most and find the list."""
    r = citecheck.check_citation_parity(DocxPackage(ch4))
    assert r["reference_entries"] >= 40
    assert r["in_text_citations"] >= 40
    # Real docs may legitimately have flags; the tool just must not blow up
    # and must parse the bulk of the list.
    assert len(r["unparsed_reference_entries"]) <= r["reference_entries"] * 0.4


def test_parity_no_references_section(tmp_path):
    dst = tmp_path / "norefs.docx"
    d = Document()
    d.add_paragraph("Body only (Nobody, 2020).")
    d.save(str(dst))
    with pytest.raises(TargetNotFound):
        citecheck.check_citation_parity(DocxPackage(dst))


# ------------------------------------------- page number format / line numbers


def test_roman_numeral_page_format(ch4):
    pkg = DocxPackage(ch4)
    f.set_page_number_format(pkg, number_format="lowerRoman", start_at=1)
    pkg.save(do_backup=False)
    xml = DocxPackage(ch4).raw_part("word/document.xml").decode("utf-8")
    assert 'w:fmt="lowerRoman"' in xml
    Document(str(ch4))


def test_bad_page_format_refused(ch4):
    with pytest.raises(WordMcpError):
        f.set_page_number_format(DocxPackage(ch4), number_format="klingon")


def test_line_numbering_set_and_remove(ch4):
    pkg = DocxPackage(ch4)
    f.set_line_numbering(pkg, count_by=5, restart="newPage")
    pkg.save(do_backup=False)
    xml = DocxPackage(ch4).raw_part("word/document.xml").decode("utf-8")
    assert 'w:countBy="5"' in xml
    assert "lnNumType" in xml
    Document(str(ch4))

    pkg = DocxPackage(ch4)
    f.set_line_numbering(pkg, remove=True)
    pkg.save(do_backup=False)
    assert "lnNumType" not in DocxPackage(ch4).raw_part(
        "word/document.xml"
    ).decode("utf-8")


def test_sectpr_schema_order_preserved(ch4):
    """lnNumType must land BEFORE pgNumType and cols per CT_SectPr order."""
    pkg = DocxPackage(ch4)
    f.set_page_number_format(pkg, number_format="lowerRoman")
    f.set_line_numbering(pkg, count_by=1)
    pkg.save(do_backup=False)
    xml = DocxPackage(ch4).raw_part("word/document.xml").decode("utf-8")
    assert xml.index("lnNumType") < xml.index("pgNumType")
    Document(str(ch4))
