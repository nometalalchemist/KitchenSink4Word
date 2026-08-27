"""Regression tests for the 2026-08-27 stress-test findings (v1.0.2)."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import notes, read, tables as tb, text as tx

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def doc(tmp_path):
    """Fresh doc with a table and a footnoted paragraph."""
    dst = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("Intro text before everything else here.")
    d.add_paragraph("Text with anchor one and anchor two present.")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tb.create_table(pkg, [["Header"], ["Solo cell"]], at_end=True)
    pkg.save(do_backup=False)
    return dst


# Finding 1: format_text must reach table cells


def test_format_text_finds_table_cell_text(doc):
    pkg = DocxPackage(doc)
    result = tx.format_text(pkg, find="Solo cell", formatting={"bold": True})
    pkg.save(do_backup=False)
    assert result["formatted"].get("location") == "table cell"
    d = Document(str(doc))
    cell = d.tables[0].rows[1].cells[0]
    assert any(r.bold for p in cell.paragraphs for r in p.runs)


def test_format_text_error_mentions_search_scope(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(Exception) as exc:
        tx.format_text(pkg, find="nowhere at all", formatting={"bold": True})
    assert "table cells" in str(exc.value)


# Finding 2: content deletion purges orphan note definitions


def test_delete_paragraph_purges_footnote_definitions(doc):
    pkg = DocxPackage(doc)
    notes.add_note(pkg, "footnote", anchor_text="anchor one", note_text="First.")
    notes.add_note(pkg, "footnote", anchor_text="anchor two", note_text="Second.")
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    target = next(
        p["index"]
        for p in read.get_paragraphs(pkg2)
        if "anchor one" in p["text"]
    )
    result = tx.delete_paragraphs(pkg2, target)
    pkg2.save(do_backup=False)
    assert "note_definitions_purged" in result

    v = notes.validate_notes(DocxPackage(doc))
    assert v["footnotes"]["orphan_definitions"] == []
    assert v["footnotes"]["needs_cleanup"] is False


def test_set_cells_purges_footnote_in_cell(doc):
    pkg = DocxPackage(doc)
    notes.add_note(pkg, "footnote", anchor_text="Solo cell", note_text="In cell.")
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    result = tb.set_cells(pkg2, 0, [{"row": 1, "cell": 0, "text": "overwritten"}])
    pkg2.save(do_backup=False)
    assert "note_definitions_purged" in result
    assert notes.validate_notes(DocxPackage(doc))["footnotes"]["needs_cleanup"] is False


def test_validate_notes_flags_needs_cleanup(doc):
    """External orphan (not created by our tools) -> needs_cleanup True,
    purge_orphans removes it."""
    pkg = DocxPackage(doc)
    notes.add_note(pkg, "footnote", anchor_text="anchor one", note_text="x")
    # Simulate an externally-orphaned definition: remove the body ref only.
    from word_mcp.core.package import qn

    ref = next(pkg.root().iter(qn("w:footnoteReference")))
    run = ref.getparent()
    run.getparent().remove(run)
    pkg.mark_dirty()
    v = notes.validate_notes(pkg)
    assert v["footnotes"]["needs_cleanup"] is True
    assert v["footnotes"]["ok"] is True  # not corruption
    purged = notes.purge_orphans(pkg)
    assert purged["purged"]["footnotes"]
    assert notes.validate_notes(pkg)["footnotes"]["needs_cleanup"] is False


# Finding 4: single-cell merge refused


def test_single_cell_merge_refused(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="single cell"):
        tb.merge_cells(
            pkg, 0, start_row=0, end_row=0, start_col=0, end_col=0
        )


# Finding 5: superscript + subscript conflict


def test_superscript_subscript_conflict_refused(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="mutually exclusive"):
        tx.format_text(
            pkg,
            find="Intro",
            formatting={"superscript": True, "subscript": True},
        )


def test_superscript_false_subscript_true_ok(doc):
    pkg = DocxPackage(doc)
    tx.format_text(
        pkg,
        find="Intro",
        formatting={"superscript": False, "subscript": True},
    )
    pkg.save(do_backup=False)
    Document(str(doc))


# Finding 6: max_replacements guard


def test_max_replacements_guard_aborts_atomically(doc):
    before = "\n".join(
        p["text"] for p in read.get_paragraphs(DocxPackage(doc))
    )
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="max_replacements"):
        tx.search_and_replace(
            pkg,
            [{"find": ".+", "replace": "DESTROYED", "regex": True}],
            max_replacements=1,
        )
    pkg.save(do_backup=False)
    after = "\n".join(p["text"] for p in read.get_paragraphs(DocxPackage(doc)))
    assert after == before, "guard must abort with zero changes applied"


def test_max_replacements_allows_under_limit(doc):
    pkg = DocxPackage(doc)
    result = tx.search_and_replace(
        pkg,
        [{"find": "anchor", "replace": "ANCHOR"}],
        max_replacements=10,
    )
    assert result["total"] == 2
