"""BUSINESS bundle: mail merge / templates, batch operations, form fields."""

import csv

import pytest
from docx import Document
from lxml import etree

import word_mcp.server as srv
from word_mcp.core.errors import AmbiguousTarget, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import batch as bt, forms as fm, mailmerge as mm, read as rd

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"


def _els(xml_fragment: str) -> list:
    root = etree.fromstring(
        f'<w:root xmlns:w="{W}" xmlns:w14="{W14}">{xml_fragment}</w:root>'
    )
    return list(root)


def _append_body_xml(pkg: DocxPackage, xml_fragment: str) -> None:
    body = pkg.body()
    sectpr = body.find(qn("w:sectPr"))
    for el in _els(xml_fragment):
        if sectpr is not None:
            sectpr.addprevious(el)
        else:
            body.append(el)
    pkg.mark_dirty()


def _body_text(path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


# ------------------------------------------------------------------ templates


@pytest.fixture
def template(tmp_path):
    path = tmp_path / "template.docx"
    srv.create_document(str(path))
    srv.insert_paragraphs(
        str(path),
        [
            {"text": "Dear {{name}},"},
            {"text": "Your order {{order_id}} ships to {{city}}."},
            {"text": "Regards, {{name}}"},
        ],
        at_end=True,
        backup=False,
    )
    # Deliberately split the first placeholder across runs: bolding "{{na"
    # fragments "Dear {{name}}," into three runs.
    srv.format_text(str(path), {"bold": True}, find="{{na", backup=False)
    return path


def test_template_fixture_placeholder_is_split_across_runs(template):
    para = next(
        p for p in Document(str(template)).paragraphs
        if p.text == "Dear {{name}},"
    )
    assert any(r.text == "{{na" for r in para.runs), (
        "fixture must fragment the placeholder to prove runmap-safety"
    )


def test_list_placeholders_finds_split_and_counts(template):
    res = mm.list_template_placeholders(DocxPackage(template))
    assert res["names"] == ["city", "name", "order_id"]
    by_name = {p["name"]: p for p in res["placeholders"]}
    assert by_name["name"]["count"] == 2
    assert by_name["city"]["count"] == 1
    assert by_name["name"]["locations"][0]["part"] == "word/document.xml"


def test_fill_template_runmap_safe_and_formatting_kept(template):
    pkg = DocxPackage(template)
    res = mm.fill_template(
        pkg, {"name": "Ada", "order_id": 42, "city": "Seoul"}
    )
    pkg.save(do_backup=False)
    assert res["placeholders_replaced"] == {"name": 2, "order_id": 1, "city": 1}
    text = _body_text(template)
    assert "Dear Ada," in text
    assert "Your order 42 ships to Seoul." in text
    assert "Regards, Ada" in text
    assert "{{" not in text
    para = next(
        p for p in Document(str(template)).paragraphs if p.text == "Dear Ada,"
    )
    assert any(r.bold and "Ada" in r.text for r in para.runs), (
        "replacement must inherit the first replaced run's bold"
    )


def test_fill_template_missing_error_is_atomic(template):
    pkg = DocxPackage(template)
    with pytest.raises(WordMcpError, match="city.*order_id|order_id.*city"):
        mm.fill_template(pkg, {"name": "Ada"})
    texts = "\n".join(e["text"] for e in rd.get_paragraphs(pkg))
    assert "{{name}}" in texts and "{{city}}" in texts and "{{order_id}}" in texts


def test_fill_template_missing_skip_and_empty(template):
    pkg = DocxPackage(template)
    res = mm.fill_template(pkg, {"name": "Ada"}, missing="skip")
    assert res["skipped"] == ["city", "order_id"]
    texts = "\n".join(e["text"] for e in rd.get_paragraphs(pkg))
    assert "Dear Ada," in texts and "{{city}}" in texts

    pkg2 = DocxPackage(template)
    res2 = mm.fill_template(pkg2, {"name": "Ada"}, missing="empty")
    assert sorted(res2["filled_empty"]) == ["city", "order_id"]
    texts2 = "\n".join(e["text"] for e in rd.get_paragraphs(pkg2))
    assert "Your order  ships to ." in texts2


def test_fill_template_rejects_control_chars(template):
    pkg = DocxPackage(template)
    with pytest.raises(WordMcpError, match="control characters"):
        mm.fill_template(
            pkg, {"name": "A\x07B", "order_id": "1", "city": "X"}
        )


def test_mergefields_listed_and_filled(template):
    pkg = DocxPackage(template)
    # Simple field.
    _append_body_xml(
        pkg,
        '<w:p><w:r><w:t>Simple: </w:t></w:r>'
        '<w:fldSimple w:instr=" MERGEFIELD  city  \\* MERGEFORMAT ">'
        '<w:r><w:t>&#171;city&#187;</w:t></w:r></w:fldSimple></w:p>',
    )
    # Complex field (begin / instrText / separate / result / end).
    _append_body_xml(
        pkg,
        '<w:p><w:r><w:t>Complex: </w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> MERGEFIELD name </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:t>&#171;name&#187;</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>',
    )
    pkg.save(do_backup=False)

    pkg = DocxPackage(template)
    listed = mm.list_template_placeholders(pkg)
    mf = {m["name"]: m for m in listed["mergefields"]}
    assert mf["city"]["kind"] == "simple"
    assert mf["name"]["kind"] == "complex"

    res = mm.fill_template(pkg, {"name": "Ada", "order_id": "1", "city": "Seoul"})
    pkg.save(do_backup=False)
    assert res["mergefields_replaced"] == {"city": 1, "name": 1}
    text = _body_text(template)
    assert "Simple: Seoul" in text and "Complex: Ada" in text
    xml = DocxPackage(template).raw_part("word/document.xml").decode("utf-8")
    assert "MERGEFIELD" not in xml


# ----------------------------------------------------------------- mail merge

ROWS = [
    {"name": "Ada", "order_id": "1", "city": "Seoul"},
    {"name": "Bob", "order_id": "2", "city": "Busan"},
    {"name": "Cyn", "order_id": "3", "city": "Daegu"},
]


def test_mail_merge_from_list_of_dicts(template, tmp_path):
    out = tmp_path / "out"
    res = mm.mail_merge(
        str(template), ROWS, str(out),
        filename_pattern="{row_index}_{name}.docx",
    )
    assert res["rows"] == 3
    assert len(res["outputs"]) == 3
    assert (out / "1_Ada.docx").is_file() and (out / "3_Cyn.docx").is_file()
    assert "Dear Bob," in _body_text(out / "2_Bob.docx")
    # Template itself untouched.
    assert "{{name}}" in _body_text(template)


def test_mail_merge_refuses_collisions_before_writing(template, tmp_path):
    out = tmp_path / "out"
    mm.mail_merge(
        str(template), ROWS, str(out),
        filename_pattern="{row_index}_{name}.docx",
    )
    before = sorted(f.name for f in out.iterdir())
    with pytest.raises(WordMcpError, match="refusing to overwrite"):
        mm.mail_merge(
            str(template), ROWS, str(out),
            filename_pattern="{row_index}_{name}.docx",
        )
    assert sorted(f.name for f in out.iterdir()) == before


def test_mail_merge_refuses_duplicate_output_names(template, tmp_path):
    out = tmp_path / "dupes"
    rows = [dict(r, city="Same") for r in ROWS]
    with pytest.raises(WordMcpError, match="same file"):
        mm.mail_merge(
            str(template), rows, str(out), filename_pattern="{city}.docx"
        )
    assert not any(out.iterdir())


def test_mail_merge_missing_error_refuses_before_writing(template, tmp_path):
    out = tmp_path / "missing"
    rows = [{"name": "Ada", "order_id": "1"}]  # no city
    with pytest.raises(WordMcpError, match="row 1 lacks.*city"):
        mm.mail_merge(str(template), rows, str(out))
    # Refusal precedes even the output dir creation.
    assert not out.exists() or not any(out.iterdir())


def test_mail_merge_from_csv(template, tmp_path):
    data = tmp_path / "rows.csv"
    with open(data, "w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=["name", "order_id", "city"])
        writer.writeheader()
        writer.writerows(ROWS)
    out = tmp_path / "csvout"
    res = mm.mail_merge(
        str(template), str(data), str(out),
        filename_pattern="order_{order_id}.docx",
    )
    assert res["rows"] == 3
    assert "ships to Daegu" in _body_text(out / "order_3.docx")


def test_mail_merge_sanitizes_hostile_filenames(template, tmp_path):
    out = tmp_path / "hostile"
    rows = [{"name": 'A/d:a*"?', "order_id": "1", "city": "Seoul"}]
    res = mm.mail_merge(
        str(template), rows, str(out), filename_pattern="{name}.docx"
    )
    from pathlib import Path

    assert Path(res["outputs"][0]).name == "A_d_a___.docx"
    assert Path(res["outputs"][0]).is_file()


# ---------------------------------------------------------------------- batch


def _make_doc(path, n_paragraphs):
    d = Document()
    d.add_paragraph("alpha lives here")
    for i in range(n_paragraphs - 1):
        d.add_paragraph(f"filler paragraph {i}")
    d.save(str(path))
    return str(path)


BATCH_OPS = [
    {"tool": "search_and_replace",
     "params": {"replacements": [{"find": "alpha", "replace": "beta"}]}},
    {"tool": "delete_paragraphs", "params": {"start": 5}},
    {"tool": "search_and_replace",
     "params": {"replacements": [{"find": "beta", "replace": "gamma"}]}},
]


@pytest.fixture
def batch_files(tmp_path):
    a = _make_doc(tmp_path / "a.docx", 7)
    b = _make_doc(tmp_path / "b.docx", 3)  # delete_paragraphs(5) fails here
    c = _make_doc(tmp_path / "c.docx", 7)
    return a, b, c


def test_batch_apply_stop_on_error(batch_files):
    a, b, c = batch_files
    res = bt.batch_apply([a, b, c], BATCH_OPS, stop_on_error=True, backup=False)
    assert res["saved"] == [a]
    assert res["failed"] == [b]
    assert res["not_attempted"] == [c]
    assert "gamma" in _body_text(a)
    # Failed file: all-or-nothing — op 1 succeeded in memory but nothing saved.
    assert "alpha" in _body_text(b) and "beta" not in _body_text(b)
    assert "alpha" in _body_text(c)
    b_entry = res["files"][1]
    assert b_entry["ok"] is False
    assert b_entry["operations"][1]["ok"] is False
    assert "NOT saved" in b_entry["error"]
    assert "already saved" in res["note"]


def test_batch_apply_continue_on_error(batch_files):
    a, b, c = batch_files
    res = bt.batch_apply(
        [a, b, c], BATCH_OPS, stop_on_error=False, backup=False
    )
    assert res["saved"] == [a, c]
    assert res["failed"] == [b]
    assert res["not_attempted"] == []
    assert "gamma" in _body_text(a) and "gamma" in _body_text(c)
    assert "alpha" in _body_text(b)


def test_batch_apply_one_save_one_backup_per_file(batch_files, tmp_path):
    a, _, c = batch_files
    bt.batch_apply([a, c], BATCH_OPS, backup=True)
    assert len(list(tmp_path.glob("a.bak-*.docx"))) == 1
    assert len(list(tmp_path.glob("c.bak-*.docx"))) == 1


def test_batch_apply_rejects_unknown_tool_before_work(batch_files):
    a, b, c = batch_files
    with pytest.raises(WordMcpError, match="not batchable.*search_and_replace"):
        bt.batch_apply(
            [a, b, c],
            [{"tool": "com_export_pdf", "params": {}}],
            backup=False,
        )
    assert "alpha" in _body_text(a)


def test_batch_apply_validates_files_exist_first(batch_files, tmp_path):
    a, b, c = batch_files
    ghost = str(tmp_path / "nope.docx")
    with pytest.raises(WordMcpError, match="missing files"):
        bt.batch_apply([a, ghost, c], BATCH_OPS, backup=False)
    assert "alpha" in _body_text(a), "no file may be touched on refusal"


def test_batch_apply_wires_furniture_and_properties(tmp_path):
    a = _make_doc(tmp_path / "f1.docx", 2)
    b = _make_doc(tmp_path / "f2.docx", 2)
    ops = [
        {"tool": "set_footer",
         "params": {"text": "Confidential", "alignment": "left"}},
        {"tool": "set_document_properties", "params": {"author": "Batch Bot"}},
        {"tool": "add_watermark", "params": {"text": "DRAFT"}},
    ]
    res = bt.batch_apply([a, b], ops, backup=False)
    assert res["saved"] == [a, b]
    for path in (a, b):
        doc = Document(path)
        assert doc.core_properties.author == "Batch Bot"
        assert doc.sections[0].footer.paragraphs[0].text == "Confidential"


# ---------------------------------------------------------------- form fields

_FORM_XML = (
    # Modern: plain-text control, placeholder showing.
    '<w:p><w:sdt><w:sdtPr><w:alias w:val="Full Name"/>'
    '<w:tag w:val="full_name"/><w:id w:val="1001"/><w:showingPlcHdr/>'
    '<w:text/></w:sdtPr><w:sdtContent><w:r><w:rPr><w:i/></w:rPr>'
    '<w:t>Click here to enter text.</w:t></w:r></w:sdtContent></w:sdt></w:p>'
    # Modern: checkbox (w14), unchecked.
    '<w:p><w:sdt><w:sdtPr><w:tag w:val="agree"/><w:id w:val="1002"/>'
    '<w14:checkbox><w14:checked w14:val="0"/>'
    '<w14:checkedState w14:val="2612" w14:font="MS Gothic"/>'
    '<w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/></w14:checkbox>'
    '</w:sdtPr><w:sdtContent><w:r><w:t>&#9744;</w:t></w:r></w:sdtContent>'
    '</w:sdt></w:p>'
    # Modern: dropdown, placeholder showing.
    '<w:p><w:sdt><w:sdtPr><w:tag w:val="dept"/><w:id w:val="1003"/>'
    '<w:showingPlcHdr/><w:dropDownList>'
    '<w:listItem w:displayText="Engineering" w:value="ENG"/>'
    '<w:listItem w:displayText="Sales" w:value="SLS"/></w:dropDownList>'
    '</w:sdtPr><w:sdtContent><w:r><w:t>Choose an item.</w:t></w:r>'
    '</w:sdtContent></w:sdt></w:p>'
    # Legacy: FORMTEXT, empty display.
    '<w:p><w:r><w:fldChar w:fldCharType="begin"><w:ffData>'
    '<w:name w:val="applicant"/><w:enabled/><w:textInput/></w:ffData>'
    '</w:fldChar></w:r>'
    '<w:r><w:instrText xml:space="preserve"> FORMTEXT </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    '<w:r><w:t xml:space="preserve">     </w:t></w:r>'
    '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    # Legacy: FORMCHECKBOX (no separate/result — normal for checkboxes).
    '<w:p><w:r><w:fldChar w:fldCharType="begin"><w:ffData>'
    '<w:name w:val="subscribed"/><w:checkBox><w:sizeAuto/>'
    '<w:default w:val="0"/></w:checkBox></w:ffData></w:fldChar></w:r>'
    '<w:r><w:instrText> FORMCHECKBOX </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    # Legacy: FORMDROPDOWN.
    '<w:p><w:r><w:fldChar w:fldCharType="begin"><w:ffData>'
    '<w:name w:val="color"/><w:ddList><w:result w:val="0"/>'
    '<w:listEntry w:val="Red"/><w:listEntry w:val="Blue"/></w:ddList>'
    '</w:ffData></w:fldChar></w:r>'
    '<w:r><w:instrText> FORMDROPDOWN </w:instrText></w:r>'
    '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
    '<w:r><w:t>Red</w:t></w:r>'
    '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
)


@pytest.fixture
def form_doc(tmp_path):
    path = tmp_path / "form.docx"
    srv.create_document(str(path))
    pkg = DocxPackage(path)
    _append_body_xml(pkg, _FORM_XML)
    pkg.save(do_backup=False)
    return path


def test_list_form_fields(form_doc):
    res = fm.list_form_fields(DocxPackage(form_doc))
    fields = {f["name"]: f for f in res["fields"]}
    assert res["count"] == 6
    assert fields["full_name"]["kind"] == "sdt_text"
    assert fields["full_name"]["placeholder_showing"] is True
    assert fields["full_name"]["alias"] == "Full Name"
    assert fields["agree"]["kind"] == "sdt_checkbox"
    assert fields["agree"]["value"] is False
    assert fields["dept"]["options"] == ["Engineering", "Sales"]
    assert fields["applicant"]["kind"] == "legacy_text"
    assert fields["subscribed"]["kind"] == "legacy_checkbox"
    assert fields["subscribed"]["value"] is False
    assert fields["color"]["kind"] == "legacy_dropdown"
    assert fields["color"]["value"] == "Red"
    assert fields["color"]["options"] == ["Red", "Blue"]


def test_fill_form_fields_all_kinds(form_doc):
    pkg = DocxPackage(form_doc)
    res = fm.fill_form_fields(
        pkg,
        {
            "full_name": "Ada Lovelace",
            "agree": True,
            "dept": "Sales",
            "applicant": "Bob",
            "subscribed": True,
            "color": "Blue",
        },
    )
    assert res["count"] == 6
    pkg.save(do_backup=False)

    fields = {
        f["name"]: f
        for f in fm.list_form_fields(DocxPackage(form_doc))["fields"]
    }
    assert fields["full_name"]["value"] == "Ada Lovelace"
    assert fields["full_name"]["placeholder_showing"] is False
    assert fields["agree"]["value"] is True
    assert fields["dept"]["value"] == "Sales"
    assert fields["applicant"]["value"] == "Bob"
    assert fields["subscribed"]["value"] is True
    assert fields["color"]["value"] == "Blue"
    xml = DocxPackage(form_doc).raw_part("word/document.xml").decode("utf-8")
    assert "☒" in xml, "checkbox glyph must be updated with the state"
    assert 'w14:val="1"' in xml


def test_fill_dropdown_invalid_value_refuses_atomically(form_doc):
    pkg = DocxPackage(form_doc)
    with pytest.raises(WordMcpError, match="not an option"):
        fm.fill_form_fields(
            pkg, {"full_name": "Ada", "dept": "Marketing"}
        )
    # Validation precedes mutation: full_name must be untouched too.
    fields = {f["name"]: f for f in fm.list_form_fields(pkg)["fields"]}
    assert fields["full_name"]["placeholder_showing"] is True
    with pytest.raises(WordMcpError, match="not an option"):
        fm.fill_form_fields(pkg, {"color": "Green"})


def test_fill_checkbox_requires_boolean(form_doc):
    pkg = DocxPackage(form_doc)
    with pytest.raises(WordMcpError, match="checkbox"):
        fm.fill_form_fields(pkg, {"agree": "yes"})


def test_fill_unknown_name_error_and_skip(form_doc):
    pkg = DocxPackage(form_doc)
    with pytest.raises(WordMcpError, match="no form field matches.*ghost"):
        fm.fill_form_fields(pkg, {"ghost": "x"})
    res = fm.fill_form_fields(
        pkg, {"ghost": "x", "applicant": "Bob"}, missing="skip"
    )
    assert res["skipped_unknown"] == ["ghost"]
    assert res["filled"] == {"applicant": "legacy_text"}


def test_fill_duplicate_name_refused(tmp_path):
    path = tmp_path / "dup.docx"
    srv.create_document(str(path))
    pkg = DocxPackage(path)
    one = (
        '<w:p><w:sdt><w:sdtPr><w:tag w:val="dup"/><w:text/></w:sdtPr>'
        '<w:sdtContent><w:r><w:t>a</w:t></w:r></w:sdtContent></w:sdt></w:p>'
    )
    _append_body_xml(pkg, one + one)
    pkg.save(do_backup=False)
    pkg = DocxPackage(path)
    with pytest.raises(AmbiguousTarget, match="2 form fields match 'dup'"):
        fm.fill_form_fields(pkg, {"dup": "value"})


def test_validate_form_completeness(form_doc):
    pkg = DocxPackage(form_doc)
    res = fm.validate_form_completeness(pkg)
    assert res["complete"] is False
    unfilled = {f["name"] for f in res["unfilled"]}
    # Placeholder-showing controls and the blank legacy text are unfilled;
    # unchecked boxes only count when explicitly required.
    assert unfilled == {"full_name", "dept", "applicant"}

    req = fm.validate_form_completeness(pkg, required=["agree", "nothere"])
    assert req["complete"] is False
    assert req["missing_fields"] == ["nothere"]
    assert {f["name"] for f in req["unfilled"]} == {"agree"}

    fm.fill_form_fields(
        pkg,
        {
            "full_name": "Ada",
            "agree": True,
            "dept": "Engineering",
            "applicant": "Bob",
            "subscribed": True,
            "color": "Blue",
        },
    )
    done = fm.validate_form_completeness(pkg)
    assert done["complete"] is True
    assert fm.validate_form_completeness(pkg, required=["agree"])["complete"]
