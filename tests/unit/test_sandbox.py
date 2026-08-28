"""Path sandboxing (KS4W_ALLOWED_ROOTS): opt-in containment for every path
the server touches.

Self-contained: documents are built on the fly with python-docx. The env var
is manipulated through monkeypatch; the sandbox module re-parses whenever the
raw env value changes, so no reload tricks are needed.
"""

import os
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core import sandbox
from word_mcp.core.package import DocxPackage
from word_mcp.core.sandbox import SandboxViolation, check_path
from word_mcp.ops import backups as bk, dataio as dio, media

ENV = sandbox.ENV_VAR


# ------------------------------------------------------------------ fixtures


def _fresh_doc(where: Path, name: str = "doc.docx") -> Path:
    f = where / name
    doc = Document()
    doc.add_paragraph("Anchor paragraph so the body is never empty.")
    doc.save(str(f))
    return f


@pytest.fixture()
def unsandboxed(monkeypatch):
    monkeypatch.delenv(ENV, raising=False)


@pytest.fixture()
def root(tmp_path, monkeypatch):
    """One allowed root at tmp_path/inside; a sibling escape dir beside it."""
    inside = tmp_path / "inside"
    inside.mkdir()
    (tmp_path / "outside").mkdir()
    monkeypatch.setenv(ENV, str(inside))
    return inside


# ------------------------------------------------------- unset = unrestricted


class TestUnset:
    def test_check_path_is_identity_when_unset(self, unsandboxed):
        weird = r"C:\definitely\not\a\real\place\..\thing.docx"
        assert check_path(weird, "test") == weird
        assert sandbox.active() is False

    def test_empty_value_means_unrestricted(self, monkeypatch):
        monkeypatch.setenv(ENV, "")
        assert sandbox.active() is False
        monkeypatch.setenv(ENV, f" {os.pathsep} ")
        assert sandbox.active() is False

    def test_docx_open_anywhere_when_unset(self, unsandboxed, tmp_path):
        f = _fresh_doc(tmp_path)
        pkg = DocxPackage(f)
        assert pkg.body() is not None


# --------------------------------------------------------------- containment


class TestContainment:
    def test_inside_root_allowed(self, root):
        p = check_path(root / "sub" / "a.docx", "test")
        assert os.path.normcase(p).startswith(os.path.normcase(str(root)))

    def test_root_itself_allowed(self, root):
        assert check_path(root, "test")

    def test_outside_blocked(self, root, tmp_path):
        with pytest.raises(SandboxViolation):
            check_path(tmp_path / "outside" / "a.docx", "test")

    def test_traversal_escape_blocked(self, root, tmp_path):
        sneaky = root / "sub" / ".." / ".." / "outside" / "a.docx"
        with pytest.raises(SandboxViolation):
            check_path(sneaky, "test")

    def test_prefix_collision_blocked(self, tmp_path, monkeypatch):
        docs = tmp_path / "Documents"
        docs2 = tmp_path / "Documents2"
        docs.mkdir()
        docs2.mkdir()
        monkeypatch.setenv(ENV, str(docs))
        assert check_path(docs / "a.docx", "test")
        with pytest.raises(SandboxViolation):
            check_path(docs2 / "a.docx", "test")

    def test_case_differences_match(self, root):
        if os.path.normcase("A") != os.path.normcase("a"):
            pytest.skip("case-sensitive filesystem")
        assert check_path(str(root).upper() + os.sep + "a.docx", "test")

    def test_multi_root(self, tmp_path, monkeypatch):
        r1 = tmp_path / "r1"
        r2 = tmp_path / "r2"
        r1.mkdir()
        r2.mkdir()
        monkeypatch.setenv(ENV, os.pathsep.join([str(r1), str(r2)]))
        assert check_path(r1 / "a.docx", "test")
        assert check_path(r2 / "b.docx", "test")
        with pytest.raises(SandboxViolation):
            check_path(tmp_path / "elsewhere.docx", "test")

    def test_unc_refused_for_local_roots(self, root):
        with pytest.raises(SandboxViolation) as exc:
            check_path(r"\\some-server\share\doc.docx", "test")
        assert "UNC" in str(exc.value)

    def test_extended_length_unc_refused(self, root):
        with pytest.raises(SandboxViolation):
            check_path(r"\\?\UNC\some-server\share\doc.docx", "test")

    def test_extended_length_local_normalized(self, root):
        p = check_path("\\\\?\\" + str(root / "a.docx"), "test")
        assert not p.startswith("\\\\?\\")

    def test_nonexistent_create_target_inside_allowed(self, root):
        target = root / "new" / "deeper" / "created.docx"
        assert not target.parent.exists()
        got = check_path(target, "test")
        assert os.path.normcase(got).endswith("created.docx")

    def test_nonexistent_create_target_outside_blocked(self, root, tmp_path):
        with pytest.raises(SandboxViolation):
            check_path(tmp_path / "outside" / "new" / "created.docx", "test")

    def test_error_names_roots_and_env_var(self, root, tmp_path):
        with pytest.raises(SandboxViolation) as exc:
            check_path(tmp_path / "outside" / "a.docx", "open document")
        msg = str(exc.value)
        assert ENV in msg
        assert str(root) in msg
        assert "outside" in msg  # the offending path appears
        assert "open document" in msg  # the purpose appears

    def test_env_change_is_seen_without_reload(self, tmp_path, monkeypatch):
        a = tmp_path / "a"
        b = tmp_path / "b"
        a.mkdir()
        b.mkdir()
        monkeypatch.setenv(ENV, str(a))
        with pytest.raises(SandboxViolation):
            check_path(b / "x.docx", "test")
        monkeypatch.setenv(ENV, str(b))
        assert check_path(b / "x.docx", "test")


# ---------------------------------------------------------- junction escapes


class TestJunction:
    def test_junction_escape_blocked(self, root, tmp_path):
        """A junction inside the root pointing outside must not smuggle the
        target back in: realpath resolves it before containment."""
        target = tmp_path / "outside" / "secret"
        target.mkdir()
        link = root / "jump"
        proc = subprocess.run(
            ["cmd", "/c", "mklink", "/J", str(link), str(target)],
            capture_output=True,
            text=True,
        )
        if proc.returncode != 0 or not link.exists():
            pytest.skip(f"cannot create junction here: {proc.stderr.strip()}")
        try:
            with pytest.raises(SandboxViolation):
                check_path(link / "a.docx", "test")
        finally:
            link.rmdir()  # removes the junction, not the target

    @pytest.mark.skipif(sys.platform != "win32", reason="symlink test uses os.symlink")
    def test_symlink_escape_blocked(self, root, tmp_path):
        target = tmp_path / "outside" / "linked"
        target.mkdir()
        link = root / "sym"
        try:
            os.symlink(str(target), str(link), target_is_directory=True)
        except OSError:
            pytest.skip("symlink creation not permitted (no developer mode)")
        try:
            with pytest.raises(SandboxViolation):
                check_path(link / "a.docx", "test")
        finally:
            link.rmdir()


# ------------------------------------------------------ enforcement plumbing


class TestEnforcement:
    def test_docx_open_outside_blocked(self, root, tmp_path):
        f = _fresh_doc(tmp_path / "outside")
        with pytest.raises(SandboxViolation):
            DocxPackage(f)

    def test_docx_open_inside_allowed_and_saves(self, root):
        f = _fresh_doc(root)
        pkg = DocxPackage(f)
        pkg.save()

    def test_save_dest_outside_blocked(self, root, tmp_path):
        f = _fresh_doc(root)
        pkg = DocxPackage(f)
        with pytest.raises(SandboxViolation):
            pkg.save(tmp_path / "outside" / "escape.docx")

    def test_create_snapshot_dest_dir_outside_blocked(self, root, tmp_path):
        f = _fresh_doc(root)
        with pytest.raises(SandboxViolation):
            bk.create_snapshot(str(f), dest_dir=str(tmp_path / "outside"))

    def test_create_snapshot_inside_allowed(self, root):
        f = _fresh_doc(root)
        out = bk.create_snapshot(str(f))
        assert Path(out["snapshot"]).exists()

    def test_snapshot_source_outside_blocked(self, root, tmp_path):
        f = _fresh_doc(tmp_path / "outside")
        with pytest.raises(SandboxViolation):
            bk.create_snapshot(str(f), dest_dir=str(root))

    def test_export_table_output_outside_blocked(self, root, tmp_path):
        f = root / "tables.docx"
        doc = Document()
        t = doc.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "a"
        t.cell(0, 1).text = "b"
        doc.save(str(f))
        pkg = DocxPackage(f)
        with pytest.raises(SandboxViolation):
            dio.export_table(
                pkg, 0, format="csv",
                output_path=str(tmp_path / "outside" / "t.csv"),
            )

    def test_import_data_file_outside_blocked(self, root, tmp_path):
        stolen = tmp_path / "outside" / "data.csv"
        stolen.write_text("a,b\n1,2\n", encoding="utf-8")
        with pytest.raises(SandboxViolation):
            dio._load_rows(str(stolen))

    def test_add_image_read_outside_blocked(self, root, tmp_path):
        f = _fresh_doc(root)
        pkg = DocxPackage(f)
        with pytest.raises(SandboxViolation):
            media.add_image(
                pkg, str(tmp_path / "outside" / "img.png"), at_end=True
            )

    def test_restore_backup_legacy_source_outside_blocked(self, root, tmp_path):
        f = _fresh_doc(root)
        outside_bak = _fresh_doc(tmp_path / "outside", "doc.bak-x.docx")
        with pytest.raises(SandboxViolation):
            bk.restore_backup(str(f), str(outside_bak))

    def test_list_backups_outside_blocked(self, root, tmp_path):
        with pytest.raises(SandboxViolation):
            bk.list_backups(directory=str(tmp_path / "outside"))

    def test_purge_outside_blocked(self, root, tmp_path):
        with pytest.raises(SandboxViolation):
            bk.purge_backups("legacy", directory=str(tmp_path / "outside"))

    def test_split_document_output_dir_outside_blocked(self, root, tmp_path):
        f = root / "split.docx"
        doc = Document()
        doc.add_heading("One", level=1)
        doc.add_paragraph("body")
        doc.add_heading("Two", level=1)
        doc.add_paragraph("body 2")
        doc.save(str(f))
        with pytest.raises(SandboxViolation):
            dio.split_document(str(f), str(tmp_path / "outside" / "parts"))

    def test_normal_edit_cycle_unaffected_inside_root(self, root):
        """A full open-modify-save cycle inside the root works with the
        sandbox on, including the backup slot rotation beside the doc."""
        f = _fresh_doc(root)
        pkg = DocxPackage(f)
        pkg.save()  # first save: rotates slots under root/.ks4w-backups
        pkg2 = DocxPackage(f)
        assert pkg2.body() is not None
