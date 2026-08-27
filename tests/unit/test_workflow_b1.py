"""Wave B bundle 1 gate: front-matter assembly, chapter-aware running
headers (STYLEREF), document diagnostics.

Docs are built with the server's own file-based tools; the new ops are
exercised directly (they are not registered in server.py yet — the
integration snippet does that). Deliberately broken docs for the
diagnostics tests are hand-damaged with lxml surgery.
"""

import copy

import pytest
from docx import Document
from lxml import etree

import word_mcp.server as srv
from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import chapterheaders as ch
from word_mcp.ops import diagnostics as dg
from word_mcp.ops import frontmatter as fm
from word_mcp.ops import toc as toc_ops
from word_mcp.ops.furniture import _sect_prs

_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def make_chapter_doc(tmp_path, name="doc.docx", chapters=3):
    path = str(tmp_path / name)
    srv.create_document(path)
    for i in range(1, chapters + 1):
        srv.add_heading(path, f"Chapter {i}", level=1, at_end=True, backup=False)
        srv.insert_paragraphs(
            path,
            [{"text": f"Body text for chapter {i}, long enough to matter."}],
            at_end=True,
            backup=False,
            live="off",
        )
    return path


def edit(path, fn):
    """Apply an op the way server._edit does: load, mutate, atomic save."""
    pkg = DocxPackage(path)
    result = fn(pkg)
    pkg.save(do_backup=False)
    return result


FRONT_SPEC = {
    "sections": [
        {"kind": "title_page", "lines": ["A Study of Things", "by", "An Author"]},
        {"kind": "blank_or_copyright", "lines": ["Copyright 2026. All rights reserved."]},
        {
            "kind": "abstract",
            "title": "Abstract",
            "text": "This study examines things.\n\nIt finds that things vary.",
        },
        {"kind": "toc"},
        {"kind": "list_of_figures"},
    ],
    "page_numbering": {
        "front": "lowerRoman",
        "body": "decimal",
        "body_restart_at": 1,
    },
}


# ---------------------------------------------------------------- front matter


def test_front_matter_full_assembly(tmp_path):
    path = make_chapter_doc(tmp_path)
    outline_before = [h["text"] for h in srv.get_outline(path, live="off")]
    assert outline_before == ["Chapter 1", "Chapter 2", "Chapter 3"]

    result = edit(
        path, lambda pkg: fm.assemble_front_matter(pkg, copy.deepcopy(FRONT_SPEC))
    )
    assert result["sections_total"] == 2
    assert result["front_matter"] == {
        "section_index": 0,
        "number_format": "lowerRoman",
        "start_at": 1,
    }
    assert result["body"]["number_format"] == "decimal"
    assert [e["kind"] for e in result["inserted"]] == [
        "title_page", "blank_or_copyright", "abstract", "toc", "list_of_figures",
    ]
    assert not any("skipped" in e for e in result["inserted"])

    pkg = DocxPackage(path)
    sects = _sect_prs(pkg)
    assert len(sects) == 2
    pg0 = sects[0].find(qn("w:pgNumType"))
    assert pg0.get(qn("w:fmt")) == "lowerRoman"
    assert pg0.get(qn("w:start")) == "1"
    pg1 = sects[1].find(qn("w:pgNumType"))
    assert pg1.get(qn("w:fmt")) == "decimal"
    assert pg1.get(qn("w:start")) == "1"
    # The front-matter section closes with a nextPage break.
    assert sects[0].find(qn("w:type")).get(qn("w:val")) == "nextPage"

    # TOC and List of Figures both present.
    info = toc_ops.read_toc(pkg)
    assert info["present"]
    kinds = {t["kind"] for t in info["tocs"]}
    assert kinds == {"main", "caption_list"}
    settings = pkg.raw_part("word/settings.xml").decode("utf-8")
    assert "updateFields" in settings

    # Front matter order: title page first, centered; body content intact.
    paras = srv.get_text(path, live="off")
    assert paras[0]["index"] == 0
    assert paras[0]["text"] == "A Study of Things"
    first_p = pkg.body().find(qn("w:p"))
    jc = first_p.find(f"{qn('w:pPr')}/{qn('w:jc')}")
    assert jc is not None and jc.get(qn("w:val")) == "center"
    outline_after = [h["text"] for h in srv.get_outline(path, live="off")]
    assert outline_after == ["Abstract", "Chapter 1", "Chapter 2", "Chapter 3"]
    body_texts = [p["text"] for p in paras]
    assert "Body text for chapter 3, long enough to matter." in body_texts

    Document(path)  # Word-library round-trip


def test_front_matter_refuses_existing_roman_then_force(tmp_path):
    path = make_chapter_doc(tmp_path)
    srv.set_page_number_format(
        path, section=0, number_format="lowerRoman", backup=False
    )
    before = open(path, "rb").read()
    with pytest.raises(WordMcpError, match="front matter appears to exist"):
        edit(path, lambda pkg: fm.assemble_front_matter(pkg, copy.deepcopy(FRONT_SPEC)))
    # Atomicity: refusal changed nothing on disk.
    assert open(path, "rb").read() == before

    spec = copy.deepcopy(FRONT_SPEC)
    spec["force"] = True
    result = edit(path, lambda pkg: fm.assemble_front_matter(pkg, spec))
    assert result["sections_total"] == 2
    Document(path)


def test_front_matter_refuses_leading_toc_then_force_skips_toc(tmp_path):
    path = make_chapter_doc(tmp_path)
    srv.insert_toc(path, at_start=True, update_on_open=False, backup=False)
    with pytest.raises(WordMcpError, match="front matter appears to exist"):
        edit(path, lambda pkg: fm.assemble_front_matter(pkg, copy.deepcopy(FRONT_SPEC)))

    spec = copy.deepcopy(FRONT_SPEC)
    spec["force"] = True
    result = edit(path, lambda pkg: fm.assemble_front_matter(pkg, spec))
    # The duplicate TOC request is skipped and reported, not doubled.
    toc_entry = next(e for e in result["inserted"] if e["kind"] == "toc")
    assert "skipped" in toc_entry
    pkg = DocxPackage(path)
    mains = [t for t in toc_ops.read_toc(pkg)["tocs"] if t["kind"] == "main"]
    assert len(mains) == 1
    Document(path)


def test_front_matter_bad_spec_refused(tmp_path):
    path = make_chapter_doc(tmp_path, chapters=1)
    with pytest.raises(WordMcpError, match="unknown kind"):
        edit(
            path,
            lambda pkg: fm.assemble_front_matter(
                pkg, {"sections": [{"kind": "dedication_page"}]}
            ),
        )
    with pytest.raises(WordMcpError, match="page_numbering"):
        edit(
            path,
            lambda pkg: fm.assemble_front_matter(
                pkg,
                {
                    "sections": [{"kind": "toc"}],
                    "page_numbering": {"front": "romanish"},
                },
            ),
        )


# ------------------------------------------------------------ chapter headers


def test_chapter_headers_basic(tmp_path):
    path = make_chapter_doc(tmp_path)
    result = edit(
        path, lambda pkg: ch.setup_chapter_headers(pkg, level=1, include_number=True)
    )
    assert [e["section"] for e in result["sections"]] == [0]
    assert result["style_referenced"] == "heading 1"
    assert any("\\n" in code for code in result["field_codes"])
    assert result["parts_created"], "a header part should have been created"

    pkg = DocxPackage(path)
    part = result["parts_written"][0]
    xml = pkg.raw_part(part).decode("utf-8")
    assert "STYLEREF" in xml
    # first_page_blank default: titlePg flipped on the section.
    assert _sect_prs(pkg)[0].find(qn("w:titlePg")) is not None
    Document(path)

    report = ch.validate_chapter_headers(pkg)
    assert report["ok"]
    assert report["chapter_level"] == 1
    assert report["sections"][0]["styleref_fields"]
    assert report["sections_missing_chapter_header"] == []


def test_chapter_headers_refused_without_headings(tmp_path):
    path = str(tmp_path / "plain.docx")
    srv.create_document(path)
    srv.insert_paragraphs(
        path, [{"text": "Just prose, no headings."}], at_end=True,
        backup=False, live="off",
    )
    with pytest.raises(TargetNotFound, match="no body headings"):
        edit(path, lambda pkg: ch.setup_chapter_headers(pkg))


def test_chapter_headers_watermark_survives(tmp_path):
    path = make_chapter_doc(tmp_path)
    srv.add_watermark(path, "DRAFT", backup=False)
    result = edit(path, lambda pkg: ch.setup_chapter_headers(pkg, level=1))
    assert result["watermark_preserved_in"], "watermark part should be reported"

    pkg = DocxPackage(path)
    part = result["parts_written"][0]
    xml = pkg.raw_part(part).decode("utf-8")
    assert "PowerPlusWaterMarkObject" in xml, "watermark destroyed by header write"
    assert "STYLEREF" in xml
    Document(path)


def test_chapter_headers_scope_and_validate_gaps(tmp_path):
    # Two sections, each with a level-1 heading.
    path = make_chapter_doc(tmp_path, chapters=1)
    last = srv.get_text(path, live="off")[-1]["index"]
    srv.add_section_break(path, after_index=last, backup=False)
    srv.add_heading(path, "Chapter 2", level=1, at_end=True, backup=False)
    srv.insert_paragraphs(
        path, [{"text": "Second chapter body."}], at_end=True,
        backup=False, live="off",
    )
    pkg = DocxPackage(path)
    assert len(_sect_prs(pkg)) == 2

    # Explicit scope: only section 1 gets its own chapter header; section 0
    # has no header at all -> reported as the gap.
    edit(path, lambda pkg: ch.setup_chapter_headers(pkg, level=1, scope=[1]))
    report = ch.validate_chapter_headers(DocxPackage(path))
    assert report["chapter_level"] == 1
    assert report["sections_missing_chapter_header"] == [0]
    assert not report["ok"]
    assert report["sections"][1]["styleref_fields"]

    # auto scope fills the gap.
    result = edit(path, lambda pkg: ch.setup_chapter_headers(pkg, level=1))
    assert [e["section"] for e in result["sections"]] == [0, 1]
    report = ch.validate_chapter_headers(DocxPackage(path))
    assert report["ok"]
    assert report["sections_missing_chapter_header"] == []
    Document(path)


def test_chapter_headers_bad_scope(tmp_path):
    path = make_chapter_doc(tmp_path, chapters=1)
    with pytest.raises(TargetNotFound, match="out of range"):
        edit(path, lambda pkg: ch.setup_chapter_headers(pkg, scope=[7]))


# ---------------------------------------------------------------- diagnostics


def test_diagnostics_healthy_doc(tmp_path):
    path = make_chapter_doc(tmp_path)
    report = dg.diagnose_document(DocxPackage(path))
    assert report["ok"]
    assert not [p for p in report["problems"] if p["severity"] == "error"]
    profile = report["info"]["size_profile"]
    assert profile["part_count"] > 0
    assert profile["total_bytes"] > 0
    assert len(profile["largest_parts"]) > 0
    assert "notes" in report["info"]
    assert "fields" in report["info"]


def test_diagnostics_dangling_relationship(tmp_path):
    path = make_chapter_doc(tmp_path, name="dangling.docx", chapters=1)
    pkg = DocxPackage(path)
    rels_root = pkg.root("word/_rels/document.xml.rels")
    rel = etree.SubElement(rels_root, f"{{{_REL_NS}}}Relationship")
    rel.set("Id", "rId9901")
    rel.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
    )
    rel.set("Target", "media/does-not-exist.png")
    pkg.mark_dirty("word/_rels/document.xml.rels")
    pkg.save(do_backup=False)

    report = dg.diagnose_document(DocxPackage(path))
    assert not report["ok"]
    hits = [
        p
        for p in report["problems"]
        if p["category"] == "relationships" and "does-not-exist.png" in p["detail"]
    ]
    assert hits and hits[0]["severity"] == "error"


def test_diagnostics_undefined_pstyle(tmp_path):
    path = make_chapter_doc(tmp_path, name="ghoststyle.docx", chapters=1)
    pkg = DocxPackage(path)
    p = pkg.body().find(qn("w:p"))
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.Element(qn("w:pPr"))
        p.insert(0, ppr)
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        pstyle = etree.Element(qn("w:pStyle"))
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), "GhostStyle99")
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    report = dg.diagnose_document(DocxPackage(path))
    hits = [
        p
        for p in report["problems"]
        if p["category"] == "styles" and "GhostStyle99" in p["detail"]
    ]
    assert hits and hits[0]["severity"] == "warning"
    # Latent bug class, not corruption: the doc still opens, so ok holds.
    assert report["ok"]


def test_diagnostics_unbalanced_fldchar(tmp_path):
    path = make_chapter_doc(tmp_path, name="unbalanced.docx", chapters=1)
    pkg = DocxPackage(path)
    p = pkg.body().find(qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    etree.SubElement(r, qn("w:fldChar")).set(qn("w:fldCharType"), "begin")
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    report = dg.diagnose_document(DocxPackage(path))
    assert not report["ok"]
    hits = [
        p
        for p in report["problems"]
        if p["category"] == "fields" and p.get("location") == "word/document.xml"
    ]
    assert hits and hits[0]["severity"] == "error"
    assert "1 begin vs 0 end" in hits[0]["detail"]


def test_diagnostics_orphan_part(tmp_path):
    path = make_chapter_doc(tmp_path, name="orphan.docx", chapters=1)
    pkg = DocxPackage(path)
    pkg.set_raw_part(
        "word/orphanpart.xml", b'<?xml version="1.0"?><orphan/>'
    )
    pkg.save(do_backup=False)

    report = dg.diagnose_document(DocxPackage(path))
    hits = [
        p
        for p in report["problems"]
        if p["category"] == "orphan_parts"
        and p.get("location") == "word/orphanpart.xml"
    ]
    assert hits


def test_diagnostics_never_raises_on_missing_styles_part(tmp_path):
    """Degrade-to-reporting: removing whole parts must not crash diagnose."""
    path = make_chapter_doc(tmp_path, name="weird.docx", chapters=1)
    pkg = DocxPackage(path)
    # Simulate a stripped-down producer: no styles.xml, no numbering.
    pkg._raw.pop("word/styles.xml", None)
    pkg._order.remove("word/styles.xml")
    pkg._trees.pop("word/styles.xml", None)
    report = dg.diagnose_document(pkg)  # must not raise
    assert isinstance(report["ok"], bool)
    assert isinstance(report["problems"], list)


# ------------------------------------------------- workflow interaction check


def test_front_matter_then_chapter_headers(tmp_path):
    """The bundle's own pieces compose: front matter first, then chapter
    headers on the body section only."""
    path = make_chapter_doc(tmp_path)
    edit(path, lambda pkg: fm.assemble_front_matter(pkg, copy.deepcopy(FRONT_SPEC)))
    # Auto scope: the Abstract heading lives in section 0, chapters in
    # section 1, so both sections contain level-1 headings.
    result = edit(path, lambda pkg: ch.setup_chapter_headers(pkg, level=1, scope=[1]))
    assert [e["section"] for e in result["sections"]] == [1]
    report = dg.diagnose_document(DocxPackage(path))
    assert report["ok"]
    Document(path)
