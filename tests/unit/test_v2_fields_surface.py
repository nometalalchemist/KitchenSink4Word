"""STAGED v2-surface rewrite (Wave D): bookmarks, hyperlinks, captions
(2.7 slice) plus their delete_element branches.

insert_bookmark / insert_hyperlink / insert_caption renames and
delete_element types bookmark / hyperlink / caption per
integration/v2_briefs/wave_D.md. Skips until the Phase 2 integrator
registers the v2 tools. Ops-level coverage stays in
tests/unit/test_fields.py and tests/unit/test_v2_delete_ops.py. Original
v1 test file untouched.
"""

from pathlib import Path

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import AmbiguousTarget, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import fields

pytestmark = pytest.mark.skipif(
    not hasattr(srv, "insert_bookmark"),
    reason="v2 surface not yet registered (staged for the Phase 2 "
    "integrator)",
)


def fresh(tmp_path: Path, name="doc.docx") -> str:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Alpha paragraph one.")
    doc.add_paragraph("Beta paragraph two.")
    doc.save(str(f))
    return str(f)


def ok(result) -> bool:
    return not (isinstance(result, dict) and result.get("ok") is False)


def body_text(path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


# ------------------------------------------------------------------ bookmark


def test_bookmark_insert_then_delete_element(tmp_path):
    f = fresh(tmp_path)
    assert ok(srv.insert_bookmark(f, name="mark_a", anchor_text="Beta"))
    assert [
        b["name"] for b in fields.list_bookmarks(DocxPackage(f))
    ] == ["mark_a"]

    r = srv.delete_element(f, type="bookmark", id="mark_a")
    assert ok(r)
    assert fields.list_bookmarks(DocxPackage(f)) == []
    assert "Beta paragraph two." in body_text(f)


def test_delete_element_internal_bookmark_refuses(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(WordMcpError):
        srv.delete_element(f, type="bookmark", id="_Toc999")


# ----------------------------------------------------------------- hyperlink


def test_hyperlink_insert_then_delete_by_url(tmp_path):
    f = fresh(tmp_path)
    assert ok(
        srv.insert_hyperlink(
            f, anchor_text="Alpha", url="https://one.example"
        )
    )
    r = srv.delete_element(f, type="hyperlink", id="https://one.example")
    assert ok(r)
    pkg = DocxPackage(f)
    assert not list(pkg.root().iter(qn("w:hyperlink")))
    assert "Alpha paragraph one." in body_text(f)


def test_hyperlink_ambiguous_delete_refuses_with_matches(tmp_path):
    f = fresh(tmp_path)
    srv.insert_hyperlink(
        f, anchor_text="paragraph", occurrence=1, url="https://one.example"
    )
    srv.insert_hyperlink(
        f, anchor_text="paragraph", occurrence=2, url="https://two.example"
    )
    with pytest.raises(AmbiguousTarget) as exc_info:
        srv.delete_element(
            f, type="hyperlink", location={"search": {"text": "paragraph"}}
        )
    assert len(exc_info.value.matches) == 2

    r2 = srv.delete_element(
        f,
        type="hyperlink",
        location={"search": {"text": "paragraph", "occurrence": 2}},
    )
    assert ok(r2)
    pkg = DocxPackage(f)
    assert len(list(pkg.root().iter(qn("w:hyperlink")))) == 1


# ------------------------------------------------------------------- caption


def make_table_doc(tmp_path) -> str:
    f = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Intro.")
    doc.add_table(rows=2, cols=2)
    doc.save(str(f))
    return str(f)


def test_caption_insert_then_delete_element(tmp_path):
    f = make_table_doc(tmp_path)
    assert ok(
        srv.insert_caption(f, text="Results summary", table_index=0)
    )
    assert "Results summary" in body_text(f)

    # The caption paragraph sits where the table was (above it).
    pkg = DocxPackage(f)
    cap_index = next(
        i
        for i, p in enumerate(pkg.body().findall(qn("w:p")))
        if "Results summary" in "".join(t.text or "" for t in p.iter(qn("w:t")))
    )
    r = srv.delete_element(
        f, type="caption", location={"paragraph": cap_index,
                                     "position": "replace"}
    )
    assert ok(r)
    assert "Results summary" not in body_text(f)


def test_delete_element_caption_refuses_prose(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(WordMcpError):
        srv.delete_element(
            f, type="caption",
            location={"paragraph": 0, "position": "replace"},
        )
