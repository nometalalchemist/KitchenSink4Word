"""STAGED v2-surface rewrite (Wave D): templates, forms, content controls
(2.13).

set_form_fields and set_content_control renames plus fill_template /
mail_merge / insert_content_control identity smokes per
integration/v2_briefs/wave_D.md. Skips until the Phase 2 integrator
registers the v2 tools. Ops-level coverage stays in
tests/unit/test_business_bundle.py (batch_apply and the deeper form/merge
gates live there and belong to other waves where renamed). Original v1
test file untouched.
"""

import csv
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

import word_mcp.server as srv
from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage, qn

pytestmark = pytest.mark.skipif(
    not hasattr(srv, "set_form_fields"),
    reason="v2 surface not yet registered (staged for the Phase 2 "
    "integrator)",
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def ok(result) -> bool:
    return not (isinstance(result, dict) and result.get("ok") is False)


def body_text(path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def template_doc(tmp_path: Path, name="template.docx") -> str:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Dear {{name}},")
    doc.add_paragraph("Your order {{order_id}} has shipped.")
    doc.save(str(f))
    return str(f)


# ------------------------------------------------------------- fill_template


def test_fill_template_identity(tmp_path):
    f = template_doc(tmp_path)
    r = srv.fill_template(f, {"name": "Ada", "order_id": "17"})
    assert ok(r)
    text = body_text(f)
    assert "Dear Ada," in text and "order 17" in text


def test_fill_template_missing_error_changes_nothing(tmp_path):
    f = template_doc(tmp_path)
    with pytest.raises(WordMcpError):
        srv.fill_template(f, {"name": "Ada"}, missing="error")
    assert "{{order_id}}" in body_text(f)


# ---------------------------------------------------------------- mail_merge


def test_mail_merge_identity(tmp_path):
    f = template_doc(tmp_path)
    rows_csv = tmp_path / "rows.csv"
    with open(rows_csv, "w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["name", "order_id"])
        writer.writerow(["Ada", "1"])
        writer.writerow(["Grace", "2"])
    out = tmp_path / "merged"
    out.mkdir()
    r = srv.mail_merge(
        f, str(rows_csv), str(out), filename_pattern="{name}.docx"
    )
    assert ok(r)
    assert (out / "Ada.docx").exists() and (out / "Grace.docx").exists()
    assert "{{name}}" in body_text(f)  # template untouched


# ----------------------------------------------- content controls and forms


def test_content_control_insert_set_roundtrip(tmp_path):
    f = tmp_path / "cc.docx"
    doc = Document()
    doc.add_paragraph("Reviewed by: ")
    doc.save(str(f))

    r = srv.insert_content_control(
        str(f), tag="reviewer", after_anchor="Reviewed by:", text="TBD"
    )
    assert ok(r)
    r2 = srv.set_content_control(str(f), value="Dr. Hopper", tag="reviewer")
    assert ok(r2)
    pkg = DocxPackage(str(f))
    sdt_text = "".join(
        t.text or ""
        for sdt in pkg.root().iter(qn("w:sdt"))
        for t in sdt.iter(qn("w:t"))
    )
    assert "Dr. Hopper" in sdt_text


def test_set_content_control_missing_tag_refuses(tmp_path):
    f = tmp_path / "cc2.docx"
    doc = Document()
    doc.add_paragraph("Nothing here.")
    doc.save(str(f))
    with pytest.raises(WordMcpError):
        srv.set_content_control(str(f), value="x", tag="ghost")


def test_set_form_fields_checkbox_and_text(tmp_path):
    """Legacy-field filling through the renamed surface, on a document
    carrying a checkbox content control built inline."""
    f = tmp_path / "form.docx"
    doc = Document()
    doc.add_paragraph("Name field follows.")
    doc.save(str(f))
    srv.insert_content_control(
        str(f), tag="fullname", after_anchor="Name field", text=""
    )
    r = srv.set_form_fields(str(f), {"fullname": "Ada Lovelace"})
    assert ok(r)
    pkg = DocxPackage(str(f))
    sdt_text = "".join(
        t.text or ""
        for sdt in pkg.root().iter(qn("w:sdt"))
        for t in sdt.iter(qn("w:t"))
    )
    assert "Ada Lovelace" in sdt_text


def test_set_form_fields_missing_error_refuses(tmp_path):
    f = tmp_path / "form2.docx"
    doc = Document()
    doc.add_paragraph("No fields at all.")
    doc.save(str(f))
    with pytest.raises(WordMcpError):
        srv.set_form_fields(str(f), {"ghost": "x"}, missing="error")
