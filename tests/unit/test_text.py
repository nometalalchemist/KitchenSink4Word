"""Phase 2 gate: text edits on real docs — fragmented-run replace, formatting,
paragraph CRUD, revision-XML preservation."""

import re
import shutil
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from word_mcp.core.errors import (
    AmbiguousTarget,
    TargetNotFound,
    WordMcpError,
)
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import read, text as t

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


@pytest.fixture
def ejir(tmp_path):
    dst = tmp_path / "ejir_rw.docx"
    shutil.copy(CORPUS / "ejir_rw.docx", dst)
    return dst


@pytest.fixture
def ch13(tmp_path):
    dst = tmp_path / "ch1-3.docx"
    shutil.copy(CORPUS / "ch1-3.docx", dst)
    return dst


def reopen_text(path):
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


# ------------------------------------------------------------ search & replace


def test_simple_replace_roundtrip(ch4):
    pkg = DocxPackage(ch4)
    before = reopen_text(ch4)
    assert "Delta Model" in before
    result = t.search_and_replace(
        pkg, [{"find": "Delta Model", "replace": "DELTA-MODEL-X"}]
    )
    pkg.save(do_backup=False)
    after = reopen_text(ch4)
    assert result["total"] == before.count("Delta Model")
    assert "DELTA-MODEL-X" in after
    assert "Delta Model" not in after


def test_replace_across_fragmented_runs(tmp_path):
    """Build a doc whose target phrase is split across five runs with
    different formatting, exactly like Word fragments text."""
    src = tmp_path / "frag.docx"
    doc = Document()
    p = doc.add_paragraph()
    for piece, bold in (("The Del", False), ("ta ", True), ("Mo", False),
                        ("d", True), ("el governs.", False)):
        run = p.add_run(piece)
        run.bold = bold
    doc.save(str(src))

    pkg = DocxPackage(src)
    result = t.search_and_replace(
        pkg, [{"find": "Delta Model", "replace": "Alpha Framework"}]
    )
    pkg.save(do_backup=False)
    assert result["total"] == 1
    reopened = Document(str(src))
    assert reopened.paragraphs[0].text == "The Alpha Framework governs."
    # Formatting outside the match must survive: "el governs." was not bold.
    runs = reopened.paragraphs[0].runs
    assert runs[0].text.startswith("The ")


def test_batch_replace_counts(ch13):
    pkg = DocxPackage(ch13)
    body = reopen_text(ch13)
    result = t.search_and_replace(
        pkg,
        [
            {"find": "alliance", "replace": "ALLIANCE"},
            {"find": "Korea", "replace": "KOREA"},
        ],
    )
    pkg.save(do_backup=False)
    after = reopen_text(ch13)
    assert result["replaced"]["alliance"] == body.count("alliance")
    assert result["replaced"]["Korea"] == body.count("Korea")
    assert "alliance" not in after
    assert after.count("ALLIANCE") >= result["replaced"]["alliance"]


def test_replacement_containing_find_string_terminates(ch4):
    """'alliance' -> 'alliance-structure' must replace each original occurrence
    exactly once, never re-matching its own output (the infinite-growth bug)."""
    pkg = DocxPackage(ch4)
    before = reopen_text(ch4)
    n = before.count("alliance")
    result = t.search_and_replace(
        pkg, [{"find": "alliance", "replace": "alliance-structure"}]
    )
    pkg.save(do_backup=False)
    after = reopen_text(ch4)
    assert result["replaced"]["alliance"] == n
    assert after.count("alliance-structure") == n
    assert after.count("alliance-structure-structure") == 0


def test_identity_replacement_terminates(ch4):
    pkg = DocxPackage(ch4)
    before = reopen_text(ch4)
    result = t.search_and_replace(pkg, [{"find": "the", "replace": "the"}])
    pkg.save(do_backup=False)
    assert reopen_text(ch4) == before
    assert result["replaced"]["the"] == before.count("the")


def test_regex_replace(ch4):
    pkg = DocxPackage(ch4)
    result = t.search_and_replace(
        pkg,
        [{"find": r"\(19(\d\d)\)", "replace": r"(19\1!)", "regex": True}],
    )
    pkg.save(do_backup=False)
    after = reopen_text(ch4)
    if result["total"]:
        assert re.search(r"\(19\d\d!\)", after)
    assert not re.search(r"\(19\d\d\)", after)


def test_replace_preserves_revision_xml(ejir):
    """Editing a doc full of tracked changes must leave every revision
    element untouched (count, authors, deleted text)."""
    pkg = DocxPackage(ejir)
    revs_before = read.get_tracked_changes(pkg)
    # Replace a word that exists in plain text.
    t.search_and_replace(pkg, [{"find": "Mitzen", "replace": "MITZEN"}])
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ejir)
    revs_after = read.get_tracked_changes(pkg2)
    assert len(revs_after) == len(revs_before)
    assert [r["type"] for r in revs_after] == [r["type"] for r in revs_before]
    assert [r["author"] for r in revs_after] == [r["author"] for r in revs_before]
    # Deleted-text payloads specifically must be untouched.
    dels_before = [r["text"] for r in revs_before if r["type"] == "deletion"]
    dels_after = [r["text"] for r in revs_after if r["type"] == "deletion"]
    assert dels_before == dels_after


def test_replace_skips_deleted_text(ejir):
    """A find string that only exists inside w:del content must not match."""
    pkg = DocxPackage(ejir)
    dels = [
        r["text"]
        for r in read.get_tracked_changes(pkg)
        if r["type"] == "deletion" and len(r["text"].strip()) > 12
    ]
    if not dels:
        pytest.skip("no long deletions")
    # Find a deleted string absent from visible text.
    visible = "\n".join(e["text"] for e in read.get_paragraphs(pkg))
    target = next((d for d in dels if d not in visible), None)
    if target is None:
        pytest.skip("all deleted strings also appear in visible text")
    result = t.search_and_replace(
        pkg, [{"find": target, "replace": "SHOULD-NOT-APPEAR"}]
    )
    assert result["total"] == 0


# ------------------------------------------------------------ paragraph CRUD


def test_insert_and_delete_paragraphs(ch4):
    pkg = DocxPackage(ch4)
    n_before = read.get_document_info(pkg)["paragraphs"]
    t.insert_paragraphs(
        pkg,
        [{"text": "MARKER ONE"}, {"text": "MARKER TWO", "style": "Heading1"}],
        after_index=5,
    )
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg2)
    assert paras[6]["text"] == "MARKER ONE"
    assert paras[7]["text"] == "MARKER TWO"
    assert paras[7].get("style") == "Heading1"
    assert len(paras) == n_before + 2

    t.delete_paragraphs(pkg2, 6, 7)
    pkg2.save(do_backup=False)
    pkg3 = DocxPackage(ch4)
    assert read.get_document_info(pkg3)["paragraphs"] == n_before
    assert "MARKER" not in reopen_text(ch4)


def test_insert_after_anchor(ch4):
    pkg = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg)
    long_para = max(paras, key=lambda p: len(p["text"]))
    anchor = long_para["text"][:60]
    t.insert_paragraphs(pkg, [{"text": "ANCHOR-INSERTED"}], after_anchor=anchor)
    pkg.save(do_backup=False)
    after = read.get_paragraphs(DocxPackage(ch4))
    inserted = [p for p in after if p["text"] == "ANCHOR-INSERTED"]
    assert len(inserted) == 1
    assert inserted[0]["index"] == long_para["index"] + 1


def test_ambiguous_anchor_refused(ch13):
    pkg = DocxPackage(ch13)
    with pytest.raises((AmbiguousTarget, TargetNotFound)):
        t.insert_paragraphs(pkg, [{"text": "X"}], after_anchor="the")


def test_replace_paragraph_text_keeps_style(ch13):
    pkg = DocxPackage(ch13)
    outline = read.get_outline(pkg)
    target = outline[3]
    t.replace_paragraph_text(pkg, target["paragraph_index"], "REWRITTEN HEADING")
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch13)
    new_outline = read.get_outline(pkg2)
    match = [h for h in new_outline if h["text"] == "REWRITTEN HEADING"]
    assert len(match) == 1
    assert match[0]["level"] == target["level"]


def test_delete_refuses_section_paragraph(ch13):
    pkg = DocxPackage(ch13)
    sect_paras = [
        idx
        for kind, idx, el in read.body_items(pkg)
        if kind == "paragraph"
        and el.find(f"{qn('w:pPr')}/{qn('w:sectPr')}") is not None
    ]
    if not sect_paras:
        pytest.skip("no mid-document section paragraphs")
    with pytest.raises(WordMcpError):
        t.delete_paragraphs(pkg, sect_paras[0])


# ---------------------------------------------------------------- formatting


def test_format_substring_bold(ch4):
    pkg = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg)
    long_para = max(paras, key=lambda p: len(p["text"]))
    word = long_para["text"].split()[4]
    t.format_text(
        pkg,
        paragraph_index=long_para["index"],
        find=word,
        formatting={"bold": True, "color": "FF0000"},
    )
    pkg.save(do_backup=False)
    doc = Document(str(ch4))
    para = doc.paragraphs[long_para["index"]]
    bold_runs = [r for r in para.runs if r.bold]
    assert bold_runs
    assert word in "".join(r.text for r in bold_runs)
    # Text unchanged overall.
    assert para.text == long_para["text"]


def test_paragraph_format_batch(ch4):
    pkg = DocxPackage(ch4)
    t.set_paragraph_format(
        pkg,
        [1, 2, 3],
        {"alignment": "center", "space_after_pt": 12},
    )
    pkg.save(do_backup=False)
    doc = Document(str(ch4))
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for i in (1, 2, 3):
        assert doc.paragraphs[i].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_apply_style_by_name_or_id(ch13):
    pkg = DocxPackage(ch13)
    target = next(
        p["index"] for p in read.get_paragraphs(pkg) if len(p["text"]) > 30
    )
    result = t.apply_style(pkg, [target], "Heading2")
    assert result["style_id"]
    pkg.save(do_backup=False)
    outline = read.get_outline(DocxPackage(ch13))
    assert any(
        h["paragraph_index"] == target and h["level"] == 2 for h in outline
    )


def test_apply_unknown_style_refused(ch4):
    pkg = DocxPackage(ch4)
    with pytest.raises(TargetNotFound):
        t.apply_style(pkg, [0], "NoSuchStyleAnywhere")


# ------------------------------------------------------------- edit integrity


@pytest.mark.parametrize(
    "doc_name", ["ch4.docx", "ch1-3.docx", "ejir_rw.docx", "niu.docx"]
)
def test_edited_doc_still_valid_package(doc_name, tmp_path):
    """After a mixed edit batch, the doc must remain a structurally valid
    package with all untouched parts byte-identical."""
    src = tmp_path / doc_name
    shutil.copy(CORPUS / doc_name, src)
    with zipfile.ZipFile(src) as zf:
        before = {n: zf.read(n) for n in zf.namelist()}

    pkg = DocxPackage(src)
    t.search_and_replace(pkg, [{"find": "the", "replace": "THE"}])
    t.insert_paragraphs(pkg, [{"text": "END MARKER"}], at_end=True)
    pkg.save(do_backup=False)

    with zipfile.ZipFile(src) as zf:
        after = {n: zf.read(n) for n in zf.namelist()}
    assert set(before) == set(after)
    changed = {n for n in before if before[n] != after[n]}
    assert "word/document.xml" in changed
    unexpected = changed - {"word/document.xml"}
    assert not unexpected, f"parts changed that should not have: {unexpected}"
    Document(str(src))  # opens cleanly
