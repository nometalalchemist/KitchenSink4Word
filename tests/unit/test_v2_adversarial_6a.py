"""v2 Phase 6a adversarial round: regression tests for every confirmed
finding. Attack log: integration/adversarial_v2_log.md; round report in
Agent Results. Fixtures are synthetic (python-docx); tests drive the ops
layer or the in-process fastmcp client, matching the attack that found
each defect.
"""

from __future__ import annotations

import hashlib
from pathlib import Path

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import batch as bt
from word_mcp.ops import read as rd


def _build(path: Path, n: int = 6) -> Path:
    d = Document()
    d.add_heading("Alpha", 1)
    for i in range(n):
        d.add_paragraph(f"Paragraph number {i} with unique token{i}.")
    d.save(str(path))
    return path


def _texts(pkg: DocxPackage) -> list[str]:
    return [rd.paragraph_text(el)
            for k, _i, el in rd.body_items(pkg) if k == "paragraph"]


def _anchor_for(path: Path, needle: str) -> str:
    out = srv.get_document_view(str(path), stamp_anchors=True)
    for line in out["view"].split("\n"):
        if needle in line and line.startswith("["):
            return line[1:line.index("]")]
    raise AssertionError(f"no anchor for {needle!r}")


# ---------------------------------------------------------------- finding 1


def test_duplicate_delete_anchors_delete_once(tmp_path):
    """F1: {'op':'delete','anchors':[a,a]} deleted the anchored paragraph
    AND its neighbor (duplicate index -> two single-item runs -> the second
    bottom-up pass hit the shifted paragraph). Dedupe at validation."""
    doc = _build(tmp_path / "f1.docx")
    a = _anchor_for(doc, "token2")
    out = srv.apply_edits(str(doc), [
        {"op": "delete", "anchors": [a, a]},
    ])
    assert out["changed"]["0"]["deleted"] == 1
    texts = _texts(DocxPackage(doc))
    assert not any("token2" in t for t in texts)
    assert any("token3" in t for t in texts), "neighbor was deleted"


def test_duplicate_delete_mixed_spellings_delete_once(tmp_path):
    """Same defect via two spellings of one anchor (short + long suffix)."""
    doc = _build(tmp_path / "f1b.docx")
    a = _anchor_for(doc, "token2")
    pkg = DocxPackage(doc)
    from word_mcp.ops import view as vw

    info = vw.resolve_anchor(pkg, a)
    full = [rec["digest"] for rec in info["map"]["paragraphs"]
            if rec["el"] is info["el"]][0]
    out = srv.apply_edits(str(doc), [
        {"op": "delete", "anchors": [a, full]},
    ])
    assert out["changed"]["0"]["deleted"] == 1
    texts = _texts(DocxPackage(doc))
    assert any("token3" in t for t in texts)


def test_validate_plan_dedupes_delete_indices(tmp_path):
    """The live batch consumes plan['indices'] directly, so the dedupe
    must land in validation, not only in the file-mode apply."""
    doc = _build(tmp_path / "f1c.docx")
    a = _anchor_for(doc, "token1")
    pkg = DocxPackage(doc)
    plans = bt.validate_edits(pkg, [{"op": "delete", "anchors": [a, a, a]}])
    assert plans[0]["indices"] == [2]  # heading + token0 precede it


# --------------------------------------------------- steering message fixes


def test_delete_table_anchor_names_delete_element(tmp_path):
    doc = tmp_path / "f2.docx"
    d = Document()
    d.add_paragraph("lead")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "x"
    d.save(str(doc))
    out = srv.get_document_view(str(doc), stamp_anchors=True)
    import re

    tanchor = re.search(r"\[t:([0-9a-f]+)\]", out["view"]).group(1)
    with pytest.raises(WordMcpError) as exc:
        srv.apply_edits(str(doc), [{"op": "delete",
                                    "anchor": f"t:{tanchor}"}])
    assert "delete_element" in str(exc.value)


# ------------------------------------------- envelope ok-collision findings


def test_validate_top_level_is_passed_not_ok(tmp_path):
    """F3: validate returned a domain ok:false for a failing battery; on
    the wire that is indistinguishable from a failed call with no error
    object. The domain flag is now 'passed'."""
    doc = _build(tmp_path / "f3.docx")
    res = srv.validate(str(doc), checks=["core"])
    assert res["passed"] is True
    assert "ok" not in res


def test_diagnose_document_top_level_is_healthy_not_ok(tmp_path):
    """F4: same collision on diagnose_document; re-keyed to 'healthy' at
    the server (ops v1 shape untouched)."""
    doc = _build(tmp_path / "f4.docx")
    res = srv.diagnose_document(str(doc))
    assert isinstance(res["healthy"], bool)
    assert "ok" not in res
    from word_mcp.ops import diagnostics as dg

    raw = dg.diagnose_document(DocxPackage(doc))
    assert "ok" in raw  # v1 ops shape preserved


# --------------------------------------------- table structure guard fixes


def _table_doc(tmp_path, name="t.docx"):
    doc = tmp_path / name
    d = Document()
    t = d.add_table(rows=3, cols=3)
    t.cell(0, 0).text = "x"
    d.save(str(doc))
    return doc


@pytest.mark.parametrize("count", [0, -3, True])
def test_mts_insert_count_must_be_positive_int(tmp_path, count):
    """F5: count=0/-3 rode through modify_table_structure to a silent
    zero-row no-op reported as success."""
    doc = _table_doc(tmp_path)
    with pytest.raises(WordMcpError, match="count must be"):
        srv.modify_table_structure(
            str(doc), 0, "insert", target="rows", at=0, count=count)


def test_merge_inverted_rectangle_names_inversion(tmp_path):
    """F6: an inverted merge rectangle said 'row range out of bounds'
    (both rows were in bounds; the range was inverted)."""
    doc = _table_doc(tmp_path)
    with pytest.raises(WordMcpError, match="inverted"):
        srv.modify_table_structure(
            str(doc), 0, "merge",
            range={"start_row": 1, "end_row": 0,
                   "start_col": 0, "end_col": 0})
    with pytest.raises(WordMcpError, match="inverted"):
        srv.modify_table_structure(
            str(doc), 0, "merge",
            range={"start_row": 0, "end_row": 0,
                   "start_col": 2, "end_col": 1})


def test_manage_note_refuses_both_note_id_and_position(tmp_path):
    """F7: edit/delete with BOTH note_id and position was accepted and
    note_id silently won (mis-target risk when they disagree)."""
    doc = _build(tmp_path / "f7.docx")
    srv.manage_note(
        str(doc), "insert", note_type="footnote", text="a note",
        location={"search": {"text": "token1"}})
    for action in ("edit", "delete"):
        with pytest.raises(WordMcpError, match="exactly ONE"):
            srv.manage_note(
                str(doc), action, note_type="footnote",
                text="x" if action == "edit" else None,
                note_id="1", position=1)
