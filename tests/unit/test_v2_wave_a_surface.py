"""V2 STAGING (Wave A, NEW file): server-surface coverage for the Wave A
consolidations that had no v1 entry-point tests: list_elements (18-type
dispatch), validate (14-check dispatch, options routing, loud refusals),
and the absorptions (get_text textbox, find_text formatting, word_count
exclusions, manage_backups snapshot, set_document_properties
update_fields_on_open, insert_document location object).

Runs only against the rebuilt v2 server.py; the integrator swaps this
directory in with the rest of tests/v2_staging/. Synthetic documents
only (python-docx builders plus ops-level setup); no Word, no COM, no
real-document set needed. Per-branch payloads are the v1 ops shapes
verbatim, per the Wave A brief (integration/v2_briefs/wave_A.md).
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import fields as fl
from word_mcp.ops import notes as nt

ALL_TYPES = [
    "tables", "images", "charts", "equations", "bookmarks", "sources",
    "sections", "section_blocks", "footnotes", "endnotes", "fields",
    "reference_fields", "form_fields", "content_controls",
    "template_placeholders", "index_entries", "lists", "toc",
]


def _fresh(tmp_path: Path, name: str, texts: list[str]) -> Path:
    f = tmp_path / name
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    doc.save(str(f))
    return f


@pytest.fixture
def plain_doc(tmp_path):
    return _fresh(tmp_path, "plain.docx", ["Alpha one.", "Beta two.",
                                           "Gamma three."])


@pytest.fixture
def rich_doc(tmp_path):
    """One table, one bookmark, one footnote: enough to see non-empty
    items from several branches."""
    f = tmp_path / "rich.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    doc.add_paragraph("Anchor paragraph text.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "H1"
    doc.save(str(f))
    pkg = DocxPackage(str(f))
    fl.add_bookmark(pkg, "wave_a_mark", anchor_text="Anchor paragraph")
    nt.add_note(pkg, "footnote", anchor_text="Intro paragraph",
                note_text="A footnote body.")
    pkg.save(do_backup=False)
    return f


# ------------------------------------------------------------ list_elements


def test_list_elements_dispatches_every_type(plain_doc):
    for t in ALL_TYPES:
        res = srv.list_elements(str(plain_doc), type=t)
        assert res["type"] == t, res
        assert isinstance(res["items"], list), t
        assert res["count"] == len(res["items"]), t


def test_list_elements_tables_items_keep_v1_shape(rich_doc):
    res = srv.list_elements(str(rich_doc), type="tables")
    assert res["count"] == 1
    item = res["items"][0]
    # v1 list_tables item shape verbatim: index, rows, grid_columns,
    # has_merges, header_preview.
    assert item["rows"] == 2 and item["grid_columns"] == 2
    assert item["index"] == 0


def test_list_elements_bookmarks_and_footnotes(rich_doc):
    marks = srv.list_elements(str(rich_doc), type="bookmarks")
    assert any(b["name"] == "wave_a_mark" for b in marks["items"])
    notes = srv.list_elements(str(rich_doc), type="footnotes")
    assert notes["count"] == 1
    # v1 list_footnotes item shape verbatim: id, position, text.
    assert notes["items"][0]["text"] == "A footnote body."
    assert notes["items"][0]["position"] == 1


def test_list_elements_toc_siblings_carried(plain_doc):
    res = srv.list_elements(str(plain_doc), type="toc")
    # v1 read_toc dict keys carried as siblings of items.
    assert res["present"] is False
    assert res["items"] == []


def test_list_elements_unknown_type_refuses(plain_doc):
    with pytest.raises(WordMcpError, match="template_placeholders"):
        srv.list_elements(str(plain_doc), type="figures")


def test_list_elements_name_filter_on_bookmarks(rich_doc):
    hit = srv.list_elements(
        str(rich_doc), type="bookmarks", filter={"name": "wave_a"}
    )
    assert hit["count"] == 1
    miss = srv.list_elements(
        str(rich_doc), type="bookmarks", filter={"name": "zzz_absent"}
    )
    assert miss["count"] == 0


def test_list_elements_inapplicable_filter_refuses(rich_doc):
    # Per the brief's filter matrix, tables supports neither filter key.
    with pytest.raises(WordMcpError):
        srv.list_elements(
            str(rich_doc), type="tables", filter={"name": "H1"}
        )


# ----------------------------------------------------------------- validate


def test_validate_default_is_core(plain_doc):
    res = srv.validate(str(plain_doc))
    assert set(res["results"]) == {"core"}
    core = res["results"]["core"]
    assert core["passed"] is True
    # v1 validate_document findings shape verbatim.
    f = core["findings"]
    assert f["package_ok"] is True
    assert f["fields_balanced"] is True
    assert res["passed"] is True
    assert "ok" not in res  # the envelope owns top-level ok (6a fix)


def test_validate_battery_runs_multiple_checks(rich_doc):
    # citation_parity is exercised separately: its v1 op refuses documents
    # without a References/Bibliography heading (v1-verbatim behavior).
    checks = ["core", "notes", "accessibility", "captions",
              "cross_references", "reference_fields"]
    res = srv.validate(str(rich_doc), checks=checks)
    assert set(res["results"]) == set(checks)
    for name in checks:
        entry = res["results"][name]
        assert isinstance(entry["passed"], bool), name
        assert isinstance(entry["findings"], dict), name
    # v1 shapes verbatim inside findings.
    assert "summary" in res["results"]["accessibility"]["findings"]
    # v1 validate_notes shape verbatim: per-kind ok flags.
    assert "ok" in res["results"]["notes"]["findings"]["footnotes"]


def test_validate_unknown_check_refuses(plain_doc):
    with pytest.raises(WordMcpError, match="chapter_headers"):
        srv.validate(str(plain_doc), checks=["speling"])


def test_validate_template_requires_options(plain_doc):
    with pytest.raises(WordMcpError, match="template"):
        srv.validate(str(plain_doc), checks=["template"])


def test_validate_template_takes_v1_rules_dict(plain_doc):
    rules = {"fonts": {"allowed": ["Calibri", "Times New Roman"]}}
    res = srv.validate(
        str(plain_doc), checks=["template"], options={"template": rules}
    )
    findings = res["results"]["template"]["findings"]
    assert "compliant" in findings  # v1 shape verbatim


def test_validate_redaction_requires_targets(plain_doc):
    with pytest.raises(WordMcpError, match="targets"):
        srv.validate(str(plain_doc), checks=["redaction"])
    res = srv.validate(
        str(plain_doc), checks=["redaction"],
        options={"redaction": {"targets": [{"find": "sekret"}]}},
    )
    assert res["results"]["redaction"]["findings"]["clean"] is True
    assert res["results"]["redaction"]["passed"] is True


def test_validate_options_for_unrequested_check_refuse(plain_doc):
    with pytest.raises(WordMcpError):
        srv.validate(
            str(plain_doc), checks=["core"],
            options={"redaction": {"targets": [{"find": "x"}]}},
        )


def test_validate_forms_options_routed(plain_doc):
    res = srv.validate(
        str(plain_doc), checks=["forms"],
        options={"forms": {"required": ["missing_field_name"]}},
    )
    findings = res["results"]["forms"]["findings"]
    assert findings["complete"] is False
    assert "missing_field_name" in findings["missing_fields"]
    assert res["results"]["forms"]["passed"] is False
    assert res["passed"] is False


# -------------------------------------------------------------- absorptions


def test_get_text_textbox_mode_shape(plain_doc):
    res = srv.get_text(str(plain_doc), textbox=True, live="off")
    # v1 get_textbox_text dict verbatim (no boxes in this doc).
    assert res["count"] == 0
    assert res["boxes"] == []


def test_get_text_textbox_exclusive_with_body_params(plain_doc):
    with pytest.raises(WordMcpError):
        srv.get_text(str(plain_doc), textbox=True, contains="Alpha",
                     live="off")


def test_find_text_formatting_mode(tmp_path):
    f = tmp_path / "fmt.docx"
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Bold body text")
    run.bold = True
    doc.add_paragraph("Plain body text")
    doc.save(str(f))
    res = srv.find_text(str(f), formatting={"bold": True}, live="off")
    texts = [m["text"] for m in res["matches"]]
    assert "Bold body text" in texts
    assert res["total"] == 1


def test_find_text_formatting_refuses_regex(plain_doc):
    with pytest.raises(WordMcpError):
        srv.find_text(str(plain_doc), query="Alpha", regex=True,
                      formatting={"bold": True}, live="off")


def test_find_text_plain_still_requires_query(plain_doc):
    with pytest.raises(WordMcpError):
        srv.find_text(str(plain_doc), live="off")


def test_word_count_exclusions_invariant(plain_doc):
    res = srv.word_count(str(plain_doc), exclusions=["headings"],
                         live="off")
    # v1 word_count_with_exclusions shape verbatim.
    assert res["total"] == res["included"] + res["excluded_total"]
    assert res["exclusions_applied"] == ["headings"]


def test_word_count_unknown_zone_refuses(plain_doc):
    with pytest.raises(WordMcpError, match="block_quotes"):
        srv.word_count(str(plain_doc), exclusions=["sidebars"], live="off")


def test_word_count_plain_mode_unchanged(plain_doc):
    res = srv.word_count(str(plain_doc), live="off")
    # v1 word_count shape verbatim: totals dict, no exclusions keys.
    assert res["totals"]["words"] > 0
    assert "excluded_total" not in res


def test_manage_backups_snapshot_action(plain_doc, tmp_path):
    (tmp_path / "snaps").mkdir()
    res = srv.manage_backups(
        action="snapshot", file_path=str(plain_doc),
        dest_dir=str(tmp_path / "snaps"),
    )
    snap = Path(res["snapshot"])
    assert snap.exists()
    assert snap.name.endswith("plain.docx")
    assert snap.name[:9].rstrip("_").isdigit() or snap.name[:8].isdigit()
    # Source untouched, still opens.
    Document(str(plain_doc))


def test_manage_backups_snapshot_needs_file_path():
    with pytest.raises(WordMcpError):
        srv.manage_backups(action="snapshot")


def test_set_document_properties_update_fields_flag(plain_doc):
    srv.set_document_properties(
        str(plain_doc), title="Flagged", update_fields_on_open=True,
        backup=False,
    )
    with zipfile.ZipFile(plain_doc) as z:
        settings = z.read("word/settings.xml").decode("utf-8")
    assert "updateFields" in settings
    doc = Document(str(plain_doc))
    assert doc.core_properties.title == "Flagged"


# --------------------------------------------- insert_document via location


def test_insert_document_default_appends_at_end(tmp_path):
    src = _fresh(tmp_path, "src.docx", ["Source body."])
    tgt = _fresh(tmp_path, "tgt.docx", ["Target one.", "Target two."])
    srv.insert_document(str(tgt), str(src), backup=False)
    texts = [p["text"] for p in srv.get_text(str(tgt), live="off")]
    assert texts[-1] == "Source body."


def test_insert_document_location_paragraph(tmp_path):
    src = _fresh(tmp_path, "src.docx", ["Inserted body."])
    tgt = _fresh(tmp_path, "tgt.docx", ["Target one.", "Target two."])
    srv.insert_document(
        str(tgt), str(src), location={"paragraph": 0}, backup=False
    )
    texts = [p["text"] for p in srv.get_text(str(tgt), live="off")]
    assert texts.index("Inserted body.") == texts.index("Target one.") + 1


def test_insert_document_ambiguous_search_refuses(tmp_path):
    from word_mcp.core.errors import AmbiguousTarget

    src = _fresh(tmp_path, "src.docx", ["Inserted body."])
    tgt = _fresh(tmp_path, "tgt.docx", ["Dup text.", "Dup text."])
    with pytest.raises(AmbiguousTarget):
        srv.insert_document(
            str(tgt), str(src),
            location={"search": {"text": "Dup text"}}, backup=False,
        )


def test_insert_document_rejects_unsupported_position(tmp_path):
    src = _fresh(tmp_path, "src.docx", ["Inserted body."])
    tgt = _fresh(tmp_path, "tgt.docx", ["Target one."])
    with pytest.raises(WordMcpError):
        srv.insert_document(
            str(tgt), str(src),
            location={"paragraph": 0, "position": "replace"}, backup=False,
        )


# ------------------------------------------------- retired names are gone


def test_v1_names_not_registered():
    retired = [
        "create_snapshot", "set_update_fields_flag", "get_textbox_text",
        "find_formatted", "word_count_with_exclusions", "read_toc",
        "get_lists", "list_tables", "list_images", "list_charts",
        "list_equations", "list_bookmarks", "list_sources",
        "list_sections", "list_section_blocks", "list_footnotes",
        "list_endnotes", "list_fields", "list_reference_fields",
        "list_form_fields", "list_content_controls",
        "list_template_placeholders", "list_index_entries",
        "validate_document", "validate_captions",
        "validate_chapter_headers", "validate_form_completeness",
        "check_defined_terms", "check_brand_compliance",
        "check_template_compliance", "audit_accessibility",
    ]
    import asyncio

    shipped = {t.name for t in asyncio.run(srv.mcp.list_tools())}
    still_there = [n for n in retired if n in shipped]
    assert not still_there, f"retired v1 names still registered: {still_there}"
