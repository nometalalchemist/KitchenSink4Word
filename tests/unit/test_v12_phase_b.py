"""v1.2 Phase B: format long-tail, gridBefore tables, table features, columns."""

import shutil
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from word_mcp.core.errors import UnsupportedStructure, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import furniture as fu, read, tables as tb, text as tx

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def doc(tmp_path):
    dst = tmp_path / "d.docx"
    d = Document()
    d.add_paragraph("First paragraph with MixedCase words here. second sentence.")
    d.add_paragraph("Another paragraph entirely.")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tb.create_table(
        pkg,
        [["Name", "Score"], ["delta", "30"], ["alpha", "10"], ["charlie", "20"]],
        at_end=True,
    )
    pkg.save(do_backup=False)
    return dst


# ------------------------------------------------------- char format long-tail


def test_new_char_format_keys(doc):
    pkg = DocxPackage(doc)
    tx.format_text(
        pkg,
        find="MixedCase",
        formatting={
            "small_caps": True,
            "double_strike": True,
            "char_spacing_pt": 1.5,
            "kerning_pt": 12,
            "language": "en-US",
            "east_asian_language": "ko-KR",
        },
    )
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    for frag in (
        "smallCaps", "dstrike", '<w:spacing w:val="30"/>', 'w:kern w:val="24"',
        'w:eastAsia="ko-KR"',
    ):
        assert frag in xml, frag
    Document(str(doc))


def test_hidden_text_key(doc):
    pkg = DocxPackage(doc)
    tx.format_text(pkg, find="Another", formatting={"hidden": True})
    pkg.save(do_backup=False)
    assert "vanish" in DocxPackage(doc).raw_part("word/document.xml").decode()


# ------------------------------------------------------ para format long-tail


def test_borders_shading_tabs(doc):
    pkg = DocxPackage(doc)
    tx.set_paragraph_format(
        pkg,
        [0],
        {
            "borders": "all",
            "shading": "FFF2CC",
            "tab_stops": [
                {"position_pt": 216, "alignment": "center"},
                {"position_pt": 432, "alignment": "right", "leader": "dot"},
            ],
            "widow_control": True,
            "keep_lines_together": True,
            "page_break_before": True,
        },
    )
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    for frag in (
        "pBdr", 'w:fill="FFF2CC"', 'w:leader="dot"', 'w:pos="8640"',
        "keepLines", "pageBreakBefore",
    ):
        assert frag in xml, frag
    # Schema order: pBdr before shd before tabs before spacing/jc.
    assert xml.index("pBdr") < xml.index('w:fill="FFF2CC"')
    Document(str(doc))


def test_ppr_schema_order_with_jc(doc):
    """jc added first, then borders: pBdr must still land BEFORE jc."""
    pkg = DocxPackage(doc)
    tx.set_paragraph_format(pkg, [1], {"alignment": "center"})
    tx.set_paragraph_format(pkg, [1], {"borders": ["top"]})
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    p = [el for k, i, el in read.body_items(pkg2) if k == "paragraph" and i == 1][0]
    ppr = p.find(qn("w:pPr"))
    names = [etree.QName(c).localname for c in ppr]
    assert names.index("pBdr") < names.index("jc")


def test_bad_tab_alignment(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError):
        tx.set_paragraph_format(
            pkg, [0], {"tab_stops": [{"position_pt": 100, "alignment": "diagonal"}]}
        )


# ------------------------------------------------------------------ change_case


def test_change_case_upper_and_title(doc):
    pkg = DocxPackage(doc)
    tx.change_case(pkg, "upper", indices=[1])
    pkg.save(do_backup=False)
    paras = read.get_paragraphs(DocxPackage(doc))
    assert paras[1]["text"] == "ANOTHER PARAGRAPH ENTIRELY."

    pkg = DocxPackage(doc)
    tx.change_case(pkg, "title", indices=[1])
    pkg.save(do_backup=False)
    paras = read.get_paragraphs(DocxPackage(doc))
    assert paras[1]["text"] == "Another Paragraph Entirely."


def test_change_case_sentence(doc):
    pkg = DocxPackage(doc)
    tx.change_case(pkg, "sentence", indices=[0])
    pkg.save(do_backup=False)
    text = read.get_paragraphs(DocxPackage(doc))[0]["text"]
    assert text == "First paragraph with mixedcase words here. Second sentence."


def test_change_case_find_scope(doc):
    pkg = DocxPackage(doc)
    tx.change_case(pkg, "upper", find="another")
    pkg.save(do_backup=False)
    text = read.get_paragraphs(DocxPackage(doc))[1]["text"]
    assert "ANOTHER" in text
    assert "paragraph entirely" in text  # rest untouched


# ------------------------------------------------------------ gridBefore tables


def make_gridbefore_table(path):
    """Build a table with Word-style gridBefore/gridAfter rows (per the
    research ground truth: 3-col grid; row 1 skips last col, row 2 skips
    first col)."""
    pkg = DocxPackage(path)
    tb.create_table(pkg, [["a", "b", "c"], ["d", "e", "f"]], at_end=True)
    pkg.save(do_backup=False)
    pkg = DocxPackage(path)
    tbl = [el for k, i, el in read.body_items(pkg) if k == "table"][-1]
    rows = tbl.findall(qn("w:tr"))
    # Row 0: drop last cell, gridAfter=1. Row 1: drop first cell, gridBefore=1.
    r0_cells = rows[0].findall(qn("w:tc"))
    rows[0].remove(r0_cells[2])
    tb._set_row_zone(tbl, rows[0], "after", 1)
    r1_cells = rows[1].findall(qn("w:tc"))
    rows[1].remove(r1_cells[0])
    tb._set_row_zone(tbl, rows[1], "before", 1)
    pkg.mark_dirty()
    pkg.save(do_backup=False)
    return len([x for x, i, e in read.body_items(pkg) if x == "table"]) - 1


def test_gridbefore_table_readable_and_consistent(doc):
    idx = make_gridbefore_table(doc)
    g = read.get_table(DocxPackage(doc), idx)
    assert g["grid_columns"] == 3
    assert g["rows"] == 2
    Document(str(doc))


def test_gridbefore_column_delete_middle(doc):
    """Deleting the middle column hits real cells in both rows."""
    idx = make_gridbefore_table(doc)
    pkg = DocxPackage(doc)
    tb.delete_columns(pkg, idx, [1])
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    g = read.get_table(pkg2, idx)
    assert g["grid_columns"] == 2
    tbl = [el for k, i, el in read.body_items(pkg2) if k == "table"][idx]
    rows = tbl.findall(qn("w:tr"))
    assert tb._row_zones(rows[0]) == (0, 1)
    assert tb._row_zones(rows[1]) == (1, 0)
    Document(str(doc))


def test_gridbefore_column_delete_in_skip_zone(doc):
    """Deleting column 0 consumes row 1's gridBefore, removes row 0's cell."""
    idx = make_gridbefore_table(doc)
    pkg = DocxPackage(doc)
    tb.delete_columns(pkg, idx, [0])
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    g = read.get_table(pkg2, idx)
    assert g["grid_columns"] == 2
    tbl = [el for k, i, el in read.body_items(pkg2) if k == "table"][idx]
    rows = tbl.findall(qn("w:tr"))
    assert tb._row_zones(rows[1]) == (0, 0)  # gridBefore consumed
    assert [c["text"] for c in g["cells"][1]] == ["e", "f"]
    Document(str(doc))


def test_gridbefore_insert_at_start(doc):
    idx = make_gridbefore_table(doc)
    pkg = DocxPackage(doc)
    tb.insert_columns(pkg, idx, at=0)
    pkg.save(do_backup=False)
    g = read.get_table(DocxPackage(doc), idx)
    assert g["grid_columns"] == 4
    Document(str(doc))


# ------------------------------------------------------------- table features


def test_apply_table_style(doc):
    pkg = DocxPackage(doc)
    tb.apply_table_style(pkg, 0, "BandedTable")
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    xml = pkg2.raw_part("word/document.xml").decode()
    assert 'w:tblStyle w:val="BandedTable"' in xml.replace("<", " ").replace(
        ">", " "
    ) or 'val="BandedTable"' in xml
    styles = pkg2.raw_part("word/styles.xml").decode()
    assert "band1Horz" in styles
    assert "TableNormal" in styles
    Document(str(doc))


def test_sort_table_text_and_numeric(doc):
    pkg = DocxPackage(doc)
    tb.sort_table(pkg, 0, column=0)
    pkg.save(do_backup=False)
    g = read.get_table(DocxPackage(doc), 0)
    assert [r[0]["text"] for r in g["cells"]] == ["Name", "alpha", "charlie", "delta"]

    pkg = DocxPackage(doc)
    tb.sort_table(pkg, 0, column=1, numeric=True, descending=True)
    pkg.save(do_backup=False)
    g = read.get_table(DocxPackage(doc), 0)
    assert [r[1]["text"] for r in g["cells"]] == ["Score", "30", "20", "10"]


def test_sort_refuses_vmerge(doc):
    pkg = DocxPackage(doc)
    tb.merge_cells(pkg, 0, start_row=1, end_row=2, start_col=0, end_col=0)
    with pytest.raises(UnsupportedStructure):
        tb.sort_table(pkg, 0, column=0)


def test_split_table(doc):
    pkg = DocxPackage(doc)
    n_before = len(read.list_tables(pkg))
    tb.split_table(pkg, 0, at_row=2)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    tabs = read.list_tables(pkg2)
    assert len(tabs) == n_before + 1
    g0 = read.get_table(pkg2, 0)
    g1 = read.get_table(pkg2, 1)
    assert g0["rows"] == 2 and g1["rows"] == 2
    assert g1["cells"][0][0]["text"] == "alpha"
    Document(str(doc))


def test_header_row_repeat_toggle(doc):
    pkg = DocxPackage(doc)
    tb.set_header_row_repeat(pkg, 0, rows=1, on=True)
    pkg.save(do_backup=False)
    assert "tblHeader" in DocxPackage(doc).raw_part("word/document.xml").decode()
    pkg = DocxPackage(doc)
    tb.set_header_row_repeat(pkg, 0, on=False)
    pkg.save(do_backup=False)
    assert "tblHeader" not in DocxPackage(doc).raw_part(
        "word/document.xml"
    ).decode()


def test_nested_table_read_write(doc):
    pkg = DocxPackage(doc)
    tbl = [el for k, i, el in read.body_items(pkg) if k == "table"][0]
    tc = tbl.findall(qn("w:tr"))[1].findall(qn("w:tc"))[1]
    inner = etree.SubElement(tc, qn("w:tbl"))
    grid = etree.SubElement(inner, qn("w:tblGrid"))
    for _ in range(2):
        etree.SubElement(grid, qn("w:gridCol")).set(qn("w:w"), "1000")
    for vals in (["n1", "n2"], ["n3", "n4"]):
        tr = etree.SubElement(inner, qn("w:tr"))
        for v in vals:
            tc2 = etree.SubElement(tr, qn("w:tc"))
            p = etree.SubElement(tc2, qn("w:p"))
            r = etree.SubElement(p, qn("w:r"))
            etree.SubElement(r, qn("w:t")).text = v
    etree.SubElement(tc, qn("w:p"))  # trailing para required in cells
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    nested = tb.get_nested_table(pkg2, 0, row=1, cell=1)
    assert nested["cells"][0][0]["text"] == "n1"
    tb.set_nested_cells(
        pkg2, 0, row=1, cell=1, edits=[{"row": 1, "cell": 1, "text": "EDITED"}]
    )
    pkg2.save(do_backup=False)
    nested = tb.get_nested_table(DocxPackage(doc), 0, row=1, cell=1)
    assert nested["cells"][1][1]["text"] == "EDITED"
    Document(str(doc))


def test_cell_text_direction(doc):
    pkg = DocxPackage(doc)
    tb.format_cells(pkg, 0, [{"row": 0, "cell": 0}], {"text_direction": "btLr"})
    pkg.save(do_backup=False)
    assert 'w:textDirection w:val="btLr"' in DocxPackage(doc).raw_part(
        "word/document.xml"
    ).decode().replace("<", " ").replace("/>", " ") or "btLr" in DocxPackage(
        doc
    ).raw_part("word/document.xml").decode()


# ---------------------------------------------------------------- columns/pages


def test_multi_column_section(doc):
    pkg = DocxPackage(doc)
    fu.set_columns(pkg, count=2, separator=True)
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert 'w:num="2"' in xml and 'w:sep="1"' in xml
    Document(str(doc))


def test_unequal_columns_validated(doc):
    pkg = DocxPackage(doc)
    fu.set_columns(pkg, count=2, widths_pt=[200, 250])
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert 'w:equalWidth="0"' in xml
    with pytest.raises(WordMcpError):
        fu.set_columns(DocxPackage(doc), count=3, widths_pt=[100, 100])


def test_page_x_of_y(doc):
    pkg = DocxPackage(doc)
    fu.add_page_numbers(pkg, x_of_y=True)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    footers = [n for n in pkg2.part_names() if n.startswith("word/footer")]
    combined = "".join(
        pkg2.raw_part(f).decode("utf-8", errors="replace") for f in footers
    )
    assert "NUMPAGES" in combined
    assert "Page " in combined
    Document(str(doc))
