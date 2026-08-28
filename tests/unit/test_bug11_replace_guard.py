"""Bug 11 (2026-08-28 Ch5 session): index-addressed replace_paragraph_text
hitting the wrong paragraph after index shifts looked like duplicate
insertion. Guard: optional expect substring refuses on mismatch with zero
mutation; replaced_text (the old text) is always returned for
verification."""

import hashlib
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import text as tx


@pytest.fixture()
def doc(tmp_path: Path) -> Path:
    p = tmp_path / "bug11.docx"
    d = Document()
    d.add_paragraph("First paragraph about goals.")
    d.add_paragraph("Second paragraph about tasks.")
    d.add_paragraph("Third paragraph about bonds.")
    d.save(str(p))
    return p


def md5(p: Path) -> str:
    return hashlib.md5(p.read_bytes()).hexdigest()


def test_expect_match_replaces_and_returns_old_text(doc: Path) -> None:
    pkg = DocxPackage(str(doc))
    result = tx.replace_paragraph_text(
        pkg, 1, "Rewritten tasks paragraph.", expect="about tasks"
    )
    pkg.save(do_backup=False)
    assert result["replaced_paragraph"] == 1
    assert result["replaced_text"] == "Second paragraph about tasks."
    out = Document(str(doc))
    assert out.paragraphs[1].text == "Rewritten tasks paragraph."
    assert len(out.paragraphs) == 3  # replaced in place, never inserted


def test_expect_mismatch_refuses_with_zero_mutation(doc: Path) -> None:
    before = md5(doc)
    pkg = DocxPackage(str(doc))
    with pytest.raises(WordMcpError) as exc:
        # simulates the stale index: caller thinks 0 is the tasks paragraph
        tx.replace_paragraph_text(
            pkg, 0, "Rewritten tasks paragraph.", expect="about tasks"
        )
    msg = str(exc.value)
    assert "does not contain the expected text" in msg
    assert "First paragraph about goals." in msg  # names what IS there
    assert "Nothing was changed" in msg
    assert md5(doc) == before


def test_no_expect_still_returns_replaced_text(doc: Path) -> None:
    pkg = DocxPackage(str(doc))
    result = tx.replace_paragraph_text(pkg, 2, "New third.")
    pkg.save(do_backup=False)
    assert result["replaced_text"] == "Third paragraph about bonds."
    assert Document(str(doc)).paragraphs[2].text == "New third."
