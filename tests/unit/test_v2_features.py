"""v1.1 features: tracked-change writing, caption lists, note conversion.
(COM compare is tested in the gate script — it needs real Word.)"""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import notes, read, revisions as rv, tables as tb, text as tx, toc

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


@pytest.fixture
def niu(tmp_path):
    dst = tmp_path / "niu.docx"
    shutil.copy(CORPUS / "niu.docx", dst)
    return dst


def visible(path):
    return "\n".join(e["text"] for e in read.get_paragraphs(DocxPackage(path)))


# ------------------------------------------------------- tracked replace


def test_tracked_replace_shows_as_revision(ch4):
    pkg = DocxPackage(ch4)
    result = tx.search_and_replace(
        pkg,
        [{"find": "Delta Model", "replace": "Delta Framework"}],
        track=True,
        author="Test Author",
    )
    pkg.save(do_backup=False)
    assert result["tracked_as"] == "Test Author"
    assert result["total"] > 0

    pkg2 = DocxPackage(ch4)
    summary = read.revision_summary(pkg2)
    assert summary["by_author"].get("Test Author", 0) >= result["total"]
    # Visible text shows the replacement (insertions display, deletions hide).
    text = visible(ch4)
    assert "Delta Framework" in text
    assert "Delta Model" not in text
    Document(str(ch4))


def test_tracked_replace_reject_restores_original(ch4):
    before = visible(ch4)
    pkg = DocxPackage(ch4)
    tx.search_and_replace(
        pkg,
        [{"find": "Delta Model", "replace": "WRONG TERM"}],
        track=True,
        author="Reviewer",
    )
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    rv.reject_revisions(pkg2, author="Reviewer")
    pkg2.save(do_backup=False)
    assert visible(ch4) == before
    assert read.revision_summary(DocxPackage(ch4))["total"] == 0


def test_tracked_replace_accept_keeps_replacement(ch4):
    pkg = DocxPackage(ch4)
    tx.search_and_replace(
        pkg,
        [{"find": "Delta Model", "replace": "Delta Framework"}],
        track=True,
    )
    pkg.save(do_backup=False)
    before_accept = visible(ch4)

    pkg2 = DocxPackage(ch4)
    rv.accept_revisions(pkg2)
    pkg2.save(do_backup=False)
    assert visible(ch4) == before_accept
    assert read.revision_summary(DocxPackage(ch4))["total"] == 0
    assert "Delta Framework" in visible(ch4)


# ----------------------------------------- tracked insert/delete paragraphs


def test_tracked_insert_and_reject(ch4):
    before = visible(ch4)
    pkg = DocxPackage(ch4)
    tx.insert_paragraphs(
        pkg,
        [{"text": "A tracked new paragraph."}],
        after_index=3,
        track=True,
        author="Author",
    )
    pkg.save(do_backup=False)
    assert "A tracked new paragraph." in visible(ch4)
    assert read.revision_summary(DocxPackage(ch4))["by_author"]["Author"] >= 1

    pkg2 = DocxPackage(ch4)
    rv.reject_revisions(pkg2, author="Author")
    pkg2.save(do_backup=False)
    assert visible(ch4) == before


def test_tracked_delete_and_reject(ch4):
    before = visible(ch4)
    pkg = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg)
    victim = next(p for p in paras if len(p["text"]) > 80)
    result = tx.delete_paragraphs(
        pkg, victim["index"], track=True, author="Chair"
    )
    pkg.save(do_backup=False)
    assert result["deleted_tracked"] == 1
    # Text hidden but nothing physically removed.
    assert victim["text"] not in visible(ch4)
    n_paras = read.get_document_info(DocxPackage(ch4))["paragraphs"]
    assert n_paras == len(paras)

    pkg2 = DocxPackage(ch4)
    rv.reject_revisions(pkg2, author="Chair")
    pkg2.save(do_backup=False)
    assert visible(ch4) == before


def test_tracked_delete_accept_removes(ch4):
    pkg = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg)
    victim = next(p for p in paras if len(p["text"]) > 80)
    tx.delete_paragraphs(pkg, victim["index"], track=True)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    rv.accept_revisions(pkg2)
    pkg2.save(do_backup=False)
    assert victim["text"] not in visible(ch4)
    assert read.revision_summary(DocxPackage(ch4))["total"] == 0
    Document(str(ch4))


# ------------------------------------------------------- tracked cell edits


def test_tracked_set_cells_roundtrip(ch4):
    pkg = DocxPackage(ch4)
    tb.create_table(pkg, [["A", "B"], ["old text", "keep"]], at_end=True)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    idx = len(read.list_tables(pkg2)) - 1
    tb.set_cells(
        pkg2,
        idx,
        [{"row": 1, "cell": 0, "text": "new text"}],
        track=True,
        author="Editor",
    )
    pkg2.save(do_backup=False)
    g = read.get_table(DocxPackage(ch4), idx)
    assert g["cells"][1][0]["text"] == "new text"

    pkg3 = DocxPackage(ch4)
    rv.reject_revisions(pkg3, author="Editor")
    pkg3.save(do_backup=False)
    g = read.get_table(DocxPackage(ch4), idx)
    assert g["cells"][1][0]["text"] == "old text"


# ----------------------------------------------------------- caption lists


def test_insert_caption_list(ch4):
    from word_mcp.ops import fields as fl

    pkg = DocxPackage(ch4)
    tb.create_table(pkg, [["x"]], at_end=True)
    fl.add_caption(pkg, table_index=0, label="Table", text="First table")
    toc.insert_caption_list(pkg, label="Table", at_start=True, update_on_open=False)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    info = toc.read_toc(pkg2)
    assert info["present"]
    kinds = [t["kind"] for t in info["tocs"]]
    assert "caption_list" in kinds
    xml = pkg2.raw_part("word/document.xml").decode("utf-8")
    assert '\\c "Table"' in xml
    assert "List of Tables" in xml
    Document(str(ch4))


def test_duplicate_caption_list_refused(ch4):
    pkg = DocxPackage(ch4)
    toc.insert_caption_list(pkg, label="Figure", at_start=True, update_on_open=False)
    with pytest.raises(WordMcpError, match="already has"):
        toc.insert_caption_list(pkg, label="Figure", at_start=True)
    # A different label is fine, and so is a main TOC.
    toc.insert_caption_list(pkg, label="Table", at_start=True, update_on_open=False)
    toc.insert_toc(pkg, at_start=True, update_on_open=False)


# ---------------------------------------------------------- note conversion


def test_convert_single_footnote_to_endnote(niu):
    pkg = DocxPackage(niu)
    fn_before = read.list_footnotes(pkg)
    target_text = fn_before[2]["text"]
    result = notes.convert_notes(
        pkg, "footnotes_to_endnotes", position=3
    )
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(niu)
    assert result["converted"] == 1
    assert len(read.list_footnotes(pkg2)) == len(fn_before) - 1
    ens = read.list_endnotes(pkg2)
    assert len(ens) == 1
    assert ens[0]["text"] == target_text
    v = notes.validate_notes(pkg2)
    assert v["footnotes"]["ok"] and v["endnotes"]["ok"]
    Document(str(niu))


def test_convert_all_and_back(niu):
    pkg = DocxPackage(niu)
    texts_before = [n["text"] for n in read.list_footnotes(pkg)]
    result = notes.convert_notes(pkg, "footnotes_to_endnotes")
    pkg.save(do_backup=False)
    assert result["converted"] == 171

    pkg2 = DocxPackage(niu)
    assert len(read.list_footnotes(pkg2)) == 0
    ens = read.list_endnotes(pkg2)
    assert len(ens) == 171
    assert [n["text"] for n in ens] == texts_before

    notes.convert_notes(pkg2, "endnotes_to_footnotes")
    pkg2.save(do_backup=False)
    pkg3 = DocxPackage(niu)
    assert len(read.list_endnotes(pkg3)) == 0
    fns = read.list_footnotes(pkg3)
    assert [n["text"] for n in fns] == texts_before
    v = notes.validate_notes(pkg3)
    assert v["footnotes"]["ok"] and not v["footnotes"]["needs_cleanup"]
    Document(str(niu))


def test_convert_bad_direction(niu):
    pkg = DocxPackage(niu)
    with pytest.raises(WordMcpError):
        notes.convert_notes(pkg, "sideways")
