"""v2 delete_element parity-gap ops: image / chart / bookmark / hyperlink.

These four deletions did not exist in v1.6 (V2_DESIGN Section 3.2 closes
the gap). Ops-layer only; the delete_element server multiplex dispatches to
them at Phase 2 integration per integration/v2_briefs/wave_D.md.

Self-contained: documents are built on the fly with python-docx plus the
package's own ops; assertions parse our own output structurally.
"""

import copy
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from word_mcp.core.errors import (
    AmbiguousTarget,
    TargetNotFound,
    WordMcpError,
)
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import charts as ch
from word_mcp.ops import fields, media

A = media._A
R = media._R_NS
CT_NS = media._CT_NS


# ------------------------------------------------------------------- helpers


def make_png(path: Path, w=8, h=4, rgb=(200, 30, 30)):
    """Minimal valid PNG without external deps."""

    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(
            ">I", zlib.crc32(c) & 0xFFFFFFFF
        )

    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    path.write_bytes(png)
    return path


def fresh_doc(tmp_path: Path, name="doc.docx") -> Path:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("First paragraph, always present.")
    doc.add_paragraph("Second paragraph for positioning.")
    doc.save(str(f))
    return f


def doc_with_image(tmp_path: Path, name="doc.docx"):
    f = fresh_doc(tmp_path, name)
    png = make_png(tmp_path / "pic.png")
    pkg = DocxPackage(f)
    media.add_image(pkg, str(png), at_end=True)
    pkg.save(do_backup=False)
    return f


CHART_DATA = {
    "categories": ["Alpha", "Beta"],
    "series": [{"name": "S1", "values": [1, 2]}],
}


def body_paragraph_count(path: Path) -> int:
    pkg = DocxPackage(path)
    return len(pkg.body().findall(qn("w:p")))


# -------------------------------------------------------------- delete_image


def test_delete_image_removes_drawing_rel_and_part(tmp_path):
    f = doc_with_image(tmp_path)
    before_paras = body_paragraph_count(f)
    pkg = DocxPackage(f)
    assert len(media.list_images(pkg)) == 1
    result = media.delete_image(pkg, 0)
    pkg.save(do_backup=False)

    assert result["deleted_image"] == 0
    assert result["removed_paragraph"] is True
    assert result["media_part"] == "word/media/image1.png"
    assert result["media_part_removed"] is True

    pkg2 = DocxPackage(f)
    assert media.list_images(pkg2) == []
    assert not pkg2.has_part("word/media/image1.png")
    rels = pkg2.root("word/_rels/document.xml.rels")
    assert not any(
        (r.get("Type") or "").endswith("/image") for r in rels
    )
    assert body_paragraph_count(f) == before_paras - 1


def test_delete_image_out_of_range(tmp_path):
    f = doc_with_image(tmp_path)
    pkg = DocxPackage(f)
    with pytest.raises(TargetNotFound):
        media.delete_image(pkg, 5)


def test_delete_image_shared_rel_keeps_part(tmp_path):
    """Two drawings sharing one relationship id: deleting one keeps the
    relationship and the media part alive for the other."""
    f = doc_with_image(tmp_path)
    pkg = DocxPackage(f)
    body = pkg.body()
    img_p = next(
        p for p in body.findall(qn("w:p"))
        if p.find(f".//{{{A}}}blip") is not None
    )
    img_p.addnext(copy.deepcopy(img_p))
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    assert len(media.list_images(pkg)) == 2
    result = media.delete_image(pkg, 0)
    pkg.save(do_backup=False)
    assert result["media_part_removed"] is False

    pkg2 = DocxPackage(f)
    assert len(media.list_images(pkg2)) == 1
    assert pkg2.has_part("word/media/image1.png")
    rels = pkg2.root("word/_rels/document.xml.rels")
    assert any((r.get("Type") or "").endswith("/image") for r in rels)


def test_delete_image_part_kept_when_other_rels_reference_it(tmp_path):
    """A media part still targeted by another rels part (header pattern)
    survives even after the body drawing and document rel go."""
    f = doc_with_image(tmp_path)
    pkg = DocxPackage(f)
    other_rels = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<Relationships xmlns='
        b'"http://schemas.openxmlformats.org/package/2006/relationships">'
        b'<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/'
        b'officeDocument/2006/relationships/image" '
        b'Target="media/image1.png"/></Relationships>'
    )
    pkg.set_raw_part("word/_rels/header9.xml.rels", other_rels)
    result = media.delete_image(pkg, 0)
    assert result["media_part_removed"] is False
    assert pkg.has_part("word/media/image1.png")


def test_delete_image_inline_in_shared_paragraph_keeps_paragraph(tmp_path):
    """A drawing sharing its paragraph with text: only the drawing goes."""
    f = fresh_doc(tmp_path)
    png = make_png(tmp_path / "pic.png")
    pkg = DocxPackage(f)
    media.add_image(pkg, str(png), at_end=True)
    img_p = next(
        p for p in pkg.body().findall(qn("w:p"))
        if p.find(f".//{{{A}}}blip") is not None
    )
    r = etree.SubElement(img_p, qn("w:r"))
    etree.SubElement(r, qn("w:t")).text = "caption text stays"
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    before = body_paragraph_count(f)
    result = media.delete_image(pkg, 0)
    pkg.save(do_backup=False)
    assert result["removed_paragraph"] is False
    assert body_paragraph_count(f) == before
    text = "\n".join(p.text for p in Document(str(f)).paragraphs)
    assert "caption text stays" in text


# -------------------------------------------------------------- delete_chart


def test_delete_chart_removes_everything(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    ch.add_chart(pkg, "column", CHART_DATA, at_end=True)
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    listed = ch.list_charts(pkg)
    assert len(listed) == 1
    chart_part = listed[0]["part"]
    result = ch.delete_chart(pkg, 0)
    pkg.save(do_backup=False)

    assert result["deleted_chart"] == 0
    assert result["kind"] == "chart"
    assert result["part"] == chart_part
    assert chart_part in result["removed_parts"]
    assert result["removed_paragraph"] is True

    pkg2 = DocxPackage(f)
    assert ch.list_charts(pkg2) == []
    assert not pkg2.has_part(chart_part)
    assert not any(
        n.startswith("word/embeddings/") for n in pkg2.part_names()
    )
    assert not any(n.startswith("word/charts/") for n in pkg2.part_names())
    rels = pkg2.root("word/_rels/document.xml.rels")
    assert not any((r.get("Type") or "").endswith("/chart") for r in rels)
    ct = pkg2.root("[Content_Types].xml")
    assert not any(
        (o.get("PartName") or "").startswith("/word/charts/")
        for o in ct.findall(f"{{{CT_NS}}}Override")
    )


def test_delete_chart_out_of_range(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    with pytest.raises(TargetNotFound):
        ch.delete_chart(pkg, 0)


def test_delete_first_of_two_charts_keeps_second(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    ch.add_chart(pkg, "column", CHART_DATA, title="One", at_end=True)
    ch.add_chart(pkg, "line", CHART_DATA, title="Two", at_end=True)
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    result = ch.delete_chart(pkg, 0)
    pkg.save(do_backup=False)
    assert result["deleted_chart"] == 0

    pkg2 = DocxPackage(f)
    listed = ch.list_charts(pkg2)
    assert len(listed) == 1
    assert listed[0]["title"] == "Two"
    assert listed[0]["index"] == 0
    assert pkg2.has_part(listed[0]["part"])
    assert listed[0]["embedded_workbook"] is True


# ----------------------------------------------------------- delete_bookmark


def test_delete_bookmark_keeps_text(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    fields.add_bookmark(pkg, "target_one", anchor_text="Second paragraph")
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    assert [b["name"] for b in fields.list_bookmarks(pkg)] == ["target_one"]
    result = fields.delete_bookmark(pkg, "target_one")
    pkg.save(do_backup=False)
    assert result["deleted_bookmark"] == "target_one"
    assert result["markers_removed"] == 1
    assert result["dangling_references"] == 0

    pkg2 = DocxPackage(f)
    assert fields.list_bookmarks(pkg2) == []
    assert not list(pkg2.root().iter(qn("w:bookmarkStart")))
    assert not list(pkg2.root().iter(qn("w:bookmarkEnd")))
    text = "\n".join(p.text for p in Document(str(f)).paragraphs)
    assert "Second paragraph for positioning." in text


def test_delete_bookmark_reports_dangling_references(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    fields.add_bookmark(pkg, "meth_sec", anchor_text="Second paragraph")
    fields.add_cross_reference(
        pkg, after_anchor="First paragraph", to_bookmark="meth_sec"
    )
    result = fields.delete_bookmark(pkg, "meth_sec")
    assert result["dangling_references"] == 1


def test_delete_bookmark_refuses_internal_and_missing(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    with pytest.raises(WordMcpError, match="Word-internal"):
        fields.delete_bookmark(pkg, "_Toc12345")
    with pytest.raises(TargetNotFound):
        fields.delete_bookmark(pkg, "never_existed")


# ---------------------------------------------------------- delete_hyperlink


def linked_doc(tmp_path, name="doc.docx"):
    f = fresh_doc(tmp_path, name)
    pkg = DocxPackage(f)
    fields.add_hyperlink(
        pkg, anchor_text="Second paragraph", url="https://example.org/a"
    )
    pkg.save(do_backup=False)
    return f


def test_delete_hyperlink_unwraps_and_drops_rel(tmp_path):
    f = linked_doc(tmp_path)
    pkg = DocxPackage(f)
    result = fields.delete_hyperlink(pkg, anchor_text="Second paragraph")
    pkg.save(do_backup=False)
    assert result["deleted_hyperlink"] == "https://example.org/a"
    assert result["unwrapped_runs"] >= 1

    pkg2 = DocxPackage(f)
    assert not list(pkg2.root().iter(qn("w:hyperlink")))
    rels = pkg2.root("word/_rels/document.xml.rels")
    assert not any(
        (r.get("Type") or "").endswith("/hyperlink") for r in rels
    )
    text = "\n".join(p.text for p in Document(str(f)).paragraphs)
    assert "Second paragraph for positioning." in text
    # The Hyperlink character style reference is stripped from the runs.
    assert not [
        st
        for st in pkg2.root().iter(qn("w:rStyle"))
        if st.get(qn("w:val")) == "Hyperlink"
    ]


def test_delete_hyperlink_by_url(tmp_path):
    f = linked_doc(tmp_path)
    pkg = DocxPackage(f)
    result = fields.delete_hyperlink(pkg, url="https://example.org/a")
    assert result["text"] == "Second paragraph"


def test_delete_hyperlink_ambiguous_refuses_with_matches(tmp_path):
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    fields.add_hyperlink(
        pkg, anchor_text="paragraph", occurrence=1, url="https://one.example"
    )
    fields.add_hyperlink(
        pkg, anchor_text="paragraph", occurrence=2, url="https://two.example"
    )
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    with pytest.raises(AmbiguousTarget) as exc_info:
        fields.delete_hyperlink(pkg, anchor_text="paragraph")
    matches = exc_info.value.matches
    assert len(matches) == 2
    assert matches[0]["occurrence"] == 1
    assert matches[1]["url"] == "https://two.example"

    # occurrence picks deterministically in document order.
    result = fields.delete_hyperlink(pkg, anchor_text="paragraph", occurrence=2)
    pkg.save(do_backup=False)
    assert result["deleted_hyperlink"] == "https://two.example"
    pkg2 = DocxPackage(f)
    remaining = list(pkg2.root().iter(qn("w:hyperlink")))
    assert len(remaining) == 1


def test_delete_hyperlink_param_and_notfound_guards(tmp_path):
    f = linked_doc(tmp_path)
    pkg = DocxPackage(f)
    with pytest.raises(WordMcpError, match="anchor_text or url"):
        fields.delete_hyperlink(pkg)
    with pytest.raises(TargetNotFound):
        fields.delete_hyperlink(pkg, url="https://nowhere.example")
    with pytest.raises(TargetNotFound, match="occurrence"):
        fields.delete_hyperlink(
            pkg, anchor_text="Second paragraph", occurrence=4
        )


def test_delete_internal_anchor_hyperlink(tmp_path):
    """w:hyperlink with w:anchor (no relationship): unwraps cleanly, href
    reported as '#name'."""
    f = fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    fields.add_bookmark(pkg, "sec_two", anchor_text="Second paragraph")
    p = pkg.body().findall(qn("w:p"))[0]
    link = etree.SubElement(p, qn("w:hyperlink"))
    link.set(qn("w:anchor"), "sec_two")
    r = etree.SubElement(link, qn("w:r"))
    etree.SubElement(r, qn("w:t")).text = "jump link"
    pkg.mark_dirty()
    pkg.save(do_backup=False)

    pkg = DocxPackage(f)
    result = fields.delete_hyperlink(pkg, url="#sec_two")
    pkg.save(do_backup=False)
    assert result["deleted_hyperlink"] == "#sec_two"
    pkg2 = DocxPackage(f)
    assert not list(pkg2.root().iter(qn("w:hyperlink")))
    text = "\n".join(p.text for p in Document(str(f)).paragraphs)
    assert "jump link" in text


# -------------------------------------------------- delete_content_control


def cc_doc(tmp_path: Path, name="cc.docx") -> Path:
    """A document with one tagged plain-text content control."""
    from word_mcp.ops import forms as fm

    f = fresh_doc(tmp_path, name)
    pkg = DocxPackage(f)
    fm.insert_content_control(
        pkg, tag="reviewer", after_anchor="First paragraph", text="TBD"
    )
    pkg.save(do_backup=False)
    return f


def test_delete_content_control_by_tag(tmp_path):
    from word_mcp.ops import forms as fm

    f = cc_doc(tmp_path)
    pkg = DocxPackage(f)
    assert fm.list_content_controls(pkg)["count"] == 1
    result = fm.delete_content_control(pkg, tag="reviewer")
    pkg.save(do_backup=False)
    assert result["deleted_content_control"] == "reviewer"

    pkg2 = DocxPackage(f)
    assert fm.list_content_controls(pkg2)["count"] == 0
    assert not list(pkg2.root().iter(qn("w:sdt")))
    text = "\n".join(p.text for p in Document(str(f)).paragraphs)
    assert "First paragraph, always present." in text


def test_delete_content_control_by_index(tmp_path):
    from word_mcp.ops import forms as fm

    f = cc_doc(tmp_path)
    pkg = DocxPackage(f)
    idx = fm.list_content_controls(pkg)["controls"][0]["index"]
    result = fm.delete_content_control(pkg, index=idx)
    assert result["deleted_content_control"] == "reviewer"
    assert fm.list_content_controls(pkg)["count"] == 0


def test_delete_content_control_guards(tmp_path):
    from word_mcp.ops import forms as fm

    f = cc_doc(tmp_path)
    pkg = DocxPackage(f)
    with pytest.raises(WordMcpError, match="exactly one"):
        fm.delete_content_control(pkg)
    with pytest.raises(WordMcpError, match="exactly one"):
        fm.delete_content_control(pkg, tag="reviewer", index=0)
    with pytest.raises(TargetNotFound):
        fm.delete_content_control(pkg, tag="ghost")
    with pytest.raises(TargetNotFound):
        fm.delete_content_control(pkg, index=99)


def test_delete_content_control_locked_refused(tmp_path):
    from word_mcp.ops import forms as fm

    f = cc_doc(tmp_path)
    pkg = DocxPackage(f)
    sdt = next(iter(pkg.root().iter(qn("w:sdt"))))
    pr = sdt.find(qn("w:sdtPr"))
    lock = etree.SubElement(pr, qn("w:lock"))
    lock.set(qn("w:val"), "sdtLocked")
    with pytest.raises(WordMcpError, match="locked"):
        fm.delete_content_control(pkg, tag="reviewer")
    # nothing changed
    assert fm.list_content_controls(pkg)["count"] == 1
