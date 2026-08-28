"""DATA-PLUMBING bundle: table <-> CSV/JSON, extract_images, split_document.

Self-contained — every document is built on the fly with python-docx plus the
package's own ops; no external fixtures needed.
"""

import csv
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import (
    dataio as dio,
    media,
    notes,
    read as rd,
    tables as tb,
)


# ------------------------------------------------------------------- fixtures


def _fresh_doc(tmp_path: Path, name: str) -> Path:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Anchor paragraph so the body is never empty.")
    doc.save(str(f))
    return f


def _png(width: int, height: int, rgb: tuple[int, int, int]) -> bytes:
    """Minimal valid truecolor PNG, stdlib only."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + bytes(rgb) * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _merged_table_pkg(tmp_path: Path) -> DocxPackage:
    """3x3 table, col 0 rows 0-1 merged vertically, row 2 cols 1-2 merged
    horizontally."""
    f = _fresh_doc(tmp_path, "merged.docx")
    pkg = DocxPackage(f)
    tb.create_table(
        pkg,
        [["A", "B", "C"], ["D", "E", "F"], ["G", "H", "I"]],
        at_end=True,
        header_row=False,
    )
    tb.merge_cells(pkg, 0, start_row=0, end_row=1, start_col=0, end_col=0)
    tb.merge_cells(pkg, 0, start_row=2, end_row=2, start_col=1, end_col=2)
    return pkg


EXPECTED_MERGED_GRID = [
    ["A\nD", "B", "C"],
    ["", "E", "F"],
    ["G", "H\nI", ""],
]


def _body_text(path: Path) -> str:
    pkg = DocxPackage(path)
    return "".join(t.text or "" for t in pkg.body().iter(qn("w:t")))


# ------------------------------------------------------------------- exporting


def test_export_merged_table_inline_json(tmp_path):
    pkg = _merged_table_pkg(tmp_path)
    result = dio.export_table(pkg, 0, format="json")
    assert result["data"] == EXPECTED_MERGED_GRID
    assert result["rows"] == 3 and result["columns"] == 3
    assert {"row": 0, "col": 0, "rowspan": 2, "colspan": 1} in result["merges"]
    assert {"row": 2, "col": 1, "rowspan": 1, "colspan": 2} in result["merges"]
    assert len(result["merges"]) == 2


def test_export_merged_table_csv_file_and_refusal(tmp_path):
    pkg = _merged_table_pkg(tmp_path)
    out = tmp_path / "table.csv"
    result = dio.export_table(pkg, 0, format="csv", output_path=str(out))
    assert result["output_path"] == str(out)
    assert result["merges"]  # merge topology as a report field alongside CSV
    with open(out, encoding="utf-8", newline="") as fh:
        assert [list(r) for r in csv.reader(fh)] == EXPECTED_MERGED_GRID
    # An existing output file is refused, not overwritten.
    with pytest.raises(WordMcpError, match="already exists"):
        dio.export_table(pkg, 0, format="csv", output_path=str(out))


def test_export_bad_format_refused(tmp_path):
    pkg = _merged_table_pkg(tmp_path)
    with pytest.raises(WordMcpError, match="csv or json"):
        dio.export_table(pkg, 0, format="xlsx")


def test_export_nested_table_flattened(tmp_path):
    f = _fresh_doc(tmp_path, "nested.docx")
    pkg = DocxPackage(f)
    tb.create_table(pkg, [["host", "x"], ["y", "z"]], at_end=True, header_row=False)
    tb.create_table(pkg, [["n1", "n2"]], at_end=True, header_row=False)
    # Move the second table inside cell (0,0) of the first (nested table).
    tbls = [el for k, _, el in rd.body_items(pkg) if k == "table"]
    inner = tbls[1]
    inner.getparent().remove(inner)
    tc = tbls[0].findall(qn("w:tr"))[0].findall(qn("w:tc"))[0]
    tc.append(inner)
    tc.append(etree.Element(qn("w:p")))  # Word requires a trailing paragraph
    pkg.mark_dirty()

    result = dio.export_table(pkg, 0, format="json")
    assert "n1" in result["data"][0][0] and "n2" in result["data"][0][0]
    assert result["nested_table_cells"][0]["row"] == 0
    assert result["nested_table_cells"][0]["col"] == 0


# ------------------------------------------------------------------- importing


def test_roundtrip_export_import_new_table(tmp_path):
    src_pkg = _merged_table_pkg(tmp_path)
    exported = dio.export_table(src_pkg, 0, format="json")

    dest = _fresh_doc(tmp_path, "dest.docx")
    dest_pkg = DocxPackage(dest)
    result = dio.import_table(dest_pkg, exported["data"], has_header=False)
    assert result["mode"] == "created"
    got = rd.get_table(dest_pkg, 0)
    texts = [[c["text"] for c in row] for row in got["cells"]]
    assert texts == EXPECTED_MERGED_GRID


def test_import_from_csv_file(tmp_path):
    csv_path = tmp_path / "data.csv"
    grid = [["Name", "Score"], ["alpha", "1"], ["beta", "2"]]
    with open(csv_path, "w", encoding="utf-8", newline="") as fh:
        csv.writer(fh).writerows(grid)

    f = _fresh_doc(tmp_path, "csvdest.docx")
    pkg = DocxPackage(f)
    result = dio.import_table(pkg, str(csv_path), at_end=True)
    assert result["mode"] == "created"
    pkg.save(do_backup=False)
    got = rd.get_table(DocxPackage(f), 0)
    assert [[c["text"] for c in row] for row in got["cells"]] == grid


def test_import_inline_overwrite_existing(tmp_path):
    f = _fresh_doc(tmp_path, "overwrite.docx")
    pkg = DocxPackage(f)
    tb.create_table(pkg, [["a", "b"], ["c", "d"]], at_end=True, header_row=False)
    result = dio.import_table(
        pkg, [["w", "x"], ["y", "z"]], table_index=0
    )
    assert result["mode"] == "overwritten"
    got = rd.get_table(pkg, 0)
    assert [[c["text"] for c in row] for row in got["cells"]] == [
        ["w", "x"],
        ["y", "z"],
    ]


def test_import_dimension_mismatch_refused(tmp_path):
    f = _fresh_doc(tmp_path, "mismatch.docx")
    pkg = DocxPackage(f)
    tb.create_table(pkg, [["a", "b"], ["c", "d"]], at_end=True, header_row=False)
    data = [["1", "2", "3"], ["4", "5", "6"], ["7", "8", "9"]]
    with pytest.raises(WordMcpError) as exc:
        dio.import_table(pkg, data, table_index=0)
    # Both shapes are listed in the refusal.
    assert "2x2" in str(exc.value) and "3x3" in str(exc.value)


def test_import_covered_position_value_refused(tmp_path):
    pkg = _merged_table_pkg(tmp_path)
    bad = [
        ["A2\nD2", "B2", "C2"],
        ["LOST", "E2", "F2"],  # (1,0) is covered by the vertical merge
        ["G2", "H2", ""],
    ]
    with pytest.raises(WordMcpError, match="merge-covered"):
        dio.import_table(pkg, bad, table_index=0)
    # Valid version (covered positions empty) succeeds and hits the anchors.
    ok = [
        ["A2\nD2", "B2", "C2"],
        ["", "E2", "F2"],
        ["G2", "H2", ""],
    ]
    dio.import_table(pkg, ok, table_index=0)
    assert dio.export_table(pkg, 0, format="json")["data"] == ok


def test_import_bad_payloads_refused(tmp_path):
    f = _fresh_doc(tmp_path, "badpayload.docx")
    pkg = DocxPackage(f)
    with pytest.raises(WordMcpError):
        dio.import_table(pkg, ["not", "2d"])
    txt = tmp_path / "data.txt"
    txt.write_text("a,b\n", encoding="utf-8")
    with pytest.raises(WordMcpError, match=r"\.csv or \.json"):
        dio.import_table(pkg, str(txt))


# -------------------------------------------------------------- extract_images


def _doc_with_two_images(tmp_path: Path) -> tuple[Path, bytes, bytes]:
    png_a = _png(4, 2, (255, 0, 0))
    png_b = _png(3, 5, (0, 0, 255))
    (tmp_path / "a.png").write_bytes(png_a)
    (tmp_path / "b.png").write_bytes(png_b)
    f = _fresh_doc(tmp_path, "imgs.docx")
    pkg = DocxPackage(f)
    media.add_image(pkg, str(tmp_path / "a.png"), at_end=True)
    media.add_image(pkg, str(tmp_path / "b.png"), at_end=True)
    pkg.save(do_backup=False)
    return f, png_a, png_b


def test_extract_images_two_pngs(tmp_path):
    f, png_a, png_b = _doc_with_two_images(tmp_path)
    out_dir = tmp_path / "extracted"
    result = dio.extract_images(DocxPackage(f), str(out_dir))
    assert result["extracted"] == 2
    files = sorted(p.name for p in out_dir.iterdir())
    assert files == ["image0.png", "image1.png"]
    assert (out_dir / "image0.png").read_bytes() == png_a
    assert (out_dir / "image1.png").read_bytes() == png_b
    by_index = {e["index"]: e for e in result["images"]}
    assert (by_index[0]["width_px"], by_index[0]["height_px"]) == (4, 2)
    assert (by_index[1]["width_px"], by_index[1]["height_px"]) == (3, 5)
    # Each image reports where it appears in the body.
    assert "body_paragraph" in by_index[0]["appears"]


def test_extract_images_collision_refused_no_partial_write(tmp_path):
    f, _a, _b = _doc_with_two_images(tmp_path)
    out_dir = tmp_path / "collide"
    out_dir.mkdir()
    # Only the SECOND name collides; the first must still not be written.
    (out_dir / "image1.png").write_bytes(b"pre-existing")
    with pytest.raises(WordMcpError, match="already exist"):
        dio.extract_images(DocxPackage(f), str(out_dir))
    assert not (out_dir / "image0.png").exists()
    assert (out_dir / "image1.png").read_bytes() == b"pre-existing"


def test_extract_images_prefix(tmp_path):
    f, _a, _b = _doc_with_two_images(tmp_path)
    out_dir = tmp_path / "prefixed"
    dio.extract_images(DocxPackage(f), str(out_dir), prefix="fig")
    assert sorted(p.name for p in out_dir.iterdir()) == [
        "fig0.png",
        "fig1.png",
    ]


# -------------------------------------------------------------- split_document


def _three_chapter_doc(tmp_path: Path) -> Path:
    """Front matter, then three chapters: image in ch1, footnote in ch2,
    table in ch3."""
    f = tmp_path / "book.docx"
    doc = Document()
    doc.add_paragraph("This is the front matter paragraph.")
    doc.add_heading("Chapter One", 1)
    doc.add_paragraph("Alpha body text lives here.")
    doc.add_heading("Chapter Two", 1)
    doc.add_paragraph("Beta body text lives here.")
    doc.add_heading("Chapter Three", 1)
    doc.add_paragraph("Gamma body text lives here.")
    doc.save(str(f))

    png = tmp_path / "split_img.png"
    png.write_bytes(_png(4, 4, (0, 255, 0)))
    pkg = DocxPackage(f)
    media.add_image(pkg, str(png), after_anchor="Alpha body text lives here.")
    notes.add_note(
        pkg,
        "footnote",
        anchor_text="Beta body text",
        note_text="the beta footnote",
    )
    tb.create_table(pkg, [["H1", "H2"], ["a", "b"]], at_end=True)
    pkg.save(do_backup=False)
    return f


def test_split_three_chapters(tmp_path):
    src = _three_chapter_doc(tmp_path)
    out_dir = tmp_path / "split"
    result = dio.split_document(str(src), str(out_dir), level=1)

    expected = [
        "00_front_matter.docx",
        "01_Chapter One.docx",
        "02_Chapter Two.docx",
        "03_Chapter Three.docx",
    ]
    assert sorted(p.name for p in out_dir.iterdir()) == expected
    assert result["sections"] == 4
    outputs = [out_dir / name for name in expected]

    # Every output is a valid standalone document.
    for path in outputs:
        Document(str(path))  # python-docx opens it
        vpkg = DocxPackage(path)  # round-trips through the package layer
        vn = notes.validate_notes(vpkg)
        for kind in ("footnotes", "endnotes"):
            assert vn[kind]["ok"], (path.name, vn)
            assert not vn[kind]["needs_cleanup"], (path.name, vn)

    # The ch2 footnote lives only in ch2's file.
    counts = [len(rd.list_footnotes(DocxPackage(p))) for p in outputs]
    assert counts == [0, 0, 1, 0]
    assert "the beta footnote" in [
        n["text"] for n in rd.list_footnotes(DocxPackage(outputs[2]))
    ][0]

    # The ch1 image part travels only with ch1.
    has_media = [
        any(name.startswith("word/media/") for name in DocxPackage(p).part_names())
        for p in outputs
    ]
    assert has_media == [False, True, False, False]

    # The ch3 table travels only with ch3.
    tables_per = [
        rd.get_document_info(DocxPackage(p))["tables"] for p in outputs
    ]
    assert tables_per == [0, 0, 0, 1]

    # Concatenated body text equals the original: nothing lost, nothing added.
    assert "".join(_body_text(p) for p in outputs) == _body_text(src)

    # The source file itself was never touched.
    assert _body_text(src).startswith("This is the front matter paragraph.")


def test_split_output_collision_refused(tmp_path):
    src = _three_chapter_doc(tmp_path)
    out_dir = tmp_path / "collide_split"
    out_dir.mkdir()
    (out_dir / "01_Chapter One.docx").write_bytes(b"pre-existing")
    with pytest.raises(WordMcpError, match="already exist"):
        dio.split_document(str(src), str(out_dir), level=1)
    # Refusal happened before any write: no other output appeared.
    assert sorted(p.name for p in out_dir.iterdir()) == ["01_Chapter One.docx"]
    assert (out_dir / "01_Chapter One.docx").read_bytes() == b"pre-existing"


def test_split_no_headings_refused(tmp_path):
    f = _fresh_doc(tmp_path, "flat.docx")
    with pytest.raises(WordMcpError, match="no headings"):
        dio.split_document(str(f), str(tmp_path / "nowhere"), level=1)
    assert not (tmp_path / "nowhere").exists()


def test_split_filename_from_index(tmp_path):
    src = _three_chapter_doc(tmp_path)
    out_dir = tmp_path / "by_index"
    dio.split_document(str(src), str(out_dir), level=1, filename_from="index")
    assert sorted(p.name for p in out_dir.iterdir()) == [
        "00_front_matter.docx",
        "01.docx",
        "02.docx",
        "03.docx",
    ]


def test_split_bad_args_refused(tmp_path):
    src = _three_chapter_doc(tmp_path)
    with pytest.raises(WordMcpError, match="heading or index"):
        dio.split_document(str(src), str(tmp_path / "x"), filename_from="title")
    with pytest.raises(WordMcpError, match="1..9"):
        dio.split_document(str(src), str(tmp_path / "x"), level=0)
