"""Wave C staging: the v2 table surface (V2_DESIGN Section 2.8) exercised
through the registered server tools.

Self-skipping until the integrator lands the consolidated registrations
(keyed on modify_table_structure existing on the server module), so the file
is safe to collect on the v1 surface today and becomes the Wave C gate the
moment the v2 registrations land. All fixtures are built inline with
python-docx; nothing here depends on shared test documents. The ops-layer
suite in tests/unit/test_tables.py is untouched and keeps covering the
internals (the ops functions keep their v1 signatures).
"""

from __future__ import annotations

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import TargetNotFound, WordMcpError

pytestmark = pytest.mark.skipif(
    not hasattr(srv, "modify_table_structure"),
    reason="v2 table surface not yet integrated into server.py",
)

DATA = [
    ["H1", "H2", "H3", "H4"],
    ["a1", "b1", "c1", "d1"],
    ["a2", "b2", "c2", "d2"],
    ["a3", "b3", "c3", "d3"],
]


def _base_doc(tmp_path, name="doc.docx"):
    dst = tmp_path / name
    d = Document()
    d.add_paragraph("Intro paragraph.")
    d.add_heading("Data", level=1)
    d.add_paragraph("Closing paragraph.")
    d.save(str(dst))
    return dst


@pytest.fixture
def doc(tmp_path):
    """Base doc plus a known 4x4 table appended by the v2 create_table."""
    dst = _base_doc(tmp_path)
    srv.create_table(str(dst), DATA, backup=False)
    return dst


@pytest.fixture
def nested_doc(tmp_path):
    """A body table whose (1, 1) cell hosts a 2x2 nested table."""
    dst = tmp_path / "nested.docx"
    d = Document()
    d.add_paragraph("Host document.")
    host = d.add_table(rows=2, cols=2)
    host.style = "Table Grid"
    for r in range(2):
        for c in range(2):
            host.cell(r, c).text = f"host{r}{c}"
    inner = host.cell(1, 1).add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            inner.cell(r, c).text = f"in{r}{c}"
    d.save(str(dst))
    return dst


def _grid(path, index=0, **kw):
    return srv.get_table(str(path), index, **kw)


def _texts(g, row):
    return [c["text"] for c in g["cells"][row]]


# ------------------------------------------------------------- create_table


def test_create_table_default_end(doc):
    g = _grid(doc)
    assert g["rows"] == 4
    assert g["grid_columns"] == 4
    assert _texts(g, 0) == ["H1", "H2", "H3", "H4"]


def test_create_table_reports_shape(tmp_path):
    dst = _base_doc(tmp_path)
    result = srv.create_table(str(dst), DATA, backup=False)
    assert result["created"] == {"rows": 4, "columns": 4}
    assert "saved" in result


def test_create_table_at_paragraph_location(tmp_path):
    dst = _base_doc(tmp_path)
    result = srv.create_table(
        str(dst), [["x", "y"]], location={"paragraph": 0}, backup=False
    )
    assert result["created"]["columns"] == 2
    assert len(Document(str(dst)).tables) == 1


def test_create_table_after_heading_location(tmp_path):
    dst = _base_doc(tmp_path)
    srv.create_table(
        str(dst), [["x", "y"]],
        location={"after_heading": {"text": "Data"}}, backup=False,
    )
    assert len(Document(str(dst)).tables) == 1


def test_create_table_unsupported_position_refused(tmp_path):
    dst = _base_doc(tmp_path)
    with pytest.raises(WordMcpError):
        srv.create_table(
            str(dst), [["x"]],
            location={"paragraph": 0, "position": "before"}, backup=False,
        )


def test_create_table_location_not_found_refused(tmp_path):
    dst = _base_doc(tmp_path)
    with pytest.raises((TargetNotFound, WordMcpError)):
        srv.create_table(
            str(dst), [["x"]],
            location={"after_heading": {"text": "No Such Heading"}},
            backup=False,
        )


# --------------------------------------------------- modify_table_structure


def test_insert_rows_branch(doc):
    srv.modify_table_structure(
        str(doc), 0, action="insert", target="rows", at=4, count=2,
        backup=False,
    )
    assert _grid(doc)["rows"] == 6


def test_insert_rows_copy_format(doc):
    srv.modify_table_structure(
        str(doc), 0, action="insert", target="rows", at=1,
        copy_format_from=1, backup=False,
    )
    g = _grid(doc)
    assert g["rows"] == 5
    assert _texts(g, 2) == ["a1", "b1", "c1", "d1"]


def test_insert_columns_branch(doc):
    srv.modify_table_structure(
        str(doc), 0, action="insert", target="columns", at=2, backup=False,
    )
    g = _grid(doc)
    assert g["grid_columns"] == 5
    assert _texts(g, 0) == ["H1", "H2", "", "H3", "H4"]


def test_delete_rows_branch(doc):
    srv.modify_table_structure(
        str(doc), 0, action="delete", target="rows", start=1, end=2,
        backup=False,
    )
    g = _grid(doc)
    assert g["rows"] == 2
    assert _texts(g, 1) == ["a3", "b3", "c3", "d3"]


def test_delete_columns_branch(doc):
    srv.modify_table_structure(
        str(doc), 0, action="delete", target="columns", columns=[1],
        backup=False,
    )
    g = _grid(doc)
    assert g["grid_columns"] == 3
    assert _texts(g, 0) == ["H1", "H3", "H4"]


def test_merge_then_unmerge_roundtrip(doc):
    srv.modify_table_structure(
        str(doc), 0, action="merge",
        range={"start_row": 1, "end_row": 2, "start_col": 0, "end_col": 1},
        backup=False,
    )
    assert _grid(doc)["has_merges"] is True
    srv.modify_table_structure(
        str(doc), 0, action="unmerge", row=1, cell=0, backup=False,
    )
    assert _grid(doc)["has_merges"] is False


def test_split_branch(doc):
    srv.modify_table_structure(
        str(doc), 0, action="split", at_row=2, backup=False,
    )
    assert len(Document(str(doc)).tables) == 2
    assert _grid(doc, 0)["rows"] == 2
    assert _grid(doc, 1)["rows"] == 2


def test_unknown_action_refused(doc):
    with pytest.raises(WordMcpError):
        srv.modify_table_structure(
            str(doc), 0, action="resize", backup=False,
        )


def test_insert_without_target_refused(doc):
    with pytest.raises(WordMcpError):
        srv.modify_table_structure(
            str(doc), 0, action="insert", at=1, backup=False,
        )


def test_merge_without_range_refused(doc):
    with pytest.raises(WordMcpError):
        srv.modify_table_structure(str(doc), 0, action="merge", backup=False)


def test_cross_action_param_bleed_refused(doc):
    # width_pt belongs to insert+columns, not insert+rows.
    with pytest.raises(WordMcpError):
        srv.modify_table_structure(
            str(doc), 0, action="insert", target="rows", at=1,
            width_pt=100.0, backup=False,
        )
    # at_row belongs to split, not delete.
    with pytest.raises(WordMcpError):
        srv.modify_table_structure(
            str(doc), 0, action="delete", target="rows", start=1,
            at_row=2, backup=False,
        )


# ---------------------------------------------------- set_table_properties


def test_style_property(doc):
    result = srv.set_table_properties(str(doc), 0, style="BandedTable",
                                      backup=False)
    assert "saved" in result


def test_column_widths_property(doc):
    srv.set_table_properties(
        str(doc), 0, column_widths=[100.0, 100.0, 100.0, 100.0],
        backup=False,
    )
    assert _grid(doc)["grid_widths_twips"] == [2000, 2000, 2000, 2000]


def test_header_row_repeat_int_and_off(doc):
    result = srv.set_table_properties(str(doc), 0, header_row_repeat=2,
                                      backup=False)
    assert result["header_rows"] == 2
    result = srv.set_table_properties(str(doc), 0, header_row_repeat=False,
                                      backup=False)
    assert result["header_rows"] == 0


def test_combined_properties_single_call(doc):
    result = srv.set_table_properties(
        str(doc), 0, style="TableGrid",
        column_widths=[120.0, 120.0, 120.0, 120.0],
        header_row_repeat=True, backup=False,
    )
    assert result["header_rows"] == 1
    assert _grid(doc)["grid_widths_twips"] == [2400, 2400, 2400, 2400]


def test_no_properties_refused(doc):
    with pytest.raises(WordMcpError):
        srv.set_table_properties(str(doc), 0, backup=False)


# ---------------------------------------------------------------- set_cells


def test_scatter_edits(doc):
    srv.set_cells(
        str(doc), 0,
        edits=[
            {"row": 1, "cell": 0, "text": "X"},
            {"row": 3, "cell": 3, "text": "Y"},
        ],
        backup=False, live="off",
    )
    g = _grid(doc)
    assert g["cells"][1][0]["text"] == "X"
    assert g["cells"][3][3]["text"] == "Y"


def test_block_mode(doc):
    srv.set_cells(
        str(doc), 0,
        block={"origin": {"row": 1, "cell": 1},
               "values": [["P", "Q"], ["R", "S"]]},
        backup=False,
    )
    g = _grid(doc)
    assert _texts(g, 1) == ["a1", "P", "Q", "d1"]
    assert _texts(g, 2) == ["a2", "R", "S", "d2"]


def test_nested_mode(nested_doc):
    srv.set_cells(
        str(nested_doc), 0,
        edits=[{"row": 0, "cell": 1, "text": "edited"}],
        nested={"row": 1, "cell": 1}, backup=False,
    )
    g = _grid(nested_doc, 0, nested={"row": 1, "cell": 1})
    assert g["cells"][0][1]["text"] == "edited"


def test_both_modes_refused(doc):
    with pytest.raises(WordMcpError):
        srv.set_cells(
            str(doc), 0,
            edits=[{"row": 0, "cell": 0, "text": "x"}],
            block={"origin": {"row": 0, "cell": 0}, "values": [["x"]]},
            backup=False,
        )


def test_neither_mode_refused(doc):
    with pytest.raises(WordMcpError):
        srv.set_cells(str(doc), 0, backup=False)


def test_nested_plus_block_refused(nested_doc):
    with pytest.raises(WordMcpError):
        srv.set_cells(
            str(nested_doc), 0,
            block={"origin": {"row": 0, "cell": 0}, "values": [["x"]]},
            nested={"row": 1, "cell": 1}, backup=False,
        )


def test_nested_plus_track_refused(nested_doc):
    with pytest.raises(WordMcpError):
        srv.set_cells(
            str(nested_doc), 0,
            edits=[{"row": 0, "cell": 0, "text": "x"}],
            nested={"row": 1, "cell": 1}, track=True, backup=False,
        )


# ---------------------------------------------------------------- get_table


def test_get_table_nested_read(nested_doc):
    g = _grid(nested_doc, 0, nested={"row": 1, "cell": 1})
    assert g["rows"] == 2
    assert g["cells"][0][0]["text"] == "in00"
    assert g["nested_in"] == {"table": 0, "row": 1, "cell": 1}


def test_get_table_nested_index_out_of_range(nested_doc):
    with pytest.raises((TargetNotFound, WordMcpError)):
        _grid(nested_doc, 0, nested={"row": 1, "cell": 1, "index": 5})


def test_get_table_nested_unknown_key_refused(nested_doc):
    with pytest.raises(WordMcpError):
        _grid(nested_doc, 0, nested={"row": 1, "col": 1})


def test_get_table_no_nested_table_in_cell(nested_doc):
    with pytest.raises((TargetNotFound, WordMcpError)):
        _grid(nested_doc, 0, nested={"row": 0, "cell": 0})


# ------------------------------------------- identity tools, v2 smoke pass


def test_delete_table_identity(doc):
    srv.delete_table(str(doc), 0, backup=False)
    with pytest.raises((TargetNotFound, WordMcpError)):
        _grid(doc)


def test_format_cells_identity(doc):
    result = srv.format_cells(
        str(doc), 0, targets=[{"row": 0}], formatting={"bold": True},
        backup=False,
    )
    assert "saved" in result


def test_sort_table_identity(doc):
    srv.set_cells(
        str(doc), 0,
        edits=[
            {"row": 1, "cell": 0, "text": "zz"},
            {"row": 3, "cell": 0, "text": "aa"},
        ],
        backup=False, live="off",
    )
    srv.sort_table(str(doc), 0, column=0, backup=False)
    g = _grid(doc)
    # Lexicographic ascending: "a2" < "aa" < "zz"; header row stays put.
    assert [g["cells"][r][0]["text"] for r in (1, 2, 3)] == ["a2", "aa", "zz"]


def test_sort_refused_on_vertical_merge(doc):
    srv.modify_table_structure(
        str(doc), 0, action="merge",
        range={"start_row": 1, "end_row": 2, "start_col": 0, "end_col": 0},
        backup=False,
    )
    with pytest.raises(WordMcpError):
        srv.sort_table(str(doc), 0, column=1, backup=False)


def test_export_import_roundtrip(doc, tmp_path):
    exported = srv.export_table(str(doc), 0, format="json")
    assert exported["data"][0] == ["H1", "H2", "H3", "H4"]
    dst = _base_doc(tmp_path, "target.docx")
    srv.import_table(str(dst), exported["data"], at_end=True, backup=False)
    g = _grid(dst)
    assert g["rows"] == 4
    assert _texts(g, 0) == ["H1", "H2", "H3", "H4"]


def test_import_shape_mismatch_refused(doc):
    with pytest.raises(WordMcpError):
        srv.import_table(str(doc), [["only", "two"]], table_index=0,
                         backup=False)


def test_copy_table_identity(doc, tmp_path):
    dst = _base_doc(tmp_path, "target.docx")
    result = srv.copy_table(str(dst), str(doc), 0, at_end=True, backup=False)
    assert "saved" in result
    g = _grid(dst)
    assert g["rows"] == 4
    assert _texts(g, 0) == ["H1", "H2", "H3", "H4"]
