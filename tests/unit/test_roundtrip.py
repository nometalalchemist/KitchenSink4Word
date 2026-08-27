"""Phase 0 gate: every corpus document loads, saves, and round-trips intact.

'Intact' means: every part the code did not touch is byte-identical, and the
output still opens in python-docx. The Word repair-prompt check runs separately
in tests/word_validator.py (COM, slow, phase-gate only).
"""

import zipfile
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.package import DocxPackage

CORPUS = Path(__file__).resolve().parents[1] / "corpus"
CORPUS_DOCS = sorted(CORPUS.glob("*.docx"))


def corpus_ids():
    return [p.name for p in CORPUS_DOCS]


@pytest.mark.parametrize("doc_path", CORPUS_DOCS, ids=corpus_ids())
def test_load(doc_path):
    pkg = DocxPackage(doc_path)
    assert pkg.has_part("word/document.xml")
    assert pkg.body() is not None


@pytest.mark.parametrize("doc_path", CORPUS_DOCS, ids=corpus_ids())
def test_untouched_roundtrip_is_byte_identical_per_part(doc_path, tmp_path):
    pkg = DocxPackage(doc_path)
    out = tmp_path / doc_path.name
    pkg.save(out, do_backup=False)

    with zipfile.ZipFile(doc_path) as zin, zipfile.ZipFile(out) as zout:
        assert zin.namelist() == zout.namelist(), "part order or set changed"
        for name in zin.namelist():
            assert zin.read(name) == zout.read(name), f"part changed: {name}"


@pytest.mark.parametrize("doc_path", CORPUS_DOCS, ids=corpus_ids())
def test_roundtrip_output_opens_in_python_docx(doc_path, tmp_path):
    pkg = DocxPackage(doc_path)
    out = tmp_path / doc_path.name
    pkg.save(out, do_backup=False)
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


@pytest.mark.parametrize("doc_path", CORPUS_DOCS, ids=corpus_ids())
def test_touched_but_unchanged_part_stays_semantically_stable(doc_path, tmp_path):
    """Parse document.xml, mark it dirty WITHOUT editing, save: output must
    still be well-formed, open in python-docx, and contain identical text."""
    pkg = DocxPackage(doc_path)
    pkg.tree("word/document.xml")
    pkg.mark_dirty("word/document.xml")
    out = tmp_path / doc_path.name
    pkg.save(out, do_backup=False)

    orig_text = "\n".join(p.text for p in Document(str(doc_path)).paragraphs)
    new_text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert orig_text == new_text
