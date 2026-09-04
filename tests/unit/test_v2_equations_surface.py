"""STAGED v2-surface rewrite (Wave D): equations (2.7 slice).

insert_equation rename (location object for display placement, anchor_text
kept for inline) and delete_element(type:"equation") per
integration/v2_briefs/wave_D.md. Skips until the Phase 2 integrator
registers the v2 tools. The LaTeX fidelity gates stay in
tests/unit/test_equations.py (ops layer, unchanged). Original v1 test file
untouched.
"""

from pathlib import Path

import pytest
from docx import Document
from lxml import etree

import word_mcp.server as srv
from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import equations as eq

pytestmark = pytest.mark.skipif(
    not hasattr(srv, "insert_equation"),
    reason="v2 surface not yet registered (staged for the Phase 2 "
    "integrator)",
)

M = eq.M_NS


def fresh(tmp_path: Path, name="doc.docx") -> str:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("The model is defined as follows.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(f))
    return str(f)


def ok(result) -> bool:
    return not (isinstance(result, dict) and result.get("ok") is False)


def m_tags(path):
    root = etree.fromstring(
        DocxPackage(path).raw_part("word/document.xml")
    )
    return {
        etree.QName(el).localname
        for el in root.iter()
        if etree.QName(el).namespace == M
    }


def test_insert_equation_display_at_location(tmp_path):
    f = fresh(tmp_path)
    r = srv.insert_equation(
        f, r"\frac{a+b}{c}", display=True, location={"paragraph": 0}
    )
    assert ok(r)
    assert "oMathPara" in m_tags(f)
    listed = eq.list_equations(DocxPackage(f))
    assert listed["equation_count"] == 1


def test_insert_equation_inline_after_anchor(tmp_path):
    f = fresh(tmp_path)
    r = srv.insert_equation(
        f, r"x^2", display=False, anchor_text="defined as"
    )
    assert ok(r)
    tags = m_tags(f)
    assert "oMath" in tags and "oMathPara" not in tags


def test_insert_equation_bad_latex_refuses_unmodified(tmp_path):
    f = fresh(tmp_path)
    before = DocxPackage(f).raw_part("word/document.xml")
    with pytest.raises(WordMcpError):
        srv.insert_equation(f, r"\frac{", display=True)
    assert DocxPackage(f).raw_part("word/document.xml") == before


def test_delete_element_equation(tmp_path):
    f = fresh(tmp_path)
    srv.insert_equation(f, r"\sqrt{x}", display=True)
    r = srv.delete_element(f, type="equation", id=0)
    assert ok(r)
    assert eq.list_equations(DocxPackage(f))["equation_count"] == 0


def test_delete_element_equation_out_of_range(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises((TargetNotFound, WordMcpError)):
        srv.delete_element(f, type="equation", id=0)
