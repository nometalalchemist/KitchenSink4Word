"""STAGED v2-surface rewrite (Wave D): charts (2.7 slice).

insert_chart / set_chart_data renames and delete_element(type:"chart") per
integration/v2_briefs/wave_D.md. Skips until the Phase 2 integrator
registers the v2 tools. The deep structural chart gates stay in
tests/unit/test_charts.py (ops layer, unchanged); this file checks the
renamed surface routes to them. Original v1 test file untouched.
"""

from pathlib import Path

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage

pytestmark = pytest.mark.skipif(
    not hasattr(srv, "insert_chart"),
    reason="v2 surface not yet registered (staged for the Phase 2 "
    "integrator)",
)

DATA = {
    "categories": ["Alpha", "Beta", "Gamma"],
    "series": [
        {"name": "S1", "values": [4.3, 2.5, 3.5]},
        {"name": "S2", "values": [1, 2, 3]},
    ],
}


def fresh(tmp_path: Path, name="doc.docx") -> str:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Anchor paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(f))
    return str(f)


def ok(result) -> bool:
    return not (isinstance(result, dict) and result.get("ok") is False)


def test_insert_chart_with_location(tmp_path):
    f = fresh(tmp_path)
    r = srv.insert_chart(
        f, "column", DATA, title="Staged", location={"paragraph": 0}
    )
    assert ok(r)
    pkg = DocxPackage(f)
    assert pkg.has_part("word/charts/chart1.xml")
    assert any(n.startswith("word/embeddings/") for n in pkg.part_names())


def test_insert_chart_default_location_is_end(tmp_path):
    f = fresh(tmp_path)
    assert ok(srv.insert_chart(f, "line", DATA))
    assert DocxPackage(f).has_part("word/charts/chart1.xml")


def test_set_chart_data_roundtrip(tmp_path):
    f = fresh(tmp_path)
    srv.insert_chart(f, "column", DATA)
    new_data = {
        "categories": ["One", "Two", "Three", "Four"],
        "series": [
            {"name": "S1", "values": [1, 2, 3, 4]},
            {"name": "S2", "values": [4, 3, 2, 1]},
        ],
    }
    r = srv.set_chart_data(f, chart_id=0, data=new_data)
    assert ok(r)
    xml = DocxPackage(f).raw_part("word/charts/chart1.xml").decode("utf-8")
    assert "Four" in xml


def test_set_chart_data_series_count_change_refuses(tmp_path):
    f = fresh(tmp_path)
    srv.insert_chart(f, "column", DATA)
    bad = {"categories": ["A"], "series": [{"name": "Only", "values": [1]}]}
    with pytest.raises(WordMcpError):
        srv.set_chart_data(f, chart_id=0, data=bad)


def test_delete_element_chart_removes_parts(tmp_path):
    f = fresh(tmp_path)
    srv.insert_chart(f, "pie", {
        "categories": ["A", "B"],
        "series": [{"name": "S", "values": [3, 7]}],
    })
    r = srv.delete_element(f, type="chart", id=0)
    assert ok(r)
    pkg = DocxPackage(f)
    assert not any(n.startswith("word/charts/") for n in pkg.part_names())
    assert not any(
        n.startswith("word/embeddings/") for n in pkg.part_names()
    )


def test_delete_element_chart_out_of_range(tmp_path):
    f = fresh(tmp_path)
    with pytest.raises(TargetNotFound):
        srv.delete_element(f, type="chart", id=3)
