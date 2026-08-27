"""v1.2 Phase C: move_section, template transfer, properties, styles, alt text."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import AmbiguousTarget, TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.ops import media, read, structure as sx, tables as tb, template as tp, text as tx, toc

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def doc(tmp_path):
    """Structured doc: three H1 sections with content, one with a table."""
    dst = tmp_path / "d.docx"
    d = Document()
    d.add_heading("Alpha Section", 1)
    d.add_paragraph("Alpha body one.")
    d.add_paragraph("Alpha body two.")
    d.add_heading("Beta Section", 1)
    d.add_paragraph("Beta body.")
    d.add_heading("Beta Child", 2)
    d.add_paragraph("Beta child body.")
    d.add_heading("Gamma Section", 1)
    d.add_paragraph("Gamma body.")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tb.create_table(pkg, [["t1", "t2"]], after_anchor="Beta body.")
    pkg.save(do_backup=False)
    return dst


@pytest.fixture
def ch13(tmp_path):
    dst = tmp_path / "ch1-3.docx"
    shutil.copy(CORPUS / "ch1-3.docx", dst)
    return dst


def outline_texts(path):
    return [h["text"] for h in read.get_outline(DocxPackage(path))]


# ------------------------------------------------------------- move_section


def test_move_section_to_end(doc):
    pkg = DocxPackage(doc)
    r = sx.move_section(pkg, "Alpha Section", at_end=True)
    pkg.save(do_backup=False)
    assert r["elements"] == 3
    assert outline_texts(doc) == [
        "Beta Section", "Beta Child", "Gamma Section", "Alpha Section",
    ]
    text = "\n".join(p["text"] for p in read.get_paragraphs(DocxPackage(doc)))
    assert text.index("Gamma body") < text.index("Alpha body one")
    Document(str(doc))


def test_move_section_carries_subsections_and_tables(doc):
    """Beta's block includes its H2 child AND the table."""
    pkg = DocxPackage(doc)
    r = sx.move_section(pkg, "Beta Section", at_end=True)
    pkg.save(do_backup=False)
    assert r["elements"] == 6  # H1, body, table, spacer para, H2, child body
    assert outline_texts(doc) == [
        "Alpha Section", "Gamma Section", "Beta Section", "Beta Child",
    ]
    # Table travelled with the section.
    pkg2 = DocxPackage(doc)
    items = read.body_items(pkg2)
    kinds = [k for k, _, _ in items]
    table_pos = kinds.index("table")
    gamma_pos = next(
        i for i, (k, _, el) in enumerate(items)
        if k == "paragraph" and "Gamma Section" in read.paragraph_text(el)
    )
    assert table_pos > gamma_pos
    Document(str(doc))


def test_move_section_before_heading(doc):
    pkg = DocxPackage(doc)
    sx.move_section(pkg, "Gamma Section", before_heading="Alpha Section")
    pkg.save(do_backup=False)
    assert outline_texts(doc)[0] == "Gamma Section"


def test_move_into_self_refused(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="inside the moved section"):
        sx.move_section(pkg, "Beta Section", before_heading="Beta Child")


def test_move_refuses_field_cut(doc):
    pkg = DocxPackage(doc)
    toc.insert_toc(pkg, after_index=1, update_on_open=False)
    pkg.save(do_backup=False)
    # TOC sdt sits inside Alpha's block region; moving Alpha must NOT corrupt
    # — sdt travels whole, so this should actually succeed cleanly.
    pkg2 = DocxPackage(doc)
    sx.move_section(pkg2, "Alpha Section", at_end=True)
    pkg2.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert xml.count('w:fldCharType="begin"') == xml.count('w:fldCharType="end"')
    Document(str(doc))


def test_list_section_blocks(doc):
    blocks = sx.list_section_blocks(DocxPackage(doc))
    by_heading = {b["heading"]: b["elements"] for b in blocks}
    assert by_heading["Alpha Section"] == 3
    assert by_heading["Beta Section"] == 6
    assert by_heading["Beta Child"] == 2


def test_move_real_chapter_section(ch13):
    """Move a real dissertation section and keep the document healthy."""
    pkg = DocxPackage(ch13)
    blocks = [
        b for b in sx.list_section_blocks(pkg)
        if b["elements"] and 3 < b["elements"] < 40 and b["level"] == 2
    ]
    if len(blocks) < 2:
        pytest.skip("not enough movable level-2 sections")
    victim = blocks[0]["heading"]
    n_before = read.get_document_info(pkg)["paragraphs"]
    sx.move_section(pkg, victim, at_end=True)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch13)
    assert read.get_document_info(pkg2)["paragraphs"] == n_before
    assert outline_texts(ch13)[-1] == victim
    Document(str(ch13))


# ---------------------------------------------------------- template transfer


def test_apply_template_restyles_by_name(tmp_path, doc):
    """Reference doc's Heading1 has a distinct color; after transfer the
    target's headings must reference a style that exists (no dangling)."""
    ref = tmp_path / "ref.docx"
    d = Document()
    d.add_heading("Ref heading", 1)
    d.save(str(ref))
    rp = DocxPackage(ref)
    sx.define_style(
        rp, style_id="Heading1", name="heading 1", style_type="paragraph",
        character_formatting={"color": "8B0000", "size_pt": 18, "bold": True},
    )
    rp.save(do_backup=False)

    pkg = DocxPackage(doc)
    result = tp.apply_template(pkg, str(ref))
    pkg.save(do_backup=False)
    assert result["dangling_style_refs"] == []
    assert "word/styles.xml" in result["parts_copied"]
    styles = DocxPackage(doc).raw_part("word/styles.xml").decode()
    assert "8B0000" in styles
    # Outline preserved (headings still resolve).
    assert len(outline_texts(doc)) == 4
    Document(str(doc))


def test_apply_template_clones_custom_styles(tmp_path, doc):
    """A custom style used in the target but absent from the reference gets
    cloned, not dropped."""
    pkg = DocxPackage(doc)
    sx.define_style(
        pkg, style_id="MySpecial", name="My Special",
        character_formatting={"italic": True},
    )
    tx.apply_style(pkg, [1], "MySpecial")
    pkg.save(do_backup=False)

    ref = tmp_path / "ref2.docx"
    Document().save(str(ref))
    pkg2 = DocxPackage(doc)
    result = tp.apply_template(pkg2, str(ref))
    pkg2.save(do_backup=False)
    assert "MySpecial" in result["styles_cloned"]
    assert result["dangling_style_refs"] == []
    Document(str(doc))


def test_apply_template_real_reference(ch13, tmp_path, doc):
    """Use the real dissertation as the style reference."""
    pkg = DocxPackage(doc)
    result = tp.apply_template(pkg, str(ch13))
    pkg.save(do_backup=False)
    assert result["dangling_style_refs"] == []
    Document(str(doc))


# ------------------------------------------------- properties, styles, alt text


def test_set_document_properties(doc):
    pkg = DocxPackage(doc)
    sx.set_document_properties(
        pkg, title="The Delta Model", author="Nykolus Alvut",
        keywords="alliances, legitimacy, authority",
    )
    pkg.save(do_backup=False)
    d = Document(str(doc))
    assert d.core_properties.title == "The Delta Model"
    assert d.core_properties.author == "Nykolus Alvut"
    assert "legitimacy" in d.core_properties.keywords


def test_define_and_apply_character_style(doc):
    pkg = DocxPackage(doc)
    sx.define_style(
        pkg, style_id="Emphatic", name="Emphatic", style_type="character",
        based_on="DefaultParagraphFont",
        character_formatting={"bold": True, "color": "C00000"},
    )
    sx.apply_character_style(pkg, find="Alpha body one", style="Emphatic")
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert 'w:rStyle w:val="Emphatic"' in xml.replace("<", " ") or "Emphatic" in xml
    Document(str(doc))


def test_character_style_missing_refused(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(TargetNotFound, match="define_style"):
        sx.apply_character_style(pkg, find="Alpha", style="NoSuch")


def test_define_style_validation(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError):
        sx.define_style(pkg, style_id="1bad", name="x")
    with pytest.raises(WordMcpError):
        sx.define_style(
            pkg, style_id="Ok", name="x", style_type="character",
            paragraph_formatting={"alignment": "center"},
        )


def test_image_alt_text(doc, tmp_path):
    from tests.unit.test_furniture_toc_media import make_png

    png = make_png(tmp_path / "a.png")
    pkg = DocxPackage(doc)
    media.add_image(pkg, str(png), at_end=True, width_pt=100)
    sx.set_image_alt_text(
        pkg, 0, description="A test rectangle", title="Test image"
    )
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert 'descr="A test rectangle"' in xml
    Document(str(doc))
