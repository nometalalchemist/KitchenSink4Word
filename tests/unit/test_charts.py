"""CHARTS bundle: add_chart / list_charts / update_chart_data.

Self-contained: documents are built on the fly with python-docx and the
package's own ops. Validation here is structural (parse our own output,
assert parts/rels/content-types/cache consistency); the open-in-real-Word
render gate runs in the deferred validation round.
"""

import copy
import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from word_mcp.core.errors import (
    TargetNotFound,
    UnsupportedStructure,
    WordMcpError,
)
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import charts as ch
from word_mcp.ops import media

C = ch._C
A = ch._A
R = ch._R_NS
CT = ch._CT_NS
REL = ch._REL_NS
CX = ch._CX


# ------------------------------------------------------------------- fixtures


def _fresh_doc(tmp_path: Path, name: str = "doc.docx") -> Path:
    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Anchor paragraph so the body is never empty.")
    doc.add_paragraph("Second paragraph for positioning tests.")
    doc.save(str(f))
    return f


CAT_DATA = {
    "categories": ["Alpha", "Beta", "Gamma"],
    "series": [
        {"name": "S1", "values": [4.3, 2.5, 3.5]},
        {"name": "S2", "values": [1, 2, 3]},
    ],
}

SCATTER_DATA = {
    "series": [
        {"name": "P1", "x": [1, 2, 3], "y": [4.3, 2.5, 3.5]},
        {"name": "P2", "x": [1.5, 2.5, 3.5], "y": [1, 2, 3]},
    ]
}


def _add_and_save(tmp_path, chart_type, data, name="doc.docx", **kw):
    f = _fresh_doc(tmp_path, name)
    pkg = DocxPackage(f)
    result = ch.add_chart(pkg, chart_type, data, **kw)
    pkg.save(do_backup=False)
    return f, result


def _zip_part(path: Path, part: str) -> bytes:
    with zipfile.ZipFile(path) as zf:
        return zf.read(part)


def _chart_root(path: Path, part: str = "word/charts/chart1.xml"):
    return etree.fromstring(_zip_part(path, part))


def _local(el) -> str:
    return etree.QName(el).localname


def _f_text(wrap) -> str:
    for ref in wrap:
        f = ref.find(f"{{{C}}}f")
        if f is not None:
            return f.text
    return None


def _cache_values(wrap):
    """(ptCount, [pt values]) for the first ref/lit child under a data wrap."""
    for holder in wrap.iter():
        if _local(holder) in ("numCache", "strCache", "numLit", "strLit"):
            ptc = holder.find(f"{{{C}}}ptCount")
            pts = [
                pt.find(f"{{{C}}}v").text
                for pt in holder.findall(f"{{{C}}}pt")
            ]
            return int(ptc.get("val")), pts
    return None, []


def _sheet_cells(path: Path, wb_part: str) -> dict:
    """cell ref -> value from the embedded workbook's sheet1."""
    with zipfile.ZipFile(io.BytesIO(_zip_part(path, wb_part))) as zf:
        sheet = etree.fromstring(zf.read("xl/worksheets/sheet1.xml"))
    sml = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    out = {}
    for cell in sheet.iter(f"{{{sml}}}c"):
        t = cell.find(f"{{{sml}}}is/{{{sml}}}t")
        v = cell.find(f"{{{sml}}}v")
        out[cell.get("r")] = t.text if t is not None else v.text
    return out


# ------------------------------------------------------------------ add_chart


def test_add_column_chart_full_structure(tmp_path):
    f, result = _add_and_save(tmp_path, "column", CAT_DATA, title="Totals")
    assert result["chart_added"] == "word/charts/chart1.xml"
    assert result["type"] == "column"
    assert result["series"] == 2
    assert result["points"] == 3
    assert result["chart_index"] == 0

    with zipfile.ZipFile(f) as zf:
        names = set(zf.namelist())
    assert "word/charts/chart1.xml" in names
    assert "word/charts/_rels/chart1.xml.rels" in names
    assert "word/embeddings/Microsoft_Excel_Worksheet1.xlsx" in names

    # Content types: Override for the chart part, Default for xlsx.
    ct = etree.fromstring(_zip_part(f, "[Content_Types].xml"))
    overrides = {
        o.get("PartName"): o.get("ContentType")
        for o in ct.findall(f"{{{CT}}}Override")
    }
    assert overrides["/word/charts/chart1.xml"] == ch._CHART_CT
    defaults = {
        d.get("Extension"): d.get("ContentType")
        for d in ct.findall(f"{{{CT}}}Default")
    }
    assert defaults["xlsx"] == ch._XLSX_CT

    # document.xml.rels carries the chart relationship.
    doc_rels = etree.fromstring(_zip_part(f, "word/_rels/document.xml.rels"))
    chart_rels = [
        r for r in doc_rels if r.get("Type") == ch._CHART_REL
    ]
    assert len(chart_rels) == 1
    assert chart_rels[0].get("Target") == "charts/chart1.xml"

    # Chart part rels: package rel to the workbook.
    crels = etree.fromstring(
        _zip_part(f, "word/charts/_rels/chart1.xml.rels")
    )
    pkg_rels = [r for r in crels if r.get("Type") == ch._PACKAGE_REL]
    assert len(pkg_rels) == 1
    assert (
        pkg_rels[0].get("Target")
        == "../embeddings/Microsoft_Excel_Worksheet1.xlsx"
    )

    root = _chart_root(f)
    # chartSpace order: chart before externalData.
    kids = [_local(k) for k in root]
    assert kids.index("chart") < kids.index("externalData")
    ext = root.find(f"{{{C}}}externalData")
    assert ext.get(f"{{{R}}}id") == "rId1"
    assert ext.find(f"{{{C}}}autoUpdate").get("val") == "0"

    chart = root.find(f"{{{C}}}chart")
    # Title first, then plotArea; title text round-trips.
    chart_kids = [_local(k) for k in chart]
    assert chart_kids[0] == "title"
    assert chart_kids.index("title") < chart_kids.index("plotArea")
    assert "".join(
        t.text or "" for t in chart.find(f"{{{C}}}title").iter(f"{{{A}}}t")
    ) == "Totals"

    bar = chart.find(f"{{{C}}}plotArea/{{{C}}}barChart")
    assert bar.find(f"{{{C}}}barDir").get("val") == "col"
    assert bar.find(f"{{{C}}}grouping").get("val") == "clustered"
    sers = bar.findall(f"{{{C}}}ser")
    assert len(sers) == 2

    # Series 0: order of children, tx/cat/val f + caches.
    ser = sers[0]
    ser_kids = [_local(k) for k in ser]
    assert ser_kids == ["idx", "order", "tx", "cat", "val"]
    assert _f_text(ser.find(f"{{{C}}}tx")) == "Sheet1!$B$1"
    assert _f_text(ser.find(f"{{{C}}}cat")) == "Sheet1!$A$2:$A$4"
    assert _f_text(ser.find(f"{{{C}}}val")) == "Sheet1!$B$2:$B$4"
    n, cats = _cache_values(ser.find(f"{{{C}}}cat"))
    assert n == 3 and cats == ["Alpha", "Beta", "Gamma"]
    n, vals = _cache_values(ser.find(f"{{{C}}}val"))
    assert n == 3 and vals == ["4.3", "2.5", "3.5"]
    # Series 1 lands in column C.
    assert _f_text(sers[1].find(f"{{{C}}}val")) == "Sheet1!$C$2:$C$4"
    # No per-series spPr by default (theme accents apply).
    assert ser.find(f"{{{C}}}spPr") is None

    # Axis pairing: barChart axIds match the axes, crossAx swapped.
    ax_ids = [a.get("val") for a in bar.findall(f"{{{C}}}axId")]
    plot = chart.find(f"{{{C}}}plotArea")
    cat_ax = plot.find(f"{{{C}}}catAx")
    val_ax = plot.find(f"{{{C}}}valAx")
    assert cat_ax.find(f"{{{C}}}axId").get("val") == ax_ids[0]
    assert val_ax.find(f"{{{C}}}axId").get("val") == ax_ids[1]
    assert cat_ax.find(f"{{{C}}}crossAx").get("val") == ax_ids[1]
    assert val_ax.find(f"{{{C}}}crossAx").get("val") == ax_ids[0]
    ax_kids = [_local(k) for k in cat_ax]
    assert ax_kids == ["axId", "scaling", "delete", "axPos", "crossAx"]


def test_add_bar_chart_horizontal(tmp_path):
    f, _ = _add_and_save(tmp_path, "bar", CAT_DATA)
    bar = _chart_root(f).find(f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart")
    assert bar.find(f"{{{C}}}barDir").get("val") == "bar"


def test_add_line_chart(tmp_path):
    f, _ = _add_and_save(tmp_path, "line", CAT_DATA)
    plot = _chart_root(f).find(f"{{{C}}}chart/{{{C}}}plotArea")
    line = plot.find(f"{{{C}}}lineChart")
    assert line.find(f"{{{C}}}grouping").get("val") == "standard"
    for ser in line.findall(f"{{{C}}}ser"):
        assert ser.find(f"{{{C}}}smooth").get("val") == "0"
        assert [_local(k) for k in ser][-1] == "smooth"
    assert len(line.findall(f"{{{C}}}axId")) == 2
    assert plot.find(f"{{{C}}}catAx") is not None
    assert plot.find(f"{{{C}}}valAx") is not None


def test_add_pie_chart_no_axes(tmp_path):
    data = {
        "categories": ["A", "B", "C"],
        "series": [{"name": "Share", "values": [50, 30, 20]}],
    }
    f, result = _add_and_save(tmp_path, "pie", data)
    assert result["series"] == 1
    plot = _chart_root(f).find(f"{{{C}}}chart/{{{C}}}plotArea")
    pie = plot.find(f"{{{C}}}pieChart")
    assert pie.find(f"{{{C}}}varyColors").get("val") == "1"
    assert pie.find(f"{{{C}}}firstSliceAng") is not None
    # No axId in the pie model, no axes in plotArea.
    assert pie.findall(f"{{{C}}}axId") == []
    assert plot.find(f"{{{C}}}catAx") is None
    assert plot.find(f"{{{C}}}valAx") is None


def test_add_scatter_chart(tmp_path):
    f, result = _add_and_save(tmp_path, "scatter", SCATTER_DATA)
    assert result["points"] == 3
    plot = _chart_root(f).find(f"{{{C}}}chart/{{{C}}}plotArea")
    sc = plot.find(f"{{{C}}}scatterChart")
    assert sc.find(f"{{{C}}}scatterStyle").get("val") == "lineMarker"
    sers = sc.findall(f"{{{C}}}ser")
    assert len(sers) == 2
    # Column pairs: series 0 in A/B, series 1 in C/D.
    assert _f_text(sers[0].find(f"{{{C}}}xVal")) == "Sheet1!$A$2:$A$4"
    assert _f_text(sers[0].find(f"{{{C}}}yVal")) == "Sheet1!$B$2:$B$4"
    assert _f_text(sers[1].find(f"{{{C}}}xVal")) == "Sheet1!$C$2:$C$4"
    assert _f_text(sers[1].find(f"{{{C}}}yVal")) == "Sheet1!$D$2:$D$4"
    n, xs = _cache_values(sers[1].find(f"{{{C}}}xVal"))
    assert n == 3 and xs == ["1.5", "2.5", "3.5"]
    # Both axes are value axes.
    assert len(plot.findall(f"{{{C}}}valAx")) == 2
    assert plot.find(f"{{{C}}}catAx") is None


def test_embedded_workbook_layout(tmp_path):
    f, result = _add_and_save(tmp_path, "column", CAT_DATA)
    cells = _sheet_cells(f, result["embedded_workbook"])
    assert "A1" not in cells  # A1 stays blank, Word-style
    assert cells["B1"] == "S1" and cells["C1"] == "S2"
    assert cells["A2"] == "Alpha" and cells["A4"] == "Gamma"
    assert cells["B2"] == "4.3" and cells["C4"] == "3"


def test_embedded_workbook_layout_scatter(tmp_path):
    f, result = _add_and_save(tmp_path, "scatter", SCATTER_DATA)
    cells = _sheet_cells(f, result["embedded_workbook"])
    assert cells["B1"] == "P1" and cells["D1"] == "P2"
    assert cells["A2"] == "1" and cells["B2"] == "4.3"
    assert cells["C2"] == "1.5" and cells["D4"] == "3"


def test_no_title_no_legend(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA, legend=False)
    chart = _chart_root(f).find(f"{{{C}}}chart")
    assert chart.find(f"{{{C}}}title") is None
    assert chart.find(f"{{{C}}}legend") is None


def test_legend_default_on(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    chart = _chart_root(f).find(f"{{{C}}}chart")
    leg = chart.find(f"{{{C}}}legend")
    assert leg is not None
    assert leg.find(f"{{{C}}}legendPos").get("val") == "b"


def test_colors_emit_sppr(tmp_path):
    f, _ = _add_and_save(
        tmp_path, "column", CAT_DATA, colors=["#4472C4", "ED7D31"]
    )
    bar = _chart_root(f).find(f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart")
    sers = bar.findall(f"{{{C}}}ser")
    clrs = [
        next(s.find(f"{{{C}}}spPr").iter(f"{{{A}}}srgbClr")).get("val")
        for s in sers
    ]
    assert clrs == ["4472C4", "ED7D31"]
    # spPr sits after tx, before cat (fixed sequence).
    kids = [_local(k) for k in sers[0]]
    assert kids.index("tx") < kids.index("spPr") < kids.index("cat")


def test_colors_line_chart_strokes_the_line(tmp_path):
    f, _ = _add_and_save(
        tmp_path, "line", CAT_DATA, colors=["112233", "445566"]
    )
    line = _chart_root(f).find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}lineChart"
    )
    sppr = line.find(f"{{{C}}}ser/{{{C}}}spPr")
    assert sppr.find(f"{{{A}}}ln/{{{A}}}solidFill/{{{A}}}srgbClr") is not None


def test_rows_input_shape(tmp_path):
    rows = [["", "S1", "S2"], ["Alpha", 4.3, 1], ["Beta", "2.5", 2]]
    f, result = _add_and_save(tmp_path, "column", rows)
    assert result["series"] == 2 and result["points"] == 2
    ser = _chart_root(f).find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart/{{{C}}}ser"
    )
    n, vals = _cache_values(ser.find(f"{{{C}}}val"))
    assert vals == ["4.3", "2.5"]


def test_csv_path_input(tmp_path):
    csv_file = tmp_path / "data.csv"
    csv_file.write_text(
        "Label,North,South\nQ1,10,20\nQ2,15,25\n", encoding="utf-8"
    )
    f, result = _add_and_save(tmp_path, "line", str(csv_file))
    assert result["series"] == 2 and result["points"] == 2
    line = _chart_root(f).find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}lineChart"
    )
    names = [ch._series_name(s) for s in line.findall(f"{{{C}}}ser")]
    assert names == ["North", "South"]


def test_json_path_dict_input(tmp_path):
    import json

    j = tmp_path / "data.json"
    j.write_text(json.dumps(CAT_DATA), encoding="utf-8")
    f, result = _add_and_save(tmp_path, "column", str(j))
    assert result["series"] == 2 and result["points"] == 3


def test_positioning_after_anchor(tmp_path):
    f = _fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    ch.add_chart(pkg, "column", CAT_DATA, after_anchor="Anchor paragraph")
    body = pkg.body()
    paras = body.findall(qn("w:p"))
    # Chart paragraph directly follows the anchor (index 1 of 3).
    assert paras[1].find(f".//{{{C}}}chart") is not None
    assert paras[0].find(f".//{{{C}}}chart") is None


def test_two_charts_unique_parts_rids_docpr(tmp_path):
    f = _fresh_doc(tmp_path)
    png = tmp_path / "pic.png"
    import struct
    import zlib

    def chunk(tag, data):
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 4, 4, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x10\x20\x30" * 4 for _ in range(4))
    png.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    pkg = DocxPackage(f)
    media.add_image(pkg, str(png), at_end=True)
    r1 = ch.add_chart(pkg, "column", CAT_DATA)
    r2 = ch.add_chart(pkg, "pie", {
        "categories": ["A", "B"],
        "series": [{"name": "S", "values": [1, 2]}],
    })
    pkg.save(do_backup=False)
    assert r1["chart_added"] == "word/charts/chart1.xml"
    assert r2["chart_added"] == "word/charts/chart2.xml"
    assert r2["embedded_workbook"].endswith("Worksheet2.xlsx")
    assert r1["chart_index"] == 0 and r2["chart_index"] == 1
    # docPr ids unique across image + charts.
    wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    ids = [
        el.get("id") for el in pkg.root().iter(f"{{{wp}}}docPr")
    ]
    assert len(ids) == 3 and len(set(ids)) == 3
    # Both chart rels resolve.
    for part in ("word/charts/chart1.xml", "word/charts/chart2.xml"):
        ch._check_chart_closure(pkg, part)


def test_size_defaults_and_overrides(tmp_path):
    f, result = _add_and_save(tmp_path, "column", CAT_DATA)
    assert result["width_pt"] == 432.0 and result["height_pt"] == 252.0
    f2, r2 = _add_and_save(
        tmp_path, "column", CAT_DATA, name="doc2.docx",
        width_pt=300, height_pt=200,
    )
    assert r2["width_pt"] == 300.0 and r2["height_pt"] == 200.0


# ------------------------------------------------------- add_chart refusals


def _pkg(tmp_path, name="r.docx"):
    return DocxPackage(_fresh_doc(tmp_path, name))


def test_refuse_bad_chart_type(tmp_path):
    with pytest.raises(WordMcpError, match="unsupported chart_type"):
        ch.add_chart(_pkg(tmp_path), "doughnut", CAT_DATA)


def test_refuse_ragged_rows(tmp_path):
    rows = [["", "S1"], ["A", 1], ["B", 1, 2]]
    with pytest.raises(WordMcpError, match="ragged"):
        ch.add_chart(_pkg(tmp_path), "column", rows)


def test_refuse_ragged_dict(tmp_path):
    data = {
        "categories": ["A", "B", "C"],
        "series": [{"name": "S1", "values": [1, 2]}],
    }
    with pytest.raises(WordMcpError, match="ragged"):
        ch.add_chart(_pkg(tmp_path), "column", data)


def test_refuse_empty_data(tmp_path):
    with pytest.raises(WordMcpError):
        ch.add_chart(_pkg(tmp_path), "column", [])
    with pytest.raises(WordMcpError):
        ch.add_chart(_pkg(tmp_path), "column", {"categories": [], "series": []})


def test_refuse_non_numeric_value(tmp_path):
    data = {
        "categories": ["A"],
        "series": [{"name": "S1", "values": ["not-a-number"]}],
    }
    with pytest.raises(WordMcpError, match="non-numeric"):
        ch.add_chart(_pkg(tmp_path), "column", data)


def test_refuse_scatter_string_x(tmp_path):
    data = {"series": [{"name": "S", "x": ["Alpha", "Beta"], "y": [1, 2]}]}
    with pytest.raises(WordMcpError, match="line chart"):
        ch.add_chart(_pkg(tmp_path), "scatter", data)


def test_refuse_pie_multi_series(tmp_path):
    with pytest.raises(WordMcpError, match="pie charts show exactly one"):
        ch.add_chart(_pkg(tmp_path), "pie", CAT_DATA)


def test_refuse_positioning_conflict(tmp_path):
    with pytest.raises(WordMcpError, match="at most one"):
        ch.add_chart(
            _pkg(tmp_path), "column", CAT_DATA, after_index=0, at_end=True
        )


def test_refuse_colors_short_and_invalid(tmp_path):
    with pytest.raises(WordMcpError, match="1 entries.*2 series"):
        ch.add_chart(_pkg(tmp_path), "column", CAT_DATA, colors=["4472C4"])
    with pytest.raises(WordMcpError, match="invalid color"):
        ch.add_chart(
            _pkg(tmp_path), "column", CAT_DATA, colors=["red", "blue"]
        )
    with pytest.raises(WordMcpError, match="pie"):
        ch.add_chart(
            _pkg(tmp_path),
            "pie",
            {"categories": ["A"], "series": [{"name": "S", "values": [1]}]},
            colors=["4472C4"],
        )


def test_refuse_bad_size(tmp_path):
    with pytest.raises(WordMcpError, match="width_pt"):
        ch.add_chart(_pkg(tmp_path), "column", CAT_DATA, width_pt=0)


def test_refusal_leaves_file_unchanged(tmp_path):
    f = _fresh_doc(tmp_path)
    before = f.read_bytes()
    pkg = DocxPackage(f)
    with pytest.raises(WordMcpError):
        ch.add_chart(pkg, "column", [["", "S1"], ["A", "x"]])
    assert f.read_bytes() == before


# ---------------------------------------------------------------- list_charts


def test_list_charts(tmp_path):
    f = _fresh_doc(tmp_path)
    pkg = DocxPackage(f)
    ch.add_chart(pkg, "column", CAT_DATA, title="Totals")
    ch.add_chart(pkg, "scatter", SCATTER_DATA)
    out = ch.list_charts(pkg)
    assert len(out) == 2
    first, second = out
    assert first["index"] == 0
    assert first["type"] == "barChart"
    assert first["title"] == "Totals"
    assert first["series"] == ["S1", "S2"]
    assert first["points"] == 3
    assert first["embedded_workbook"] is True
    assert first["supported_for_update"] is True
    assert first["reason"] is None
    assert first["width_pt"] == 432.0
    assert second["type"] == "scatterChart"
    assert second["series"] == ["P1", "P2"]


def _append_chartex_drawing(pkg):
    """Minimal cx: (modern chart) drawing hooked into the body."""
    body = pkg.body()
    p = etree.SubElement(body, qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    drawing = etree.SubElement(r, qn("w:drawing"))
    wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    inline = etree.SubElement(drawing, f"{{{wp}}}inline")
    extent = etree.SubElement(inline, f"{{{wp}}}extent")
    extent.set("cx", "5486400")
    extent.set("cy", "3200400")
    graphic = etree.SubElement(inline, f"{{{A}}}graphic")
    gdata = etree.SubElement(graphic, f"{{{A}}}graphicData")
    gdata.set("uri", CX)
    cx_chart = etree.SubElement(
        gdata, f"{{{CX}}}chart", nsmap={"cx": CX, "r": R}
    )
    cx_chart.set(f"{{{R}}}id", "rId99")


def test_list_charts_labels_chartex(tmp_path):
    pkg = _pkg(tmp_path)
    ch.add_chart(pkg, "column", CAT_DATA)
    _append_chartex_drawing(pkg)
    out = ch.list_charts(pkg)
    assert len(out) == 2
    assert out[1]["type"] == "chartex"
    assert out[1]["supported_for_update"] is False
    assert "chartex" in out[1]["reason"]


def test_list_charts_flags_unsupported_type(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    part = "word/charts/chart1.xml"
    root = pkg.root(part)
    bar = root.find(f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart")
    bar.tag = f"{{{C}}}radarChart"
    out = ch.list_charts(pkg)
    assert out[0]["supported_for_update"] is False
    assert "radarChart" in out[0]["reason"]


def test_list_charts_flags_combo(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    root = pkg.root("word/charts/chart1.xml")
    plot = root.find(f"{{{C}}}chart/{{{C}}}plotArea")
    bar = plot.find(f"{{{C}}}barChart")
    bar.addnext(copy.deepcopy(bar))
    plot[2].tag = f"{{{C}}}lineChart"  # second group, different type
    out = ch.list_charts(pkg)
    assert out[0]["supported_for_update"] is False
    assert "combo" in out[0]["reason"]


# ---------------------------------------------------------- update_chart_data


NEW_CAT_DATA = {
    "categories": ["One", "Two", "Three", "Four"],
    "series": [
        {"name": "S1", "values": [10, 20, 30, 40]},
        {"name": "S2", "values": [5, 6, 7, 8]},
    ],
}


def test_update_roundtrip_on_own_chart(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    result = ch.update_chart_data(
        pkg, 0, NEW_CAT_DATA, series_names=["North", "South"]
    )
    pkg.save(do_backup=False)
    assert result["points_before"] == 3
    assert result["points_after"] == 4
    assert result["embedded_workbook"] == "regenerated"

    root = _chart_root(f)
    bar = root.find(f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart")
    sers = bar.findall(f"{{{C}}}ser")
    n, cats = _cache_values(sers[0].find(f"{{{C}}}cat"))
    assert n == 4 and cats == ["One", "Two", "Three", "Four"]
    n, vals = _cache_values(sers[1].find(f"{{{C}}}val"))
    assert n == 4 and vals == ["5", "6", "7", "8"]
    assert _f_text(sers[0].find(f"{{{C}}}cat")) == "Sheet1!$A$2:$A$5"
    assert _f_text(sers[1].find(f"{{{C}}}val")) == "Sheet1!$C$2:$C$5"
    n, names = _cache_values(sers[0].find(f"{{{C}}}tx"))
    assert names == ["North"]
    # Workbook regenerated with the new data.
    cells = _sheet_cells(f, "word/embeddings/Microsoft_Excel_Worksheet1.xlsx")
    assert cells["A5"] == "Four" and cells["C5"] == "8"
    assert cells["B1"] == "North" and cells["C1"] == "South"


def test_update_scatter(tmp_path):
    f, _ = _add_and_save(tmp_path, "scatter", SCATTER_DATA)
    pkg = DocxPackage(f)
    new = {
        "series": [
            {"name": "P1", "x": [1, 2], "y": [9, 8]},
            {"name": "P2", "x": [3, 4], "y": [7, 6]},
        ]
    }
    result = ch.update_chart_data(pkg, 0, new)
    pkg.save(do_backup=False)
    assert result["points_after"] == 2
    sc = _chart_root(f).find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}scatterChart"
    )
    sers = sc.findall(f"{{{C}}}ser")
    n, xs = _cache_values(sers[1].find(f"{{{C}}}xVal"))
    assert n == 2 and xs == ["3", "4"]
    assert _f_text(sers[1].find(f"{{{C}}}yVal")) == "Sheet1!$D$2:$D$3"


def test_update_pie(tmp_path):
    data = {
        "categories": ["A", "B"],
        "series": [{"name": "Share", "values": [60, 40]}],
    }
    f, _ = _add_and_save(tmp_path, "pie", data)
    pkg = DocxPackage(f)
    new = {
        "categories": ["X", "Y", "Z"],
        "series": [{"name": "Share", "values": [10, 20, 70]}],
    }
    result = ch.update_chart_data(pkg, 0, new)
    assert result["points_after"] == 3
    pie = pkg.root("word/charts/chart1.xml").find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}pieChart"
    )
    n, vals = _cache_values(pie.find(f"{{{C}}}ser/{{{C}}}val"))
    assert vals == ["10", "20", "70"]


def test_update_preserves_c14_extlst(tmp_path):
    """Word-round-tripped charts carry c14/c16 extLst content in c:ser and
    inside caches; in-place surgery must keep every bit of it."""
    c14 = "http://schemas.microsoft.com/office/drawing/2007/8/2/chart"
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    part = "word/charts/chart1.xml"
    root = pkg.root(part)
    ser = root.find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart/{{{C}}}ser"
    )
    # Synthetic Word-style extension content: extLst on the ser...
    ext_lst = etree.SubElement(ser, f"{{{C}}}extLst")
    ext = etree.SubElement(ext_lst, f"{{{C}}}ext")
    ext.set("uri", "{6F2FDCE9-48DA-4B69-8628-5D25D57E5C99}")
    style = etree.SubElement(ext, f"{{{c14}}}style", nsmap={"c14": c14})
    style.set("val", "102")
    # ...and one inside the numCache.
    cache = ser.find(f"{{{C}}}val/{{{C}}}numRef/{{{C}}}numCache")
    cache_ext_lst = etree.SubElement(cache, f"{{{C}}}extLst")
    cache_ext = etree.SubElement(cache_ext_lst, f"{{{C}}}ext")
    cache_ext.set("uri", "{TEST-CACHE-EXT}")
    pkg.mark_dirty(part)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(f)
    ch.update_chart_data(pkg2, 0, NEW_CAT_DATA)
    pkg2.save(do_backup=False)
    root2 = _chart_root(f)
    ser2 = root2.find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart/{{{C}}}ser"
    )
    kept = ser2.find(f"{{{C}}}extLst/{{{C}}}ext")
    assert kept is not None
    assert kept.find(f"{{{c14}}}style").get("val") == "102"
    cache2 = ser2.find(f"{{{C}}}val/{{{C}}}numRef/{{{C}}}numCache")
    cache_kids = [_local(k) for k in cache2]
    # extLst survives AND sits last, after the rewritten points.
    assert cache_kids[-1] == "extLst"
    assert cache2.find(f"{{{C}}}extLst/{{{C}}}ext").get("uri") == "{TEST-CACHE-EXT}"
    n, vals = _cache_values(ser2.find(f"{{{C}}}val"))
    assert n == 4 and vals == ["10", "20", "30", "40"]


def test_update_without_workbook_reports_none(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    # Simulate a PHPWord-style chart: no externalData, no package rel.
    part = "word/charts/chart1.xml"
    root = pkg.root(part)
    ext = root.find(f"{{{C}}}externalData")
    root.remove(ext)
    pkg.mark_dirty(part)
    rels_part = "word/charts/_rels/chart1.xml.rels"
    rels_root = pkg.root(rels_part)
    for rel in list(rels_root):
        rels_root.remove(rel)
    pkg.mark_dirty(rels_part)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(f)
    result = ch.update_chart_data(pkg2, 0, NEW_CAT_DATA)
    assert result["embedded_workbook"] == "none"
    n, vals = _cache_values(
        pkg2.root(part).find(
            f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart/{{{C}}}ser/{{{C}}}val"
        )
    )
    assert n == 4


# ------------------------------------------------- update_chart_data refusals


def test_update_refuses_series_count_change(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    one_series = {
        "categories": ["A", "B"],
        "series": [{"name": "Only", "values": [1, 2]}],
    }
    with pytest.raises(UnsupportedStructure, match="series count"):
        ch.update_chart_data(pkg, 0, one_series)


def test_update_refuses_radar(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    root = pkg.root("word/charts/chart1.xml")
    root.find(
        f"{{{C}}}chart/{{{C}}}plotArea/{{{C}}}barChart"
    ).tag = f"{{{C}}}radarChart"
    with pytest.raises(UnsupportedStructure, match="radarChart"):
        ch.update_chart_data(pkg, 0, CAT_DATA)


def test_update_refuses_combo(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    plot = pkg.root("word/charts/chart1.xml").find(
        f"{{{C}}}chart/{{{C}}}plotArea"
    )
    bar = plot.find(f"{{{C}}}barChart")
    bar.addnext(copy.deepcopy(bar))
    with pytest.raises(UnsupportedStructure, match="combo"):
        ch.update_chart_data(pkg, 0, CAT_DATA)


def test_update_refuses_multilvl_categories(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    root = pkg.root("word/charts/chart1.xml")
    for ser in root.iter(f"{{{C}}}ser"):
        cat = ser.find(f"{{{C}}}cat")
        for child in list(cat):
            cat.remove(child)
        etree.SubElement(cat, f"{{{C}}}multiLvlStrRef")
    with pytest.raises(UnsupportedStructure, match="multi-level"):
        ch.update_chart_data(pkg, 0, CAT_DATA)


def test_update_refuses_chartex(tmp_path):
    pkg = _pkg(tmp_path)
    _append_chartex_drawing(pkg)
    with pytest.raises(UnsupportedStructure, match="modern chart"):
        ch.update_chart_data(pkg, 0, CAT_DATA)


def test_update_refuses_ragged(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    bad = {
        "categories": ["A", "B"],
        "series": [
            {"name": "S1", "values": [1, 2]},
            {"name": "S2", "values": [1]},
        ],
    }
    with pytest.raises(WordMcpError, match="ragged"):
        ch.update_chart_data(pkg, 0, bad)


def test_update_refuses_bad_index(tmp_path):
    pkg = _pkg(tmp_path)
    with pytest.raises(TargetNotFound, match="out of range"):
        ch.update_chart_data(pkg, 0, CAT_DATA)


def test_update_refuses_bad_series_names_count(tmp_path):
    f, _ = _add_and_save(tmp_path, "column", CAT_DATA)
    pkg = DocxPackage(f)
    with pytest.raises(WordMcpError, match="series_names"):
        ch.update_chart_data(pkg, 0, CAT_DATA, series_names=["Only one"])


# ------------------------------------------------------------------- plumbing


def test_col_letter():
    assert ch._col_letter(0) == "A"
    assert ch._col_letter(25) == "Z"
    assert ch._col_letter(26) == "AA"
    assert ch._col_letter(27) == "AB"
    assert ch._col_letter(701) == "ZZ"
    assert ch._col_letter(702) == "AAA"


def test_fmt_num():
    assert ch._fmt_num(4.0) == "4"
    assert ch._fmt_num(4.3) == "4.3"
    assert ch._fmt_num(-2.0) == "-2"
    assert ch._fmt_num(0.1) == "0.1"


def test_special_characters_survive(tmp_path):
    data = {
        "categories": ["A & B", "<Test>", 'Quo"te'],
        "series": [{"name": "R&D <spend>", "values": [1, 2, 3]}],
    }
    f, _ = _add_and_save(tmp_path, "column", data, title="P&L <2026>")
    root = _chart_root(f)
    chart = root.find(f"{{{C}}}chart")
    assert "".join(
        t.text or "" for t in chart.find(f"{{{C}}}title").iter(f"{{{A}}}t")
    ) == "P&L <2026>"
    ser = chart.find(f"{{{C}}}plotArea/{{{C}}}barChart/{{{C}}}ser")
    _, cats = _cache_values(ser.find(f"{{{C}}}cat"))
    assert cats == ["A & B", "<Test>", 'Quo"te']
    assert ch._series_name(ser) == "R&D <spend>"
    cells = _sheet_cells(f, "word/embeddings/Microsoft_Excel_Worksheet1.xlsx")
    assert cells["A2"] == "A & B" and cells["B1"] == "R&D <spend>"


def test_registration_snippet_importable():
    """The integration snippet registers cleanly on the shared mcp instance
    (paste-readiness smoke test; mirrors the dataio bundle's convention)."""
    import importlib
    import sys

    root = Path(__file__).resolve().parents[2]
    if not (root / "integration" / "charts_registrations.py").exists():
        pytest.skip("integration/ staging dir not present (gitignored)")
    sys.path.insert(0, str(root))
    try:
        mod = importlib.import_module("integration.charts_registrations")
        assert hasattr(mod, "add_chart")
        assert hasattr(mod, "list_charts")
        assert hasattr(mod, "update_chart_data")
    finally:
        sys.path.remove(str(root))
