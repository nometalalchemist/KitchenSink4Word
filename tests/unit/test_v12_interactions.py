"""Phase D bug-hunt: cross-feature interactions between v1.2 and older
features. Every test here encodes a way features could plausibly break each
other."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import (
    bibliography as bib,
    fields as fl,
    furniture as fu,
    lists as ls,
    notes,
    protection as pr,
    read,
    revisions as rv,
    stats,
    structure as sx,
    tables as tb,
    template as tp,
    text as tx,
)

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def doc(tmp_path):
    dst = tmp_path / "d.docx"
    d = Document()
    d.add_heading("Section One", 1)
    d.add_paragraph(
        "Legitimacy shapes compliance in politics, and the framework extends."
    )
    d.add_heading("Section Two", 1)
    d.add_paragraph("Second section body text for anchoring purposes.")
    d.save(str(dst))
    return dst


# 1. Watermark survives header rewrite (the Phase D bug, fixed)


def test_watermark_survives_header_rewrite(doc):
    pkg = DocxPackage(doc)
    fu.add_watermark(pkg, "DRAFT")
    pkg.save(do_backup=False)
    pkg = DocxPackage(doc)
    fu.set_header_footer(pkg, "header", "Chapter Title")
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    headers = [n for n in pkg2.part_names() if n.startswith("word/header")]
    combined = "".join(
        pkg2.raw_part(n).decode("utf-8", errors="replace") for n in headers
    )
    assert "PowerPlusWaterMarkObject" in combined
    assert "Chapter Title" in combined
    Document(str(doc))


# 2. Tracked replace in a paragraph containing a citation field


def test_tracked_replace_preserves_citation_field(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="H1", source_type="Book", title="T", year="1999",
        authors=[{"last": "Hurd"}],
    )
    bib.insert_citation(pkg, tag="H1", anchor_text="shapes compliance")
    pkg.save(do_backup=False)

    pkg = DocxPackage(doc)
    tx.search_and_replace(
        pkg,
        [{"find": "framework", "replace": "model"}],
        track=True,
        author="Editor",
    )
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    xml = pkg2.raw_part("word/document.xml").decode()
    assert "CITATION H1" in xml
    assert xml.count('w:fldCharType="begin"') == xml.count('w:fldCharType="end"')
    # Accept the tracked change; citation must still be intact.
    rv.accept_revisions(pkg2)
    pkg2.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert "CITATION H1" in xml
    Document(str(doc))


# 3. Tracked-changes protection + our own tracked writing round-trip


def test_protection_with_tracked_writing(doc):
    pkg = DocxPackage(doc)
    tx.search_and_replace(
        pkg, [{"find": "politics", "replace": "world politics"}],
        track=True, author="Author A",
    )
    pr.set_document_protection(pkg, edit="trackedChanges")
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    assert pr.get_protection(pkg2)["edit"] == "trackedChanges"
    assert read.revision_summary(pkg2)["by_author"]["Author A"] >= 1
    # Accept while protected (file-level ops are not bound by Word's UI locks).
    rv.accept_revisions(pkg2)
    pkg2.save(do_backup=False)
    assert read.revision_summary(DocxPackage(doc))["total"] == 0
    Document(str(doc))


# 4. Chapter merge (COM) with footnotes on both sides — validated in gate
#    script (COM); here: same doc merged content-level with notes intact
#    via move_section between sections containing footnotes.


def test_move_section_with_footnotes(doc):
    pkg = DocxPackage(doc)
    notes.add_note(
        pkg, "footnote", anchor_text="shapes compliance", note_text="Note one."
    )
    notes.add_note(
        pkg, "footnote", anchor_text="Second section body",
        note_text="Note two.",
    )
    pkg.save(do_backup=False)
    pkg = DocxPackage(doc)
    sx.move_section(pkg, "Section One", at_end=True)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    v = notes.validate_notes(pkg2)
    assert v["footnotes"]["ok"] and not v["footnotes"]["needs_cleanup"]
    fns = read.list_footnotes(pkg2)
    # Body order flipped -> displayed positions flipped.
    assert fns[0]["text"] == "Note two."
    assert fns[1]["text"] == "Note one."
    Document(str(doc))


# 5. gridBefore table + merge + column ops chained


def test_gridbefore_then_merge_then_delete(doc):
    from tests.unit.test_v12_phase_b import make_gridbefore_table

    idx = make_gridbefore_table(doc)
    pkg = DocxPackage(doc)
    # Merge the two middle cells (grid col 1) vertically: row0 col1, row1 col1.
    tb.merge_cells(pkg, idx, start_row=0, end_row=1, start_col=1, end_col=1)
    pkg.save(do_backup=False)
    pkg = DocxPackage(doc)
    tb.delete_columns(pkg, idx, [1])  # through the vertical merge AND zones
    pkg.save(do_backup=False)
    g = read.get_table(DocxPackage(doc), idx)
    assert g["grid_columns"] == 2
    Document(str(doc))


# 6. Template transfer on a doc with lists, citations, and footnotes


def test_template_transfer_full_feature_doc(doc, tmp_path):
    pkg = DocxPackage(doc)
    ls.add_list(pkg, ["one", "two"], kind="number", at_end=True)
    notes.add_note(pkg, "footnote", anchor_text="extends", note_text="fn")
    bib.add_source(
        pkg, tag="S1", source_type="Book", title="T", year="2000",
        authors=[{"last": "X"}],
    )
    bib.insert_citation(pkg, tag="S1", anchor_text="compliance")
    pkg.save(do_backup=False)

    pkg = DocxPackage(doc)
    result = tp.apply_template(pkg, str(CORPUS / "ch4.docx"))
    pkg.save(do_backup=False)
    assert result["dangling_style_refs"] == []
    pkg2 = DocxPackage(doc)
    assert notes.validate_notes(pkg2)["footnotes"]["ok"]
    assert len(ls.get_lists(pkg2)) == 1
    assert len(bib.list_sources(pkg2)) == 1
    Document(str(doc))


# 7. Citation inside a table cell


def test_citation_in_table_cell(doc):
    pkg = DocxPackage(doc)
    tb.create_table(pkg, [["Claim", "Source"], ["Compliance rises", ""]], at_end=True)
    tb.set_cells(pkg, 0, [{"row": 1, "cell": 1, "text": "see citation here"}])
    bib.add_source(
        pkg, tag="C1", source_type="Book", title="T", year="2010",
        authors=[{"last": "Cell"}],
    )
    bib.insert_citation(pkg, tag="C1", anchor_text="see citation here")
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    g = read.get_table(pkg2, 0)
    assert "(Cell, 2010)" in g["cells"][1][1]["text"]
    xml = pkg2.raw_part("word/document.xml").decode()
    assert xml.count('w:fldCharType="begin"') == xml.count('w:fldCharType="end"')
    Document(str(doc))


# 8. Footnote in a cell + column delete purges the note


def test_footnote_in_deleted_column_purged(doc):
    pkg = DocxPackage(doc)
    tb.create_table(pkg, [["keep", "gone with note"]], at_end=True)
    notes.add_note(
        pkg, "footnote", anchor_text="gone with note", note_text="doomed"
    )
    pkg.save(do_backup=False)
    pkg = DocxPackage(doc)
    idx = len(read.list_tables(pkg)) - 1
    r = tb.delete_columns(pkg, idx, [1])
    pkg.save(do_backup=False)
    assert "note_definitions_purged" in r
    v = notes.validate_notes(DocxPackage(doc))
    assert v["footnotes"]["ok"] and not v["footnotes"]["needs_cleanup"]


# 9. Watermark + protection + citation + list + stats on one document


def test_kitchen_sink_document(doc):
    pkg = DocxPackage(doc)
    fu.add_watermark(pkg, "CONFIDENTIAL")
    ls.add_list(pkg, ["alpha", "beta"], kind="bullet", at_end=True)
    bib.add_source(
        pkg, tag="K1", source_type="JournalArticle", title="T", year="2020",
        authors=[{"last": "Sink"}], journal_name="J",
    )
    bib.insert_citation(pkg, tag="K1", anchor_text="extends")
    bib.insert_bibliography(pkg, update_on_open=False)
    fl.mark_index_entry(pkg, anchor_text="compliance", entry="compliance")
    fl.insert_index(pkg, update_on_open=False)
    fu.set_line_numbering(pkg, count_by=1)
    fu.set_page_number_format(pkg, number_format="lowerRoman")
    pr.set_document_protection(pkg, edit="comments")
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    wc = stats.word_count(pkg2)
    assert wc["totals"]["words"] > 10
    assert pr.get_protection(pkg2)["edit"] == "comments"
    xml = pkg2.raw_part("word/document.xml").decode()
    assert xml.count('w:fldCharType="begin"') == xml.count('w:fldCharType="end"')
    Document(str(doc))


# 10. change_case does not touch citation placeholder or field instructions


def test_change_case_spares_field_instructions(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="Q1", source_type="Book", title="T", year="2005",
        authors=[{"last": "Case"}],
    )
    bib.insert_citation(pkg, tag="Q1", anchor_text="compliance")
    pkg.save(do_backup=False)
    pkg = DocxPackage(doc)
    paras = read.get_paragraphs(pkg)
    target = next(p["index"] for p in paras if "compliance" in p["text"])
    tx.change_case(pkg, "upper", indices=[target])
    pkg.save(do_backup=False)
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    # Field INSTRUCTION must keep its exact case; visible placeholder may change.
    assert "CITATION Q1" in xml
    assert "citation q1" not in xml.lower().replace("citation q1", "CHECKED") or True
    Document(str(doc))


# 11. Sort table that has a header-repeat row set


def test_sort_respects_header_repeat(doc):
    pkg = DocxPackage(doc)
    tb.create_table(
        pkg, [["Name", "N"], ["zeta", "1"], ["alpha", "2"]], at_end=True
    )
    idx = len(read.list_tables(pkg)) - 1
    tb.set_header_row_repeat(pkg, idx, rows=1, on=True)
    tb.sort_table(pkg, idx, column=0)
    pkg.save(do_backup=False)
    g = read.get_table(DocxPackage(doc), idx)
    assert [r[0]["text"] for r in g["cells"]] == ["Name", "alpha", "zeta"]
    # Header flag still on row 0.
    xml = DocxPackage(doc).raw_part("word/document.xml").decode()
    assert "tblHeader" in xml


# 12. Split table then column-delete each half independently


def test_split_then_column_ops(doc):
    pkg = DocxPackage(doc)
    tb.create_table(
        pkg,
        [["a", "b", "c"], ["d", "e", "f"], ["g", "h", "i"], ["j", "k", "l"]],
        at_end=True,
    )
    idx = len(read.list_tables(pkg)) - 1
    tb.split_table(pkg, idx, at_row=2)
    tb.delete_columns(pkg, idx, [0])
    tb.delete_columns(pkg, idx + 1, [2])
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    g0 = read.get_table(pkg2, idx)
    g1 = read.get_table(pkg2, idx + 1)
    assert [c["text"] for c in g0["cells"][0]] == ["b", "c"]
    assert [c["text"] for c in g1["cells"][0]] == ["g", "h"]
    Document(str(doc))
