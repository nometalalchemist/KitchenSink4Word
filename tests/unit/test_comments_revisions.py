"""Phase 7 gate: comment authoring + revision accept/reject on the real
committee-revised EJIR doc (126 revisions by the chair) and the
committee-commented outline (33 comments, threaded)."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.package import DocxPackage
from word_mcp.ops import comments as cm, read, revisions as rv

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def ejir(tmp_path):
    dst = tmp_path / "ejir_rw.docx"
    shutil.copy(CORPUS / "ejir_rw.docx", dst)
    return dst


@pytest.fixture
def outline(tmp_path):
    dst = tmp_path / "outline.docx"
    shutil.copy(CORPUS / "outline.docx", dst)
    return dst


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


def visible_text(path):
    pkg = DocxPackage(path)
    return "\n".join(e["text"] for e in read.get_paragraphs(pkg))


# ------------------------------------------------------------------ revisions


def test_accept_all_preserves_visible_text(ejir):
    before = visible_text(ejir)
    pkg = DocxPackage(ejir)
    n_revs = read.revision_summary(pkg)["total"]
    result = rv.accept_revisions(pkg)
    pkg.save(do_backup=False)

    assert result["revisions_resolved"] > 0
    pkg2 = DocxPackage(ejir)
    assert read.revision_summary(pkg2)["total"] == 0
    after = visible_text(ejir)
    assert after == before, "accepting all must not change what is displayed"
    Document(str(ejir))


def test_reject_all_restores_original(ejir):
    pkg = DocxPackage(ejir)
    revs = read.get_tracked_changes(pkg)
    some_insertion = next(
        (r["text"] for r in revs if r["type"] == "insertion" and len(r["text"].strip()) > 15),
        None,
    )
    some_deletion = next(
        (r["text"] for r in revs if r["type"] == "deletion" and len(r["text"].strip()) > 15),
        None,
    )
    result = rv.reject_revisions(pkg)
    pkg.save(do_backup=False)

    assert result["revisions_resolved"] > 0
    pkg2 = DocxPackage(ejir)
    assert read.revision_summary(pkg2)["total"] == 0
    after = visible_text(ejir)
    if some_deletion:
        assert some_deletion in after, "rejected deletion must restore its text"
    if some_insertion:
        assert some_insertion not in after, "rejected insertion must vanish"
    Document(str(ejir))


def test_author_filter_no_match(ejir):
    pkg = DocxPackage(ejir)
    result = rv.accept_revisions(pkg, author="Nobody Real")
    assert result["revisions_resolved"] == 0
    assert read.revision_summary(pkg)["total"] > 0


def test_author_filter_match(ejir):
    pkg = DocxPackage(ejir)
    authors = list(read.revision_summary(pkg)["by_author"])
    result = rv.accept_revisions(pkg, author=authors[0])
    pkg.save(do_backup=False)
    assert result["revisions_resolved"] > 0
    remaining = read.revision_summary(DocxPackage(ejir))["by_author"]
    assert authors[0] not in remaining


# ------------------------------------------------------------------- comments


def test_add_comment_full_stack(ch4):
    """ch4 has no comments infrastructure; everything must be created."""
    pkg = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg)
    lp = max(paras, key=lambda p: len(p["text"]))
    anchor = " ".join(lp["text"].split()[5:9])
    result = cm.add_comment(
        pkg,
        anchor_text=anchor,
        text="This claim needs a citation.",
        author="Nykolus Alvut",
    )
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    for part in (
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsIds.xml",
        "word/people.xml",
    ):
        assert pkg2.has_part(part), part
    got = read.get_comments(pkg2)
    assert len(got) == 1
    assert got[0]["author"] == "Nykolus Alvut"
    assert got[0]["text"] == "This claim needs a citation."
    assert got[0]["anchored_text"] == anchor
    assert got[0]["resolved"] is False
    Document(str(ch4))


def test_reply_resolve_roundtrip(ch4):
    pkg = DocxPackage(ch4)
    paras = read.get_paragraphs(pkg)
    lp = max(paras, key=lambda p: len(p["text"]))
    anchor = " ".join(lp["text"].split()[5:9])
    r1 = cm.add_comment(
        pkg, anchor_text=anchor, text="Original point.", author="Reviewer"
    )
    cm.reply_to_comment(
        pkg, comment_id=r1["comment_id"], text="Addressed in rev 2.", author="Author"
    )
    cm.resolve_comment(pkg, comment_id=r1["comment_id"])
    pkg.save(do_backup=False)

    got = read.get_comments(DocxPackage(ch4))
    assert len(got) == 2
    root = next(c for c in got if c["text"] == "Original point.")
    reply = next(c for c in got if c["text"] == "Addressed in rev 2.")
    assert reply["reply_to"] == root["id"]
    assert root["resolved"] is True
    Document(str(ch4))


def test_delete_comment_cascades_replies(outline):
    pkg = DocxPackage(outline)
    all_comments = read.get_comments(pkg)
    threaded = [c for c in all_comments if c["reply_to"] is not None]
    if not threaded:
        pytest.skip("no threaded comments in outline")
    parent_id = threaded[0]["reply_to"]
    expected_gone = {parent_id} | {
        c["id"] for c in all_comments if c["reply_to"] == parent_id
    }
    result = cm.delete_comment(pkg, comment_id=parent_id)
    pkg.save(do_backup=False)

    remaining = {c["id"] for c in read.get_comments(DocxPackage(outline))}
    assert not (expected_gone & remaining)
    assert set(result["deleted_comments"]) >= expected_gone
    Document(str(outline))


def test_delete_leaf_comment_only(outline):
    pkg = DocxPackage(outline)
    all_comments = read.get_comments(pkg)
    reply_ids = {c["reply_to"] for c in all_comments if c["reply_to"]}
    leaves = [
        c for c in all_comments if c["reply_to"] and c["id"] not in reply_ids
    ]
    if not leaves:
        pytest.skip("no leaf replies")
    victim = leaves[0]
    n = len(all_comments)
    cm.delete_comment(pkg, comment_id=victim["id"])
    pkg.save(do_backup=False)
    remaining = read.get_comments(DocxPackage(outline))
    assert len(remaining) == n - 1
    assert all(c["id"] != victim["id"] for c in remaining)


def test_comment_on_committee_doc_alongside_revisions(ejir):
    """Comment + accept revisions on the same doc: both features coexist."""
    pkg = DocxPackage(ejir)
    n_comments = len(read.get_comments(pkg))
    paras = read.get_paragraphs(pkg)
    lp = max(paras, key=lambda p: len(p["text"]))
    anchor = " ".join(lp["text"].split()[2:5])
    cm.add_comment(pkg, anchor_text=anchor, text="Post-review note.", author="NA")
    rv.accept_revisions(pkg)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ejir)
    assert len(read.get_comments(pkg2)) == n_comments + 1
    assert read.revision_summary(pkg2)["total"] == 0
    Document(str(ejir))
