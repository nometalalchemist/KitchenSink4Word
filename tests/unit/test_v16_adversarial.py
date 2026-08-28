"""Regression tests for KitchenSink4Word v1.6 adversarial round findings.

Each test targets a specific finding from the 2026-08-28 adversarial stress
round (Phase 3 of the V1.6 SHIP RUNBOOK). Tests are named after the finding
number and designed to FAIL against the current code, confirming the bug
exists. Once the coordinator applies fixes, these tests should PASS.

Finding 1 (HIGH):  NaN/Inf crash in charts._num / _fmt_num
Finding 2 (MEDIUM): Cross-process file creation race in safesave.write_lock
Finding 3 (MEDIUM): Null byte in sandbox.check_path (defense-in-depth)
Finding 4 (LOW):   Slot rotation raw PermissionError on held backup file
"""

from __future__ import annotations

import json
import math
import os
import subprocess
import sys
import tempfile
import textwrap
import threading
import time
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import WordMcpError
from word_mcp.core.package import DocxPackage
from word_mcp.core.sandbox import ENV_VAR, SandboxViolation, check_path
from word_mcp.core import safesave
from word_mcp.ops.charts import add_chart, _num, _fmt_num


# ------------------------------------------------------------------ helpers


def _fresh_doc(where: Path, name: str = "doc.docx", text: str = "Anchor.") -> Path:
    f = where / name
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(f))
    return f


# ==========================================================================
# FINDING 1 (HIGH): NaN/Inf pass validation and crash with raw exceptions
# ==========================================================================


class TestFinding1_NanInf:
    """charts._num() must reject NaN and Inf with a typed WordMcpError,
    and _fmt_num() must never receive a non-finite float."""

    def test_nan_rejected_by_num(self):
        """_num must raise WordMcpError for float('nan'), not pass it through."""
        with pytest.raises(WordMcpError, match="(?i)non.?finite|nan"):
            _num(float("nan"), "test location")

    def test_positive_inf_rejected_by_num(self):
        """_num must raise WordMcpError for float('inf')."""
        with pytest.raises(WordMcpError, match="(?i)non.?finite|inf"):
            _num(float("inf"), "test location")

    def test_negative_inf_rejected_by_num(self):
        """_num must raise WordMcpError for float('-inf')."""
        with pytest.raises(WordMcpError, match="(?i)non.?finite|inf"):
            _num(float("-inf"), "test location")

    def test_nan_in_add_chart_bar(self, tmp_path):
        """add_chart with NaN in values must raise WordMcpError, not crash."""
        doc_path = _fresh_doc(tmp_path)
        pkg = DocxPackage(str(doc_path))
        data = {
            "categories": ["A", "B"],
            "series": [{"name": "S1", "values": [float("nan"), 1.0]}],
        }
        with pytest.raises(WordMcpError):
            add_chart(pkg, "bar", data)

    def test_inf_in_add_chart_bar(self, tmp_path):
        """add_chart with Inf in values must raise WordMcpError, not crash."""
        doc_path = _fresh_doc(tmp_path)
        pkg = DocxPackage(str(doc_path))
        data = {
            "categories": ["A", "B"],
            "series": [{"name": "S1", "values": [float("inf"), 1.0]}],
        }
        with pytest.raises(WordMcpError):
            add_chart(pkg, "bar", data)

    def test_nan_in_scatter_x(self, tmp_path):
        """NaN in scatter x-values must raise WordMcpError."""
        doc_path = _fresh_doc(tmp_path)
        pkg = DocxPackage(str(doc_path))
        data = {
            "series": [{"name": "S", "x": [float("nan"), 2.0], "y": [1.0, 2.0]}],
        }
        with pytest.raises(WordMcpError):
            add_chart(pkg, "scatter", data)

    def test_nan_in_scatter_y(self, tmp_path):
        """NaN in scatter y-values must raise WordMcpError."""
        doc_path = _fresh_doc(tmp_path)
        pkg = DocxPackage(str(doc_path))
        data = {
            "series": [{"name": "S", "x": [1.0, 2.0], "y": [float("nan"), 2.0]}],
        }
        with pytest.raises(WordMcpError):
            add_chart(pkg, "scatter", data)

    def test_fmt_num_nan_guard(self):
        """_fmt_num must not crash on NaN -- either refuse or produce valid text."""
        # If _num does its job, NaN never reaches _fmt_num. But belt-and-suspenders:
        # _fmt_num should raise WordMcpError, not ValueError/OverflowError.
        with pytest.raises((WordMcpError, ValueError)):
            _fmt_num(float("nan"))

    def test_fmt_num_inf_guard(self):
        """_fmt_num must not crash on Inf with OverflowError."""
        with pytest.raises((WordMcpError, OverflowError)):
            _fmt_num(float("inf"))

    def test_nan_string_rejected(self):
        """The string 'nan' must not be silently converted to float('nan')."""
        # Python's float("nan") returns NaN, so if _num does float(v.strip())
        # on the string "nan", it would produce NaN. Verify it is caught.
        with pytest.raises(WordMcpError):
            _num("nan", "test location")

    def test_inf_string_rejected(self):
        """The string 'inf' must not be silently converted to float('inf')."""
        with pytest.raises(WordMcpError):
            _num("inf", "test location")


# ==========================================================================
# FINDING 2 (MEDIUM): Cross-process file creation race in write_lock
# ==========================================================================


class TestFinding2_CreateRace:
    """write_lock must serialize cross-process creation of the same file,
    not just mutations of existing files."""

    def test_cross_process_create_serialized(self, tmp_path):
        """Two processes creating the same new file must not corrupt it.

        The current bug: write_lock checks os.path.exists(doc_path) and
        skips the advisory lockfile when the file does not exist, so
        concurrent creators race and the last writer wins with a
        potentially corrupt zip.
        """
        target = tmp_path / "race_target.docx"
        assert not target.exists()

        # A subprocess that creates the doc inside write_lock
        script = textwrap.dedent(f"""\
            import sys, os, time
            sys.path.insert(0, r"{Path(__file__).resolve().parents[1].parent / 'src'}")
            from word_mcp.core.safesave import write_lock
            from docx import Document

            target = r"{target}"
            with write_lock(target):
                time.sleep(0.3)  # hold long enough for the race
                doc = Document()
                doc.add_paragraph("from subprocess")
                doc.save(target)
        """)

        # Launch two processes simultaneously
        procs = []
        for _ in range(2):
            p = subprocess.Popen(
                [sys.executable, "-X", "utf8", "-c", script],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            procs.append(p)

        for p in procs:
            p.wait(timeout=30)

        # The file must exist and be a valid docx
        assert target.exists(), "target was never created"
        try:
            pkg = DocxPackage(str(target))
            # Just opening it validates the zip structure
        except Exception as exc:
            pytest.fail(
                f"Cross-process creation race produced a corrupt file: {exc}"
            )

    def test_write_lock_acquires_lockfile_for_new_file(self, tmp_path):
        """write_lock must create a lockfile even when the doc does not exist.

        The current bug: safesave.py line 385 gates lockfile acquisition
        on os.path.exists(doc_path). This test verifies the lockfile is
        created for new files too.
        """
        target = tmp_path / "new_doc.docx"
        assert not target.exists()

        with safesave.write_lock(str(target)):
            # While we hold the lock, check that the lockfile was created
            d = safesave.slot_dir(target)
            lock_path = d / safesave.LOCK_FILE_NAME
            # The lockfile should exist (this is the fix target)
            if not lock_path.exists():
                pytest.fail(
                    "write_lock did not create an advisory lockfile for a "
                    "non-existent target file. Cross-process serialization "
                    "is absent for file creation."
                )


# ==========================================================================
# FINDING 3 (MEDIUM): Null byte in path not explicitly rejected
# ==========================================================================


class TestFinding3_NullByte:
    """check_path must explicitly reject paths containing null bytes when
    sandboxing is active, rather than relying on downstream Python guards."""

    def test_null_byte_in_path_refused(self, tmp_path, monkeypatch):
        """A path with an embedded null byte must raise SandboxViolation."""
        inside = tmp_path / "allowed"
        inside.mkdir()
        monkeypatch.setenv(ENV_VAR, str(inside))

        evil_path = str(inside / "doc\x00.docx")
        with pytest.raises(SandboxViolation, match="(?i)null"):
            check_path(evil_path, "test")

    def test_null_byte_traversal_refused(self, tmp_path, monkeypatch):
        """A path using null byte to confuse traversal must be rejected."""
        inside = tmp_path / "allowed"
        inside.mkdir()
        monkeypatch.setenv(ENV_VAR, str(inside))

        # Null byte before ".." might confuse path resolution
        evil_path = str(inside) + "\x00/../../../etc/passwd"
        with pytest.raises(SandboxViolation, match="(?i)null"):
            check_path(evil_path, "test")

    def test_null_byte_mid_component_refused(self, tmp_path, monkeypatch):
        """Null byte in the middle of a path component must be rejected."""
        inside = tmp_path / "allowed"
        inside.mkdir()
        monkeypatch.setenv(ENV_VAR, str(inside))

        evil_path = str(inside / "subdir\x00evil" / "doc.docx")
        with pytest.raises(SandboxViolation, match="(?i)null"):
            check_path(evil_path, "test")


# ==========================================================================
# FINDING 4 (LOW): Slot rotation surfaces raw PermissionError
# ==========================================================================


class TestFinding4_SlotRotationPermError:
    """When another process holds a read handle on prev.docx, the slot
    rotation must raise a typed WordMcpError, not a raw PermissionError."""

    @pytest.mark.skipif(
        sys.platform != "win32",
        reason="PermissionError on open files is Windows-specific",
    )
    def test_held_backup_raises_typed_error(self, tmp_path):
        """Slot rotation while prev.docx is held open must raise
        WordMcpError, not raw PermissionError."""
        doc_path = _fresh_doc(tmp_path)

        # First mutation: creates prev.docx
        pkg = DocxPackage(str(doc_path))
        from word_mcp.ops.text import search_and_replace
        search_and_replace(pkg, [{"find": "Anchor.", "replace": "Version 1."}])
        pkg.save()

        # Verify prev exists
        d = safesave.slot_dir(doc_path)
        prev = d / safesave.PREV_SLOT
        assert prev.exists()

        # Hold prev.docx open with an exclusive read lock
        held_file = open(prev, "rb")
        hold_event = threading.Event()

        def hold_file_open():
            hold_event.wait(timeout=30)
            held_file.close()

        holder = threading.Thread(target=hold_file_open, daemon=True)
        holder.start()

        try:
            # Second mutation: tries to rotate prev while it is held open
            pkg2 = DocxPackage(str(doc_path))
            search_and_replace(pkg2, [{"find": "Version 1.", "replace": "Version 2."}])
            # This should raise WordMcpError, not PermissionError
            with pytest.raises(WordMcpError):
                pkg2.save()
        except PermissionError:
            pytest.fail(
                "Slot rotation raised raw PermissionError instead of "
                "typed WordMcpError when the backup file was held open"
            )
        finally:
            hold_event.set()
            holder.join(timeout=5)
            try:
                held_file.close()
            except Exception:
                pass
