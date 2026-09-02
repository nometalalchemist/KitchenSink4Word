"""v2 Phase 3: the document-view/batch-edit layer (V2_DESIGN Section 9).

Covers: view projection shapes at every detail level, anchor stability
across unrelated edits, the stamp_anchors opt-in (and that plain reads
never mutate: file hash asserted unchanged), the volatile-anchor fallback,
apply_edits happy paths for all 8 op types, whole-batch stale refusal
(hash-verified nothing mutated), changed-map chaining, markdown mapping,
and the anchor selector through resolve_location including STALE_ANCHOR.

All file-mode fixtures are synthetic (python-docx built). The live test
spawns its own Word instance (marked live, auto-skipped without Word).
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

import pytest
from docx import Document

import word_mcp.server as srv
from word_mcp.core.errors import (
    StaleAnchor,
    TargetNotFound,
    UnsupportedStructure,
    WordMcpError,
)
from word_mcp.core.locate import resolve_location
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import batch as bt
from word_mcp.ops import read as rd
from word_mcp.ops import view as vw


# ------------------------------------------------------------------ fixtures


PARAS_AFTER_METHODS = "Cell data follows in the table."


def _build(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Introduction", 1)                      # 0
    doc.add_paragraph("The alliance framework matters.")    # 1
    doc.add_paragraph("Second paragraph with alliance terms.")  # 2
    doc.add_heading("Methods", 1)                           # 3
    doc.add_paragraph(PARAS_AFTER_METHODS)                  # 4
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "a"
    t.cell(0, 1).text = "b"
    t.cell(1, 0).text = "c"
    t.cell(1, 1).text = "d"
    doc.add_paragraph("Closing paragraph.")                 # 5
    doc.add_paragraph("Format me with centered text.")      # 6
    doc.add_paragraph("Style me as a heading.")             # 7
    doc.add_paragraph("Format the word target here.")       # 8
    doc.save(str(path))
    return path


@pytest.fixture()
def plain_doc(tmp_path) -> Path:
    """Unstamped: python-docx emits no w14:paraId, so anchors are volatile."""
    return _build(tmp_path / "plain.docx")


@pytest.fixture()
def stamped_doc(tmp_path) -> Path:
    path = _build(tmp_path / "stamped.docx")
    out = srv.get_document_view(str(path), stamp_anchors=True)
    assert out["stamped"] > 0
    return path


def _md5(path: Path) -> str:
    return hashlib.md5(Path(path).read_bytes()).hexdigest()


_BLOCK_RE = re.compile(r"^\[([0-9a-f]{4,40})\] (#*\s*)?(.*)$")


def _anchors_by_text(view_text: str) -> dict[str, str]:
    """paragraph text -> anchor, from a rendered view."""
    out = {}
    for line in view_text.split("\n"):
        m = _BLOCK_RE.match(line)
        if m:
            out[m.group(3)] = m.group(1)
    return out


def _table_anchor(view_text: str) -> str:
    m = re.search(r"\[t:([0-9a-f]{4,40})\]", view_text)
    assert m, "no table anchor in view"
    return m.group(1)


# ------------------------------------------------------------ view: shapes


def test_view_volatile_projection_shape(plain_doc):
    out = vw.get_document_view(DocxPackage(plain_doc))
    assert out["anchor_mode"] == "volatile"
    assert out["volatile_anchors"] == out["blocks"] == 10  # 9 paras + 1 table
    text = out["view"]
    # header: anchor legend, volatile caution, outline map
    assert "VOLATILE" in text
    assert "Outline: 1 Introduction" in text
    # headings project with # prefixes, prose blocks with anchors
    anchors = _anchors_by_text(text)
    assert "Introduction" in anchors
    assert "The alliance framework matters." in anchors
    # tables are pipe tables under a t: anchor
    assert re.search(r"\[t:[0-9a-f]+\] table 2x2", text)
    assert "| a | b |" in text


def test_view_detail_structure(plain_doc):
    out = vw.get_document_view(DocxPackage(plain_doc), detail="structure")
    lines = [ln for ln in out["view"].split("\n") if _BLOCK_RE.match(ln)]
    assert len(lines) == 2  # only the two headings
    assert "# Introduction (2 paras)" in lines[0]
    assert out["blocks"] == 2


def test_view_scope_outline_and_paragraphs(plain_doc):
    pkg = DocxPackage(plain_doc)
    intro = vw.get_document_view(pkg, scope={"outline": "1"})
    text = intro["view"]
    assert "The alliance framework matters." in text
    assert "Methods" not in text  # scoped views carry no outline header
    sl = vw.get_document_view(pkg, scope={"paragraphs": {"start": 1, "end": 3}})
    assert sl["blocks"] == 2
    assert "Introduction" not in sl["view"]
    with pytest.raises(TargetNotFound):
        vw.get_document_view(pkg, scope={"outline": "9.9"})
    with pytest.raises(WordMcpError):
        vw.get_document_view(pkg, scope={"bogus": 1})


def test_view_include_tables_off_and_wide_stub(tmp_path):
    path = _build(tmp_path / "wide.docx")
    out = vw.get_document_view(DocxPackage(path), include={"tables": False})
    assert not re.search(r"\[t:[0-9a-f]", out["view"])  # legend aside, no tables
    doc = Document(str(path))
    wide = doc.add_table(rows=1, cols=14)
    wide.cell(0, 0).text = "w0"
    doc.save(str(path))
    out2 = vw.get_document_view(DocxPackage(path))
    assert "stub only" in out2["view"]
    assert "get_table" in out2["view"]


def test_view_detail_full_revision_markers(tmp_path):
    path = _build(tmp_path / "full.docx")
    srv.search_and_replace(
        str(path), [{"find": "alliance framework", "replace": "coalition"}],
        track=True, live="off",
    )
    out = vw.get_document_view(DocxPackage(path), detail="full")
    assert "{++coalition++}" in out["view"]
    assert "{--alliance framework--}" in out["view"]
    assert "Markers:" in out["view"]
    # default detail stays clean of markers
    plain = vw.get_document_view(DocxPackage(path))
    assert "{++" not in plain["view"]


# --------------------------------------------- reads pure, stamping opt-in


def test_plain_view_never_mutates(plain_doc):
    before = _md5(plain_doc)
    srv.get_document_view(str(plain_doc))
    srv.get_document_view(str(plain_doc), detail="full")
    srv.get_document_view(str(plain_doc), detail="structure")
    assert _md5(plain_doc) == before


def test_stamp_anchors_is_explicit_and_durable(plain_doc):
    before = _md5(plain_doc)
    out = srv.get_document_view(str(plain_doc), stamp_anchors=True)
    assert out["stamped"] > 0
    assert out["anchor_mode"] == "paraId"
    assert "saved" in out  # went through the normal backup/save cycle
    assert _md5(plain_doc) != before
    # second stamped call: nothing left to stamp, pure read, no mutation
    mid = _md5(plain_doc)
    again = srv.get_document_view(str(plain_doc), stamp_anchors=True)
    assert "stamped" not in again
    assert _md5(plain_doc) == mid


def test_anchor_stability_across_unrelated_edits(stamped_doc):
    v1 = srv.get_document_view(str(stamped_doc))
    a1 = _anchors_by_text(v1["view"])
    target = a1["The alliance framework matters."]
    srv.set_paragraph_text(
        str(stamped_doc), {"search": {"text": "Closing paragraph."}},
        "Closing paragraph, edited.", live="off",
    )
    v2 = srv.get_document_view(str(stamped_doc))
    a2 = _anchors_by_text(v2["view"])
    assert a2["The alliance framework matters."] == target
    assert a2["Introduction"] == a1["Introduction"]
    assert _table_anchor(v2["view"]) == _table_anchor(v1["view"])


def test_volatile_anchors_resolve_but_shift(plain_doc):
    v1 = srv.get_document_view(str(plain_doc))
    anchor = _anchors_by_text(v1["view"])["Closing paragraph."]
    pkg = DocxPackage(plain_doc)
    r = resolve_location(pkg, {"anchor": anchor})
    assert r.paragraph_index == 5
    assert r.matched["volatile"] is True


# ------------------------------------------------------- apply_edits: happy


def test_apply_edits_all_eight_ops(stamped_doc):
    v = srv.get_document_view(str(stamped_doc))
    a = _anchors_by_text(v["view"])
    t = _table_anchor(v["view"])
    md = (
        "### Data Sources\n"
        "One plain paragraph.\n"
        "- first item\n"
        "- second item\n"
        "\n"
        "| X | Y |\n"
        "|---|---|\n"
        "| 1 | 2 |\n"
    )
    edits = [
        {"op": "insert", "location": {"anchor": a["Methods"]},
         "markdown": md},
        {"op": "replace", "anchor": a["The alliance framework matters."],
         "find": "alliance", "text": "coalition"},
        {"op": "set_text", "anchor": a["Second paragraph with alliance terms."],
         "text": "Entirely new second paragraph."},
        {"op": "set_style", "anchor": a["Style me as a heading."],
         "style": "Heading 2"},
        {"op": "format", "anchor": a["Format the word target here."],
         "find": "target", "formatting": {"bold": True}},
        {"op": "set_paragraph_format", "anchor": a["Format me with centered text."],
         "format": {"alignment": "center"}},
        {"op": "set_cell", "anchor": f"t:{t}:r1c1", "text": "Header!"},
        {"op": "delete", "anchor": a["Closing paragraph."]},
    ]
    out = srv.apply_edits(str(stamped_doc), edits, live="off")
    assert out["applied"] == 8
    assert set(out["changed"]) == {str(i) for i in range(8)}
    assert out["changed"]["0"]["inserted_paragraphs"] >= 4
    assert out["changed"]["0"]["inserted_tables"] == 1
    assert out["changed"]["0"]["anchors"]  # durable anchors for chaining
    assert out["changed"]["0"]["table_anchors"]
    assert out["changed"]["1"]["replaced"] == 1
    assert out["changed"]["7"]["deleted"] == 1
    assert "saved" in out

    texts = [p["text"] for p in srv.get_text(str(stamped_doc), live="off")]
    assert "The coalition framework matters." in texts
    assert "Entirely new second paragraph." in texts
    assert "Closing paragraph." not in texts
    assert "Data Sources" in texts
    # the markdown table landed after Methods, BEFORE the original table,
    # so it is body table 0 and the original shifted to 1; the set_cell op
    # still hit the original (element identity, not index)
    tbl_new = srv.get_table(str(stamped_doc), 0)
    assert tbl_new["cells"][0][0]["text"] == "X"
    assert tbl_new["cells"][1][1]["text"] == "2"
    tbl_orig = srv.get_table(str(stamped_doc), 1)
    assert tbl_orig["cells"][0][0]["text"] == "Header!"
    assert out["changed"]["6"]["table"] == 1  # current index reported
    # heading landed in the outline; list paragraphs carry real numbering
    outline = [h["text"] for h in srv.get_outline(str(stamped_doc), live="off")]
    assert "Data Sources" in outline
    pkg = DocxPackage(stamped_doc)
    numbered = [
        el for k, i, el in rd.body_items(pkg)
        if k == "paragraph" and el.find(
            f"{qn('w:pPr')}/{qn('w:numPr')}"
        ) is not None
    ]
    assert len(numbered) == 2
    # styled paragraph
    styled = [p for p in srv.get_text(str(stamped_doc), live="off")
              if p["text"] == "Style me as a heading."]
    assert styled and styled[0]["style"].startswith("Heading")


def test_changed_map_chaining(stamped_doc):
    out1 = srv.apply_edits(
        str(stamped_doc),
        [{"op": "insert", "location": {"search": {"text": "Closing paragraph."}},
          "markdown": "Chained paragraph, first form."}],
        live="off",
    )
    new_anchor = out1["changed"]["0"]["anchors"][0]
    out2 = srv.apply_edits(
        str(stamped_doc),
        [{"op": "set_text", "anchor": new_anchor,
          "text": "Chained paragraph, rewritten without re-viewing."}],
        live="off",
    )
    assert out2["applied"] == 1
    texts = [p["text"] for p in srv.get_text(str(stamped_doc), live="off")]
    assert "Chained paragraph, rewritten without re-viewing." in texts


def test_replace_all_occurrences_and_occurrence_pick(stamped_doc):
    v = srv.get_document_view(str(stamped_doc))
    a = _anchors_by_text(v["view"])
    anchor = a["Second paragraph with alliance terms."]
    srv.apply_edits(
        str(stamped_doc),
        [{"op": "set_text", "anchor": anchor, "text": "one two one two one"}],
        live="off",
    )
    out = srv.apply_edits(
        str(stamped_doc),
        [{"op": "replace", "anchor": anchor, "find": "one", "text": "ONE",
          "occurrence": 2}],
        live="off",
    )
    assert out["changed"]["0"]["replaced"] == 1
    texts = [p["text"] for p in srv.get_text(str(stamped_doc), live="off")]
    assert "one two ONE two one" in texts
    out2 = srv.apply_edits(
        str(stamped_doc),
        [{"op": "replace", "anchor": anchor, "find": "one", "text": "1"}],
        live="off",
    )
    assert out2["changed"]["0"]["replaced"] == 2


# ---------------------------------------------------- apply_edits: refusals


def test_whole_batch_stale_refusal_nothing_mutated(stamped_doc):
    before = _md5(stamped_doc)
    v = srv.get_document_view(str(stamped_doc))
    a = _anchors_by_text(v["view"])
    with pytest.raises(WordMcpError) as ei:
        srv.apply_edits(
            str(stamped_doc),
            [
                {"op": "set_text", "anchor": a["Closing paragraph."],
                 "text": "would be valid"},
                {"op": "set_text", "anchor": "abcdef123456",
                 "text": "stale target"},
            ],
            live="off",
        )
    exc = ei.value
    assert getattr(exc, "code", None) == "STALE_ANCHOR"
    assert "nothing was applied" in str(exc)
    assert exc.detail["failures"][0]["index"] == 1
    assert _md5(stamped_doc) == before  # untouched, hash-verified


def test_unsupported_markdown_refuses_whole_batch(stamped_doc):
    before = _md5(stamped_doc)
    with pytest.raises(WordMcpError) as ei:
        srv.apply_edits(
            str(stamped_doc),
            [{"op": "insert", "location": {"paragraph": 0},
              "markdown": "![figure](chart.png)"}],
            live="off",
        )
    assert getattr(ei.value, "code", None) == "UNSUPPORTED_CONTENT"
    assert "insert_image" in str(ei.value)
    assert _md5(stamped_doc) == before


def test_bad_op_and_param_refusals(stamped_doc):
    before = _md5(stamped_doc)
    v = srv.get_document_view(str(stamped_doc))
    a = _anchors_by_text(v["view"])
    with pytest.raises(WordMcpError, match="not an apply_edits op"):
        srv.apply_edits(
            str(stamped_doc), [{"op": "explode", "anchor": "aaaa"}],
            live="off",
        )
    with pytest.raises(WordMcpError, match="missing required"):
        srv.apply_edits(
            str(stamped_doc),
            [{"op": "replace", "anchor": a["Introduction"]}], live="off",
        )
    with pytest.raises(WordMcpError, match="CELL anchor"):
        srv.apply_edits(
            str(stamped_doc),
            [{"op": "set_cell", "anchor": a["Introduction"], "text": "x"}],
            live="off",
        )
    with pytest.raises(WordMcpError, match="atomic"):
        srv.apply_edits(
            str(stamped_doc),
            [{"op": "delete", "anchor": a["Introduction"]}],
            atomic=False, live="off",
        )
    assert _md5(stamped_doc) == before


def test_paragraph_op_refuses_table_anchor(stamped_doc):
    v = srv.get_document_view(str(stamped_doc))
    t = _table_anchor(v["view"])
    with pytest.raises(WordMcpError) as ei:
        srv.apply_edits(
            str(stamped_doc),
            [{"op": "set_text", "anchor": f"t:{t}", "text": "nope"}],
            live="off",
        )
    assert "set_cell" in str(ei.value)


# ------------------------------------------------------------ markdown unit


def test_parse_markdown_segments():
    segs = bt.parse_markdown(
        "# Title\npara one\npara two\n\n1. first\n2. second\n- bullet\n"
        "| A | B |\n| 1 | 2 |\n"
    )
    kinds = [s["kind"] for s in segs]
    assert kinds == ["paragraphs", "paragraphs", "list", "list", "table"]
    assert segs[0]["items"][0] == {"text": "Title", "level": 1}
    assert segs[2]["list_kind"] == "number"
    assert segs[3]["list_kind"] == "bullet"
    assert segs[4]["data"] == [["A", "B"], ["1", "2"]]
    with pytest.raises(UnsupportedStructure):
        bt.parse_markdown("```python\nprint('hi')\n```")
    with pytest.raises(UnsupportedStructure):
        bt.parse_markdown("> a quote")
    with pytest.raises(UnsupportedStructure):
        bt.parse_markdown("---")
    with pytest.raises(WordMcpError):
        bt.parse_markdown("   ")
    esc = bt.parse_markdown("| a \\| b | c |")
    assert esc[0]["data"] == [["a | b", "c"]]


# ------------------------------------------- anchor selector via locate


def test_anchor_selector_resolves_through_location(stamped_doc):
    pkg = DocxPackage(stamped_doc)
    view = vw.get_document_view(pkg)
    a = _anchors_by_text(view["view"])
    r = resolve_location(pkg, {"anchor": a["Methods"], "position": "before"})
    assert r.selector == "anchor"
    assert r.paragraph_index == 3
    assert r.position == "before"
    assert "volatile" not in r.matched


def test_anchor_selector_stale_and_table_refusals(stamped_doc):
    pkg = DocxPackage(stamped_doc)
    with pytest.raises(StaleAnchor) as ei:
        resolve_location(pkg, {"anchor": "abcdef123456"})
    assert "get_document_view" in str(ei.value)
    t = _table_anchor(vw.get_document_view(pkg)["view"])
    with pytest.raises(UnsupportedStructure):
        resolve_location(pkg, {"anchor": f"t:{t}"})
    with pytest.raises(UnsupportedStructure):
        resolve_location(pkg, {"anchor": f"t:{t}:r1c1"})


def test_anchor_rides_positional_tools(stamped_doc):
    """Section 9.4 binding 1: anchors work in every positional tool."""
    v = srv.get_document_view(str(stamped_doc))
    a = _anchors_by_text(v["view"])
    srv.set_paragraph_text(
        str(stamped_doc), {"anchor": a["Closing paragraph."]},
        "Rewritten through an anchor location.", live="off",
    )
    texts = [p["text"] for p in srv.get_text(str(stamped_doc), live="off")]
    assert "Rewritten through an anchor location." in texts


def test_cell_anchor_bounds_refuse(stamped_doc):
    v = srv.get_document_view(str(stamped_doc))
    t = _table_anchor(v["view"])
    pkg = DocxPackage(stamped_doc)
    with pytest.raises(TargetNotFound):
        vw.resolve_anchor(pkg, f"t:{t}:r9c1")
    with pytest.raises(TargetNotFound):
        vw.resolve_anchor(pkg, f"t:{t}:r1c9")


# --------------------------------------------------------------- round trip


def test_round_trip_view_edit_review(stamped_doc):
    """Gate 2: view -> apply_edits -> re-view keeps untouched anchors
    stable and issues coherent anchors for inserted paragraphs."""
    v1 = srv.get_document_view(str(stamped_doc))
    a1 = _anchors_by_text(v1["view"])
    out = srv.apply_edits(
        str(stamped_doc),
        [
            {"op": "replace", "anchor": a1["The alliance framework matters."],
             "find": "matters", "text": "endures"},
            {"op": "insert", "location": {"anchor": a1["Methods"]},
             "markdown": "Fresh paragraph after Methods."},
        ],
        live="off",
    )
    inserted = out["changed"]["1"]["anchors"][0]
    v2 = srv.get_document_view(str(stamped_doc))
    a2 = _anchors_by_text(v2["view"])
    assert v2["anchor_mode"] == "paraId"
    # untouched paragraphs keep their anchors
    for text in ("Introduction", "Methods", "Closing paragraph.",
                 "Second paragraph with alliance terms."):
        assert a2[text] == a1[text], text
    # the edited paragraph keeps its anchor too (paraId survives run edits)
    assert a2["The alliance framework endures."] == \
        a1["The alliance framework matters."]
    # the inserted paragraph's reported anchor is the one the view shows
    assert a2["Fresh paragraph after Methods."] == inserted


# --------------------------------------------------------------------- live


def _word_available() -> bool:
    from test_live_core import _word_available as f

    return f()


@pytest.mark.live
def test_apply_edits_live_one_undo_group(tmp_path_factory):
    """The live route: whole batch in ONE run_live session (one undo
    group), fixture-spawned Word only, PID-precise cleanup via
    quit_instance_holding."""
    if not _word_available():
        pytest.skip("Word not available")
    import pythoncom
    import win32com.client

    path = tmp_path_factory.mktemp("live_batch") / "batch.docx"
    _build(path)
    srv.get_document_view(str(path), stamp_anchors=True)
    v = srv.get_document_view(str(path))
    a = _anchors_by_text(v["view"])

    pythoncom.CoInitialize()
    app = win32com.client.DispatchEx("Word.Application")
    app.Visible = True
    app.Documents.Open(str(path))
    try:
        result = srv.apply_edits(
            str(path),
            [
                {"op": "replace",
                 "anchor": a["The alliance framework matters."],
                 "find": "alliance", "text": "coalition"},
                {"op": "set_text", "anchor": a["Closing paragraph."],
                 "text": "Closing paragraph, live-edited."},
                {"op": "insert", "location": {"anchor": a["Methods"]},
                 "markdown": "### Live Sub\nLive-inserted paragraph."},
            ],
        )
        assert result["applied"] == 3
        assert result["live"] is True
        assert result["undo_grouped"] is True  # ONE Ctrl+Z step
        live_read = srv.get_text(str(path))
        texts = [p["text"] for p in live_read]
        assert "The coalition framework matters." in texts
        assert "Closing paragraph, live-edited." in texts
        assert "Live-inserted paragraph." in texts
    finally:
        app = None
        from test_live_core import quit_instance_holding

        try:
            quit_instance_holding(str(path))
        finally:
            pythoncom.CoUninitialize()
