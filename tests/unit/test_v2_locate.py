"""v2 Phase 1: the location-object resolver (core/locate.py).

Covers V2_DESIGN Section 6: every selector's happy path, every refusal
path, the merge-test regression (heading text recurring in body prose must
refuse with the match list, never resolve first-match), occurrence
disambiguation, match modes, position modifiers, range resolution
including inversion, the fresh-document index-0 rule (WS0), entity and
curly-quote hints (Bug 11 / L6), one-selector-only enforcement, and how
ambiguity refusals surface through envelope.py.

All fixtures are synthetic (python-docx built); nothing here needs the
real-document set or a live Word.
"""

from __future__ import annotations

import pytest
from docx import Document
from lxml import etree

from word_mcp import envelope
from word_mcp.core.errors import (
    AmbiguousTarget,
    RangeOutOfBounds,
    StaleAnchor,
    TargetNotFound,
    UnsupportedStructure,
    WordMcpError,
    WordNotRunning,
)
from word_mcp.core.locate import (
    POSITIONS,
    SELECTORS,
    ResolvedLocation,
    ResolvedRange,
    is_range_spec,
    resolve_location,
    resolve_range,
)
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import fields as ff
from word_mcp.ops import read as rd


# ------------------------------------------------------------------ fixtures


def _pkg(tmp_path, name, build) -> DocxPackage:
    path = tmp_path / name
    doc = Document()
    build(doc)
    doc.save(str(path))
    return DocxPackage(path)


def _set_outline_lvl(pkg: DocxPackage, para_index: int, level0: int) -> None:
    """Give a paragraph a direct w:outlineLvl (the L8/NSU template pattern:
    a heading with no Heading style)."""
    paras = [el for k, i, el in rd.body_items(pkg) if k == "paragraph"]
    p = paras[para_index]
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.Element(qn("w:pPr"))
        p.insert(0, ppr)
    etree.SubElement(ppr, qn("w:outlineLvl")).set(qn("w:val"), str(level0))


@pytest.fixture
def doc(tmp_path) -> DocxPackage:
    """Structured document: two chapters, sub-headings, prose that reuses
    heading text (the merge-test shape), repeated prose, quotes, entities.

    Body paragraph map:
      0  H1 "Introduction"
      1  prose "Opening prose with a target word."
      2  H1 "Discussion"
      3  prose "First discussion paragraph mentions alpha & beta."
      4  H2 "Methods"
      5  prose "as noted in the Discussion section above, results vary."
      6  H2 "Results"
      7  prose "The word target appears here too."
      8  prose "Word autocorrects to ‘curly’ quotes."
      9  H1 "Conclusion"
    """

    def build(d):
        d.add_heading("Introduction", level=1)
        d.add_paragraph("Opening prose with a target word.")
        d.add_heading("Discussion", level=1)
        d.add_paragraph("First discussion paragraph mentions alpha & beta.")
        d.add_heading("Methods", level=2)
        d.add_paragraph("as noted in the Discussion section above, results vary.")
        d.add_heading("Results", level=2)
        d.add_paragraph("The word target appears here too.")
        d.add_paragraph("Word autocorrects to ‘curly’ quotes.")
        d.add_heading("Conclusion", level=1)

    return _pkg(tmp_path, "structured.docx", build)


@pytest.fixture
def fresh(tmp_path) -> DocxPackage:
    """A document exactly as create_document produces it: the body holds no
    w:p at all, only the trailing sectPr (the WS0 precondition)."""
    path = tmp_path / "fresh.docx"
    Document().save(str(path))
    pkg = DocxPackage(path)
    assert pkg.body().find(qn("w:p")) is None
    return pkg


# ------------------------------------------------------- paragraph selector


def test_paragraph_happy_path(doc):
    loc = resolve_location(doc, {"paragraph": 3})
    assert isinstance(loc, ResolvedLocation)
    assert loc.paragraph_index == 3
    assert loc.selector == "paragraph"
    assert loc.position == "after"  # the default
    assert "alpha & beta" in loc.matched["text"]
    assert loc.char_start is None and loc.char_end is None


def test_paragraph_position_modifiers(doc):
    for pos in POSITIONS:
        loc = resolve_location(doc, {"paragraph": 1, "position": pos})
        assert loc.position == pos


def test_paragraph_bad_position(doc):
    with pytest.raises(WordMcpError, match="position"):
        resolve_location(doc, {"paragraph": 1, "position": "above"})


def test_paragraph_out_of_range(doc):
    with pytest.raises(TargetNotFound, match="out of range.*document has 10"):
        resolve_location(doc, {"paragraph": 10})


def test_paragraph_negative_refuses(doc):
    with pytest.raises(WordMcpError, match=">= 0"):
        resolve_location(doc, {"paragraph": -1})


def test_paragraph_non_int_refuses(doc):
    with pytest.raises(WordMcpError, match="integer"):
        resolve_location(doc, {"paragraph": "3"})
    with pytest.raises(WordMcpError, match="integer"):
        resolve_location(doc, {"paragraph": True})


def test_fresh_document_index_0(fresh):
    """WS0: index 0 addresses the displayed default empty paragraph, and
    resolving it never mutates the package."""
    loc = resolve_location(fresh, {"paragraph": 0})
    assert loc.paragraph_index == 0
    assert loc.matched["implicit"] is True
    assert fresh.body().find(qn("w:p")) is None  # no materialization on read


def test_fresh_document_index_1_refuses(fresh):
    with pytest.raises(TargetNotFound, match="only index 0"):
        resolve_location(fresh, {"paragraph": 1})


# --------------------------------------------------- after_heading selector


def test_after_heading_unique_exact(doc):
    loc = resolve_location(doc, {"after_heading": {"text": "Introduction"}})
    assert loc.paragraph_index == 0
    assert loc.matched["level"] == 1
    assert loc.matched["outline"] == "1"
    assert loc.matched["detected_via"] == "heading_style"


def test_after_heading_outline_lvl_heading(tmp_path):
    """L8/NSU: a Normal-styled paragraph with a direct outlineLvl IS a
    heading to the resolver."""

    def build(d):
        d.add_paragraph("Styled Like Body But A Heading")
        d.add_paragraph("Ordinary prose.")

    pkg = _pkg(tmp_path, "outlinelvl.docx", build)
    _set_outline_lvl(pkg, 0, 0)
    loc = resolve_location(
        pkg, {"after_heading": {"text": "Styled Like Body But A Heading"}}
    )
    assert loc.paragraph_index == 0
    assert loc.matched["detected_via"] == "outline_level"
    assert loc.matched["level"] == 1


def test_merge_test_regression_prose_recurrence_refuses(doc):
    """THE merge-test regression: 'Discussion' is a heading (paragraph 2)
    AND recurs in body prose (paragraph 5). Without occurrence this MUST
    refuse with the full match list, never resolve first-match."""
    with pytest.raises(AmbiguousTarget) as ei:
        resolve_location(doc, {"after_heading": {"text": "Discussion"}})
    matches = ei.value.matches
    assert [m["paragraph"] for m in matches] == [2, 5]
    heading_match = matches[0]
    assert heading_match["outline"] == "2"
    prose_match = matches[1]
    assert "outline" not in prose_match
    assert "Discussion section" in prose_match["context"]
    assert "occurrence" in str(ei.value)


def test_after_heading_occurrence_resolves_past_prose(doc):
    """occurrence counts matching HEADINGS in document order; explicit
    disambiguation resolves even though prose recurrence exists."""
    loc = resolve_location(
        doc, {"after_heading": {"text": "Discussion", "occurrence": 1}}
    )
    assert loc.paragraph_index == 2
    assert loc.matched["outline"] == "2"


def test_after_heading_two_headings_same_text(tmp_path):
    def build(d):
        d.add_heading("Summary", level=1)
        d.add_paragraph("prose one")
        d.add_heading("Summary", level=1)
        d.add_paragraph("prose two")

    pkg = _pkg(tmp_path, "twosummaries.docx", build)
    with pytest.raises(AmbiguousTarget) as ei:
        resolve_location(pkg, {"after_heading": {"text": "Summary"}})
    assert [m["paragraph"] for m in ei.value.matches] == [0, 2]
    loc = resolve_location(
        pkg, {"after_heading": {"text": "Summary", "occurrence": 2}}
    )
    assert loc.paragraph_index == 2
    assert loc.matched["outline"] == "2"


def test_after_heading_occurrence_out_of_range(doc):
    with pytest.raises(TargetNotFound, match="occurrence 3 out of range"):
        resolve_location(
            doc, {"after_heading": {"text": "Discussion", "occurrence": 3}}
        )


def test_after_heading_occurrence_zero_refuses(doc):
    with pytest.raises(WordMcpError, match="1-based"):
        resolve_location(
            doc, {"after_heading": {"text": "Discussion", "occurrence": 0}}
        )


def test_after_heading_match_contains(doc):
    loc = resolve_location(
        doc, {"after_heading": {"text": "Intro", "match": "contains"}}
    )
    assert loc.paragraph_index == 0


def test_after_heading_exact_does_not_substring(doc):
    with pytest.raises(TargetNotFound):
        resolve_location(doc, {"after_heading": {"text": "Intro"}})


def test_after_heading_bad_match_mode(doc):
    with pytest.raises(WordMcpError, match="match must be one of"):
        resolve_location(
            doc, {"after_heading": {"text": "Discussion", "match": "fuzzy"}}
        )


def test_after_heading_prose_only_not_found_names_prose(doc):
    """Text present only in prose: NOT_FOUND, and the message points at the
    prose location and the right selectors instead."""
    with pytest.raises(TargetNotFound) as ei:
        resolve_location(doc, {"after_heading": {"text": "results vary"}})
    msg = str(ei.value)
    assert "no heading matched" in msg
    assert "paragraph 5" in msg
    assert "search or paragraph" in msg


def test_after_heading_entity_hint(doc):
    """Bug 11: an XML-entity query gets the plain-text hint."""
    with pytest.raises(TargetNotFound, match=r"'&' not '&amp;'"):
        resolve_location(doc, {"after_heading": {"text": "alpha &amp; beta"}})


def test_after_heading_close_match_suggestion(doc):
    with pytest.raises(TargetNotFound, match="close heading matches"):
        resolve_location(doc, {"after_heading": {"text": "Concluson"}})


def test_after_heading_spec_shape(doc):
    with pytest.raises(WordMcpError, match="requires 'text'"):
        resolve_location(doc, {"after_heading": {"occurrence": 1}})
    with pytest.raises(WordMcpError, match="unknown key"):
        resolve_location(
            doc, {"after_heading": {"text": "Discussion", "bogus": 1}}
        )
    with pytest.raises(WordMcpError, match="takes an object"):
        resolve_location(doc, {"after_heading": "Discussion"})


# --------------------------------------------------------- outline selector


def test_outline_happy_path(doc):
    loc = resolve_location(doc, {"outline": "2.1"})
    assert loc.paragraph_index == 4
    assert loc.matched["text"] == "Methods"
    loc = resolve_location(doc, {"outline": "2.2"})
    assert loc.paragraph_index == 6
    assert loc.matched["text"] == "Results"
    loc = resolve_location(doc, {"outline": "3"})
    assert loc.paragraph_index == 9
    assert loc.matched["text"] == "Conclusion"


def test_outline_unknown_path_lists_outline(doc):
    with pytest.raises(TargetNotFound) as ei:
        resolve_location(doc, {"outline": "7.7"})
    msg = str(ei.value)
    assert "document outline" in msg
    assert "'Methods'" in msg


def test_outline_bad_shape(doc):
    with pytest.raises(WordMcpError, match="numbered path"):
        resolve_location(doc, {"outline": "3.2a"})
    with pytest.raises(WordMcpError, match="numbered path"):
        resolve_location(doc, {"outline": 3})


def test_outline_no_headings(tmp_path):
    pkg = _pkg(tmp_path, "plain.docx", lambda d: d.add_paragraph("just prose"))
    with pytest.raises(TargetNotFound, match="no headings detected"):
        resolve_location(pkg, {"outline": "1"})


# -------------------------------------------------------- bookmark selector


def test_bookmark_happy_path(doc):
    ff.add_bookmark(doc, "methods_mark", anchor_text="results vary")
    loc = resolve_location(doc, {"bookmark": "methods_mark"})
    assert loc.paragraph_index == 5
    assert loc.matched["bookmark"] == "methods_mark"


def test_bookmark_not_found_close_match(doc):
    ff.add_bookmark(doc, "methods_mark", anchor_text="results vary")
    with pytest.raises(TargetNotFound, match="methods_mark"):
        resolve_location(doc, {"bookmark": "method_mark"})


def test_bookmark_case_sensitivity_hint(doc):
    ff.add_bookmark(doc, "MethodsMark", anchor_text="results vary")
    with pytest.raises(TargetNotFound, match="case-sensitive"):
        resolve_location(doc, {"bookmark": "methodsmark"})


def test_bookmark_bad_value(doc):
    with pytest.raises(WordMcpError, match="bookmark name"):
        resolve_location(doc, {"bookmark": 7})


def test_bookmark_inside_table_refuses(tmp_path):
    def build(d):
        d.add_paragraph("before the table")
        table = d.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "cell text"

    pkg = _pkg(tmp_path, "tbl.docx", build)
    tbl = pkg.body().find(qn("w:tbl"))
    cell_p = tbl.find(f".//{qn('w:p')}")
    bs = etree.SubElement(cell_p, qn("w:bookmarkStart"))
    bs.set(qn("w:id"), "9")
    bs.set(qn("w:name"), "in_table")
    with pytest.raises(UnsupportedStructure, match="inside a table"):
        resolve_location(pkg, {"bookmark": "in_table"})


def test_bookmark_at_body_level_attaches_to_neighbor(tmp_path):
    def build(d):
        d.add_paragraph("first")
        d.add_paragraph("second")

    pkg = _pkg(tmp_path, "bodybm.docx", build)
    body = pkg.body()
    first_p = body.find(qn("w:p"))
    bs = etree.Element(qn("w:bookmarkStart"))
    bs.set(qn("w:id"), "5")
    bs.set(qn("w:name"), "between")
    first_p.addnext(bs)
    loc = resolve_location(pkg, {"bookmark": "between"})
    assert loc.paragraph_index == 1
    assert loc.matched["adjacency"] == "following"


# ---------------------------------------------------------- search selector


def test_search_unique_resolves_with_span(doc):
    loc = resolve_location(doc, {"search": {"text": "alpha & beta"}})
    assert loc.paragraph_index == 3
    assert loc.char_start is not None
    text = loc.matched["match"]
    assert text == "alpha & beta"


def test_search_ambiguous_refuses_with_matches(doc):
    with pytest.raises(AmbiguousTarget) as ei:
        resolve_location(doc, {"search": {"text": "target"}})
    matches = ei.value.matches
    assert [m["paragraph"] for m in matches] == [1, 7]
    assert all("context" in m for m in matches)


def test_search_occurrence_disambiguates(doc):
    loc = resolve_location(doc, {"search": {"text": "target", "occurrence": 2}})
    assert loc.paragraph_index == 7


def test_search_occurrence_out_of_range(doc):
    with pytest.raises(TargetNotFound, match="only 2 match"):
        resolve_location(doc, {"search": {"text": "target", "occurrence": 5}})


def test_search_match_case_default_insensitive(doc):
    loc = resolve_location(doc, {"search": {"text": "OPENING PROSE"}})
    assert loc.paragraph_index == 1


def test_search_match_case_true_respects_case(doc):
    with pytest.raises(TargetNotFound) as ei:
        resolve_location(
            doc, {"search": {"text": "OPENING PROSE", "match_case": True}}
        )
    assert "ignoring case" in str(ei.value)
    assert "paragraph 1" in str(ei.value)


def test_search_entity_hint(doc):
    """Bug 11: searching the XML entity form misses the literal '&' text
    and the refusal says exactly why."""
    with pytest.raises(TargetNotFound, match=r"'&' not '&amp;'"):
        resolve_location(doc, {"search": {"text": "alpha &amp; beta"}})


def test_search_curly_quote_hint(doc):
    """L6: straight-quote query vs curly-quote document text."""
    with pytest.raises(TargetNotFound) as ei:
        resolve_location(doc, {"search": {"text": "'curly' quotes"}})
    msg = str(ei.value)
    assert "curly and straight quotes" in msg
    assert "paragraph 8" in msg


def test_search_heading_match_carries_outline(tmp_path):
    def build(d):
        d.add_heading("Unique Heading", level=1)
        d.add_paragraph("Unique Heading mentioned again in prose.")

    pkg = _pkg(tmp_path, "srch.docx", build)
    with pytest.raises(AmbiguousTarget) as ei:
        resolve_location(pkg, {"search": {"text": "Unique Heading"}})
    matches = ei.value.matches
    assert matches[0]["paragraph"] == 0
    assert matches[0]["outline"] == "1"
    assert "outline" not in matches[1]


def test_search_spec_shape(doc):
    with pytest.raises(WordMcpError, match="non-empty"):
        resolve_location(doc, {"search": {"text": ""}})
    with pytest.raises(WordMcpError, match="unknown key"):
        resolve_location(doc, {"search": {"text": "x", "regex": True}})
    with pytest.raises(WordMcpError, match="must be a bool"):
        resolve_location(doc, {"search": {"text": "x", "match_case": "yes"}})


# ----------------------------------------------- anchor and cursor selectors


def test_anchor_unresolvable_is_stale_with_review_hint(doc):
    """An anchor id that matches nothing raises StaleAnchor (closed code
    STALE_ANCHOR) and points at get_document_view (Phase 3 layer)."""
    from word_mcp.core.errors import StaleAnchor

    with pytest.raises(StaleAnchor) as ei:
        resolve_location(doc, {"anchor": "abcdef123456"})
    assert "get_document_view" in str(ei.value)


def test_anchor_bad_shape(doc):
    with pytest.raises(WordMcpError, match="anchor id"):
        resolve_location(doc, {"anchor": 42})
    with pytest.raises(WordMcpError, match="malformed anchor"):
        resolve_location(doc, {"anchor": "not-hex!"})


def test_cursor_without_reader_refuses(doc):
    with pytest.raises(WordNotRunning, match="live Word session"):
        resolve_location(doc, {"cursor": True})


def test_cursor_with_reader_resolves(doc):
    loc = resolve_location(doc, {"cursor": True}, cursor_reader=lambda: 5)
    assert loc.paragraph_index == 5
    assert loc.selector == "cursor"
    assert "results vary" in loc.matched["text"]


def test_cursor_reader_out_of_sync(doc):
    with pytest.raises(TargetNotFound, match="out of sync"):
        resolve_location(doc, {"cursor": True}, cursor_reader=lambda: 99)


def test_cursor_value_must_be_true(doc):
    with pytest.raises(WordMcpError, match="cursor"):
        resolve_location(doc, {"cursor": 1}, cursor_reader=lambda: 0)


# --------------------------------------------- one-selector-only enforcement


def test_zero_selectors_refuses(doc):
    with pytest.raises(WordMcpError, match="exactly one selector"):
        resolve_location(doc, {})
    with pytest.raises(WordMcpError, match="exactly one selector"):
        resolve_location(doc, {"position": "after"})


def test_two_selectors_refuses(doc):
    with pytest.raises(WordMcpError, match="exactly one selector"):
        resolve_location(doc, {"paragraph": 1, "outline": "1"})


def test_unknown_key_refuses(doc):
    with pytest.raises(WordMcpError, match="unknown location key"):
        resolve_location(doc, {"paragraf": 1})


def test_non_dict_location_refuses(doc):
    with pytest.raises(WordMcpError, match="location must be an object"):
        resolve_location(doc, 3)


def test_range_passed_as_location_gets_pointer(doc):
    with pytest.raises(WordMcpError, match="looks like a .start, end. range"):
        resolve_location(doc, {"start": {"paragraph": 1}, "end": {"paragraph": 2}})


# ------------------------------------------------------------------- ranges


def test_range_happy_path(doc):
    rng = resolve_range(
        doc, {"start": {"paragraph": 2}, "end": {"outline": "2.2"}}
    )
    assert isinstance(rng, ResolvedRange)
    assert rng.start_index == 2
    assert rng.end_index == 6
    assert rng.start.selector == "paragraph"
    assert rng.end.selector == "outline"


def test_range_same_paragraph_ok(doc):
    rng = resolve_range(
        doc, {"start": {"paragraph": 3}, "end": {"paragraph": 3}}
    )
    assert rng.start_index == rng.end_index == 3


def test_range_inverted_refuses(doc):
    with pytest.raises(RangeOutOfBounds, match="inverted range"):
        resolve_range(doc, {"start": {"paragraph": 6}, "end": {"paragraph": 2}})


def test_range_inverted_char_spans_refuse(tmp_path):
    def build(d):
        d.add_paragraph("first-token middle words second-token tail")

    pkg = _pkg(tmp_path, "span.docx", build)
    with pytest.raises(RangeOutOfBounds, match="char"):
        resolve_range(
            pkg,
            {
                "start": {"search": {"text": "second-token"}},
                "end": {"search": {"text": "first-token"}},
            },
        )


def test_range_missing_end_refuses(doc):
    with pytest.raises(WordMcpError, match="missing 'end'"):
        resolve_range(doc, {"start": {"paragraph": 1}})


def test_range_unknown_key_refuses(doc):
    with pytest.raises(WordMcpError, match="unknown range key"):
        resolve_range(
            doc,
            {"start": {"paragraph": 1}, "end": {"paragraph": 2}, "step": 1},
        )


def test_range_endpoint_refusals_propagate(doc):
    with pytest.raises(AmbiguousTarget):
        resolve_range(
            doc,
            {"start": {"search": {"text": "target"}}, "end": {"paragraph": 9}},
        )


def test_is_range_spec():
    assert is_range_spec({"start": {"paragraph": 1}, "end": {"paragraph": 2}})
    assert not is_range_spec({"paragraph": 1})
    assert not is_range_spec({"search": {"text": "x"}, "start": 0})
    assert not is_range_spec("3.2")


# ------------------------------------------------------ envelope integration


def test_ambiguity_surfaces_matches_through_envelope(doc):
    """Phase 2 contract: envelope.refusal on the resolver's AmbiguousTarget
    emits the Section 6.2 shape (code, matches list, standard hint)."""
    with pytest.raises(AmbiguousTarget) as ei:
        resolve_location(doc, {"after_heading": {"text": "Discussion"}})
    out = envelope.refusal(ei.value)
    assert out["ok"] is False
    assert out["error"]["code"] == "AMBIGUOUS_LOCATION"
    assert [m["paragraph"] for m in out["error"]["matches"]] == [2, 5]
    assert out["error"]["hint"]


def test_range_out_of_bounds_code(doc):
    with pytest.raises(RangeOutOfBounds) as ei:
        resolve_range(doc, {"start": {"paragraph": 6}, "end": {"paragraph": 2}})
    out = envelope.refusal(ei.value)
    assert out["error"]["code"] == "RANGE_OUT_OF_BOUNDS"
    assert out["error"]["code"] in envelope.CLOSED_CODES


def test_stale_anchor_code_prestaged():
    out = envelope.refusal(StaleAnchor("anchor a3f9 no longer resolves"))
    assert out["error"]["code"] == "STALE_ANCHOR"
    assert "re-run" in out["error"]["hint"]


def test_not_found_code_through_envelope(doc):
    with pytest.raises(TargetNotFound) as ei:
        resolve_location(doc, {"paragraph": 10})
    out = envelope.refusal(ei.value)
    assert out["error"]["code"] == "NOT_FOUND"


def test_cursor_refusal_maps_to_app_not_running(doc):
    with pytest.raises(WordNotRunning) as ei:
        resolve_location(doc, {"cursor": True})
    out = envelope.refusal(ei.value)
    assert out["error"]["code"] == "APP_NOT_RUNNING"


def test_selector_vocabulary_is_the_design_list():
    assert SELECTORS == (
        "paragraph",
        "after_heading",
        "outline",
        "bookmark",
        "search",
        "anchor",
        "cursor",
    )
    assert POSITIONS == ("before", "after", "replace", "start", "end")
