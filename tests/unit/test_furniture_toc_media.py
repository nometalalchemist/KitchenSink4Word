"""Phase 5 gate: headers/footers, sections, page numbers, images, TOC."""

import shutil
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import furniture as f, media, read, toc

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


def make_png(path: Path, w=80, h=40, rgb=(200, 30, 30)):
    """Minimal valid PNG without external deps."""
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(
            ">I", zlib.crc32(c) & 0xFFFFFFFF
        )

    raw = b"".join(
        b"\x00" + bytes(rgb) * w for _ in range(h)
    )
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    path.write_bytes(png)
    return path


@pytest.fixture
def ch13(tmp_path):
    dst = tmp_path / "ch1-3.docx"
    shutil.copy(CORPUS / "ch1-3.docx", dst)
    return dst


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


@pytest.fixture
def unitar(tmp_path):
    dst = tmp_path / "unitar.docx"
    shutil.copy(CORPUS / "unitar.docx", dst)
    return dst


# ------------------------------------------------------------------------ TOC


def test_real_word_toc_detected_and_deletable(ch13):
    """ch1-3 carries a genuine Word-created TOC: detect it, read its cached
    entries, delete it cleanly."""
    pkg = DocxPackage(ch13)
    info = toc.read_toc(pkg)
    assert info["present"]
    assert "TOC" in info["instruction"]
    assert info["cached_entries"], "a real updated TOC has cached entries"
    n_tocs = len(info["tocs"])

    toc.delete_toc(pkg, which=0)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch13)
    info2 = toc.read_toc(pkg2)
    assert len(info2["tocs"]) == n_tocs - 1
    xml = pkg2.raw_part("word/document.xml").decode("utf-8")
    assert xml.count('w:fldCharType="begin"') == xml.count('w:fldCharType="end"')
    Document(str(ch13))


def test_insert_read_delete_toc(ch4):
    pkg = DocxPackage(ch4)
    assert not toc.read_toc(pkg)["present"]
    result = toc.insert_toc(pkg, at_start=True, update_on_open=True)
    pkg.save(do_backup=False)
    assert result["toc_inserted"]

    pkg2 = DocxPackage(ch4)
    info = toc.read_toc(pkg2)
    assert info["present"]
    assert "TOC" in info["instruction"]
    settings = pkg2.raw_part("word/settings.xml").decode("utf-8")
    assert "updateFields" in settings
    xml = pkg2.raw_part("word/document.xml").decode("utf-8")
    assert xml.count('w:fldCharType="begin"') == xml.count('w:fldCharType="end"')
    Document(str(ch4))

    toc.delete_toc(pkg2)
    pkg2.save(do_backup=False)
    pkg3 = DocxPackage(ch4)
    assert not toc.read_toc(pkg3)["present"]
    Document(str(ch4))


def test_double_toc_refused(ch4):
    pkg = DocxPackage(ch4)
    toc.insert_toc(pkg, at_start=True, update_on_open=False)
    with pytest.raises(WordMcpError):
        toc.insert_toc(pkg, at_start=True)


def test_toc_styles_injected(ch4):
    """ch4 lacks TOC styles entirely; insertion must create them."""
    pkg = DocxPackage(ch4)
    toc.insert_toc(pkg, at_start=True, update_on_open=False)
    pkg.save(do_backup=False)
    styles = DocxPackage(ch4).raw_part("word/styles.xml").decode("utf-8")
    for sid in ("TOCHeading", "TOC1", "TOC2", "TOC3"):
        assert sid in styles


def test_delete_paragraphs_refuses_cutting_toc(ch4):
    from word_mcp.ops import text as t

    pkg = DocxPackage(ch4)
    toc.insert_toc(pkg, at_start=True, update_on_open=False)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch4)
    # The TOC SDT sits before paragraph 0; its inner paragraphs are not body
    # paragraphs, so instead simulate: bare-field TOC would be paragraphs.
    # Build a bare complex field across paragraphs at the end.
    from lxml import etree

    body = pkg2.body()
    sect = body.find(qn("w:sectPr"))
    p1 = etree.Element(qn("w:p"))
    r1 = etree.SubElement(p1, qn("w:r"))
    etree.SubElement(r1, qn("w:fldChar")).set(qn("w:fldCharType"), "begin")
    p2 = etree.Element(qn("w:p"))
    r2 = etree.SubElement(p2, qn("w:r"))
    etree.SubElement(r2, qn("w:fldChar")).set(qn("w:fldCharType"), "end")
    sect.addprevious(p1)
    sect.addprevious(p2)
    pkg2.mark_dirty()
    n = read.get_document_info(pkg2)["paragraphs"]
    with pytest.raises(WordMcpError):
        t.delete_paragraphs(pkg2, n - 2)  # would cut the field in half
    t.delete_paragraphs(pkg2, n - 2, n - 1)  # whole field: fine


# ------------------------------------------------------------ headers/footers


def test_set_header_and_footer(ch4):
    pkg = DocxPackage(ch4)
    f.set_header_footer(pkg, "header", "The Delta Model — Chapter 4")
    f.set_header_footer(pkg, "footer", "", include_page_number=True)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(ch4)
    hf = f.get_headers_footers(pkg2)
    assert any(
        "Delta Model" in h["text"] for h in hf["headers"]
    ), hf["headers"]
    assert any(x["has_page_number_field"] for x in hf["footers"])
    Document(str(ch4))


def test_first_page_header(ch4):
    pkg = DocxPackage(ch4)
    f.set_header_footer(pkg, "header", "First page only", ref_type="first")
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch4)
    xml = pkg2.raw_part("word/document.xml").decode("utf-8")
    assert "titlePg" in xml
    sections = f.list_sections(pkg2)
    assert any(
        h["type"] == "first" for s in sections for h in s["headers"]
    )


def test_add_page_numbers_with_start(ch4):
    pkg = DocxPackage(ch4)
    f.add_page_numbers(pkg, prefix="Page ", start_at=7)
    pkg.save(do_backup=False)
    xml = DocxPackage(ch4).raw_part("word/document.xml").decode("utf-8")
    assert 'w:start="7"' in xml


# ------------------------------------------------------------------- sections


def test_list_sections_real_doc(ch13):
    sections = f.list_sections(DocxPackage(ch13))
    assert len(sections) >= 1
    assert "page_width_pt" in sections[0]


def test_set_section_properties(ch4):
    pkg = DocxPackage(ch4)
    f.set_section_properties(
        pkg,
        orientation="landscape",
        margins_pt={"left": 90, "right": 90},
    )
    pkg.save(do_backup=False)
    s = f.list_sections(DocxPackage(ch4))[0]
    assert s["orientation"] == "landscape"
    assert s["page_width_pt"] > s["page_height_pt"]
    assert s["margins_pt"]["left"] == 90
    Document(str(ch4))


def test_add_section_break(ch4):
    pkg = DocxPackage(ch4)
    n_before = len(f.list_sections(pkg))
    f.add_section_break(pkg, after_index=10, break_type="nextPage")
    pkg.save(do_backup=False)
    assert len(f.list_sections(DocxPackage(ch4))) == n_before + 1
    Document(str(ch4))


# --------------------------------------------------------------------- images


def test_add_list_resize_image(ch4, tmp_path):
    png = make_png(tmp_path / "test.png")
    pkg = DocxPackage(ch4)
    result = media.add_image(pkg, str(png), at_end=True, width_pt=200)
    pkg.save(do_backup=False)
    assert result["width_pt"] == 200

    pkg2 = DocxPackage(ch4)
    imgs = media.list_images(pkg2)
    assert len(imgs) == 1
    assert imgs[0]["width_pt"] == 200
    # Aspect 80x40 -> height 100.
    assert imgs[0]["height_pt"] == 100

    media.resize_image(pkg2, 0, width_pt=100)
    pkg2.save(do_backup=False)
    imgs = media.list_images(DocxPackage(ch4))
    assert imgs[0]["width_pt"] == 100
    assert imgs[0]["height_pt"] == 50
    Document(str(ch4))


def test_list_images_real_doc(unitar):
    imgs = media.list_images(DocxPackage(unitar))
    assert len(imgs) == 2
    assert all(i["target"] for i in imgs)


def test_replace_image_same_type_guard(unitar, tmp_path):
    png = make_png(tmp_path / "new.png")
    pkg = DocxPackage(unitar)
    imgs = media.list_images(pkg)
    target_ext = imgs[0]["target"].rsplit(".", 1)[1].lower()
    if target_ext == "png":
        media.replace_image(pkg, 0, str(png))
        pkg.save(do_backup=False)
        Document(str(unitar))
    else:
        with pytest.raises(WordMcpError):
            media.replace_image(pkg, 0, str(png))


def test_unsupported_image_type(ch4, tmp_path):
    bad = tmp_path / "x.svg"
    bad.write_text("<svg/>")
    pkg = DocxPackage(ch4)
    with pytest.raises(WordMcpError):
        media.add_image(pkg, str(bad), at_end=True)
