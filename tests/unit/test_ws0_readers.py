"""Workstream 0 (v1.6): empty-document indexing fix + reader upgrades.

Covers the SES resume-build feedback items:
1. index 0 addresses the implicit paragraph of a fresh create_document doc
   (python-docx bodies hold only the trailing sectPr; Word still displays
   one empty paragraph) — insert/delete/replace/format all work at index 0;
3. get_paragraph_format reads effective paragraph formatting with
   inheritance-source attribution mirroring find_formatted's matched_via;
4. get_styles reports each style's explicit paragraph/character formatting
   in define_style's input shape (read-one-define-one template cloning);
5. set_section_properties returns full section state in every response;
6. get_text always reports the effective style, including Normal.
"""

from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import (
    furniture as fu,
    read as rd,
    structure as sx,
    stylefind as sf,
    text as tx,
)


@pytest.fixture
def fresh(tmp_path):
    """A document exactly as create_document produces it: python-docx blank,
    whose body holds no w:p at all — only the trailing sectPr."""
    path = tmp_path / "fresh.docx"
    Document().save(str(path))
    pkg = DocxPackage(path)
    assert pkg.body().find(qn("w:p")) is None  # precondition of the bug
    return pkg


@pytest.fixture
def populated(tmp_path):
    path = tmp_path / "populated.docx"
    doc = Document()
    for text in ("Alpha", "Bravo", "Charlie"):
        doc.add_paragraph(text)
    doc.save(str(path))
    return DocxPackage(path)


# ------------------------------------------------- item 1: empty-doc indexing


def test_fresh_doc_insert_before_index_0(fresh):
    r = tx.insert_paragraphs(
        fresh, [{"text": "One"}, {"text": "Two"}], before_index=0
    )
    assert r["inserted"] == 2
    texts = [p["text"] for p in rd.get_paragraphs(fresh)]
    # The implicit empty paragraph stays after the insertion (Word behavior).
    assert texts == ["One", "Two", ""]


def test_fresh_doc_insert_after_index_0(fresh):
    tx.insert_paragraphs(fresh, [{"text": "Tail"}], after_index=0)
    texts = [p["text"] for p in rd.get_paragraphs(fresh)]
    assert texts == ["", "Tail"]


def test_fresh_doc_delete_index_0(fresh):
    r = tx.delete_paragraphs(fresh, 0)
    assert r["deleted"] == 1
    assert "empty paragraph" in r["note"]
    # The document still shows one (fresh) empty paragraph.
    assert [p["text"] for p in rd.get_paragraphs(fresh)] == [""]


def test_fresh_doc_replace_paragraph_0(fresh):
    tx.replace_paragraph_text(fresh, 0, "Filled in")
    assert [p["text"] for p in rd.get_paragraphs(fresh)] == ["Filled in"]


def test_fresh_doc_set_paragraph_format_0(fresh):
    tx.set_paragraph_format(fresh, [0], {"alignment": "center"})
    got = sf.get_paragraph_format(fresh, 0)["paragraphs"][0]
    assert got["format"]["alignment"] == {
        "value": "center", "source": "explicit",
    }


def test_fresh_doc_saves_valid_package(fresh, tmp_path):
    """The materialized paragraph survives the validated save round-trip."""
    tx.insert_paragraphs(fresh, [{"text": "Persisted"}], before_index=0)
    fresh.save(do_backup=False)
    again = DocxPackage(fresh.path)
    assert [p["text"] for p in rd.get_paragraphs(again)] == ["Persisted", ""]


def test_delete_all_leaves_one_empty_paragraph(populated):
    r = tx.delete_paragraphs(populated, 0, 2)
    assert r["deleted"] == 3 and "empty paragraph" in r["note"]
    assert [p["text"] for p in rd.get_paragraphs(populated)] == [""]


def test_out_of_range_error_names_valid_indices(populated):
    with pytest.raises(TargetNotFound, match=r"valid indices 0-2"):
        tx.replace_paragraph_text(populated, 7, "x")


# --------------------------------------- item 3: get_paragraph_format reader


def test_paragraph_format_sources(populated):
    sx.define_style(
        populated,
        style_id="Spaced",
        name="Spaced",
        paragraph_formatting={
            "alignment": "justify",
            "space_after_pt": 6,
            "line_spacing": 1.15,
            "indent_left_pt": 18,
            "first_line_indent_pt": -18,
            "keep_with_next": True,
        },
    )
    # Apply the style to paragraph 1, with an explicit override on top.
    p1 = [el for k, i, el in rd.body_items(populated) if k == "paragraph"][1]
    from lxml import etree

    ppr = etree.Element(qn("w:pPr"))
    p1.insert(0, ppr)
    etree.SubElement(ppr, qn("w:pStyle")).set(qn("w:val"), "Spaced")
    tx.set_paragraph_format(populated, [1], {"space_after_pt": 12})

    fmt = sf.get_paragraph_format(populated, 1)["paragraphs"][0]
    assert fmt["style"] == "Spaced"
    f = fmt["format"]
    assert f["space_after_pt"] == {"value": 12, "source": "explicit"}
    assert f["alignment"] == {"value": "justify", "source": "paragraph_style"}
    assert f["line_spacing"] == {
        "value": 1.15, "rule": "auto", "source": "paragraph_style",
    }
    assert f["indent_left_pt"]["source"] == "paragraph_style"
    assert f["first_line_indent_pt"]["value"] == -18
    assert f["keep_with_next"] == {"value": True, "source": "paragraph_style"}
    # Nothing defines these anywhere -> Word baseline, attributed as such.
    assert f["page_break_before"] == {"value": False, "source": "word_default"}
    assert f["widow_control"]["value"] is True


def test_paragraph_format_range_and_defaults(populated):
    out = sf.get_paragraph_format(populated, 0, 2)
    assert [p["index"] for p in out["paragraphs"]] == [0, 1, 2]
    for p in out["paragraphs"]:
        assert set(p["format"]) == {
            "alignment", "space_before_pt", "space_after_pt", "line_spacing",
            "indent_left_pt", "indent_right_pt", "first_line_indent_pt",
            "keep_with_next", "widow_control", "page_break_before",
            "outline_level",  # added in the v1.6 sweep (Feature Gap #2)
        }
        assert p["style"]  # effective style always reported
        for prop in p["format"].values():
            assert prop["source"] in (
                "explicit", "paragraph_style", "document_defaults",
                "word_default",
                "none",  # outline_level only: no baseline exists
            )


def test_paragraph_format_errors(populated):
    with pytest.raises(WordMcpError, match=r"valid indices 0-2"):
        sf.get_paragraph_format(populated, 0, 99)
    with pytest.raises(WordMcpError, match="start must be >= 0"):
        sf.get_paragraph_format(populated, -1)
    with pytest.raises(WordMcpError, match="end must be >= start"):
        sf.get_paragraph_format(populated, 2, 1)


# ------------------------------------ item 4: get_styles formatting objects


def test_get_styles_reports_explicit_formatting(populated):
    sx.define_style(
        populated,
        style_id="CloneMe",
        name="Clone Me",
        character_formatting={
            "font": "Times New Roman", "size_pt": 10, "bold": True,
            "color": "1F4E79",
        },
        paragraph_formatting={
            "alignment": "justify", "space_before_pt": 3, "space_after_pt": 6,
            "line_spacing": 1.15, "indent_left_pt": 18,
            "first_line_indent_pt": -18, "keep_with_next": True,
            "widow_control": False,
        },
    )
    styles = {s["id"]: s for s in rd.list_styles(populated)}
    got = styles["CloneMe"]
    assert got["paragraph_formatting"] == {
        "alignment": "justify", "space_before_pt": 3, "space_after_pt": 6,
        "line_spacing": 1.15, "indent_left_pt": 18,
        "first_line_indent_pt": -18, "keep_with_next": True,
        "widow_control": False,
    }
    assert got["character_formatting"] == {
        "font": "Times New Roman", "size_pt": 10, "bold": True,
        "color": "1F4E79",
    }
    # Read-one-define-one: the reported shapes feed define_style unchanged.
    sx.define_style(
        populated,
        style_id="CloneOfClone",
        name="Clone Of Clone",
        character_formatting=got["character_formatting"],
        paragraph_formatting=got["paragraph_formatting"],
    )
    styles = {s["id"]: s for s in rd.list_styles(populated)}
    assert (
        styles["CloneOfClone"]["paragraph_formatting"]
        == got["paragraph_formatting"]
    )
    assert (
        styles["CloneOfClone"]["character_formatting"]
        == got["character_formatting"]
    )


def test_get_styles_no_synthesized_inheritance(populated):
    """A style defining nothing of its own carries no formatting objects,
    even though it inherits plenty via based_on."""
    sx.define_style(
        populated, style_id="BareChild", name="Bare Child", based_on="Normal"
    )
    styles = {s["id"]: s for s in rd.list_styles(populated)}
    assert "paragraph_formatting" not in styles["BareChild"]
    assert "character_formatting" not in styles["BareChild"]
    assert styles["BareChild"]["based_on"] == "Normal"


# --------------------------------- item 5: set_section_properties full state


def test_set_section_properties_reads_state_without_changes(populated):
    r = fu.set_section_properties(populated, section=0)
    assert r["changed"] == []
    state = r["state"]
    assert state["page_width_pt"] == 612.0
    assert state["page_height_pt"] == 792.0
    assert state["orientation"] == "portrait"
    assert set(state["margins_pt"]) == {
        "top", "bottom", "left", "right", "header", "footer", "gutter",
    }


def test_set_section_properties_returns_post_change_state(populated):
    r = fu.set_section_properties(
        populated, section=0, orientation="landscape", margins_pt={"top": 36}
    )
    assert set(r["changed"]) == {"orientation", "margins"}
    assert r["state"]["orientation"] == "landscape"
    assert r["state"]["page_width_pt"] == 792.0  # swapped with height
    assert r["state"]["margins_pt"]["top"] == 36.0


def test_set_section_properties_out_of_range(populated):
    with pytest.raises(TargetNotFound, match=r"valid indices 0-0"):
        fu.set_section_properties(populated, section=5)


# ----------------------------------- item 6: get_text always reports style


def test_get_text_always_includes_style(populated):
    sx.define_style(populated, style_id="Special", name="Special")
    p0 = [el for k, i, el in rd.body_items(populated) if k == "paragraph"][0]
    from lxml import etree

    ppr = etree.Element(qn("w:pPr"))
    p0.insert(0, ppr)
    etree.SubElement(ppr, qn("w:pStyle")).set(qn("w:val"), "Special")

    paras = rd.get_paragraphs(populated)
    assert all("style" in p for p in paras)
    assert paras[0]["style"] == "Special"
    assert paras[1]["style"] == "Normal"  # default, no longer omitted


def test_get_text_style_on_fresh_docs_and_slices(populated):
    sliced = rd.get_paragraphs(populated, 1, 3)
    assert all(p["style"] == "Normal" for p in sliced)


# ---------------------- format_text response metadata (lxml proxy keepalive)


def test_format_text_sequential_indices_stay_correct(tmp_path):
    """Repeated sequential format_text calls must report the TRUE paragraph
    index for body matches and 'table cell' only for real table-cell matches.
    Regression: the find-only path keyed its index map on id() of ephemeral
    lxml proxies with no keepalive, so recycled proxy addresses produced
    off-by-N indices and spurious 'table cell' locations (the root cause
    behind the 'parallel format_text returns wrong paragraph index' field
    report — broken even without concurrency)."""
    path = tmp_path / "fmt.docx"
    doc = Document()
    n = 12
    for i in range(n):
        doc.add_paragraph(f"unique marker {i} alpha")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "cell target text"
    doc.save(str(path))

    import word_mcp.server as srv

    for i in range(n):
        r = srv.format_text(
            str(path), {"bold": True}, find=f"unique marker {i} alpha",
            live="off", backup=False,
        )
        loc = r["formatted"]
        assert loc.get("paragraph") == i, (
            f"call {i}: reported {loc}, expected paragraph {i}"
        )
        assert "location" not in loc
    r = srv.format_text(
        str(path), {"italic": True}, find="cell target text",
        live="off", backup=False,
    )
    assert r["formatted"].get("location") == "table cell"
