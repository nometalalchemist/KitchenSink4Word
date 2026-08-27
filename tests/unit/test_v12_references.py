"""v1.2 Phase A: bibliography system, index, watermark, protection."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound, WordMcpError
from word_mcp.core.package import DocxPackage, qn
from word_mcp.ops import (
    bibliography as bib,
    fields as fl,
    furniture as fu,
    protection as pr,
    read,
)

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def doc(tmp_path):
    dst = tmp_path / "d.docx"
    d = Document()
    d.add_paragraph(
        "Legitimacy shapes compliance in international politics, and alliance "
        "theory follows the same dynamics across cases."
    )
    d.add_paragraph("A second paragraph discusses power transition here.")
    d.save(str(dst))
    return dst


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


# ---------------------------------------------------------------- bibliography


def test_add_source_creates_store_infrastructure(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg,
        tag="Hur99",
        source_type="JournalArticle",
        title="Legitimacy and Authority in International Politics",
        year="1999",
        authors=[{"last": "Hurd", "first": "Ian"}],
        journal_name="International Organization",
        volume="53",
        issue="2",
        pages="379-408",
    )
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    assert pkg2.has_part("customXml/item1.xml")
    assert pkg2.has_part("customXml/itemProps1.xml")
    assert pkg2.has_part("customXml/_rels/item1.xml.rels")
    ct = pkg2.raw_part("[Content_Types].xml").decode()
    assert "customXmlProperties+xml" in ct
    rels = pkg2.raw_part("word/_rels/document.xml.rels").decode()
    assert "relationships/customXml" in rels
    props = pkg2.raw_part("customXml/itemProps1.xml").decode()
    assert "bibliography" in props
    # python-docx's default template ships a pre-existing bibliography store
    # (legacy "/APA.XSL"); _find_store must REUSE it, not duplicate it.
    item = pkg2.raw_part("customXml/item1.xml").decode()
    assert 'StyleName="APA"' in item
    assert not pkg2.has_part("customXml/item2.xml")

    sources = bib.list_sources(pkg2)
    assert len(sources) == 1
    assert sources[0]["tag"] == "Hur99"
    assert sources[0]["authors"] == ["Hurd"]
    Document(str(doc))


def test_duplicate_tag_refused(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="A1", source_type="Book", title="T", year="2000",
        authors=[{"last": "X"}],
    )
    with pytest.raises(WordMcpError, match="already exists"):
        bib.add_source(
            pkg, tag="A1", source_type="Book", title="T2", year="2001",
            authors=[{"last": "Y"}],
        )


def test_corporate_author_and_bad_type(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="DoD26", source_type="Report", title="Posture Statement",
        year="2026", authors=[{"corporate": "U.S. Department of Defense"}],
        institution="DoD",
    )
    assert bib.list_sources(pkg)[0]["authors"] == ["U.S. Department of Defense"]
    with pytest.raises(WordMcpError, match="source_type"):
        bib.add_source(pkg, tag="Z9", source_type="Podcast", title="x")


def test_citation_and_bibliography_fields(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="Hur99", source_type="JournalArticle",
        title="Legitimacy and Authority", year="1999",
        authors=[{"last": "Hurd", "first": "Ian"}],
        journal_name="International Organization",
    )
    r = bib.insert_citation(
        pkg, tag="Hur99", anchor_text="shapes compliance", pages="381"
    )
    assert r["placeholder"] == "(Hurd, 1999)"
    bib.insert_bibliography(pkg, update_on_open=False)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    xml = pkg2.raw_part("word/document.xml").decode()
    assert "CITATION Hur99" in xml
    assert "\\p 381" in xml
    assert "<w:citation/>" in xml  # sdt marker
    assert "BIBLIOGRAPHY" in xml
    assert xml.count('w:fldCharType="begin"') == xml.count(
        'w:fldCharType="end"'
    )
    styles = pkg2.raw_part("word/styles.xml").decode()
    assert "Bibliography" in styles
    Document(str(doc))


def test_citation_unknown_tag_refused(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(TargetNotFound, match="add_source"):
        bib.insert_citation(pkg, tag="Nope", anchor_text="Legitimacy")


def test_delete_source_cited_guard(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="K1", source_type="Book", title="T", year="1990",
        authors=[{"last": "Kelman"}],
    )
    bib.insert_citation(pkg, tag="K1", anchor_text="alliance theory")
    with pytest.raises(WordMcpError, match="cited"):
        bib.delete_source(pkg, "K1")
    bib.delete_source(pkg, "K1", force=True)
    assert bib.list_sources(pkg) == []


def test_set_style(doc):
    pkg = DocxPackage(doc)
    bib.add_source(
        pkg, tag="A1", source_type="Book", title="T", year="2000",
        authors=[{"last": "X"}],
    )
    bib.set_bibliography_style(pkg, "Chicago")
    pkg.save(do_backup=False)
    item = DocxPackage(doc).raw_part("customXml/item1.xml").decode()
    assert 'StyleName="Chicago"' in item
    with pytest.raises(WordMcpError):
        bib.set_bibliography_style(pkg, "Klingon Law Review")


# --------------------------------------------------------------------- index


def test_index_entries_and_field(doc):
    pkg = DocxPackage(doc)
    fl.mark_index_entry(
        pkg, anchor_text="Legitimacy", entry="legitimacy", bold_page=True
    )
    fl.mark_index_entry(
        pkg,
        anchor_text="alliance theory",
        entry="alliance",
        subentry="theory of",
    )
    fl.mark_index_entry(
        pkg, anchor_text="power transition", entry="power", see="alliance"
    )
    entries = fl.list_index_entries(pkg)
    assert len(entries) == 3
    assert entries[1]["subentry"] == "theory of"

    fl.insert_index(pkg, columns=2, update_on_open=False)
    pkg.save(do_backup=False)

    pkg2 = DocxPackage(doc)
    xml = pkg2.raw_part("word/document.xml").decode()
    assert 'XE "legitimacy" \\b' in xml
    assert 'XE "alliance:theory of"' in xml
    assert '\\t "See alliance"' in xml
    assert 'INDEX \\c "2"' in xml
    # XE fields have no separate (ground truth).
    assert xml.count('w:fldCharType="separate"') == 1  # only the INDEX field
    assert xml.count('w:fldCharType="begin"') == xml.count(
        'w:fldCharType="end"'
    )
    Document(str(doc))


def test_index_without_entries_refused(doc):
    pkg = DocxPackage(doc)
    with pytest.raises(WordMcpError, match="mark_index_entry"):
        fl.insert_index(pkg)


def test_index_entry_colon_escaping(doc):
    pkg = DocxPackage(doc)
    fl.mark_index_entry(
        pkg, anchor_text="second paragraph", entry="ratio 1:2"
    )
    entries = fl.list_index_entries(pkg)
    assert entries[0]["entry"] == "ratio 1:2"
    assert entries[0]["subentry"] is None


# ------------------------------------------------------------------ watermark


def test_watermark_add_and_remove(ch4):
    pkg = DocxPackage(ch4)
    r = fu.add_watermark(pkg, "DRAFT")
    pkg.save(do_backup=False)
    assert r["header_parts"] >= 1

    pkg2 = DocxPackage(ch4)
    found = False
    for name in pkg2.part_names():
        if name.startswith("word/header"):
            raw = pkg2.raw_part(name).decode("utf-8", errors="replace")
            if "PowerPlusWaterMarkObject" in raw:
                assert "_x0000_t136" in raw  # shapetype present
                assert 'string="DRAFT"' in raw
                found = True
    assert found
    Document(str(ch4))

    pkg3 = DocxPackage(ch4)
    fu.remove_watermark(pkg3)
    pkg3.save(do_backup=False)
    pkg4 = DocxPackage(ch4)
    for name in pkg4.part_names():
        if name.startswith("word/header"):
            assert "PowerPlusWaterMarkObject" not in pkg4.raw_part(name).decode(
                "utf-8", errors="replace"
            )


def test_watermark_on_headerless_doc(doc):
    pkg = DocxPackage(doc)
    fu.add_watermark(pkg, "CONFIDENTIAL", diagonal=False)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(doc)
    headers = [n for n in pkg2.part_names() if n.startswith("word/header")]
    assert headers, "header part must be created"
    assert "headerReference" in pkg2.raw_part("word/document.xml").decode()
    Document(str(doc))


def test_remove_watermark_none_found(doc):
    with pytest.raises(TargetNotFound):
        fu.remove_watermark(DocxPackage(doc))


# ----------------------------------------------------------------- protection


def test_protection_hash_matches_word_test_vector():
    """The research ground-truth vector from an actual Word file."""
    import base64

    salt = base64.b64decode("gXSXBV92kNcsLGB5MXcJ7g==")
    h, s = pr.word_protection_hash("test123", salt=salt, spin_count=100000)
    assert h == (
        "HUnRnKPzRdHVVr3eFRv+SwUClrg3769Bw4NuvqHizgoabUzTfacuMVs83iZRtGUKW"
        "ltBOlaqCBtBYGGLISgq4g=="
    )


def test_protection_tracked_changes_mode(ch4):
    pkg = DocxPackage(ch4)
    pr.set_document_protection(pkg, edit="trackedChanges")
    pkg.save(do_backup=False)
    state = pr.get_protection(DocxPackage(ch4))
    assert state["protected"] and state["edit"] == "trackedChanges"
    assert state["password_protected"] is False
    settings = DocxPackage(ch4).raw_part("word/settings.xml").decode()
    assert 'w:edit="trackedChanges"' in settings
    Document(str(ch4))


def test_protection_with_password(ch4):
    pkg = DocxPackage(ch4)
    pr.set_document_protection(pkg, edit="readOnly", password="Delta2026")
    pkg.save(do_backup=False)
    state = pr.get_protection(DocxPackage(ch4))
    assert state["password_protected"] is True
    settings = DocxPackage(ch4).raw_part("word/settings.xml").decode()
    assert 'w:cryptAlgorithmSid="14"' in settings
    assert 'w:cryptSpinCount="100000"' in settings
    assert "w:hash=" in settings and "w:salt=" in settings
    Document(str(ch4))


def test_protection_remove(ch4):
    pkg = DocxPackage(ch4)
    pr.set_document_protection(pkg, edit="comments")
    pr.remove_document_protection(pkg)
    pkg.save(do_backup=False)
    assert pr.get_protection(DocxPackage(ch4))["protected"] is False


def test_protection_bad_mode(ch4):
    with pytest.raises(WordMcpError):
        pr.set_document_protection(DocxPackage(ch4), edit="lockdown")
