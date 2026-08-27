"""Regressions for the v1.2 stress findings (1H/4M/3L, all fixed)."""

import shutil
from pathlib import Path

import pytest
from docx import Document

from word_mcp.core.errors import TargetNotFound
from word_mcp.core.package import DocxPackage
from word_mcp.ops import read, tables as tb, text as tx, toc

CORPUS = Path(__file__).resolve().parents[1] / "corpus"


@pytest.fixture
def ch4(tmp_path):
    dst = tmp_path / "ch4.docx"
    shutil.copy(CORPUS / "ch4.docx", dst)
    return dst


# HIGH: tracked replace on paragraphs with lastRenderedPageBreak runs


def test_tracked_replace_on_real_chapter_with_page_breaks(ch4):
    """ch4's corpus has 48 lastRenderedPageBreak-bearing paragraphs; tracked
    replace on 'alliance' previously failed on the first one."""
    pkg = DocxPackage(ch4)
    result = tx.search_and_replace(
        pkg,
        [{"find": "alliance", "replace": "coalition"}],
        track=True,
        author="Fix Verifier",
    )
    pkg.save(do_backup=False)
    assert result["total"] > 50  # the word is everywhere in ch4
    summary = read.revision_summary(DocxPackage(ch4))
    assert summary["by_author"]["Fix Verifier"] == result["total"] * 2  # del+ins
    Document(str(ch4))


def test_format_text_on_page_break_paragraph(ch4):
    """format_text uses the same splitter; must also work now."""
    pkg = DocxPackage(ch4)
    result = tx.format_text(
        pkg, find="Mutual Defense Treaty", formatting={"bold": True}
    )
    pkg.save(do_backup=False)
    assert result["formatted"]
    Document(str(ch4))


# MEDIUM 1: sort_table column validation


def test_sort_out_of_range_column_refused(ch4, tmp_path):
    dst = tmp_path / "s.docx"
    d = Document()
    d.add_paragraph("x")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tb.create_table(pkg, [["a", "b"], ["c", "d"]], at_end=True)
    with pytest.raises(TargetNotFound, match="out of range"):
        tb.sort_table(pkg, 0, column=9)


# MEDIUM 2: unicode title case


def test_title_case_unicode(tmp_path):
    dst = tmp_path / "u.docx"
    d = Document()
    d.add_paragraph("straße épée ΑΘΗΝΑ 한국어 test")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tx.change_case(pkg, "title", indices=[0])
    pkg.save(do_backup=False)
    text = read.get_paragraphs(DocxPackage(dst))[0]["text"]
    assert text == "Straße Épée Αθηνα 한국어 Test"


# MEDIUM 3: SDT content readable


def test_get_paragraphs_sees_toc_content(ch4):
    pkg = DocxPackage(ch4)
    toc.insert_toc(pkg, at_start=True, update_on_open=False)
    pkg.save(do_backup=False)
    paras = read.get_paragraphs(DocxPackage(ch4))
    sdt_entries = [p for p in paras if p.get("in_sdt")]
    assert sdt_entries, "TOC content must be visible to the read layer"
    assert any("Table of Contents" in p["text"] for p in sdt_entries)
    # Body indices unchanged: indexed entries still start at 0 and are ints.
    indexed = [p for p in paras if p["index"] is not None]
    assert indexed[0]["index"] == 0


def test_oracle_comparison_unaffected_by_sdt(ch4):
    """python-docx skips sdt content; our indexed entries must still align."""
    pkg = DocxPackage(ch4)
    toc.insert_toc(pkg, at_start=True, update_on_open=False)
    pkg.save(do_backup=False)
    pkg2 = DocxPackage(ch4)
    ours = [p for p in read.get_paragraphs(pkg2) if p["index"] is not None]
    theirs = Document(str(ch4)).paragraphs
    assert len(ours) == len(theirs)


# LOW 1: read_toc on caption lists


def test_read_toc_caption_list_clean_instruction(ch4):
    from word_mcp.ops import fields as fl, tables as tb2

    pkg = DocxPackage(ch4)
    tb2.create_table(pkg, [["x"]], at_end=True)
    fl.add_caption(pkg, table_index=0, label="Table", text="Only table")
    toc.insert_caption_list(pkg, label="Table", at_start=True, update_on_open=False)
    pkg.save(do_backup=False)
    info = toc.read_toc(DocxPackage(ch4))
    cap = next(t for t in info["tocs"] if t["kind"] == "caption_list")
    assert "PAGEREF" not in cap["instruction"]
    assert '\\c "Table"' in cap["instruction"]


# LOW 2: numeric sort reporting


def test_numeric_sort_reports_non_numeric(tmp_path):
    dst = tmp_path / "n.docx"
    d = Document()
    d.add_paragraph("x")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    tb.create_table(
        pkg, [["h"], ["10"], ["banana"], ["2"]], at_end=True
    )
    r = tb.sort_table(pkg, 0, column=0, numeric=True)
    assert r.get("non_numeric_cells_sorted_last") == 1


# LOW 3: removing protection from an unprotected doc is an explicit no-op


def test_remove_protection_noop_is_explicit(tmp_path):
    from word_mcp.ops import protection as pr

    dst = tmp_path / "p.docx"
    d = Document()
    d.add_paragraph("x")
    d.save(str(dst))
    pkg = DocxPackage(dst)
    r = pr.remove_document_protection(pkg)
    assert r["protection_removed"] is False
    assert "nothing changed" in r["note"]
    pr.set_document_protection(pkg, edit="readOnly")
    r2 = pr.remove_document_protection(pkg)
    assert r2["protection_removed"] is True
